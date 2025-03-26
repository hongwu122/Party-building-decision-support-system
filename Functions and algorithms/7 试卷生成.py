from openpyxl import Workbook
from openpyxl import load_workbook
from openpyxl.styles import Border, Side, Font #设置字体和边框需要的模块
from docx import Document
from docx.shared import Pt
from docx.shared import Inches
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH #段落居中
import random
import os
import sys
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT  # 导入段落对齐
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT # 导入表格对齐方式
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT # 导入单元格垂直对齐
from docx.shared import Cm, Inches, Pt, RGBColor # Cm厘米，Pt磅，Inches英寸 # 导入单位转换函数
import openpyxl
from docxtpl import DocxTemplate
from docx.enum.style import WD_STYLE_TYPE


'''
经济管理与法学学院{2021}年第{八}期入党积极分子培训班
结业考试
说明：本试卷共五大题，39小题，满分100分。考试时长：100分钟。
一、单选题（共20小题；每小题1分，满分20分，每小题只有一个选项符合题意，请把正确答案填入下列表格中）
题号	1	2	3	4	5	6	7	8	9	10	总分
得分											
题号	11	12	13	14	15	16	17	18	19	20	
得分											

二、判断题（共10小题；每小题1分，满分10分；正确的打“√”，错误的打“×”，，请把正确答案填入下列表格中）
题号	21	22	23	24	25	26	27	28	29	30	总分
得分											

三、填空题（共5题；每空1分，满分15分）

四、简答题（共4小题；满分25分）

五、论述开放题（共1小题；满分30分）

正文          宋体小四      Normal
标题1         宋体小三      Heading 1
标题2         黑体小四      Heading 2
说明字体       宋体五号     Title
'''

filepath = 'mould\题库.xlsx'
year = 2021
qishu = '九'
danxuan_num = 20
panduan_num = 10
duoxuan_num = 0
tiankong_num = 5
jianda_num = 3
lunsu_num = 1
sum_num = danxuan_num+panduan_num+duoxuan_num+tiankong_num+jianda_num+lunsu_num


# 有时候我们希望读取到公式计算出来的结果，可以使用load_workbook()中的data_only属性, data_only=True
workbook = load_workbook(filepath)

danxuan_sheet = workbook.worksheets[0]
panduan_sheet = workbook.worksheets[1]
duoxuan_sheet = workbook.worksheets[2]
tiankong_sheet = workbook.worksheets[3]
jianda_sheet = workbook.worksheets[4]
lunsu_sheet = workbook.worksheets[5]

# 打开文档
test_doc = Document('mould\模板 试卷.docx')
answer_doc = Document()

for i in test_doc.styles:
    if i.type == WD_STYLE_TYPE.PARAGRAPH:
        print(i.name)

a = '经济管理与法学学院{}年第{}期入党积极分子培训班'.format(year,qishu)
b = '结业考试'
c = '说明：本试卷共五大题，{}小题，满分100分。考试时长：100分钟。'.format(sum_num)
test_doc.add_paragraph(a, style='Heading 1') # 文章标题
test_doc.add_paragraph(b, style='Heading 1') # 文章标题
test_doc.add_paragraph(c, style='Title') # 说明
# test_doc.styles['Normal'].paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE # 1.5倍行距
answer_doc.styles['Normal'].font.name = 'Times New Roman'  # 设置英文字体
answer_doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')  # 设置中文字体
questions_num = 1 # 题号

d = '一、单选题（共{}小题；每小题1分，满分{}分，每小题只有一个选项符合题意，请把正确答案填入下列表格中）'.format(danxuan_num,danxuan_num)
test_doc.add_paragraph(d, style='Heading 2') # 标题
answer_doc.add_paragraph(d)  # 答案写入标题
# 增加一个表格
# test_doc.add_table(rows=danxuan_num//10*2 , cols=12)
# for c in test_doc.tables[0].columns[0].cells:
#     c.width = Cm(1.5)
# for c in test_doc.tables[0].columns[len(test_doc.tables[0].columns)-1].cells:
#     c.width = Cm(1.5)
# for c in range(1,len(test_doc.tables[0].rows)+1):
#     if c%2 == 0:
#         test_doc.tables[0].cell(c-1,0).text = '答案'
#         test_doc.tables[0].cell(c-1,0).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
#         test_doc.tables[0].cell(c-1,0).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
#
#     else:
#         test_doc.tables[0].cell(c-1,0).text = '题号'
#         test_doc.tables[0].cell(c-1,0).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
#         test_doc.tables[0].cell(c-1,0).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
#
# test_doc.tables[0].cell(0,len(test_doc.tables[0].columns)-1).text = '总分'
# test_doc.tables[0].cell(0,len(test_doc.tables[0].columns)-1).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
# test_doc.tables[0].cell(0,len(test_doc.tables[0].columns)-1).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
# tihao_num = 1
# for r in range(1,len(test_doc.tables[0].rows)+1):
#     if r%2 != 0:
#         for c in range(2, len(test_doc.tables[0].columns)):
#             test_doc.tables[0].cell(r-1,c-1).text = str(tihao_num)
#             test_doc.tables[0].cell(r-1,c-1).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
#             test_doc.tables[0].cell(r-1,c-1).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
#             tihao_num += 1
danxuan_set = set()
answer_danxuan_para = ''
print(danxuan_sheet.max_row - 1)
while (len(danxuan_set) < danxuan_num):
    i = random.randint(2, danxuan_sheet.max_row)
    if i not in danxuan_set:
        danxuan_set.add(i)
        danxuan_para = str(questions_num) + '、' + danxuan_sheet.cell(row=i , column=4).value # 问题题目
        danxuan_paragraph = test_doc.add_paragraph(danxuan_para)  # 写入题目
        # run = danxuan_paragraph.add_run(danxuan_para).bold = True # 单句加粗
        A = str(danxuan_sheet.cell(row=i , column=5).value)
        B = str(danxuan_sheet.cell(row=i , column=6).value)
        C = str(danxuan_sheet.cell(row=i , column=7).value)
        D = str(danxuan_sheet.cell(row=i , column=8).value)
        # print(A,C,B,D,len(A)+len(B)+len(C)+len(D))
        if (len(A)+len(B)+len(C)+len(D))<=20: # 1行
            danxuan_temp = '   A、%s   B、%s   C、%s   D、%s\n' % (A,B,C,D) # 选项
        elif (len(A)+len(B))<= 26 and (len(C)+len(D))<=26: # 2行
            danxuan_temp = '   A、%s   B、%s\n   C、%s   D、%s\n' % (A,B,C,D) # 选项
        elif (len(A)+len(B))>26 or (len(C)+len(D))>26 :# 4行
            danxuan_temp = '   A、%s\n   B、%s\n   C、%s\n   D、%s\n' % (A,B,C,D) # 选项 四行
        danxuan_paragraph = test_doc.add_paragraph(danxuan_temp) # 写入选项
        answer_danxuan_para += str(questions_num) + '、' + danxuan_sheet.cell(row=i , column=9).value + '        '# 答案
        if len(danxuan_set)%5==0:
            answer_doc.add_paragraph(answer_danxuan_para)  # 写入答案
            answer_danxuan_para=''
        questions_num += 1
    else: pass

e = '二、判断题（共{}小题；每小题1分，满分{}分；正确的打“√”，错误的打“×”，，请把正确答案填入下列表格中）'.format(panduan_num,panduan_num)
test_doc.add_paragraph(e, style='Heading 2') # 标题
answer_doc.add_paragraph(e)  # 答案写入标题
panduan_set = set()
answer_panduan_para = ''
while (len(panduan_set) < panduan_num):
    i = random.randint(2, panduan_sheet.max_row)
    if i not in panduan_set:
        panduan_set.add(i)
        panduan_para = str(questions_num) + '、' + panduan_sheet.cell(row=i, column=4).value # 问题
        panduan_paragraph = test_doc.add_paragraph(panduan_para)
        answer_panduan_para += str(questions_num) + '、' + panduan_sheet.cell(row=i, column=5).value + '        '
        if len(panduan_set)%5==0:
            answer_doc.add_paragraph(answer_panduan_para)  # 写入答案
            answer_panduan_para=''
        questions_num += 1
    else: pass

f = '三、填空题（共{}题；每空1分，满分{}分）'.format(tiankong_num,10)
test_doc.add_paragraph(f, style='Heading 2') # 标题
answer_doc.add_paragraph(e)  # 答案写入标题
tiankong_set = set()
answer_tiankong_para = ''
while (len(tiankong_set) < tiankong_num):
    i = random.randint(2, tiankong_sheet.max_row)
    if i not in tiankong_set:
        tiankong_set.add(i)
        tiankong_para = str(questions_num) + '、' + tiankong_sheet.cell(row=i, column=4).value # 问题
        tiankong_paragraph = test_doc.add_paragraph(tiankong_para)
        answer_tiankong_para += str(questions_num) + '、' + tiankong_sheet.cell(row=i, column=5).value + '\n'
        questions_num += 1
    else: pass
answer_doc.add_paragraph(answer_tiankong_para)  # 写入答案

g = '四、简答题（共{}小题；满分{}分）'.format(jianda_num,25)
test_doc.add_paragraph(g, style='Heading 2') # 标题
answer_doc.add_paragraph(e)  # 答案写入标题
jianda_set = set()
answer_jianda_para = ''
while (len(jianda_set) < jianda_num):
    i = random.randint(2, jianda_sheet.max_row)
    if i not in jianda_set:
        jianda_set.add(i)
        jianda_para = str(questions_num) + '、' + jianda_sheet.cell(row=i, column=4).value + '\n\n\n\n\n\n'# 问题
        jianda_paragraph = test_doc.add_paragraph(jianda_para)
        answer_jianda_para += str(questions_num) + '、' + jianda_sheet.cell(row=i, column=5).value + '\n\n'
        questions_num += 1
    else: pass
answer_doc.add_paragraph(answer_jianda_para)  # 写入答案

h = '五、论述开放题（共{}小题；满分{}分）'.format(lunsu_num,30)
test_doc.add_paragraph(h, style='Heading 2') # 标题
answer_doc.add_paragraph(e)  # 答案写入标题
lunsu_set = set()
answer_lunsu_para = ''
while (len(lunsu_set) < lunsu_num):
    i = random.randint(2, lunsu_sheet.max_row)
    if i not in lunsu_set:
        lunsu_set.add(i)
        lunsu_para = str(questions_num) + '、' + lunsu_sheet.cell(row=i, column=4).value # 问题
        lunsu_paragraph = test_doc.add_paragraph(lunsu_para)
        answer_lunsu_para += str(questions_num) + '、' + lunsu_sheet.cell(row=i, column=5).value
        questions_num += 1
    else: pass
answer_doc.add_paragraph(answer_lunsu_para)  # 写入答案






# 保存文件
test_doc.save('测试.docx')
answer_doc.save('测试答案.docx')









# duoxuanti_para = '第二部分  多选题(10题)'
# duo_paragraph = test_doc.add_paragraph(duoxuanti_para)
# answer_doc.add_paragraph(duoxuanti_para)  # 答案
#
# duo_num = 1
# duo_set = set()
# answer_duo_para = ''
# while (len(duo_set) < 10):
#     i = random.randint(1, duoxuan_sheet.max_row - 2)
#     if i not in duo_set:
#         duo_set.add(i)
#         # duo_para = '%s%s%s\%s' % ( duo_num,duanhao, duoxuan_sheet.cell(row=i + 2, column=4).value ,huiche) # 问题
#
#         duo_para = str(duo_num) + '、' + duoxuan_sheet.cell(row=i + 2, column=4).value + '\n'  # 问题
#
#         duo_temp = '   A、%s\n   B、%s\n   C、%s\n   D、%s\n' % (
#             duoxuan_sheet.cell(row=i + 2, column=5).value, duoxuan_sheet.cell(row=i + 2, column=6).value,
#             duoxuan_sheet.cell(row=i + 2, column=7).value, duoxuan_sheet.cell(row=i + 2, column=8).value)
#         duo_para += duo_temp
#
#         if duoxuan_sheet.cell(row=i + 2, column=9).value:
#             duo_para += '   E、' + duoxuan_sheet.cell(row=i + 2, column=9).value + '\n'
#
#         duo_paragraph = test_doc.add_paragraph(duo_para)
#
#         answer_duo_para += str(duo_num) + '、' + duoxuan_sheet.cell(row=i + 2, column=10).value + '   '
#
#         duo_num += 1
#     else:
#         pass
#
# answer_doc.add_paragraph(answer_duo_para)  # 答案



