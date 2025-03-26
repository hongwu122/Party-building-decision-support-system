from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT  # 导入段落对齐
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT # 导入表格对齐方式
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT # 导入单元格垂直对齐
from docx.shared import Cm, Inches, Pt, RGBColor # Cm厘米，Pt磅，Inches英寸 # 导入单位转换函数
import openpyxl
from docxtpl import DocxTemplate
from docx.enum.style import WD_STYLE_TYPE

def write_qingshi(cookie,yeardu,pici,year_up,year,month,day,party_name,party_num,first_people,people_num,people_sheet):
    if cookie==1: # 发展对象的请示
        a = "关于建议接收{}年度经济管理与法学学院".format(yeardu)
        b = "{}发展对象的请示".format(pici)
        c = "尊敬的院党委："
        d = "经过{}等{}个学生党支部委员会充分研究讨论，确认{}等{}名同志为{}年{}半年发展对象人选，建议学院党委将{}等{}名同志列为中共党员发展" \
            "对象，名单如下（排名以班级为序）：".format(party_name,party_num,first_people,people_num,year,year_up,
                                                    first_people,people_num)
        e = "请批示。"
        f = "经济管理与法学学院学生党建工作委员会"
        g = "{}年{}月{}日".format(year,month,day)
    if cookie==2: # 预备党员的请示
        a = "关于建议接收{}年度经济管理与法学学院".format(yeardu)
        b = "{}预备党员的请示".format(pici)
        c = "尊敬的院党委："
        d = "经过{}等{}个学生党支部召开支部大会充分讨论，认为{}等{}名同志符合预备党员的条件。现拟提请学院党委接受" \
            "{}等{}名同志为中共预备党员，名单如下（排名以班级为序）：" .format(party_name,party_num,first_people,people_num,
                                                     first_people,people_num,)
        e = "请批示。"
        f = "经济管理与法学学院学生党建工作委员会"
        g = "{}年{}月{}日".format(year,month,day)
    if cookie==3: # 预备党员转正的请示
        a = "关于建议接收{}年度经济管理与法学学院".format(yeardu)
        b = "{}预备党员转正的请示".format(pici)
        c = "尊敬的院党委："
        d = "经过{}等{}个学生党支部召开支部大会充分讨论，确认{}等{}名同志为{}年{}半年预备党员转正人选，建议学院党委将" \
            "{}等{}名同志列为中共党员，名单如下（排名以班级为序）：" .format(party_name,party_num,first_people,people_num,
                                                     year,year_up,first_people,people_num,)
        e = "请批示。"
        f = "经济管理与法学学院学生党建工作委员会"
        g = "{}年{}月{}日".format(year,month,day)
    doc = Document()
    # 判断人数，来设置表格
    if 0 <= people_num<= 64: # 四号字体
        doc.styles['Normal'].paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE  # 1.5倍行距
        doc.styles['Normal'].font.size = Pt(14)
        col_width = [2.43,1.9]
        row_height = 1
    if 64 < people_num<= 104: # 小四字体
        doc.styles['Normal'].paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE # 1.5倍行距
        doc.styles['Normal'].font.size = Pt(12)
        col_width = [2.15,1.8]
        row_height = 0.9
    if 104 < people_num <= 120:  # 小四字体
        doc.styles['Normal'].paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE  # 1.5倍行距
        doc.styles['Normal'].font.size = Pt(12)
        col_width = [2.15, 1.8]
        row_height = 0.8
    if 120 < people_num<= 168: # 小四字体
        doc.styles['Normal'].paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST # 最小倍倍行距
        doc.styles['Normal'].font.size = Pt(12)
        col_width = [1.98,1.8]
        row_height = 0.55
    if 168 < people_num<= 184: # 五号字体
        doc.styles['Normal'].paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST # 最小倍倍行距
        doc.styles['Normal'].font.size = Pt(10.5)
        col_width = [1.98,1.8]
        row_height = 0.55
    if 184 < people_num:
        doc.styles['Normal'].font.size = Pt(10)
        col_width = [1.98, 1.8]
        row_height = 0.55
        print('自行调整')
    # 标题样式
    doc.styles['Header'].font.name = 'Times New Roman'  # 设置英文字体
    doc.styles['Header']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')  # 设置中文字体
    doc.styles['Header'].font.bold = True  # 加粗
    doc.styles['Header'].font.size = Pt(16)
    doc.styles['Header'].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 居中对齐
    doc.styles['Header'].paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE # 1.5倍行距
    doc.styles['Header'].paragraph_format.space_before = Pt(0)  # 段前
    doc.styles['Header'].paragraph_format.space_after = Pt(0)  # 段后
    # 普通正文央视
    doc.styles['Footer'].font.name = 'Times New Roman'  # 设置英文字体
    doc.styles['Footer']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')  # 设置中文字体
    doc.styles['Footer'].font.size = Pt(14)
    doc.styles['Footer'].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY # 两端对齐
    doc.styles['Footer'].paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE # 1.5倍行距
    doc.styles['Footer'].paragraph_format.space_before = Pt(0)  # 段前
    doc.styles['Footer'].paragraph_format.space_after = Pt(0)  # 段后
    doc.styles['Footer'].paragraph_format.first_line_indent = doc.styles['Footer'].font.size * 2  #首行缩进2字符 1 英寸=2.54 厘米
    # 表格样式
    doc.styles['Normal'].font.name = 'Times New Roman'  # 设置英文字体
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')  # 设置中文字体
    # doc.styles['Normal'].font.size = Pt(12)
    doc.styles['Normal'].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.DISTRIBUTE # 分散对齐
    # doc.styles['Normal'].paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST # 最小倍倍行距
    doc.styles['Normal'].paragraph_format.space_before = Pt(0)  # 段前
    doc.styles['Normal'].paragraph_format.space_after = Pt(0)  # 段后
    doc.styles['Normal'].paragraph_format.first_line_indent = Inches(0)  #首行缩进2字符 1 英寸=2.54 厘米

    # 标题 两段
    doc.add_paragraph(a,style='Header')
    doc.add_paragraph(b,style='Header')
    doc.add_paragraph(c,style='Footer').paragraph_format.first_line_indent=Inches(0) # 1 英寸=2.54 厘米
    doc.add_paragraph(d,style='Footer')

    table = doc.add_table(people_num//8 if people_num%8 == 0 else people_num//8+1, 8)
    table.autofit = True   # if is True 按窗口大小自动调整
    count = 0

    for row in range(len(table.rows)):
        table.rows[row].height = Cm(row_height)  # 调整行高
        for col in range(len(table.columns)):
            # print(行, 列)  # 可以查看表格输出结果
            table.cell(row, col).text = people_sheet[count]    # 写入人名
            # table.cell(行, 列).width = doc.styles['Normal'].font.size * len(people_sheet[count])
            # table.cell(行, 列).height = doc.styles['Normal'].font.size * 5
            table.cell(row, col).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER  # 上下居中（中部居中对齐）
            # table.cell(行, 列).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 水平居中（中部居中对齐）
            count += 1
            if count == people_num:
                break
        if count == people_num:
            break
    # 修正列宽
    for col in range(len(table.columns)):
        maxlist = []
        for r in range(len(table.rows)):
            try:
                maxlist.append(len(people_sheet[8*r + col]))
                # print(people_sheet[8*r + col])
            except:pass
        maxnum = max(maxlist) # 每一列的最大值
        table.cell(len(table.rows)-1, col).width = Cm( col_width[0] if maxnum==4 else col_width[1] ) # 调整列宽 2字:1.3 3字:1.8 4字:2.1
        # 要在最后一行设置列宽度，因为设置前面的，后面一行出现空格，前面设置的宽度就不生效了

    table.alignment = WD_TABLE_ALIGNMENT.CENTER  # 设置整个表格为居中对齐
    # table.autofit = True
    doc.add_paragraph(e,style='Footer')
    doc.add_paragraph("",style='Normal')
    doc.add_paragraph(f,style='Footer').alignment = WD_ALIGN_PARAGRAPH.RIGHT  # 右对齐
    doc.add_paragraph(g,style='Footer').alignment = WD_ALIGN_PARAGRAPH.RIGHT


    doc.save(str(people_num)+a+b+'.docx')



if __name__ == '__main__':
    people_sheet = '''丰一帆 谷铝江 刘澳  孙涛  王恩黎  肖淮予  张可欣  赵威昊  陈慕天  程万里  刘极  欧阳佳
                   王静羽  杨明嘉  刘耀鸿  邓晨茜  刘炎  唐姝瑶  包雪珂  黄群  王炎鑫  王玉  宋燕茹  张雅妮
                   龙芊麦  刘静  欧阳婉婷  孙梓晴  吴永余  朱新婷  周利欢  李萍  于明霞  张灿  黄文静  颜骊静
                   彭涛  熊碗君  贺娜  刘俐俐  刘满霞  谭涛  田雅婷  张瑶  郑诗怡  张盈 吴佳鑫  符金珍  胡俊彤  '''.split()
    print(people_sheet)
    print(len(people_sheet))
    yeardu=2020
    pici='第二批'
    year_up='下'
    year=2020
    month=12
    day=12
    party_name='法学支部'
    party_num=9
    first_people=people_sheet[0]
    people_num= len(people_sheet) # 这个要固定
    cookie = 1

    write_qingshi(cookie,yeardu,pici,year_up,year,month,day,party_name,party_num,first_people,people_num,people_sheet)