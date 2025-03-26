# -*- coding: utf-8 -*-
# author：鸿武
# mali：hongwu2023@163.com
# 微信公众号：努力发芽的小鸿武
# python3默认使用UTF-8格式,开头一般不用加 # -*- coding:utf-8 -*-
# 但是有的中文中还是会出现无法识别的情况，抛出Non-UTF-8 code starting with '\xe7'  错误，这时候就需要在第一行加上这句。
banben = "V2.3.60（ Alpha 版 ）"
from tkinter import messagebox
from tkinter.filedialog import *
'''###############################################################################################'''
# 创建窗口
window = Tk()
# 窗口的大小,前两、个参数是：宽、高，后面的参数是坐标
# # 设置窗口居中
width =  650  # 1300
height =  690  # 830
screenwidth = window.winfo_screenwidth()
screenheight = window.winfo_screenheight()
alignstr = '%dx%d+%d+%d' % (width, height, (screenwidth-width)/2, (screenheight-height-100)/2)
window.geometry(alignstr)
# 禁止窗口的拉伸
# window.resizable(0, 0)  # == (Flase,Flase)
# 窗口置顶
# window.attributes("-topmost", 1)  # 1==True 处于顶层
# 标题
window.title('党建决策支持系统（高校版）')
try:
    window.iconbitmap('mould\data\\logo.ico')
except:pass
# 设置背景颜色
window['background']='gray97'
messagebox.showinfo('温馨提示', '欢迎来到！系统正在初始化，大约需要5秒钟，请稍后...')

'''###############################################################################################'''
import tkinter
from tkinter.ttk import *
from tkinter import ttk
from tkinter import scrolledtext
# # import traceback
# # from ttkbootstrap import Style as ttkStyle
# # import tkintertable
# # from tkinter import *
# from datetime import datetime
import datetime
import os
import random
import time
import platform # 获取系统信息
# from tkintertable import TableCanvas, TableModel  # tkinter展示table
# from collections import Counter # 字典计数
# import win32com.client as win32 # 调用系统权限软件
# from PIL import Image as imm  # 用于tkinter界面展示图片，裁剪修图像用
# from PIL import ImageTk as imk #
# import pygame #用于播放音乐
# from faker import Faker # 随机生成个人信息
# import webbrowser 调用浏览器
# import re # 正则表达式
# import threading # 多线程
# import linecache # 获取一行数据
# import yagmail # 邮件模块

import openpyxl
from openpyxl.styles import Color, Font, Alignment, Border, Side, PatternFill, colors
# from openpyxl.utils import get_column_letter
# from pdf2docx import Converter
from docx import Document
# from docx.oxml import OxmlElement # 设置表格单元格的边框
from docx.oxml.ns import qn

from docx.enum.text import WD_PARAGRAPH_ALIGNMENT  # 导入段落对齐
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT # 导入表格对齐方式
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT # 导入单元格垂直对齐
# # from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Cm, Inches, Pt, RGBColor # Cm厘米，Pt磅，Inches英寸 # 导入单位转换函数
from docx.table import _Cell
# 其他第三方库
try:
    ## 我们对中文进行排序时，下意识的按照拼音排序，也就是字典序，
    # 但sort、sorted()都是参照字符的编码（Unicode）大小，进行排序的，排出来的结果当然是杂乱的。
    from pypinyin import pinyin, Style  # 汉字首字母排序用
except Exception as error:
        messagebox.showinfo('错误提示', '错误信息：\n{}'.format(error))
# # 统计分析要用
# import pandas as pd # 统计分析部分，excel转csv用到
# import numpy as np # 可视化用到
# try:
#     import matplotlib.pyplot as plt
#     # 支持中文
#     plt.rcParams['font.sans-serif'] = ['SimHei']  # 用来正常显示中文标签
#     plt.rcParams['axes.unicode_minus'] = False  # 用来正常显示负号
#     from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg, NavigationToolbar2Tk
#     from matplotlib.backend_bases import key_press_handler
#     from matplotlib.figure import Figure
# except:
#     messagebox.showinfo("提示","缺少matplotlib的配置文件mpl-data，统计功能不可用！")
'''#################################################################################################################'''
# 由于tkinter中没有ToolTip功能，所以自定义这个功能如下
class ToolTip(object):
    def __init__(self, widget):
        self.widget = widget
        self.tipwindow = None
        self.id = None
        self.x = self.y = 0

    def showtip(self, text):
        "Display text in tooltip window"
        self.text = text
        if self.tipwindow or not self.text:
            return
        x, y, _cx, cy = self.widget.bbox("insert")
        x = x + self.widget.winfo_rootx() + 27
        y = y + cy + self.widget.winfo_rooty() + 27
        self.tipwindow = tw = Toplevel(self.widget)
        tw.wm_overrideredirect(1)
        tw.wm_geometry("+%d+%d" % (x, y))

        label = Label(tw, text=self.text, justify=LEFT,
                         background="#ffffe0", relief=SOLID, borderwidth=1,
                         font=("tahoma", "8", "normal"))
        label.pack(ipadx=1)

    def hidetip(self):
        tw = self.tipwindow
        self.tipwindow = None
        if tw:
            tw.destroy()
def createToolTip(widget, text):
    toolTip = ToolTip(widget)

    def enter(event):
        toolTip.showtip(text)

    def leave(event):
        toolTip.hidetip()

    widget.bind('<Enter>', enter)
    widget.bind('<Leave>', leave)
# 定义边框样式
def my_border(t_border, b_border, l_border, r_border):
    border = Border(top=Side(border_style=t_border, color=colors.BLACK),
                    bottom=Side(border_style=b_border, color=colors.BLACK),
                    left=Side(border_style=l_border, color=colors.BLACK),
                    right=Side(border_style=r_border, color=colors.BLACK))
    return border
# 自动整理列宽
def auto_column(workbook_name):
    from openpyxl.utils import get_column_letter
    wb = openpyxl.load_workbook(workbook_name)
    ws = wb.active

    # 第一步：计算每列最大宽度，并存储在列表lks中。

    lks = []  # 英文变量太费劲，用汉语首字拼音代替
    for i in range(1, ws.max_column + 1):  # 每列循环
        lk = 1  # 定义初始列宽，并在每个行循环完成后重置
        for j in range(1, ws.max_row + 1):  # 每行循环
            sz = ws.cell(row=j, column=i).value  # 每个单元格内容
            if isinstance(sz, str):  # 中文占用多个字节，需要分开处理
                lk1 = len(sz.encode('gbk'))  # gbk解码一个中文两字节，utf-8一个中文三字节，gbk合适
            else:
                lk1 = len(str(sz))
            if lk < lk1:
                lk = lk1  # 借助每行循环将最大值存入lk中
            # print(lk)
        lks.append(lk)  # 将每列最大宽度加入列表。（犯了一个错，用lks = lks.append(lk)报错，append会修改列表变量，返回值none，而none不能继续用append方法）

    # 第二步：设置列宽
    for i in range(1, ws.max_column + 1):
        k = get_column_letter(i)  # 将数字转化为列名,26个字母以内也可以用[chr(i).upper() for i in range(97, 123)]，不用导入模块
        ws.column_dimensions[k].width = lks[i - 1] + 2  # 设置列宽，一般加两个字节宽度，可以根据实际情况灵活调整

    wb.close()
    wb.save(workbook_name)
'''#################################################################################################################'''
def xls_to_xlsx(path,sole=True):# 默认单个
    import win32com.client as win32
    try:
        messagebox.showinfo("小提示", "由于您选择的文件或文件夹内包含'xls'的表格文件类型，该类型文件不被本软件读取，\n"
                                   "所以接下来会尝试进行转换操作，转换成可读取的'xlsx'文件类型，如果转换失败，请尝试手动转换!")
        scr_output(scr_10,"小提示：由于您选择的文件或文件夹内包含'xls'的表格文件类型，该类型文件不被本软件读取，\n 所以接下来会尝试进行转换操作，转换成可读取的'xlsx'文件类型，如果转换失败，请尝试手动转换!")
        excel = win32.gencache.EnsureDispatch('Excel.Application')
        if sole == True: # 单个文件
            path = eval(repr(path).replace(r'\\\\', r'/'))  # 把 \\\\ 替换成 /  不然会报错  一根是转义的\
            path = eval(repr(path).replace('/', r'\\'))  # 把 / 替换成 \\  不然会报错，
            wb = excel.Workbooks.Open(path)
            wb.SaveAs(path+"x", FileFormat = 51)    #FileFormat = 51 is for .xlsx extension
            wb.Close()                               #FileFormat = 56 is for .xls extension
            excel.Application.Quit()
        if sole != True: # 整个文件夹
            xls_files = [x for x in os.listdir(path) if os.path.splitext(x)[1] == '.xls']
            for x in xls_files:
                sole_path = path + '/' + x
                # print(sole_path)
                sole_path = eval(repr(sole_path).replace(r'\\\\',r'/')) #把 \\\\ 替换成 /  不然会报错  一根是转义的\
                # print(sole_path)
                sole_path = eval(repr(sole_path).replace('/',r'\\')) #把 / 替换成 \\  不然会报错，
                # 初步认定，win32用win的单个\，其他\\和/不识别。且需要绝对路径
                # print(sole_path)
                wb = excel.Workbooks.Open(sole_path)
                wb.SaveAs(sole_path + "x", FileFormat=51)  # FileFormat = 51 is for .xlsx extension
                wb.Close()  # FileFormat = 56 is for .xls extension
                excel.Application.Quit()
        # print('成功生成xlsx文件')
        scr_output(scr_10,"成功生成xlsx文件!现在，点击开始审核即可！")

    except Exception as error:
            error = str(error)
            print('生成xlsx文件失败')
            print('错误提示', error)
            messagebox.showinfo('错误提示', '尝试把xls文件改成xlsx文件 失败！\n请自行另存为xlsx文件类型。\n错误信息：\n{}'.format(error))
# 通用————汇总表格（暂时只针对xlsx文件格式，xls格式自动转换）
def general_merge_book(path, filename, sheet_what, biaoti_row, biaotou_row, lishi_row):
    from openpyxl.utils import get_column_letter
    # general_merge_book(path, filename, sheet_what=0, biaoti_row=0, biaotou_row=1, lishi_row=False)
    # 默认sheet==0，工作表为第一个工作表
    # 默认大标题行0行，即没有大标题行。  biaoti_row==1，即第一行是大标题，  biaoti_row==2，即前两行是大标题。
    # 默认表头行在在第一行。 biaotou_row==0行，即没有表头， biaotou_row=2在第二行
    # 没有例示行，如果例示行==True，那就例示行默认在表头行下面一行
    # '''这里需要对导入数据进行转换'''
    print(sheet_what, biaoti_row, biaotou_row, lishi_row)
    scr_output(scr_10,'\n\n！！！注意：参数设置错误，很可能导致汇总数据缺失或混乱！！\n\n本次汇总的参数设置情况：'
                      '\n\n选择的工作表：第{}个，\n选择的标题行：{}，\n选择的表头行：{}，\n选择的例示行：{}\n'.format(sheet_what, biaoti_row, biaotou_row, lishi_row))

    sheet_what = int(sheet_what[-1])-1
    if biaoti_row == '无标题':
        biaoti_row = 0
    else:
        biaoti_row = int(biaoti_row[1])
    if biaotou_row == '无表头':
        biaotou_row = 0
    else:
        biaotou_row = int(biaotou_row[1])
    # if lishi_row == '有':
    #     lishi_row = True
    # if lishi_row == '没有':
    #     lishi_row = False

    if path == "":
        messagebox.showinfo("提示", "请输入正确的需汇总文件的文件夹路径！")
        scr_output(scr_10, '\n本次没有正确输入正确的需汇总文件的文件夹路径！ \n请把汇总文件全部放在一个文件夹里面，并选择该文件夹\n\n')

    if path != "":
        try:
            xls_files = [x for x in os.listdir(path) if os.path.splitext(x)[1] == '.xls']
            if xls_files != []:# 说明有xls文件
                xls_to_xlsx(path=path, sole=False) # 给路径，让其自己转换成xlsx的
                scr_output(scr_10, '\n\n检测到有{}个xls格式文件，已经自动在原路径转换生成可读取的xlsx文件类型：\n'.format(len(xls_files)))

            # # print(os.listdir('./test')) # 打印test目录下所有文件
            # print('当前工作路径', os.path.abspath('.'))  # 打印当前目录
            # # xlsx_files = [x for x in os.listdir('./test') if os.path.isfile(x) and os.path.splitext(x)[1] == '.xlsx'] # 罗列当前目录内所有xlsx文件
            xlsx_files = [x for x in os.listdir(path) if os.path.splitext(x)[1] == '.xlsx']  # 罗列当前目录内所有xlsx文件
            scr_output(scr_10, '\n\n需要统计{}个表格'.format(len(xlsx_files)))
            scr_output(scr_10, '\n\n统计表格有：\n{}'.format(xlsx_files))
            # print('需要统计', len(xlsx_files), '个表格')
            # print('统计表格有：\n', xlsx_files)  # 本目录下的xlsx文件名字列表
            # print(xlsx_files[0])
            xlsx_file_0 = path + '/' + xlsx_files[0]
            data_biaotou = []  # 用来复制表头数据
            workbook_0 = openpyxl.load_workbook(filename=xlsx_file_0)
            worksheet_0 = workbook_0.worksheets[sheet_what]
            if biaotou_row != 0:  # 表头存在，才复制
                # 复制表头数据
                for col in range(1, worksheet_0.max_column + 1):
                    data_biaotou.append(worksheet_0.cell(row=biaotou_row, column=col).value)  # 默认表头在第一行
                # data_1 = worksheet_0.cell(row=1, column=1).value
                # data_2 = worksheet_0.cell(row=2, column=1).value
                print('表头', data_biaotou)
                scr_output(scr_10, '\n\n存在表头：\n{}'.format(data_biaotou))

            data_shuju = []  # 用来复制储存数据的数据集
            num = len(xlsx_files)
            print('\n\n开始提取数据…………\n')

            # 判断储存行开始位置
            if biaotou_row == 0:  # 没有表头，储存行开始在标题行之后
                cucun_row = biaoti_row + 1
            else:  # 有表头行，要多加一行才开始储存
                cucun_row = biaoti_row + 2
            # 遍历每个文件
            for n in range(num):
                xlsx_file = path + '/' + xlsx_files[n]
                workbook_n = openpyxl.load_workbook(filename=xlsx_file)
                worksheet_n = workbook_n.worksheets[sheet_what]

                # 删除空行（涉及删除行后，行数的索引值发生变化，所以复杂了点）
                # 前三行如果是None，就判定为空行，或者假设第二列名字出现了张三这个人
                zero_row_list = []
                # 存在汇总文件中存在例示行，需要检测出来，删掉，不要储存进来我们的数据集
                if lishi_row == '有':  zero_row_list.append(biaotou_row+1) ##################################### 例示行需要判定，不能一刀切！！！！
                # 遍历寻找空行
                for row in range(1, worksheet_n.max_row + 1):
                    if (worksheet_n.cell(row, 1).value == 'None' and worksheet_n.cell(row,2).value == 'None'
                        and worksheet_n.cell(row, 3).value == 'None') or \
                        (worksheet_n.cell(row, 1).value == None and worksheet_n.cell(row,2).value == None
                        and worksheet_n.cell(row, 3).value == None) or \
                        worksheet_n.cell(row, 2).value == '张三' :
                        zero_row_list.append(row)
                delete_row = 0
                # print('\n\n开始删除{}的空行…………\n'.format(xlsx_files[n]))
                scr_output(scr_10, '\n\n开始删除{}的空行/空列和例示行…………\n'.format(xlsx_files[n]))
                for j in zero_row_list:
                    # scr_output(scr_10,'\n原本该空值在第{}行,正在删除该空值目前所在的第{}行'.format(j , j -delete_row)
                    # print('原本该空值在第{}行,正在删除该空值目前所在的第{}行'.format(j , j -delete_row))
                    worksheet_n.delete_rows(j - delete_row, 1)
                    delete_row = delete_row + 1

                # scr_output(scr_10,'\n\n正在删除空列（第十七列后的五列数）\n')
                # 删除空列，第十七列后的列数
                # worksheet_n.delete_cols(17, 5)

                # 储存数据
                for i in range(cucun_row, worksheet_n.max_row + 1):
                    list = []
                    for j in range(1, worksheet_n.max_column + 1):
                        list.append(worksheet_n.cell(row=i, column=j).value)
                    data_shuju.append(list)

            # print('数据', data_shuju)
            time.sleep(1)
            data2 = '\n\n数据：\n'
            for i in data_shuju:  # 用来输出窗口显示数据更好看
                data2 = data2 + str(i) + '\n'
            scr_output(scr_10, data2)
            # 自此，以上代码获取了全部文件的有效数据

            # # 汇总表头和数据,新建保存总表

            # 不需要添加表头，有模板表了，以下代码注释
            # data = []
            # data.append(data_biaotou)  # 添加表头
            # for l in range(len(data_shuju)):  # 添加数据
            #     data.append(data_shuju[l])
            workbook = openpyxl.load_workbook(xlsx_file_0)  # 打开模板表，以第一个文件作为模板
            worksheet = workbook.worksheets[sheet_what]
            worksheet.title = '汇总'

            # 写入大标题
            # worksheet.cell(1,1).value = worksheet_0.cell(row=1, column=1).value
            # 写入第二行表头
            # worksheet.cell(biaotou_row, 1).value = worksheet_0.cell(row=biaotou_row, column=1).value
            # worksheet.merge_cells(start_column=1, end_column=15, start_row=1, end_row=1)  # 合并单元格
            # worksheet.merge_cells(start_column=1, end_column=15, start_row=2, end_row=2)

            # 删除这个表的除了大标题行和表头行的全部没用数据，即储存行开始的数据
            for row in range(0, worksheet.max_row):  # 懒得计算，索性多删几行
                worksheet.delete_rows(cucun_row + row, 1)

            # 写入数据
            for n_row in range(0, len(data_shuju)):  # 从第三行开始写入数据
                for n_col in range(0, len(data_shuju[n_row])):
                    if data_shuju[n_row][n_col] != None:
                        worksheet.cell(row=n_row + cucun_row, column=n_col + 1, value=str(data_shuju[n_row][n_col]))
                    else:pass
            # 获取四个区域
            max_row = worksheet.max_row  # 获得最大行数
            max_column = worksheet.max_column  # 获得最大列数
            min_row = worksheet.min_row
            min_column = worksheet.min_column

            scr_output(scr_10, '\n\n数据写入总表完成…………\n')
            scr_output(scr_10, '\n\n字体调整完成…………\n')
            scr_output(scr_10, '\n\n开始给区域设置设置框线…………\n')
            scr_output(scr_10, '\n\n开始居中对齐…………\n')

            # 给区域设置设置框线
            for row in tuple(worksheet[min_row + 2:max_row]):
                for cell in row:
                    cell.border = my_border('thin', 'thin', 'thin', 'thin')
                    # 设置单元格对齐方式 Alignment(horizontal=水平对齐模式,vertical=垂直对齐模式,text_rotation=旋转角度,wrap_text=是否自动换行)
                    cell.alignment = Alignment(horizontal='center', vertical='center', text_rotation=0, wrap_text=True)
                    # 字体对象
                    cell.font = Font(name=u'宋体', bold=False, italic=False, size=11)  # bold是否加粗, italic是否斜体

            # 区域自动调整列宽
            column_widths = []  # 定义用来获取当前列最大宽度的空列表
            for i, col in enumerate(
                    worksheet.iter_cols(min_col=min_column, max_col=max_column, min_row=min_row, max_row=max_row)):
                for cell in col:
                    value = cell.value
                    if value is not None:
                        if isinstance(value, str) is False:
                            value = str(value)
                        try:
                            column_widths[i] = max(column_widths[i], len(value))
                        except IndexError:
                            column_widths.append(len(value))
            # print('column_widths', column_widths)  # 得到该列最大的一个单元格的宽度（字符串数量）
            for i, width in enumerate(column_widths):
                col_name = get_column_letter(min_column + i)  # 获取行字母表头
                value = column_widths[i] * 2  # 设置列宽为最大长度比例
                worksheet.column_dimensions[col_name].width = value

            workbook.save(filename=filename + '.xlsx')  # 保存xlsx
            print('文件夹内全部文件汇总完成')
            messagebox.showinfo('小提示', '文件夹内全部文件汇总完成 成功！')
            scr_output(scr_10, '\n\n文件夹内全部文件汇总完成 成功！\n')
            scr_output(scr_10, '\n保存的文件路径为：\n{}\n\n\n\n\n\n'.format(
                '/'.join(path.split('/')[:-1]) + '/' + pathin6_0.get() + '.xlsx'))
            # AttributeError: 'MergedCell' object attribute 'value' is read-only
            # 读取到了合并的单元格，报错

        except Exception as error:
            error = str(error)
            print('错误提示', error)
            scr_output(scr_10, '\n汇总文件 失败！\n\n\n本次错误信息：\n{}'.format(error))
            scr_output(scr_10, '\n--------文件没有成功保存--------\n\n\n\n\n\n\n')
            messagebox.showinfo('错误提示', '汇总文件 失败！\n错误信息：\n{}'.format(error))
# 自动检测/恢复默认参数设置
def auto_general_merge_book():
    # 恢复默认值设置
    # number_chosen1_10.config(state='normal')  # 设为可编辑模式
    # number_chosen2_10.config(state='normal')  # 设为可编辑模式
    # number_chosen3_10.config(state='normal')  # 设为可编辑模式
    # number_chosen4_10.config(state='normal')  # 设为可编辑模式
    # number_chosen1_10.set('Sheet1')
    # number_chosen2_10.set('无标题')
    # number_chosen3_10.set('第1行')
    # number_chosen4_10.set('没有')
    # number_chosen1_10.config(state='readonly')  # 设为只读模式
    # number_chosen2_10.config(state='readonly')  # 设为只读模式
    # number_chosen3_10.config(state='readonly')  # 设为只读模式
    # number_chosen4_10.config(state='readonly')  # 设为只读模式
    # scr_output(scr_10,'\n您恢复了默认设置，工作簿中的第1个工作表，没有标题行，表头在第1行，没有例示行\n')
    sheet_what = number_chosen1_10.get()  # 获取工作表
    biaoti_row = 0  # '无标题' # 获取初始化设置
    biaotou_row = '无表头'
    lishi_row = '没有'
    # 自动检测功能 检测标题行，检测表头，检测例示行
    path = pathin_10.get()
    if path == "":
        messagebox.showinfo("提示", "自动检测功能需要获取文件内容，\n请输入正确的需汇总文件的文件夹路径！")
        scr_output(scr_10, '\n自动检测功能需要获取文件内容,\n本次没有正确输入正确的需汇总文件的文件夹路径！ \n请把汇总文件全部放在一个文件夹里面，并选择该文件夹\n\n')
    if path != "":
        try:
            # print('当前工作路径', os.path.abspath('.'))  # 打印当前目录
            xlsx_files = [x for x in os.listdir(path) if os.path.splitext(x)[1] == '.xlsx']  # 罗列当前目录内所有xlsx文件
            scr_output(scr_10, '\n\n自动检测到，需要统计{}个表格'.format(len(xlsx_files)))
            scr_output(scr_10, '\n\n自动检测到，统计表格有：\n{}'.format(xlsx_files))
            xlsx_file_0 = path + '/' + xlsx_files[0]
            workbook = openpyxl.load_workbook(xlsx_file_0)
            worksheet = workbook.worksheets[int(sheet_what[-1]) - 1]
            # 遍历获取相关列是否存在
            for row in tuple(worksheet[1:5]):  # 在前1到5行里检测，避免出错
                for cell in row:
                    # print(cell.value)
                    if ('姓名' in str(cell.value)) or ('名字' in str(cell.value)):
                        biaotou_row = cell.row
            for row in range(1, 6):  # 在前1到5行里检测，避免出错
                if (worksheet.cell(row, 2).value == 'None' and worksheet.cell(row, 3).value == 'None'and
                    worksheet.cell(row, 4).value == 'None') or (worksheet.cell(row, 2).value == None and
                    worksheet.cell(row, 3).value == None and worksheet.cell(row, 4).value == None) :
                    biaoti_row = biaoti_row + 1
                if worksheet.cell(row, 2).value == '张三':
                    lishi_row = '有'
            if biaoti_row == 0: number_chosen2_10.set('无标题')
            else: number_chosen2_10.set('前{}行'.format(biaoti_row))
            if type(biaotou_row) == int:  number_chosen3_10.set('第{}行'.format(biaotou_row))
            else:  number_chosen3_10.set('无表头')
            number_chosen4_10.set(lishi_row)
            scr_output(scr_10, '\n\n注意：本程序的自动检测功能并不是完全准确，请手动再次确认！\n\n！！！参数设置错误，很可能导致汇总数据缺失或混乱！！\n本次汇总的参数设置情况：'
                               '\n\n选择的工作表：第{}个，\n选择的标题行：{}，\n选择的表头行：{}，\n选择的例示行：{}\n'.format(sheet_what, biaoti_row,biaotou_row, lishi_row))

        except Exception as error:
            error = str(error)
            print('错误提示', error)
            scr_output(scr_10, '\n--------自动检测参数设置 失败！--------\n\n\n本次错误信息：\n{}\n'.format(error))
            messagebox.showinfo('错误提示', '自动检测参数设置 失败！\n错误信息：\n{}'.format(error))
'''#################################################################################################################'''
# 合并系数表格
def main0_1(path, filename):
    if path == "":
        messagebox.showinfo("提示","请输入正确的各支部递交入党申请书人数表文件夹路径！")
        scr_output(scr_0,'\n本次没有正确输入正确的各支部递交入党申请书人数表文件夹路径！\n\n')

    if path != "":
        try:
            # 判断选择表头方向执行
            if number_row_0.get() == 0 and number_col_0.get() == 0:
                scr_output(scr_0, '合并系数表 失败！\n\n\n本次错误信息：\n没有选择表头在哪个方向')
                print('没有选择表头在哪个方向')
                messagebox.showinfo('错误', message='没有选择表头在哪个方向')
            else:
                if number_row_0.get() == 1 and number_col_0.get() == 0:
                    print('\n表头横向')
                    scr_output(scr_0, '\n您选择了表头横向\n\n\n')
                if number_col_0.get() == 1 and number_row_0.get() == 0:
                    messagebox.showinfo('警告提示', '表头纵向功能还没有开发，仍按照横向表头合并，敬请期待！')
                # 纵向合并没有推出，敬请期待
                print('\n表头纵向')
                scr_output(scr_0, '\n您选择了表头纵向\n\n\n')

            xls_files = [x for x in os.listdir(path) if os.path.splitext(x)[1] == '.xls']
            if xls_files != []:  # 说明有xls文件
                xls_to_xlsx(path=path, sole=False)  # 给路径，让其自己转换成xlsx的
                scr_output(scr_0, '\n\n检测到有{}个xls格式文件，已经自动转换成可读取的xlsx文件类型：\n'.format(len(xls_files)))
            # # print(os.listdir('./test')) # 打印test目录下所有文件
            # print('当前工作路径',os.path.abspath('.')) # 打印当前目录
            # # xlfs = [x for x in os.listdir('./test') if os.path.isfile(x) and os.path.splitext(x)[1] == '.xlsx'] # 罗列当前目录内所有xlsx文件
            xlfs = [x for x in os.listdir(path) if os.path.splitext(x)[1] == '.xlsx'] # 罗列当前目录内所有xlsx文件
            scr_output(scr_0,'\n\n需要统计{}个表格'.format(len(xlfs)))
            scr_output(scr_0,'\n\n统计表格有：\n{}'.format(xlfs))
            print('需要统计',len(xlfs) , '个表格')
            print ('统计表格有：\n',xlfs)  # 本目录下的xlsx文件名字列表
            print(xlfs[0])

            xl0 = path + '/' + xlfs[0]
            data0 = []#复制表头数据
            wb0 = openpyxl.load_workbook(filename = xl0)
            ws0 = wb0.active
            for i in range(1,ws0.max_column+1):
                data0.append(ws0.cell(row = 1,column = i).value)
            print('表头',data0)
            scr_output(scr_0,'\n\n表头：\n{}'.format(data0))

            data1 = []#复制数据
            num = len(xlfs)
            for n in range(num):
                xf = path + '/' + xlfs[n]
                wb1 = openpyxl.load_workbook(filename = xf)
                ws1 = wb1.active
                for i in range(2,ws1.max_row + 1):
                    list = []
                    for j in range(1,ws1.max_column + 1):
                        list.append(ws1.cell(row=i,column=j).value)
                    data1.append(list)
            print('数据',data1)
            data2 = '\n\n数据：\n'
            for i in data1:
                data2= data2 + str(i) + '\n'
            scr_output(scr_0,data2)

            # # 汇总表头和数据,新建保存总表
            data=[]
            data.append(data0)#添加表头
            for l in range(len(data1)):#添加数据
                data.append(data1[l])
            wb = openpyxl.Workbook()#新建表
            ws = wb.active
            ws.title = '汇总'
            for n_row in range(1,len(data)+1):#写入数据
                for n_col in range(1,len(data[n_row-1])+1):
                    ws.cell(row=n_row,column=n_col,value=str(data[n_row-1][n_col-1]))
            wb.save(filename=filename + '.xlsx')#保存xlsx
            print ('系数表合并完成')
            messagebox.showinfo('小提示', '合并系数表 成功！')
            scr_output(scr_0,'\n\n合并系数表 成功！\n')
            scr_output(scr_0, '\n保存的文件路径为：\n{}\n\n\n\n\n\n'.format(
                '/'.join(path.split('/')[:-1]) + '/' + pathin2_0.get() + '.xlsx'))

            # 更新下一个选项卡的路径，更便捷
            pathin3_0.set('/'.join(path.split('/')[:-1]) + '/' + pathin2_0.get() + '.xlsx')
            scr_output(scr_0, '---------------------------------------------------------------------------------------\n' \
                   '---------------------------------------------------------------------------------------\n\n\n' \
                   '为您自动选择了合并后的文件路径是：\n{}\n\n'.format(pathin3_0.get()))

        except Exception as error:
            error = str(error)
            print('错误提示', error)
            scr_output(scr_0,'\n合并系数表 失败！\n\n\n本次错误信息：\n{}'.format(error))
            scr_output(scr_0, '\n--------文件没有成功保存--------\n\n\n\n\n\n\n')
            messagebox.showinfo('错误提示', '合并系数表 失败！\n错误信息：\n{}'.format(error))
'''#################################################################################################################'''
# 生成名额分配表 新汇总表的格式调整,名额确定
def main0_2(path, filename, xishu):
    if path == "":
        messagebox.showinfo("提示" ,"请输入正确的合并系数表文件路径！")
        scr_output(scr_0,'\n本次没有正确输入正确的合并系数表文件路径！\n\n')

    if path != "":
        try:
            from openpyxl.utils import get_column_letter
            if os.path.splitext(path)[1] == '.xls':  # 说明是xls文件
                xls_to_xlsx(path=path, sole=True)  # 给路径，让其自己转换成xlsx的
                scr_output(scr_0, '\n\n检测到本文件是xls格式文件，已经自动在原路径转换生成可读取的xlsx文件类型：\n')
                path = os.path.splitext(path)[0] + '.xlsx'

            workbook = openpyxl.load_workbook(path)
            worksheet = workbook.worksheets[0]
            one_cell = worksheet['1']  # 获取第1行的数据
            # print(one_cell) # (<Cell '汇总'.A1>, <Cell '汇总'.B1>, <Cell '汇总'.C1>, <Cell '汇总'.D1>, <Cell '汇总'.E1>)
            # 添加两列单元格
            worksheet.cell(row=1, column=4, value='乘系数后结果')
            worksheet.cell(row=1, column=5, value='最终名额')

            # 先删除空行（涉及删除行后，行数的索引值发生变化，所以复杂了点）
            zero_row_list = []
            for row in range(1, worksheet.max_row + 1):
                if (worksheet.cell(row, 2).value == 'None' and worksheet.cell(row, 3).value == 'None') or \
                        (worksheet.cell(row, 2).value == None and worksheet.cell(row, 3).value == None):
                    zero_row_list.append(row)
            delete_row = 0
            scr_output(scr_0,'\n\n开始删除空行…………\n')
            for j in zero_row_list:
                scr_output(scr_0,'\n原本该空值在第{}行,正在删除该空值目前所在的第{}行'.format(j , j -delete_row))
                print('原本该空值在第{}行,正在删除该空值目前所在的第{}行'.format(j , j -delete_row))
                worksheet.delete_rows( j -delete_row ,1)
                delete_row = delete_row + 1

            # 获取四个区域
            max_row = worksheet.max_row  # 获得最大行数
            max_column = worksheet.max_column  # 获得最大列数
            min_row = worksheet.min_row
            min_column = worksheet.min_column

            scr_output(scr_0,'\n\n开始填入系数，名额等数值…………\n')
            # 填入系数，名额等数值
            for row in range(2, max_row + 1):
                b = worksheet['C' + str(row)]
                try:
                    b = int(b.value)
                except:
                    try:
                        b = float(b.value)
                    except:
                        scr_output(scr_0,'\n\n\n表格C列似乎出现非法字符串！！！！\n\n\n')
                        print('表格C列出现非法字符串！！！！')
                # 重新写入’递交入党申请书人数‘列的数值，和写入’乘系数后结果‘列的数值
                worksheet.cell(row=row, column=3, value=b)
                worksheet.cell(row=row, column=4, value='=C{}*{}'.format(row,xishu))
                # worksheet.cell(row=row, column=4, value = b*0.75)
                worksheet.cell(row=row, column=5, value='=ROUND(D{},0)'.format(row))

            scr_output(scr_0,'\n\n开始添加总和行…………\n')
            # 添加总和行
            worksheet.cell(row=max_row +1, column=1, value='总人数')
            worksheet.cell(row=max_row +1, column=3, value='=SUM(C2:C{})'.format(max_row))
            worksheet.cell(row=max_row +1, column=4, value='=SUM(D2:D{})'.format(max_row))
            worksheet.cell(row=max_row +1, column=5, value='=SUM(E2:E{})'.format(max_row))
            max_row = max_row + 1  # 最大行数加一

            # 居中对齐，通过遍历方式实现
            # # print('max_row',max_row,'\nmax_column',max_column)
            # for i in range(max_row):
            #     for j in range(max_column):
            #         worksheet.cell(row=i+1, column=j+1).alignment = Alignment(horizontal='center', vertical='center')

            scr_output(scr_0,'\n\n开始给区域设置设置框线…………\n')
            scr_output(scr_0,'\n\n开始居中对齐…………\n')
            # 给区域设置设置框线
            for row in tuple(worksheet[min_row:max_row]):
                for cell in row:
                    cell.border = my_border('thin', 'thin', 'thin', 'thin')
                    # 设置单元格对齐方式 Alignment(horizontal=水平对齐模式,vertical=垂直对齐模式,text_rotation=旋转角度,wrap_text=是否自动换行)
                    alignment = Alignment(horizontal='center', vertical='center', text_rotation=0, wrap_text=True)
                    cell.alignment = alignment

            scr_output(scr_0,'\n\n开始区域自动调整列宽…………\n')
            # 区域自动调整列宽
            column_widths = []  # 定义用来获取当前列最大宽度的空列表
            for i, col in enumerate(worksheet.iter_cols(min_col=min_column, max_col=max_column, min_row=min_row)):
                for cell in col:
                    value = cell.value
                    if value is not None:
                        if isinstance(value, str) is False:
                            value = str(value)
                        try:
                            column_widths[i] = max(column_widths[i], len(value))
                        except IndexError:
                            column_widths.append(len(value))
            # print('column_widths', column_widths)  # 得到该列最大的一个单元格的宽度（字符串数量）
            for i, width in enumerate(column_widths):
                col_name = get_column_letter(min_column + i)  # 获取行字母表头
                value = column_widths[i] * 2 # 设置列宽为最大长度比例
                worksheet.column_dimensions[col_name].width = value

            scr_output(scr_0,'\n\n开始合并支部名列…………\n')
            # 合并支部名列
            count = 0
            last_zhiburow = None
            for row in range(1, max_row + 1):
                # print(worksheet.cell(row, 1).value)
                if worksheet.cell(row, 1).value == 'None':
                    count = count + 1
                if worksheet.cell(row, 1).value != 'None' and row != 1 and row != 2:
                    # print('合并行数起止',row-1-count,row-1)
                    worksheet.merge_cells(start_column=1, end_column=1, start_row=row - 1 - count, end_row=row - 1)
                    last_zhiburow = row
                    count = 0
            # 合并最后一个支部  # 加了最后一行总人数就不用再合并了
            # worksheet.merge_cells(start_column=1, end_column=1, start_row=last_zhiburow, end_row=max_row)

            color = PatternFill("solid", fgColor="00FFFF99")
            for column in range(1, max_column +1):
                worksheet.cell(max_row, column).fill = color  # 填充单元格
                # worksheet.cell(max_row, column).font = Font(bold=True) # 字体加粗

            workbook.save(filename + '.xlsx')
            messagebox.showinfo('小提示', '生成名额分配表 成功！')
            scr_output(scr_0,'\n\n生成名额分配表 成功！\n')
            scr_output(scr_0, '\n保存的文件路径为：\n{}\n\n\n\n\n\n'.format(
                '/'.join(path.split('/')[:-1]) + '/' + pathin4_0.get() + '.xlsx'))

        except Exception as error:
            error = str(error)
            print('错误提示', error)
            scr_output(scr_0, '\n生成名额分配表 失败！\n\n\n本次错误信息：\n{}'.format(error))
            scr_output(scr_0, '\n--------文件没有成功保存--------\n\n\n\n\n\n\n')
            messagebox.showinfo('错误提示', '生成名额分配表 失败！\n错误信息：\n{}'.format(error))
'''#################################################################################################################'''
# 合并学员册
def main0_3(path, filename, qishu):
    if path == "":
        messagebox.showinfo("提示","请输入正确的各支部学员花名册文件夹路径！")
        scr_output(scr_0,'\n本次没有正确输入正确的各支部学员花名册文件夹路径！\n\n')

    if path != "":
        try:
            if number_0_0.get() == 0: # 模板表
                xls_files = [x for x in os.listdir(path) if os.path.splitext(x)[1] == '.xls']
                if xls_files != []:  # 说明有xls文件
                    xls_to_xlsx(path=path, sole=False)  # 给路径，让其自己转换成xlsx的
                    scr_output(scr_0, '\n\n检测到有{}个xls格式文件，已经自动在原路径转换生成可读取的xlsx文件类型：\n'.format(len(xls_files)))

                # # print(os.listdir('./test')) # 打印test目录下所有文件
                print('当前工作路径',os.path.abspath('.')) # 打印当前目录
                # # xlsx_files = [x for x in os.listdir('./test') if os.path.isfile(x) and os.path.splitext(x)[1] == '.xlsx'] # 罗列当前目录内所有xlsx文件
                xlsx_files = [x for x in os.listdir(path) if os.path.splitext(x)[1] == '.xlsx']  # 罗列当前目录内所有xlsx文件
                scr_output(scr_0,'\n\n需要统计{}个表格'.format(len(xlsx_files)))
                scr_output(scr_0,'\n\n统计表格有：\n{}'.format(xlsx_files))
                print('需要统计', len(xlsx_files), '个表格')
                print('统计表格有：\n', xlsx_files)  # 本目录下的xlsx文件名字列表
                print(xlsx_files[0])

                xlsx_file_0 = path + '/' + xlsx_files[0]
                data_biaotou = []  # 复制表头数据
                workbook_0 = openpyxl.load_workbook(filename=xlsx_file_0)
                worksheet_0 = workbook_0.active
                # 复制表头数据
                for i in range(1, worksheet_0.max_column + 1):
                    data_biaotou.append(worksheet_0.cell(row=3, column=i).value) # 默认表头在第三行
                # data_1 = worksheet_0.cell(row=1, column=1).value
                # data_2 = worksheet_0.cell(row=2, column=1).value
                print('表头', data_biaotou)
                scr_output(scr_0,'\n\n表头：\n{}'.format(data_biaotou))

                data_shuju = []  # 复制数据
                num = len(xlsx_files)
                print('\n\n开始提取数据…………\n')
                for n in range(num):
                    xlsx_file = path + '/' + xlsx_files[n]
                    workbook_n = openpyxl.load_workbook(filename=xlsx_file)
                    worksheet_n = workbook_n.active

                    # 删除空行（涉及删除行后，行数的索引值发生变化，所以复杂了点）
                    zero_row_list = []
                    for row in range(1, worksheet_n.max_row + 1):
                        if (worksheet_n.cell(row, 1).value == 'None' and worksheet_n.cell(row,15).value == 'None' and worksheet_n.cell(
                                row, 2).value == 'None') or \
                                (worksheet_n.cell(row, 1).value == None and worksheet_n.cell(row,15).value == None and worksheet_n.cell(
                                    row, 2).value == None) or \
                                worksheet_n.cell(row, 2).value == '张三':
                            zero_row_list.append(row)
                    delete_row = 0
                    # print('\n\n开始删除{}的空行…………\n'.format(xlsx_files[n]))
                    scr_output(scr_0,'\n\n开始删除{}的空行/空列和张三示例行…………\n'.format(xlsx_files[n]))
                    for j in zero_row_list:
                        # scr_output(scr_0,'\n原本该空值在第{}行,正在删除该空值目前所在的第{}行'.format(j , j -delete_row)
                        # print('原本该空值在第{}行,正在删除该空值目前所在的第{}行'.format(j , j -delete_row))
                        worksheet_n.delete_rows(j - delete_row, 1)
                        delete_row = delete_row + 1

                    # scr_output(scr_0,'\n\n正在删除空列（第十七列后的五列数）\n')
                    # 删除空列，第十七列后的列数，因为有些学员册后面列存在无效数据，索性删除了。。如果不想删除请使用通用合并表格功能
                    worksheet_n.delete_cols(17, 5)

                    # 储存数据
                    for i in range(4, worksheet_n.max_row + 1):
                        list = []
                        for j in range(1, worksheet_n.max_column + 1):
                            list.append(worksheet_n.cell(row=i, column=j).value)
                        data_shuju.append(list)

                # print('数据', data_shuju)
                time.sleep(1)
                data2 = '\n\n数据：\n'
                for i in data_shuju:
                    data2 = data2 + str(i) + '\n'
                scr_output(scr_0,data2)

                # # 汇总表头和数据,新建保存总表
                data = []
                # data.append(data_biaotou)  # 添加表头
                for l in range(len(data_shuju)):  # 添加数据
                    data.append(data_shuju[l])
                workbook = openpyxl.load_workbook('mould\model\模板0 学员花名册.xlsx')  # 打开模板表
                worksheet = workbook.worksheets[0]
                # worksheet.title = '汇总'
                # worksheet.cell(1,1).value = worksheet_0.cell(row=1, column=1).value
                # 写入第二行表头
                worksheet.cell(2, 1).value = worksheet_0.cell(row=2, column=1).value
                # worksheet.merge_cells(start_column=1, end_column=15, start_row=1, end_row=1)
                # worksheet.merge_cells(start_column=1, end_column=15, start_row=2, end_row=2)
                # 删除红色的示例行
                if worksheet.cell(4, 2).value == '张三':
                    worksheet.delete_rows(4, 1)
                print(data)
                # 写入数据
                for n_row in range(1, len(data) + 1):  # 从第三行开始写入数据
                    for n_col in range(1, len(data[n_row - 1]) + 1):
                        if data[n_row - 1][n_col - 1] != None: # 过滤空值
                            # if str(data[n_row - 1][n_col - 1]).isdigit(): # 字符串中全是数字
                            #     worksheet.cell(row=n_row + 3, column=n_col, value=str(int(data[n_row - 1][n_col - 1])))
                            # else:
                            va = str(data[n_row - 1][n_col - 1])
                            # worksheet.cell(row=n_row + 3, column=n_col, value=str(data[n_row - 1][n_col - 1]).rstrip('.0'))
                            # rstrip('.0')方法会删除字符串末尾连续出现的字符集中的任意一个字符。因此，如果字符串末尾是0.0，那么三个字符都会被删除，返回的结果就是空字符串。
                            worksheet.cell(row=n_row + 3, column=n_col, value=va if va[-2:] != '.0' else va[:-2])

                        else:pass

                # 获取四个区域
                max_row = worksheet.max_row  # 获得最大行数
                max_column = worksheet.max_column  # 获得最大列数
                min_row = worksheet.min_row
                min_column = worksheet.min_column

                scr_output(scr_0,'\n\n数据写入总表完成…………\n')
                worksheet.cell(2,1,'{}年第{}期入党积极分子培训班学员花名册（学生）'.format(str(datetime.datetime.now().year),qishu))
                scr_output(scr_0,'\n\n写入期数表头总表完成…………\n')
                scr_output(scr_0,'\n\n开始给区域设置设置框线…………\n')
                scr_output(scr_0,'\n\n开始居中对齐…………\n')
                # 给区域设置设置框线
                for row in tuple(worksheet[min_row+2:max_row]):
                    for cell in row:
                        cell.border = my_border('thin', 'thin', 'thin', 'thin')
                        # 设置单元格对齐方式 Alignment(horizontal=水平对齐模式,vertical=垂直对齐模式,text_rotation=旋转角度,wrap_text=是否自动换行)
                        alignment = Alignment(horizontal='center', vertical='center', text_rotation=0, wrap_text=True)
                        cell.alignment = alignment

                workbook.save(filename=filename + '.xlsx')  # 保存xlsx
                print('学员花名册合并完成')

                # # if number_0_1 == 1:
                # # 删除None值
                # scr_output(scr_0, '\n\n开始删除空值（None）…………\n')
                # workbook_new = openpyxl.load_workbook(filename + '.xlsx')  # 打开表
                # # 储存数据
                # for i in range(1, workbook_new.worksheets[0].max_row + 1):
                #     for j in range(1, workbook_new.worksheets[0].max_column + 1):
                #         if workbook_new.worksheets[0].cell(row=i, column=j).value == 'None' or None:
                #             print('这里有None')
                #             workbook_new.worksheets[0].cell(row=i, column=j, value = '.')
                #             print(workbook_new.worksheets[0].cell(row=i, column=j).value)

                messagebox.showinfo('小提示', '合并学员花名册 成功！')
                scr_output(scr_0,'\n\n合并学员花名册 成功！\n')
                scr_output(scr_0, '\n保存的文件路径为：\n{}\n\n\n\n\n\n'.format(
                    '/'.join(path.split('/')[:-1]) + '/' + pathin6_0.get() + '.xlsx'))
                # AttributeError: 'MergedCell' object attribute 'value' is read-only
                # 读取到了合并的单元格，报错
            else:
                scr_output(scr_0, '\n本次没有合并，自定义表格合并，请前往”通用功能“-->”通用表格合并“\n\n\n\n\n\n\n')
                messagebox.showinfo('小提示', '本次没有合并，自定义表格合并，请前往”通用功能“-->”通用表格合并”进行')
        except Exception as error:
            error = str(error)
            print('错误提示', error)
            scr_output(scr_0, '\n合并学员花名册 失败！\n\n\n本次错误信息：\n{}'.format(error))
            scr_output(scr_0, '\n--------文件没有成功保存--------\n\n\n\n\n\n\n')
            messagebox.showinfo('错误提示', '合并学员花名册 失败！\n错误信息：\n{}'.format(error))
'''#################################################################################################################'''
def grouping(list_name,list_class,list_zhibu,worksheet_group,zhibu,count,qishu):

    # 居中对齐
    align = Alignment(horizontal='center', vertical='center')
    # 组名
    group = worksheet_group.title
    # 写入班级名字的开始行
    start_row = 8
    # 写入标题
    worksheet_group.cell(2, 1, '经济管理与法学学院第{}期入党积极分子培训班学员分组表'.format(qishu)).alignment = align
    # 分别合并单元格
    # worksheet_exam1.merge_cells(start_row=1, start_column=1, end_row=1, end_column=21)

    # 遍历计数
    for i in list(zip(list_name, list_class, list_zhibu)):
        if i[2] == zhibu:
            count += 1

    scr_output(scr_1,'当前正在分组的是：\n支部：{}\n人数：{}\n\n'.format(zhibu,count))
    # print(list_zhibu)

    # 遍历写入
    for i in list(zip(list_name, list_class, list_zhibu)):
        # print('每个学生元组',i)   # ('黄慧娟', '商类201班', '人营支部')
        # print('名字',i[0])      # 黄慧娟
        # print('班级',i[1])      # 商类201班
        # print('支部',i[2])      # 人营支部

        if i[2] == zhibu:
            worksheet_group.cell(start_row, 3, i[1]).alignment= align
            worksheet_group.cell(start_row, 4, i[0]).alignment= align
            start_row += 1
            # if start_row == 33:   # 此时超出25个人
            #     pass

    # 写入组名和支部，并居中对齐
    worksheet_group.cell(4, 1, '{}（{}人）'.format(group,count)).alignment= align
    worksheet_group.cell(8, 2, zhibu).alignment= align
    # 分别合并组名和支部的单元格
    worksheet_group.merge_cells(start_row=4, start_column=1, end_row=count + 7, end_column=1)
    worksheet_group.merge_cells(start_row=8, start_column=2, end_row=count + 7, end_column=2)

    # 获取四个区域
    max_row = worksheet_group.max_row  # 获得最大行数
    max_column = worksheet_group.max_column  # 获得最大列数
    min_row = worksheet_group.min_row
    min_column = worksheet_group.min_column

    scr_output(scr_1, '\n开始给区域设置设置框线…………\n')
    scr_output(scr_1, '\n开始居中对齐、调整行高…………\n')
    row_count = 3
    # 给区域设置设置框线
    for row in tuple(worksheet_group[min_row+2:max_row]):
        for cell in row:
            cell.border = my_border('thin', 'thin', 'thin', 'thin')
            # 设置单元格对齐方式 Alignment(horizontal=水平对齐模式,vertical=垂直对齐模式,text_rotation=旋转角度,wrap_text=是否自动换行)
            alignment = Alignment(horizontal='center', vertical='center', text_rotation=0, wrap_text=True)
            # 字体对象
            font = Font(name=u'宋体', bold=False, italic=False, size=12)  # bold是否加粗, italic是否斜体
            cell.alignment = alignment
            cell.font = font
        # 调整每一行的行高
        worksheet_group.row_dimensions[row_count].height = 18
        row_count = row_count + 1
def fenzu_smart(path,save,qishu):

    scr_output(scr_1, '\n开始智慧分组计算\n\n\n')
    workbook = openpyxl.load_workbook(path)
    worksheet = workbook.worksheets[0]
    # 定义小组空列表
    data = []
    # 遍历本小组每个人信息，收集成列表
    minrow = worksheet.min_row  # 最小行
    maxrow = worksheet.max_row  # 最大行
    mincol = worksheet.min_column  # 最小列
    maxcol = worksheet.max_column  # 最大列

    # 遍历获取相关列是否存在
    for row in tuple(worksheet[1:5]):  # 在前1到5行里检测，避免出错
        for cell in row:
            # print(cell.value)
            if ('姓名' in str(cell.value)) or ('名字' in str(cell.value)):
                biaotou_row = cell.row
                name_col = cell.column
            if ('专业班次' in str(cell.value)) or ('班级' in str(cell.value)) or ('专业' in str(cell.value)):
                class_col = cell.column
            if ('推荐支部' in str(cell.value)) or ('支部' in str(cell.value)):
                zhibu_col = cell.column
    try:
        for i in range(minrow + 3, maxrow + 1):
            cell1 = worksheet.cell(i, name_col).value
            cell2 = worksheet.cell(i, class_col).value
            cell3 = worksheet.cell(i, zhibu_col).value
            print([cell1, cell2, cell3], end=" ")
            data.append([cell1, cell2, cell3])
        # 使用filter()函数，删除列表中的None值
        data = list(filter(None, data))
        print(data)
    except:
        scr_output(scr_1, '\n没有找到姓名或专业班次或推荐支部列，分组失败！\n')
    # data = pd.read_excel("互联网/参考 学员花名册7.xlsx",
    #                      sheet_name='Sheet1', usecols=[1, 8, 9])
    # # # 首先将pandas读取的数据转化为array
    # # data_array = np.array(data)
    # # # 然后转化为list形式
    # # data_list =data_array.tolist()
    # # print(list(data_list))
    # # data = list(data_list)[2:]
    # # # print(data[2:])
    # data = data[2:]
    # data = data.values.tolist()
    result = []
    need_change = []

    all_dic = dict()
    all_dic_up = dict()
    for i in range(len(data)):
        if '研' not in data[i][2]:
            if data[i][2] in all_dic.keys():
                all_dic.get(data[i][2]).append([data[i][0], data[i][1], data[i][2]])
            else:
                all_dic[data[i][2]] = [[data[i][0], data[i][1], data[i][2]]]
        else:
            if data[i][2] in all_dic_up.keys():
                all_dic_up.get(data[i][2]).append([data[i][0], data[i][1], data[i][2]])
            else:
                all_dic_up[data[i][2]] = [[data[i][0], data[i][1], data[i][2]]]
    # print(all_dic)  # {'工信支部': [['韩冰凌', '工管181班'], ['邱爽', '工管181班'],……、
    # print(all_dic_up) # {'研二支部': [['罗琼', '会计'], ['陈悦', '会计'], ['朱明慧', '会计'], ['蒋露曦', '会计'], ['刘起林', '会计'], ['张莎', '法律（非法学）']]}
    # print(len(all_dic))
    for i in all_dic:
        if type(all_dic[i]) is not list:
            continue
        if len(all_dic[i]) > 25:
            need_change.append(i)
        if 20 <= len(all_dic[i]) <= 25:
            result.append(all_dic[i])
            all_dic[i] = 'over'
        # print(len(all_dic[i]))
        for j in all_dic:
            if i == j or type(all_dic[j]) is not list or type(all_dic[i]) is not list:
                continue
            if i != j:
                if len(all_dic[i]) + len(all_dic[j]) <= 25:
                    all_dic[i].extend(all_dic[j])  # 注意该函数没有返回值
                    result.append(all_dic[i])
                    all_dic[i] = 'over'
                    all_dic[j] = 'over'

    for i in need_change:
        all_dic[i + '_one'] = 'over'
        result.append(all_dic[i][:20])
        all_dic[i + '_two'] = all_dic[i][20:]
        all_dic.pop(i)
    need_change = []
    for i in all_dic:
        if type(all_dic[i]) is not list:
            continue
        if len(all_dic[i]) > 25:
            need_change.append(i)
    if len(need_change) != 0:
        for i in need_change:
            all_dic[i + '_one'] = 'over'
            result.append(all_dic[i][:20])
            all_dic[i + '_two'] = all_dic[i][20:]
            all_dic.pop(i)

    for i in all_dic:
        if type(all_dic[i]) is not list:
            continue
        if len(all_dic[i]) > 25:
            need_change.append(i)
        if 20 <= len(all_dic[i]) <= 25:
            result.append(all_dic[i])
            all_dic[i] = 'over'
        # print(len(all_dic[i]))
        for j in all_dic:
            if i == j or type(all_dic[j]) is not list or type(all_dic[i]) is not list:
                continue
            if i != j:
                if 20 <= len(all_dic[i]) + len(all_dic[j]) <= 25:
                    all_dic[i].extend(all_dic[j])  # 注意该函数没有返回值
                    result.append(all_dic[i])
                    all_dic[i] = 'over'
                    all_dic[j] = 'over'
    length = 100
    min_length = ''
    for i in all_dic:
        if type(all_dic[i]) is list and len(all_dic[i]) < length:
            min_length = i
            length = len(all_dic[i])
    for i in all_dic:
        if type(all_dic[i]) is list and i != min_length and len(all_dic[min_length]) + len(
                all_dic[i]) <= 25:
            all_dic[i].extend(all_dic[min_length])
            result.append(all_dic[i])
            break
        if type(all_dic[i]) is list and i != min_length:
            all_dic[i].extend(all_dic[min_length][0:23 - len(all_dic[i])])
            result.append(all_dic[i])
            all_dic[min_length] = all_dic[min_length][23 - len(all_dic[i]):]

    ################################################################################
    need_change_up = []
    all = 0
    for i in all_dic_up:
        all += len(all_dic_up[i])
    tem = []
    if all <= 25:
        for i in all_dic_up:
            tem.extend(all_dic_up[i])
        result.append(tem)
    for i in all_dic_up:
        if type(all_dic_up[i]) is not list:
            continue
        if len(all_dic_up[i]) > 25:
            need_change_up.append(i)
        if 20 <= len(all_dic_up[i]) <= 25:
            result.append(all_dic_up[i])
            all_dic_up[i] = 'over'
        # print(len(all_dic_up[i]))
        for j in all_dic_up:
            if i == j or type(all_dic_up[j]) is not list or type(all_dic_up[i]) is not list:
                continue
            if i != j:
                if len(all_dic_up[i]) + len(all_dic_up[j]) <= 25:
                    all_dic_up[i].extend(all_dic_up[j])  # 注意该函数没有返回值
                    result.append(all_dic_up[i])
                    all_dic_up[i] = 'over'
                    all_dic_up[j] = 'over'

    for i in need_change_up:
        all_dic_up[i + '_one'] = 'over'
        result.append(all_dic_up[i][:20])
        all_dic_up[i + '_two'] = all_dic_up[i][20:]
        all_dic_up.pop(i)

    for i in all_dic_up:
        if type(all_dic_up[i]) is not list:
            continue
        if len(all_dic_up[i]) > 25:
            need_change_up.append(i)
        if 20 <= len(all_dic_up[i]) <= 25:
            result.append(all_dic_up[i])
            all_dic_up[i] = 'over'
        # print(len(all_dic_up[i]))
        for j in all_dic_up:
            if i == j or type(all_dic_up[j]) is not list or type(all_dic_up[i]) is not list:
                continue
            if i != j:
                if 20 <= len(all_dic_up[i]) + len(all_dic_up[j]) <= 25:
                    all_dic_up[i].extend(all_dic_up[j])  # 注意该函数没有返回值
                    result.append(all_dic_up[i])
                    all_dic_up[i] = 'over'
                    all_dic_up[j] = 'over'
    length = 100
    min_length = ''
    for i in all_dic_up:
        if type(all_dic_up[i]) is list and len(all_dic_up[i]) < length:
            min_length = i
            length = len(all_dic_up[i])
    for i in all_dic_up:
        if type(all_dic_up[i]) is list and i != min_length and len(all_dic_up[min_length]) + len(
                all_dic_up[i]) <= 25:
            all_dic_up[i].extend(all_dic_up[min_length])
            result.append(all_dic_up[i])
            break
        if type(all_dic_up[i]) is list and i != min_length:
            all_dic_up[i].extend(all_dic_up[min_length][0:23 - len(all_dic_up[i])])
            result.append(all_dic_up[i])
            all_dic_up[min_length] = all_dic_up[min_length][23 - len(all_dic_up[i]):]
    ##################################################################################
    # result_df = pd.DataFrame(result)
    print(result)  # [[['卢慧珍', '电商191班'], ['覃小梅', '电商191班'], ……
    # 改为：[[['卢慧珍', '电商191班','电物支部'], ['覃小梅', '电商191班','电物支部'], ……

    workbook = openpyxl.load_workbook("mould\model\模板1 分组名单表.xlsx")

    # # 字体对象
    font = Font(name=u'宋体', bold=False, italic=False, size=12)  # bold是否加粗, italic是否斜体
    align = Alignment(horizontal='center', vertical='center', text_rotation=0, wrap_text=True)
    # # 边框
    # thin = Side(border_style="thin", color="000000")  # 边框样式，颜色
    # border = Border(left=thin, right=thin, top=thin, bottom=thin)  # 边框的位置
    # worksheet_check.cell(len(list_name) + 2 + 1, 1).border = border

    # 组名
    for i in range(len(result)):
        worksheet = workbook.worksheets[i]
        count = 1
        # 写入组名和支部，并居中对齐
        group = worksheet.title
        # 写入标题
        worksheet.cell(2, 1, '经济管理与法学学院第{}期入党积极分子培训班学员分组表'.format(qishu)).alignment = align
        for j in result[i]:
            # print(j[1],j[0]) # 电商191班 卢慧珍
            worksheet.cell(7 + count, 2).value = str(j[2])  # 支部
            worksheet.cell(7 + count, 3).value = str(j[1])  # 班级
            worksheet.cell(7 + count, 4).value = str(j[0])  # 姓名
            worksheet.row_dimensions[7 + count].height = 20  # 调整行高

            worksheet.cell(4, 1, '{}（{}人）'.format(group, worksheet.max_row-7)).alignment = align
            count += 1

        # 分别合并组名和支部的单元格
        worksheet.merge_cells(start_row=4, start_column=1, end_row=worksheet.max_row, end_column=1)
        last = None
        num = 1
        next = 8
        for row in [col for col in worksheet.columns][1]:
            # print(row.value,last)
            now = row.value
            if last != now and last != 'None' and last != None and num >= 8:
                if '支部' in last:
                    worksheet.merge_cells(start_row=next, start_column=2, end_row=num, end_column=2)
                    worksheet.cell(next, 2, '{}（{}人）'.format(last, num-next)).alignment = align
                    next = num+1
            if num == worksheet.max_row:
                worksheet.merge_cells(start_row=next, start_column=2, end_row=num, end_column=2)
                worksheet.cell(next, 2, '{}（{}人）'.format(last, num - next)).alignment = align
            last = now
            # print(num,now)
            num += 1

        scr_output(scr_1, '\n开始进行第{}组……\n'.format(group))
        # scr_output(scr_1, '\n正在给区域设置设置框线、居中对齐、调整行高\n')
        # 获取四个区域
        max_row = worksheet.max_row  # 获得最大行数
        max_column = worksheet.max_column  # 获得最大列数
        min_row = worksheet.min_row
        min_column = worksheet.min_column

        # 给区域设置设置框线
        for row in tuple(worksheet[min_row + 8:max_row]):
            for cell in row:
                cell.border = my_border('thin', 'thin', 'thin', 'thin')
                # 设置单元格对齐方式 Alignment(horizontal=水平对齐模式,vertical=垂直对齐模式,text_rotation=旋转角度,wrap_text=是否自动换行)
                cell.alignment = align
                cell.font = font

    workbook.save(save)
# 生成分组名单
def main1(path, filename, qishu):
    if path == "":
        messagebox.showinfo("提示","请输入正确的学员册文件路径！")
        scr_output(scr_1, '\n本次没有正确输入正确的学员册文件路径！\n')

    if path != "":
        # try:
        # 判断分组依据执行
        if number_row_1.get() == 0 and number_col_1.get() == 0:
            scr_output(scr_1, '\n生成分组名单 失败！\n\n\n本次错误信息：\n没有选择分组依据')
            # print('没有选择分组依据')
            messagebox.showinfo('错误', message='没有选择分组依据')
            return
        else:
            if number_row_1.get() == 1 and number_col_1.get() == 0:
                # print('\n按支部分')
                scr_output(scr_1, '\n您选择了按支部分\n\n\n')
            if number_col_1.get() == 1 and number_row_1.get() == 0:
                # messagebox.showinfo('警告提示', '按核定数分功能还没有开发，仍按照支部数目分，敬请期待！')
                # print('\n按核定数分')
                scr_output(scr_1, '\n您选择了按核定数分\n\n\n')

        if os.path.splitext(path)[1] == '.xls':  # 说明是xls文件
            xls_to_xlsx(path=path, sole=True)  # 给路径，让其自己转换成xlsx的
            scr_output(scr_1, '\n\n检测到本文件是xls格式文件，已经自动在原路径转换生成可读取的xlsx文件类型：\n')
            path = os.path.splitext(path)[1] + '.xlsx'

        if number_row_1.get() == 1 and number_col_1.get() == 0: # 按支部分
            workbook_name = openpyxl.load_workbook(path)
            worksheet_name = workbook_name.worksheets[0]

            workbook_group = openpyxl.load_workbook('mould\model\模板1 分组名单表.xlsx')

            biaotou_row =None
            name_col = None
            class_col = None
            zhibu_col = None
            # 遍历获取相关列是否存在
            for row in tuple(worksheet_name[1:5]):  # 在前1到5行里检测，避免出错
                for cell in row:
                    # print(cell.value)
                    if ('姓名' in str(cell.value)) or ('名字' in str(cell.value)):
                        biaotou_row = cell.row
                        name_col = cell.column_letter
                    if ('专业班次' in str(cell.value)) or ('班级' in str(cell.value)) or ('专业' in str(cell.value)):
                        class_col = cell.column_letter
                    if ('推荐支部' in str(cell.value)) or ('支部' in str(cell.value)):
                        zhibu_col = cell.column_letter
            if biaotou_row != None and name_col != None and class_col != None and zhibu_col != None:
                # 获取学员信息表中学员的名字列表，B4-最后  # 姓名、专业班级、推荐支部，需要灵活定位，标题行和表头也需要灵活定位
                list_name = list(cell.value for cell in worksheet_name[name_col])[biaotou_row:] # 姓名、专业班级、推荐支部，需要灵活定位，标题行和表头也需要灵活定位
                list_class = list(cell.value for cell in worksheet_name[class_col])[biaotou_row:] # 姓名、专业班级、推荐支部，需要灵活定位，标题行和表头也需要灵活定位
                list_zhibu = list(cell.value for cell in worksheet_name[zhibu_col])[biaotou_row:] # 姓名、专业班级、推荐支部，需要灵活定位，标题行和表头也需要灵活定位
            else:
                # 获取学员信息表中学员的名字列表，B4-最后  # 姓名、专业班级、推荐支部，需要灵活定位，标题行和表头也需要灵活定位
                list_name = list(cell.value for cell in worksheet_name['B'])[3:] # 姓名、专业班级、推荐支部，需要灵活定位，标题行和表头也需要灵活定位
                list_class = list(cell.value for cell in worksheet_name['I'])[3:] # 姓名、专业班级、推荐支部，需要灵活定位，标题行和表头也需要灵活定位
                list_zhibu = list(cell.value for cell in worksheet_name['O'])[3:] # 姓名、专业班级、推荐支部，需要灵活定位，标题行和表头也需要灵活定位

            scr_output(scr_1, '\n开始分组\n\n\n')

            # 按支部分组执行
            # 支部去重
            # list_zhibu_group = list(set(list_zhibu))
            # 不改变顺序的去重方法
            list_zhibu_group = []
            for i in list_zhibu:
                if i not in list_zhibu_group:
                    list_zhibu_group.append(i)
            print(list_zhibu_group)
            x = 0
            for i in list_zhibu_group:
                grouping(list_name=list_name, list_class=list_class, list_zhibu=list_zhibu,worksheet_group=workbook_group.worksheets[x], zhibu=list_zhibu_group[x], count=0, qishu=qishu)
                x += 1

            workbook_group.save('/'.join(path.split('/')[:-1]) + '/' + filename + '.xlsx')
            messagebox.showinfo('小提示', '生成分组名单表 成功！')
            scr_output(scr_1, '\n\n生成分组名单位表 成功！\n')
            scr_output(scr_1, '\n保存的文件路径为：\n{}\n\n\n\n\n\n'.format(
                '/'.join(path.split('/')[:-1]) + '/' + pathin2_1.get() + '.xlsx'))

        if number_col_1.get() == 1 and number_row_1.get() == 0:  # 按核定数分
            fenzu_smart(path=path, save='/'.join(path.split('/')[:-1]) + '/' + filename + '.xlsx',qishu=qishu)

            messagebox.showinfo('小提示', '生成分组名单表 成功！')
            scr_output(scr_1, '\n\n生成分组名单位表 成功！\n')
            scr_output(scr_1, '\n保存的文件路径为：\n{}\n\n\n\n\n\n'.format(
                '/'.join(path.split('/')[:-1]) + '/' + pathin2_1.get() + '.xlsx'))

        # except Exception as error:
        #     error = str(error)
        #     print('错误提示',error)
        #     if error == '7 must be greater than 8':
        #         messagebox.showinfo('错误提示', '您选择的学员册文件应该错误了，计数为0，找不到支部名称的相关数据'.format(error))
        #         scr_output(scr_1, '\n生成分组名单表 失败！\n\n\n本次错误信息：\n{}'.format(error))
        #     else:
        #         messagebox.showinfo('错误提示', '生成分组名单表 失败！\n错误信息：\n{}'.format(error))
        #         scr_output(scr_1, '\n生成分组名单表 失败！\n\n\n本次错误信息：\n{}'.format(error))
        #     scr_output(scr_1, '\n--------文件没有成功保存--------\n\n\n\n\n\n\n')
# 生成分组名单2
def fenzu_main2(workbook_name, output_workbook_name,qishu):
    from copy import copy
    from openpyxl.worksheet.datavalidation import DataValidation
    # print(workbook_name, output_workbook_name)
    workbook_name = workbook_name + '.xlsx'
    output_workbook_name = output_workbook_name + '.xlsx'

    # 加载原始Excel文件
    try:
        workbook = openpyxl.load_workbook(workbook_name)
    except:
        scr_output(scr_1, '\n没有找到文件"{}"，请先完成第一阶段的分组名单生成，或检查正确的分组名单文件路径！\n'.format(workbook_name))
        messagebox.showinfo('小提示', '没有找到文件,请先完成第一阶段的分组名单生成，或检查正确的分组名单文件路径！')
        return
    # 创建新的Excel文件
    new_workbook = openpyxl.Workbook()
    new_sheet = new_workbook.active
    # new_sheet.title = '分组名单'
    group_num = 1

    # 初始化开始写入的行数和列数
    current_row = 3
    current_column = 1
    cur_row = 3
    down_row1 = 3
    down_row2 = 3
    # 复制每个sheet的内容到新的sheet中
    for i, sheet in enumerate(workbook.worksheets):
        group_num = i
        print(i, sheet)  # 0 <Worksheet "第一组">
        # 判断sheet是奇数个还是偶数个
        if (i + 1) % 2 == 0:
            current_column = 7
            current_row = cur_row

        else:
            current_column = 1

        cur_row = current_row
        # 从第三行开始复制
        ori_row = 3
        ori_col = 1
        if sheet.cell(8, 3).value == None:
            print("该sheet为空")
            continue
        for row in sheet.iter_rows(min_row=3, values_only=True):
            # 忽略前两列
            for value in row:
                if value != None:
                    value = value.replace(" ", '')
                new_sheet.cell(row=current_row, column=current_column, value=value)
                original_cell = sheet.cell(row=ori_row, column=ori_col)

                new_cell = new_sheet.cell(row=current_row, column=current_column, value=value)
                # 设置单元格居中对齐，宋体字体
                new_cell.font = copy(original_cell.font)
                new_cell.alignment = copy(original_cell.alignment)
                new_cell.border = copy(original_cell.border)
                new_cell.number_format = copy(original_cell.number_format)
                new_cell.fill = copy(original_cell.fill)

                # 合并单元格
                if original_cell.coordinate in sheet.merged_cells:
                    new_sheet.merge_cells(new_cell.coordinate)

                current_column += 1
                ori_col = ori_col + 1

            # 每行写完后回到第一列/第七列
            if (i + 1) % 2 == 0:
                current_column = 7
            else:
                current_column = 1
            current_row += 1
            ori_row = ori_row + 1
            ori_col = 1

        if (i + 1) % 2 != 0:
            down_row1 = current_row
        else:
            down_row2 = current_row
        # 每个sheet复制完后，将行数加上2
        current_row = max(down_row1, down_row2) + 2

    # 保存新的Excel文件
    new_workbook.save(output_workbook_name)
    scr_output(scr_1, '\n正在整理列宽！\n')
    auto_column(output_workbook_name) # 整理列宽

    scr_output(scr_1, '\n正在写入表头！\n')
    # 加载原始Excel文件
    new_workbook = openpyxl.load_workbook(output_workbook_name)
    new_sheet = new_workbook.active
    # 写入表头
    new_sheet.cell(1, 1, workbook.worksheets[0].cell(1, 1).value).alignment = Alignment(horizontal='center',
                                                                                        vertical='center',
                                                                                        text_rotation=0, wrap_text=True)
    new_sheet.cell(2, 1, workbook.worksheets[0].cell(2, 1).value).alignment = Alignment(horizontal='center',
                                                                                        vertical='center',
                                                                                        text_rotation=0, wrap_text=True)
    if group_num >= 2:
        end_column = 11
    else:
        end_column = 5
    new_sheet.merge_cells(start_row=1, start_column=1, end_row=1, end_column=end_column)
    new_sheet.merge_cells(start_row=2, start_column=1, end_row=2, end_column=end_column)

    # 保存新的Excel文件
    new_workbook.save(output_workbook_name)
    scr_output(scr_1, '\n分组名单文件“{}”生成成功！\n'.format(output_workbook_name))

'''#################################################################################################################'''
def write_name(worksheet_exam,list_name,qishu):
    # 居中对齐
    align = Alignment(horizontal='center', vertical='center')
    # 写入标题
    worksheet_exam.cell(1, 1, '经济管理与法学学院分党校第{}期入党积极分子结业考试座位表'.format(qishu)).alignment = align
    # 分别合并单元格
    # worksheet_exam1.merge_cells(start_row=1, start_column=1, end_row=1, end_column=21)

    number = len(list_name)
    # worksheet_exam.cell(1, 1, '经济管理与法学学院第一期入党积极分子考试座位表')
    worksheet_exam.cell(18, 10, '本考场总共{}人'.format(number))

    scr_output(scr_2, '\n座位表\n 在本考场中总共有{}人\n'.format(number))

    cow = 4
    col = 2
    for i in range(0,len(list_name)):
        worksheet_exam.cell(cow,col,list_name[i])
        # print(list_name[i])
        col = col + 2
        if (i+1)>=10 and (i+1)%10==0:   # i除以10的商的余数等于0  执行换行
            col = 2
            cow = cow + 1
# 生成考试表
def main2(path, filename, qishu):
    if path == "":
        messagebox.showinfo("提示","请输入正确的学员册文件路径！")
        scr_output(scr_2, '\n本次没有正确输入正确的学员册文件路径！\n')

    if path != "":
        try:
            # print(path)
            if os.path.splitext(path)[1] == '.xls':  # 说明是xls文件
                xls_to_xlsx(path=path, sole=True)  # 给路径，让其自己转换成xlsx的
                scr_output(scr_2, '\n\n检测到本文件是xls格式文件，已经自动在原路径转换生成可读取的xlsx文件类型：\n')
                path = os.path.splitext(path)[1] + '.xlsx'
            workbook_name = openpyxl.load_workbook(path)
            worksheet_name = workbook_name.worksheets[0]

            biaotou_row =None
            name_col = None
            # class_col = None
            # zhibu_col = None
            # 遍历获取相关列是否存在
            for row in tuple(worksheet_name[1:5]):  # 在前1到5行里检测，避免出错
                for cell in row:
                    # print(cell.value)
                    if ('姓名' in str(cell.value)) or ('名字' in str(cell.value)):
                        biaotou_row = cell.row
                        name_col = cell.column_letter
                    # if ('专业班次' in str(cell.value)) or ('班级' in str(cell.value)) or ('专业' in str(cell.value)):
                    #     class_col = cell.column_letter
                    # if ('推荐支部' in str(cell.value)) or ('支部' in str(cell.value)):
                    #     zhibu_col = cell.column_letter
            if biaotou_row != None and name_col != None:
                # 获取学员信息表中学员的名字列表，B4-最后  # 姓名、专业班级、推荐支部，需要灵活定位，标题行和表头也需要灵活定位
                list_name = list(cell.value for cell in worksheet_name[name_col])[biaotou_row:] # 姓名、专业班级、推荐支部，需要灵活定位，标题行和表头也需要灵活定位
                # list_class = list(cell.value for cell in worksheet_name[class_col])[biaotou_row:] # 姓名、专业班级、推荐支部，需要灵活定位，标题行和表头也需要灵活定位
                # list_zhibu = list(cell.value for cell in worksheet_name[zhibu_col])[biaotou_row:] # 姓名、专业班级、推荐支部，需要灵活定位，标题行和表头也需要灵活定位
            else:
                # 获取学员信息表中学员的名字列表，B4-最后  # 姓名、专业班级、推荐支部，需要灵活定位，标题行和表头也需要灵活定位
                list_name = list(cell.value for cell in worksheet_name['B'])[3:] # 姓名、专业班级、推荐支部，需要灵活定位，标题行和表头也需要灵活定位
                # list_class = list(cell.value for cell in worksheet_name['I'])[3:] # 姓名、专业班级、推荐支部，需要灵活定位，标题行和表头也需要灵活定位
                # list_zhibu = list(cell.value for cell in worksheet_name['O'])[3:] # 姓名、专业班级、推荐支部，需要灵活定位，标题行和表头也需要灵活定
            # # 获取学员信息表中学员的名字列表，B4-最后
            # list_name = list(cell.value for cell in worksheet_name['B'])[3:]
            # 随机打乱列表
            random.shuffle(list_name)
            print(list_name)
            scr_output(scr_2, '本次抓取到的所有学员姓名信息：\n{}\n'.format(list_name))
            scr_output(scr_2, '\n学员人数是：{}人\n\n'.format(len(list_name)))

            workbook_exam = openpyxl.load_workbook('mould\model\模板2 考试座位表.xlsx')
            worksheet_exam1 = workbook_exam.worksheets[0]
            worksheet_exam2 = workbook_exam.worksheets[1]
            worksheet_exam3 = workbook_exam.worksheets[2]
            worksheet_exam4 = workbook_exam.worksheets[3]

            worksheet_exam1.title = '201考场'
            worksheet_exam2.title = '202考场'
            worksheet_exam3.title = '203考场'
            worksheet_exam4.title = '204考场'

            # print(len(list_name))
            # print(round(len(list_name)/2,0))
            averge_num = int(round(len(list_name) / 2, 0))
            if len(list_name) < 0:
                messagebox.showinfo("提示","学员册B4单元格及其以下面单元格找不到学员名字")
                scr_output(scr_2, '\n学员册B4单元格及其以下面单元格找不到学员名字\n')


            if 130 >= len(list_name) > 0:
                scr_output(scr_2, '\n本次学员人数在0到130区间内，共生成1个考场\n本考场是130人次以内\n')
                write_name(worksheet_exam=worksheet_exam1,list_name=list_name,qishu=qishu)
                workbook_exam.remove(worksheet_exam2)
                workbook_exam.remove(worksheet_exam3)
                workbook_exam.remove(worksheet_exam4)

            if  260 >= len(list_name) > 130:
                scr_output(scr_2, '\n本次学员人数在130到260区间内，共生成2个考场\n\n')
                list_name1 = list_name[0:averge_num].copy()
                list_name2 = list_name[averge_num:]
                write_name(worksheet_exam=worksheet_exam1,list_name=list_name1,qishu=qishu)
                write_name(worksheet_exam=worksheet_exam2,list_name=list_name2,qishu=qishu)
                workbook_exam.remove(worksheet_exam3)
                workbook_exam.remove(worksheet_exam4)


            three1averge_num = int(round(len(list_name) / 3, 0))
            three2averge_num = int(round(len(list_name) / 3 * 2, 0))
            if  390 >= len(list_name) > 260:
                scr_output(scr_2, '\n本次学员人数在260到360区间内，共生成3个考场\n\n')
                list_name1 = list_name[0:three1averge_num].copy()
                list_name2 = list_name[three1averge_num:three2averge_num]
                list_name3 = list_name[three2averge_num:]
                write_name(worksheet_exam=worksheet_exam1,list_name=list_name1,qishu=qishu)
                write_name(worksheet_exam=worksheet_exam2,list_name=list_name2,qishu=qishu)
                write_name(worksheet_exam=worksheet_exam3,list_name=list_name3,qishu=qishu)
                workbook_exam.remove(worksheet_exam4)


            quarter1averge_num = int(round(len(list_name) / 4, 0))
            quarter2averge_num = int(round(len(list_name) / 4 * 2, 0))
            quarter3averge_num = int(round(len(list_name) / 4 * 3, 0))
            if  520 >= len(list_name) > 390:
                scr_output(scr_2, '\n本次学员人数在390到480区间内，共生成4个考场\n分别是120人次，120人次，120人次，第四个考场装剩余部分的人\n')
                list_name1 = list_name[0:quarter1averge_num].copy()
                list_name2 = list_name[quarter1averge_num:quarter2averge_num]
                list_name3 = list_name[quarter2averge_num:quarter3averge_num]
                list_name4 = list_name[quarter3averge_num:]
                write_name(worksheet_exam=worksheet_exam1,list_name=list_name1,qishu=qishu)
                write_name(worksheet_exam=worksheet_exam2,list_name=list_name2,qishu=qishu)
                write_name(worksheet_exam=worksheet_exam3,list_name=list_name3,qishu=qishu)
                write_name(worksheet_exam=worksheet_exam4,list_name=list_name4,qishu=qishu)


            if  len(list_name) > 520:
                messagebox.showinfo("提示","本次没有生成，学员册学员名字超过520个，超出限制")
                scr_output(scr_2, '\n本次没有生成，学员册学员名字超过520个，超出限制，请自行先调整学员表人数再生成考试座位表\n')

            workbook_exam.save(filename + '.xlsx')
            messagebox.showinfo('小提示','生成考试座位表 成功！')
            scr_output(scr_2, '\n生成考试座位表 成功！\n')
            scr_output(scr_2, '\n保存的文件路径为：\n{}\n\n\n\n\n\n'.format(
                '/'.join(path.split('/')[:-1]) + '/' + pathin2_2.get() + '.xlsx'))

        except Exception as error:
            error = str(error)
            print('错误提示',error)
            scr_output(scr_2, '\n生成考试座位表 失败！\n\n\n本次错误信息：\n{}'.format(error))
            scr_output(scr_2, '\n--------文件没有成功保存--------\n\n\n\n\n\n\n')
            messagebox.showinfo('错误提示', '生成考试座位表 失败！\n错误信息：\n{}'.format(error))
'''#################################################################################################################'''
# 生成座位表
def main3(path, filename, qishu):
    if path == "":
        messagebox.showinfo("提示","请输入正确的分组名单文件路径！")
        scr_output(scr_3, '\n本次没有正确输入正确的分组名单文件路径！\n\n')

    if path != "":
        try:
            if os.path.splitext(path)[1] == '.xls':  # 说明是xls文件
                xls_to_xlsx(path=path, sole=True)  # 给路径，让其自己转换成xlsx的
                scr_output(scr_3, '\n\n检测到本文件是xls格式文件，已经自动在原路径转换生成可读取的xlsx文件类型：\n')
                path = os.path.splitext(path)[1] + '.xlsx'
            workbook_group = openpyxl.load_workbook(path)
            workbook_seat = openpyxl.load_workbook('mould\model\模板3 培训座位表.xlsx')
            worksheet_seat = workbook_seat.worksheets[0]

            # 使用到的颜色列表
            '''
            0099CCFF	00FFFF99	00FF8080	00CCCCFF
            0000FFFF	00FF99CC	00CCFFFF	00FFCC00
            00FFCC99	0000CCFF	0099CC00	
            '''
            # # 设置填充颜色加粗
            # color1 = PatternFill("solid",fgColor="0099CCFF")
            # # 设置单元格填充颜色
            # cell.fill = list_color[0]

            color1 = PatternFill("solid",fgColor="0099CCFF")
            color2 = PatternFill("solid",fgColor="00FFFF99")
            color3 = PatternFill("solid",fgColor="00FF8080")
            color4 = PatternFill("solid",fgColor="00CCCCFF")
            color5 = PatternFill("solid",fgColor="0000FFFF")
            color6 = PatternFill("solid",fgColor="00FF99CC")
            color7 = PatternFill("solid",fgColor="00CCFFFF")
            color8 = PatternFill("solid",fgColor="00FFCC00")
            color9 = PatternFill("solid",fgColor="00FFCC99")
            color10 = PatternFill("solid",fgColor="0000CCFF")
            color11 = PatternFill("solid",fgColor="0099CC00")

            # 定义填充单元格颜色列表
            list_color = [color1,color2,color3,color4,color5,color6,color7,color8,color9,color10,color11]
            color_number = 0

            # 奇数和偶数组的区分变量
            v = True
            # 设置学生开始写入的行列
            start_row1 = 4
            start_row2 = 4
            start_col1 = 2
            start_col2 = 12

            # 设置图例颜色的开始写入行列
            start_row3 = 18
            start_col3 = 18

            print('\n\n本次分组名单，共有{}组'.format(len(workbook_group.worksheets)))
            scr_output(scr_3, '\n\n本次分组名单，共有{}组\n\n'.format(len(workbook_group.worksheets)))

            # 在工作簿中遍历每一个表（小组）
            for worksheet in workbook_group.worksheets:

                # 定义小组人名空列表
                list_name = []

                # 遍历本小组每个人名，收集成列表
                for cell in worksheet['d'][7:]:
                    list_name.append(cell.value)
                # 使用filter()函数，删除列表中的None值
                list_name = list(filter(None, list_name))

                if list_name != []:
                    # 写入本小组人名
                    if v == True:    # 奇数组
                        print('奇数',list_name)
                        print('长度', len(list_name))
                        scr_output(scr_3, '\n\n当前奇数组，本组组员分别是： {}\n本组人数有 {}人\n\n'.format(list_name,len(list_name)))

                        for i in list_name:
                            worksheet_seat.cell(start_row1, start_col1, i).fill = list_color[color_number]
                            print('当前写入学生',i)
                            scr_output(scr_3, '当前写入学生 {}\n'.format(i))

                            if start_col1  == 10:
                                start_row1 += 1
                                start_col1 = 1
                            start_col1 += 1

                    else: # v = False   # 偶数组
                        print('偶数',list_name)
                        print('长度', len(list_name))
                        scr_output(scr_3, '\n\n当前偶数组，本组组员分别是： {}\n本组人数有 {}人\n\n'.format(list_name,len(list_name)))

                        for i in list_name:
                            worksheet_seat.cell(start_row2, start_col2, i).fill = list_color[color_number]
                            print('当前写入学生',i)
                            scr_output(scr_3, '当前写入学生 {}\n'.format(i))

                            if start_col2  == 20:
                                start_row2 += 1
                                start_col2 = 11
                            start_col2 += 1

                    if  v == True:
                        v = False
                    else: # v = False
                        v = True

                else:
                    print('\n\n\n！！！！！！！！！！！！！！！！！！！！！！！！！')
                    print('本次分组，发现分组名单中有一组空组')
                    scr_output(scr_3, '\n！！！！！！！！！！！！！！！！！！！！！！！！！\n发现分组名单中有一组空组（工作表为空）\n\n\n')

                    continue

                # 填写图例的单元格颜色
                worksheet_seat.cell(start_row3,start_col3).fill = list_color[color_number]
                # 填写图例的单元格小组名字
                worksheet_seat.cell(start_row3,start_col3).value = worksheet.title

                start_col3 += 1
                if start_col3 == 22:
                    start_col3 = 18
                    start_row3 += 1

                # 填充颜色列表+1
                color_number += 1

            # 居中对齐
            align = Alignment(horizontal='center', vertical='center')
            # 写入标题
            worksheet_seat.cell(1, 1, '经济管理与法学学院分党校第{}期入党积极分子培训班座位表'.format(qishu)).alignment = align
            # 分别合并单元格
            # worksheet_seat.merge_cells(start_row=1, start_column=1, end_row=1, end_column=21)

            workbook_seat.save(filename + '.xlsx')
            messagebox.showinfo('小提示', '生成培训座位表 成功！')
            scr_output(scr_3, '\n生成培训座位表 成功！\n\n')
            scr_output(scr_3, '\n保存的文件路径为：\n{}\n\n\n\n\n\n'.format(
                '/'.join(path.split('/')[:-1]) + '/' + pathin2_3.get() + '.xlsx'))

        except Exception as error:
            error = str(error)
            print('错误提示', error)
            scr_output(scr_3, '\n生成培训座位表 失败！\n\n\n本次错误信息：\n{}'.format(error))
            scr_output(scr_3, '\n--------文件没有成功保存--------\n\n\n\n\n\n\n')
            messagebox.showinfo('错误提示', '生成培训座位表 失败！\n错误信息：\n{}'.format(error))
'''#################################################################################################################'''
# 生成考勤表
def main4(path, filename, qishu):
    scr_output(scr_3, '本次选择的文件路径是{}\n\n'.format(path))
    if path == "":
        messagebox.showinfo("提示","请输入正确的分组名单文件路径！")
        scr_output(scr_3, '\n本次没有正确输入正确的分组名单文件路径！\n\n')

    if path != "":
        try:
            folder_path = os.getcwd() + "\\" + filename  # os.getcwd()返回当前文件所在的目录绝对路径
            # # 若文件夹不存在则自动新建
            if not os.path.exists(folder_path):  #判断是否存在文件夹如果不存在则创建为文件夹
                os.makedirs(folder_path)
                print('创建文件夹成功')
                scr_output(scr_3, '创建文件夹成功\n\n')
            else:
                print('已经存在文件夹')
                scr_output(scr_3, '已经存在文件夹\n\n')

            if os.path.splitext(path)[1] == '.xls':  # 说明是xls文件
                xls_to_xlsx(path=path, sole=True)  # 给路径，让其自己转换成xlsx的
                scr_output(scr_10, '\n\n检测到本文件是xls格式文件，已经自动在原路径转换生成可读取的xlsx文件类型：\n')
                path = os.path.splitext(path)[1] + '.xlsx'
            workbook_group = openpyxl.load_workbook(path)
            workbook_check = openpyxl.load_workbook('mould\model\模板4 第X小组通讯录及考勤表.xlsx')

            worksheet_check = workbook_check.worksheets[0]



            count = 1
            print('本次分组名单，共有{}组'.format(len(workbook_group.worksheets)))
            scr_output(scr_3, '本次分组名单，共有{}组\n\n'.format(len(workbook_group.worksheets)))

            # 在工作簿中遍历每一个表（小组）
            for worksheet in workbook_group.worksheets:
                # 定义起始行变量
                start_row = 3

                # 定义小组人名空列表
                list_name = []
                # 遍历本小组每个人名，收集成列表
                for cell in worksheet['d'][7:]:
                    list_name.append(cell.value)
                # 使用filter()函数，删除列表中的None值
                list_name = list(filter(None, list_name))

                # 定义小组班级空列表
                list_class = []
                # 遍历本小组每个人名，收集成列表
                for cell in worksheet['c'][7:]:
                    list_class.append(cell.value)
                # 使用filter()函数，删除列表中的None值
                list_class = list(filter(None, list_class))


                if list_name != [] and list_class != []:
                    print('{} 本组人数:{}'.format(worksheet.title,len(list_name)))
                    scr_output(scr_3, '{} 本组人数:{}\n\n'.format(worksheet.title,len(list_name)))
                    row1 = start_row
                    row2 = start_row
                    print(list_class)

                    # 字体对象
                    font = Font(name=u'宋体', bold=False, italic=False, size=9)  # bold是否加粗, italic是否斜体
                    # 写入
                    try:
                        # 写入本小组人名
                        for i in list_name:
                            worksheet_check.cell(row1,1,row1-2).font = font  # 写入序号
                            worksheet_check.cell(row1,3,i).font = font
                            row1 +=1

                        # 写入本小组班级
                        for i in list_class:
                            worksheet_check.cell(row2, 2, i).font = font
                            row2 += 1

                    except Exception as error:
                        print('有出错，已自动忽略 ################\n','错误信息：{}\n\n'.format(str(error)))
                        scr_output(scr_3, '有出错，已自动忽略 ################\n' + '错误信息：{}\n\n'.format(str(error)))
                        # if str(error) ==  "'MergedCell' object attribute 'value' is read-only" :
                        #     print('请检查分组名单中，是否有某单组人数已经超过25人,已自动忽略超过25人的部分')
                        #     scr_output(scr_3, '请检查分组名单中，是否有某单组人数已经超过25人,已自动忽略超过25人的部分\n\n')
                        continue

                    # # 删去多余行
                    # worksheet_check.delete_rows(len(list_name)+2+1, 25 - len(list_name))  # 从第7 + count +1 行开始删，也包括这一行，删去 25 - count 行

                    worksheet_check.cell(len(list_name)+2+1,1,'学习纪律要求：按时到课，自觉维持课堂秩序 ； 认真做好笔记，按要求完成分党校布置的各项任务；'
                                                              '认真参与小组讨论，讨论会上每人必须发言；撰写心得体会（培训心得体会一篇，1500字左右；观看纪录片微心得3篇，每篇100字），'
                                                              '认真做好自我鉴定；学员无故迟到或早退两次、缺勤一次，将视为培训不合格，不予颁发结业证书。')
                    # .alignment = Alignment(wrapText=True) 自动换行

                    # 获取四个区域
                    max_row = worksheet_check.max_row  # 获得最大行数
                    max_column = worksheet_check.max_column  # 获得最大列数
                    min_row = worksheet_check.min_row
                    min_column = worksheet_check.min_column

                    # # 字体对象
                    # font = Font(name=u'宋体', bold=False, italic=False, size=9)  # bold是否加粗, italic是否斜体
                    worksheet_check.cell(len(list_name)+2+1,1).font = font
                    align = Alignment(horizontal='center', vertical='center', wrap_text=True)
                    worksheet_check.cell(len(list_name)+2+1,1).alignment = align
                    # # 边框
                    # thin = Side(border_style="thin", color="000000")  # 边框样式，颜色
                    # border = Border(left=thin, right=thin, top=thin, bottom=thin)  # 边框的位置
                    # worksheet_check.cell(len(list_name) + 2 + 1, 1).border = border

                    scr_output(scr_3, '\n\n开始给区域设置设置框线和居中对齐…………\n')
                    # 给区域设置设置框线
                    for row in tuple(worksheet_check[min_row:max_row]):
                        for cell in row:
                            cell.border = my_border('thin', 'thin', 'thin', 'thin')
                            # 设置单元格对齐方式 Alignment(horizontal=水平对齐模式,vertical=垂直对齐模式,text_rotation=旋转角度,wrap_text=是否自动换行)
                            alignment = Alignment(horizontal='center', vertical='center', text_rotation=0,
                                                  wrap_text=True)
                            cell.alignment = alignment

                    # 合并单元格
                    worksheet_check.merge_cells(start_row=len(list_name)+2+1, start_column=1, end_row=len(list_name)+2+1, end_column=14)
                    # 调整行高
                    worksheet_check.row_dimensions[len(list_name) + 2 + 1].height = 30
                    worksheet_check.row_dimensions[len(list_name) + 2 + 1].alignment = Alignment(horizontal='left', vertical='center')

                    # 居中对齐
                    align = Alignment(horizontal='center', vertical='center')
                    # 写入标题
                    worksheet_check.cell(1, 1, '经济管理与法学学院分党校第{}期入党积极分子培训班{}成员通讯录及考勤表'.format(qishu,worksheet.title)).alignment = align
                    # 分别合并单元格
                    # worksheet_seat.merge_cells(start_row=1, start_column=1, end_row=1, end_column=21)

                    # 保存
                    try:
                        workbook_check.save('{}/{} {}考勤表.xlsx'.format(filename,count,worksheet.title))
                        print('{}/{} {}考勤表.xlsx  生成成功！'.format(filename,count, worksheet.title))
                        scr_output(scr_3, '{}/{} {}考勤表.xlsx  生成成功！\n\n'.format(filename,count, worksheet.title))

                    except Exception as error:
                        if str(error) == "[Errno 13] Permission denied: '{}/{} {}考勤表.xlsx'".format(filename,count,worksheet.title):
                            print("请检查是否打开了文件,'{}/{} {}考勤表.xlsx'该文件保存失败".format(filename,count,worksheet.title))
                            scr_output(scr_3, "请检查是否打开了文件,'{}/{} {}考勤表.xlsx'该文件保存失败\n\n".format(filename,count,worksheet.title))
                        continue

                    count += 1

                    workbook_check = openpyxl.load_workbook('mould\model\模板4 第X小组通讯录及考勤表.xlsx')
                    worksheet_check = workbook_check.worksheets[0]

                else:
                    print('！！！！！！！！！！！！！！！！！！！！！！！！！')
                    print('发现分组名单中有一组空组（组内缺少班级或者姓名信息）')
                    scr_output(scr_3, '！！！！！！！！！！！！！！！！！！！！！！！！！\n发现分组名单中有一组空组（组内缺少班级或者姓名信息）\n\n')
                    continue
                scr_output(scr_3, '\n保存的文件夹路径为：\n{}\n\n\n\n\n\n'.format(
                    '/'.join(path.split('/')[:-1]) + '/' + pathin2_4.get()))

        except Exception as error:
            error = str(error)
            print('错误提示', error)
            scr_output(scr_3, '\n生成各组考勤表 失败！\n\n\n本次错误信息：{}\n\n'.format(error))
            scr_output(scr_3, '\n--------文件没有成功保存--------\n\n\n\n\n\n\n')
            messagebox.showinfo('错误提示', '生成各组考勤表 失败！\n错误信息：\n{}'.format(error))
'''#################################################################################################################'''
# 自动检测按钮
def auto_time_management(path,sheet_is=None):
    global biaotou_row, name_col, birth_col, birth2_col, first_col, positive_col, object_col, ready_col, become_col,zhengshu_bianhao_col_set
    # 重新初始化
    number_row1_5.set(0)
    number_chosen1_5.set('')
    number_row2_5.set(0)
    number_chosen2_5.set('')
    number_row3_5.set(0)
    number_chosen3_5.set('')
    number_row4_5.set(0)
    number_chosen4_5.set('')
    number_row5_5.set(0)
    number_chosen5_5.set('')
    number_row6_5.set(0)
    number_chosen6_5.set('')
    number_row7_5.set(0)
    number_chosen7_5.set('')
    number_row8_5.set(0)
    number_chosen8_5.set('')

    number_chosen9_5.set('第1行')

    # 定义基础变量，假定相关列不存在
    biaotou_row = 1
    name_col = None
    birth_col = None
    birth2_col = None
    first_col = None
    positive_col = None
    object_col = None
    ready_col = None
    become_col = None
    zhengshu_bianhao_col_set = set()

    if path == "":
        messagebox.showinfo("提示", "自动检测功能需要获取审核文件内容，\n请输入正确的审核文件路径！")
        scr_output(scr_5, '\n自动检测功能需要获取审核文件内容，\n本次没有正确输入正确的审核文件路径！\n\n')
    if path != '':
        if sheet_is == None:
            sheet_is = int(number_chosen10_5.get()[-1])

        if os.path.splitext(path)[1] == '.xls':  # 说明是xls文件
            xls_to_xlsx(path=path, sole=True)  # 给路径，让其自己转换成xlsx的
            scr_output(scr_5, '\n\n检测到本文件是xls格式文件，已经自动在原路径转换生成可读取的xlsx文件类型：\n')
            path = os.path.splitext(path)[0] + '.xlsx'
            pathin_5.set(path)
        workbook = openpyxl.load_workbook(path)
        worksheet = workbook.worksheets[sheet_is - 1]

        # 遍历获取相关列是否存在
        for row in tuple(worksheet[1:5]):  # 在前1到5行里检测，避免出错
            for cell in row:
                # print(cell.value)
                if ('姓名' in str(cell.value)) or ('名字' in str(cell.value)):
                    biaotou_row = cell.row
                    name_col = cell.column_letter
                if '出生' in str(cell.value):                      birth_col = cell.column_letter
                if '身份证' in str(cell.value):                    birth2_col = cell.column_letter
                if ('申请书时间' in str(cell.value)) or ('入党时间' in str(cell.value)):  first_col = cell.column_letter
                if '积极分子时间' in str(cell.value):               positive_col = cell.column_letter
                if '发展对象时间' in str(cell.value):               object_col = cell.column_letter
                if '预备党员时间' in str(cell.value):               ready_col = cell.column_letter
                if '转正时间' in str(cell.value):                      become_col = cell.column_letter
                if ('结业证号' in str(cell.value)) or ('证书编号' in str(cell.value)):   zhengshu_bianhao_col_set.add(cell.column_letter)

        # global list_col
        # list_col = [biaotou_row, name_col, birth_col, birth2_col, first_col, positive_col, object_col, ready_col, become_col]
        # print(list_col)  # [2, 'C', 'F', None, 'I', 'J', 'K', 'L', 'M']
        # 更新 值的显示
        if biaotou_row:
            number_chosen9_5.set('第{}行'.format(biaotou_row))
        if name_col:
            number_row1_5.set(1)
            number_chosen1_5.set(name_col)
        else:
            number_row1_5.set(0)
        if birth2_col:
            number_row2_5.set(1)
            number_chosen2_5.set(birth2_col)
        else:
            number_row2_5.set(0)
        if birth_col:
            number_row3_5.set(1)
            number_chosen3_5.set(birth_col)
        else:
            number_row3_5.set(0)
        if first_col:
            number_row4_5.set(1)
            number_chosen4_5.set(first_col)
        else:
            number_row4_5.set(0)
        if positive_col:
            number_row5_5.set(1)
            number_chosen5_5.set(positive_col)
        else:
            number_row5_5.set(0)
        if object_col:
            number_row6_5.set(1)
            number_chosen6_5.set(object_col)
        else:
            number_row6_5.set(0)
        if ready_col:
            number_row7_5.set(1)
            number_chosen7_5.set(ready_col)
        else:
            number_row7_5.set(0)
        if become_col:
            number_row8_5.set(1)
            number_chosen8_5.set(become_col)
        else:
            number_row8_5.set(0)
    # 再次更新显示明暗
    list_row_and_chosen = [(number_row1_5, number_chosen1_5), (number_row2_5, number_chosen2_5),
                           (number_row3_5, number_chosen3_5), (number_row4_5, number_chosen4_5),
                           (number_row5_5, number_chosen5_5), (number_row6_5, number_chosen6_5),
                           (number_row7_5, number_chosen7_5), (number_row8_5, number_chosen8_5)]
    for (i, j) in list_row_and_chosen:
        disabled(number_row_5=i, number_chosen_5=j)
# 一键多表
def one_key_jiance():
    auto_time_management(pathin_5.get(), sheet_is = 1)
    workbook, filename, path = main5(path=pathin_5.get(), filename=pathin2_5.get(), sheet_is = 1)
    newpath = '/'.join(path.split('/')[:-1]) + '/' + pathin2_5.get() + '.xlsx'
    workbook.save(newpath)
    try:
        auto_time_management(newpath, sheet_is = 2)
        workbook, filename, path = main5(path=newpath, filename=pathin2_5.get(), sheet_is = 2)
        newpath = '/'.join(path.split('/')[:-1]) + '/' + pathin2_5.get() + '.xlsx'
        workbook.save(newpath)
    except:
        messagebox.showinfo('提示','只有一个工作表')
        return
    try:
        auto_time_management(newpath, sheet_is = 3)
        workbook, filename, path = main5(path=newpath, filename=pathin2_5.get(), sheet_is = 3)
        newpath = '/'.join(path.split('/')[:-1]) + '/' + pathin2_5.get() + '.xlsx'
        workbook.save(newpath)
    except:
        messagebox.showinfo('提示','只有两个工作表')
        return

# 编辑公式窗口
def gongshi():
    global list_gongshi,list_gongshi_simple
    def gongshi_save():
        global list_gongshi,list_gongshi_simple
        # list_gongshi = [gongshi1.get(),gongshi2.get(),gongshi3.get(),gongshi4.get(),gongshi5.get()]
        panduan_type = True
        panduan_int_list = [textvariable_year[0].get(),textvariable_day[0].get(),
                            textvariable_year[1].get(), textvariable_day[1].get(),
                            textvariable_year[2].get(), textvariable_day[2].get(),
                            textvariable_year[3].get(), textvariable_day[3].get(),
                            textvariable_year[4].get(), textvariable_day[4].get()]
        for i in range(len(panduan_int_list)):  # 正向遍历
            # print(panduan_int_list[i])
            if panduan_int_list[i] == '':
                messagebox.showinfo('错误提示','{}值为空，不符合规范！请重新输入！'.format('有'))
                panduan_type = False
                break
            else:
                for j in panduan_int_list[i]:
                    if '0' <= j <= '9':  # 判断是不是数字
                        pass
                    else:
                        messagebox.showinfo('错误提示','{}值为非法字符，不符合规范！请重新输入！'.format('有'))
                        panduan_type = False
                        break

        y1 = textvariable_year[0].get() if int(textvariable_year[0].get()) ==0 else ''.join(textvariable_year[0].get().lstrip("0"))
        y2 = textvariable_year[1].get() if int(textvariable_year[1].get()) ==0 else ''.join(textvariable_year[1].get().lstrip("0"))
        y3 = textvariable_year[2].get() if int(textvariable_year[2].get()) ==0 else ''.join(textvariable_year[2].get().lstrip("0"))
        y4 = textvariable_year[3].get() if int(textvariable_year[3].get()) ==0 else ''.join(textvariable_year[3].get().lstrip("0"))
        y5 = textvariable_year[4].get() if int(textvariable_year[4].get()) ==0 else ''.join(textvariable_year[4].get().lstrip("0"))
        d1 = textvariable_day[0].get() if int(textvariable_day[0].get()) ==0 else ''.join(textvariable_day[0].get().lstrip("0"))
        d2 = textvariable_day[1].get() if int(textvariable_day[1].get()) ==0 else ''.join(textvariable_day[1].get().lstrip("0"))
        d3 = textvariable_day[2].get() if int(textvariable_day[2].get()) ==0 else ''.join(textvariable_day[2].get().lstrip("0"))
        d4 = textvariable_day[3].get() if int(textvariable_day[3].get()) ==0 else ''.join(textvariable_day[3].get().lstrip("0"))
        d5 = textvariable_day[4].get() if int(textvariable_day[4].get()) ==0 else ''.join(textvariable_day[4].get().lstrip("0"))
        bk1 = textvariable_baokuo[0].get()
        bk2 = textvariable_baokuo[1].get()
        bk3 = textvariable_baokuo[2].get()
        bk4 = textvariable_baokuo[3].get()
        bk5 = textvariable_baokuo[4].get()

        if panduan_type == True:
            list_gongshi = ['int(first_value) - int(birth_value) - {}0000 -{} <{} 0'.format(y1,d1,'=' if bk1=='包括当天' else '' ),
                            # 'int(positive_value)-int(first_value) -{}0000 -{} <{} 0'.format(y2,d2,'=' if bk2=='包括当天' else '' ),
                            d2,
                            'int(object_value)- int(positive_value)-{}0000-{} <{} 0'.format(y3,d3,'=' if bk3=='包括当天' else '' ),
                            'int(ready_value) - int(object_value) -{}0000 -{} <{} 0'.format(y4,d4,'=' if bk4=='包括当天' else '' ),
                            'int(become_value)- int(ready_value) - {}0000 -{} >{} 0'.format(y5,d5,'=' if bk5=='包括当天' else '' )]
            list_gongshi_simple = [y1,y2,y3,y4,y5,  d1,d2,d3,d4,d5,  bk1,bk2,bk3,bk4,bk5]
            scr_output(scr_5,'\n公式1：{}\n公式2：{}\n公式3：{}\n公式4：{}\n公式5：{}\n\n公式保存成功！\n\n\n'.format(list_gongshi[0],
                '首次递交入党申请书 到 积极分子 间隔 {}年 {}天 {}'.format(y2, d2,'（刚好满这一天也会被检查出来）' if bk2 == '包括当天' else ''),list_gongshi[2],list_gongshi[3],list_gongshi[4]))
            # print(list_gongshi)
            gongshi.destroy()

    def gongshi_default():
        global list_gongshi,list_gongshi_simple
        # gongshi1.set('int(first_value) - int(birth_value) - 180000 < 0')
        # gongshi2.set('int(positive_value) - int(first_value) -15 <= 0')
        # gongshi3.set('int(object_value) - int(positive_value) - 10000 <= 0')
        # gongshi4.set('int(ready_value) - int(object_value) <= 0')
        # gongshi5.set('int(become_value) - int(ready_value) <= 0')
        # textvariable_year[0].set('18')
        # textvariable_year[1].set('0')
        # textvariable_year[2].set('1')
        # textvariable_year[3].set('0')
        # textvariable_year[4].set('1')
        # textvariable_day[0].set('0')
        # textvariable_day[1].set('15')
        # textvariable_day[2].set('0')
        # textvariable_day[3].set('0')
        # textvariable_day[4].set('0')
        # textvariable_baokuo[0].set('不包括当天')
        # textvariable_baokuo[1].set('包括当天')
        # textvariable_baokuo[2].set('包括当天')
        # textvariable_baokuo[3].set('包括当天')
        # textvariable_baokuo[4].set('不包括当天')
        list_gongshi_simple = ['18','0','1','0','1',    '0','180','0','0','0',   '包括当天','包括当天','包括当天','包括当天','不包括当天']
        for i in range(5):
            textvariable_year[i].set(list_gongshi_simple[0+i])
            textvariable_day[i].set(list_gongshi_simple[5+i])
            textvariable_baokuo[i].set(list_gongshi_simple[10+i])
        list_gongshi = ['int(first_value) - int(birth_value) - {}0000 -{} <{} 0'.format(list_gongshi_simple[0],list_gongshi_simple[5],'=' if list_gongshi_simple[10]=='包括当天' else '' ),
                        # 'int(positive_value)-int(first_value)- {}0000 -{} <{} 0'.format(list_gongshi_simple[1],list_gongshi_simple[6],'=' if list_gongshi_simple[11]=='包括当天' else '' ),
                        list_gongshi_simple[6],
                        'int(object_value)-int(positive_value)-{}0000 -{} <{} 0'.format(list_gongshi_simple[2],list_gongshi_simple[7],'=' if list_gongshi_simple[12]=='包括当天' else '' ),
                        'int(ready_value)- int(object_value) - {}0000 -{} <{} 0'.format(list_gongshi_simple[3],list_gongshi_simple[8],'=' if list_gongshi_simple[13]=='包括当天' else '' ),
                        'int(become_value)- int(ready_value) - {}0000 -{} >{} 0'.format(list_gongshi_simple[4],list_gongshi_simple[9],'=' if list_gongshi_simple[14]=='包括当天' else '' )]
        scr_output(scr_5,'\n公式已恢复默认！\n')

    # gongshi = Tk()
    gongshi = Toplevel(window)
    gongshi.geometry("500x190+700+310")
    try:
        gongshi.iconbitmap('mould\data\logo.ico')
    except:pass
    # 窗口置顶
    gongshi.attributes("-topmost", 1)  # 1==True 处于顶层
    # 禁止窗口的拉伸
    gongshi.resizable(0, 0)
    # 窗口的标题
    gongshi.title("公式编辑窗口")

    # 定义变量
    textvariable_year = [StringVar(),StringVar(),StringVar(),StringVar(),StringVar()]
    textvariable_day = [StringVar(),StringVar(),StringVar(),StringVar(),StringVar()]
    textvariable_baokuo = [StringVar(),StringVar(),StringVar(),StringVar(),StringVar()]
    text = ['公式1：出生日期-->申请入党  间隔', '公式2：申请入党-->积极分子  间隔', '公式3：积极分子-->发展对象  间隔',
            '公式4：发展对象-->预备党员  间隔', '公式5：预备党员-->转正时间  间隔']
    text2 = ['年 +','天']
    list_introduce = ['出生日期 和 首次递交入党申请书两个时间点之间关系','判断首次递交入党申请书 和 确认为入党积极分子两个时间点之间关系',
                      '确认为入党积极分子 到 列为发展对象两个时间点之间关系','列为发展对象 到 发展为预备党员两个时间点之间关系',
                      '发展为预备党员 到 预备党员转正两个时间点之间关系']
    a = '\n\n\n\n\n\n\n\n\n\n\n\n'
    for i in range(0, 5):
        label_gongshi = ttk.Label(gongshi, text=text[i])#标题
        label_gongshi.place(x=10, y=10 + 30 * i)
        createToolTip(label_gongshi, '{}这里是一条判断{}的公式'.format(a,list_introduce[i]))  # Add Tooltip
        a = a[:len(a)-2]
        # 多少年
        entry2_gongshi = ttk.Entry(gongshi, textvariable=textvariable_year[i])  # 输入框    # entry不能和grid连写，否则会报错
        entry2_gongshi.place(x=222, y=10 + 30 * i, width=50)
        label2_gongshi = ttk.Label(gongshi, text=text2[0])
        label2_gongshi.place(x=280, y=10 + 30 * i)
        if i == 1:
            entry2_gongshi.config(state='readonly')
        # 多少天
        entry3_gongshi = ttk.Entry(gongshi, textvariable=textvariable_day[i])  # 输入框    # entry不能和grid连写，否则会报错
        entry3_gongshi.place(x=312, y=10 + 30 * i, width=50)
        label3_gongshi = ttk.Label(gongshi, text=text2[1])
        label3_gongshi.place(x=370, y=10 + 30 * i)
        # 包不包括当天
        Combobox_gongshi = ttk.Combobox(gongshi, width=8, textvariable=textvariable_baokuo[i])
        Combobox_gongshi['values'] = ['包括当天','不包括当天']
        Combobox_gongshi.place(x=400, y=10 + 30 * i, width=80)
        Combobox_gongshi.current(0)  # 设置初始显示值，值为元组['values']的下标
        Combobox_gongshi.config(state='readonly')  # 设为只读模式
        createToolTip(Combobox_gongshi, '                                          ”包括当天“的意思就是，在刚好满足条件的这一天，还是会被审核出来标记')  # Add Tooltip
        if i == 1:
            Combobox_gongshi.config(state='disable')
        for i in range(5):
            textvariable_year[i].set(list_gongshi_simple[0+i])
            textvariable_day[i].set(list_gongshi_simple[5+i])
            textvariable_baokuo[i].set(list_gongshi_simple[10+i])


    button_gongshi = ttk.Button(gongshi, text="保存参数", command=gongshi_save)
    button_gongshi.place(x=250, y=160)

    button_gongshi = ttk.Button(gongshi, text="恢复默认", command=gongshi_default)
    button_gongshi.place(x=120, y=160)

    # 显示窗口(消息循环)
    gongshi.mainloop()
# 判断是否符合八位数字合法日期的要求
def data_legal(data):
    data = str(data)
    # 判断是8位数的数字
    if len(data) != 8:
        return  False
    for i in data:  # 正向遍历
        if '0' <= i <= '9':  # 判断是不是数字
            pass
        else:
            return False
    # if len(data)==8 and num_sole == True:
    #     print('是8位数的数字')

    # 判断日期是否合法：输入格式：2021-3-9
    runnian = False  # 判断是否为闰年

    # 记录 30天和 31天的月份 (type是集合类型)
    month_31 = {1, 3, 5, 7, 8, 10, 12}
    month_30 = {4, 6, 9, 11}

    # 将输入数据分割
    year, month, day = int(data[:4]), int(data[4:6]), int(data[6:8])
    # print(year, month, day)

    # 判断年份，以及是否为闰年
    if year % 4 == 0 and year % 100 != 0 or year % 400 == 0:
        runnian = True
    if 1900 > year or year > datetime.datetime.now().year:
        return False

    # 判断月份
    if 1 > month or month > 12:
        return False

    # 判断日期
    if month in month_31:
        if 1 > day or day > 31:
            return False
    elif month in month_30:
        if 1 > day or day > 30:
            return False
    elif month == 2:
        if not runnian and (1 > day or day > 28):
            return False
        elif runnian and (1 > day or day > 29):
            return False
    # if num_legal == True:
    #     print('日期合法')
    return True
# 检测名字是否合法
def name_legal(data):
    for i in data:
        if '\u4e00' > i or i > '\u9fff':
            return False
    return True
# 检测证书编号（结业证号是否合法）
def zhengshu_bianhao_legal(data):
    data = str(data)
    if len(data) != 9: #九位数
        return False
    for i in data:  # 正向遍历
        if '0' <= i <= '9':  # 判断是不是数字
            pass
        else:
            return False
    if str(data[0]+data[1]) != '20':
        return False
    if data[4] == data[5] == data[6] == data[7] == data[8]:
        return False
    return True
# 判断身份证号是否合法
def sfz_ID_legal(data):
    # print(data,type(data))
    if type(data) != int and  type(data) != str :
        return False
    if len(str(data)) != 18:
        return False
    y1 = data[:17]
    if y1.isdigit():
        y2 = list(data)
        w = [7, 9, 10, 5, 8, 4, 2, 1, 6, 3, 7, 9, 10, 5, 8, 4, 2]
        count = False
        for i in range(0, 17):
            count = count + (int(w[i]) * int(y2[i]))
        z = count % 11
        m = [1, 0, "X", 9, 8, 7, 6, 5, 4, 3, 2]
        s = None
        for i in range(0, 10):
            if int(i) == z:
                s = m[i]
        if str(s) == str(data[17]):
            return True
        else:
            return False
    else:
        return False
def main5_save(workbook,filename,path): # filename= pathin2_5.get()
    # workbook.save(filename + '.xlsx')
    newpath = '/'.join(path.split('/')[:-1]) + '/' + pathin2_5.get() + '.xlsx'
    workbook.save(newpath)
    messagebox.showinfo('小提示', '生成审核标注文件 成功！')
    scr_output(scr_5, '\n生成审核标注文件 成功！\n\n')
    scr_output(scr_5, '\n保存的文件路径为：\n{}\n\n\n\n\n\n'.format(newpath))
    # except Exception as error:
    #     error = str(error)
    #     print('错误提示', error)
    #     scr_output(scr_5, '\n生成审核标注表 失败！\n\n\n本次错误信息：{}\n\n'.format(error))
    #     scr_output(scr_5, '\n--------文件没有成功保存--------\n\n\n\n\n\n\n')
    #     messagebox.showinfo('错误提示', '生成审核标注表 失败！\n错误信息：\n{}'.format(error))

def is_half_year(start_date, end_date,setting):
    # import datetime
    """
    判断两个时间是否满半年
    :param start_date: 开始时间，格式为YYYYMMDD
    :param end_date: 结束时间，格式为YYYYMMDD
    :return: True表示满半年，False表示不满半年
    """
    start_year, start_month, start_day = int(start_date[:4]), int(start_date[4:6]), int(start_date[6:])
    end_year, end_month, end_day = int(end_date[:4]), int(end_date[4:6]), int(end_date[6:])
    start = datetime.date(start_year, start_month, start_day)
    end = datetime.date(end_year, end_month, end_day)
    interval = end - start
    half_year = datetime.timedelta(days=int(setting))
    return interval >= half_year

def main5(path, filename,sheet_is = None):
    from collections import Counter
    list_col = [int(number_chosen9_5.get()[1]), number_chosen1_5.get(), number_chosen3_5.get(), number_chosen2_5.get(), number_chosen4_5.get(),
                number_chosen5_5.get(), number_chosen6_5.get(), number_chosen7_5.get(),number_chosen8_5.get(), int(number_chosen10_5.get()[-1])]
    scr_output(scr_5,'\n您选择了工作簿中的第{}个工作表，表头在第{}行\n姓名是 {}列，\n出生年月是 {}列，\n身份证是 {}列，\n首次递交入党申请时间是 {}列，\n确认为积极分子时间是 {}列，'
                 '\n列为发展对象时间 {}列，\n发展为预备党员时间是 {}列，\n党员转正时间是 {}列'.format(int(number_chosen10_5.get()[-1]),int(number_chosen9_5.get()[1]),
                number_chosen1_5.get(), number_chosen3_5.get(), number_chosen2_5.get(), number_chosen4_5.get(),
                number_chosen5_5.get(), number_chosen6_5.get(), number_chosen7_5.get(),number_chosen8_5.get()))
    try:
        if list(zhengshu_bianhao_col_set) != []:
            scr_output(scr_5,'\n检测到有结业证号列的存在，其列在列表{}中\n'.format(list(zhengshu_bianhao_col_set)))
    except:  scr_output(scr_5,'\n本次没有执行自动检测！\n')

    gongshi1 = list_gongshi[0]
    gongshi2 = list_gongshi[1]
    gongshi3 = list_gongshi[2]
    gongshi4 = list_gongshi[3]
    gongshi5 = list_gongshi[4]
    if path == "":
        messagebox.showinfo("提示","请输入正确的审核文件路径！")
        scr_output(scr_5,'\n本次没有正确输入正确的审核文件路径！\n\n')
    if path != '':
        # try:
        if sheet_is == None:
            sheet_is = list_col[9]  # 即 =  int(number_chosen10_5.get()[-1])

        if os.path.splitext(path)[1] == '.xls':  # 说明是xls文件
            xls_to_xlsx(path=path, sole=True)  # 给路径，让其自己转换成xlsx的
            scr_output(scr_5, '\n\n检测到本文件是xls格式文件，已经自动在原路径转换生成可读取的xlsx文件类型：\n')
            path = os.path.splitext(path)[1] + '.xlsx'
            pathin_5.set(path)

        workbook = openpyxl.load_workbook(path)
        worksheet = workbook.worksheets[sheet_is - 1]

        # 定义颜色
        color = PatternFill("solid", fgColor="00FFFF99") #单一错误，黄色
        color2 = PatternFill("solid", fgColor="00CCFFFF") # 相关性错误，蓝色

        # 定义基础变量，假定相关列不存在
        biaotou_row = list_col[0]
        name_col = list_col[1]
        birth_col = list_col[2]
        birth2_col = list_col[3]
        first_col = list_col[4]
        positive_col = list_col[5]
        object_col = list_col[6]
        ready_col = list_col[7]
        become_col = list_col[8]
        # 新增学号检测入学时间与首次入党时间 # 暂不做
        # xuehao_col = xxx
        if not name_col:
            scr_output(scr_5,'警告：{}'.format('本表没有检测到姓名列,本次检测可能出错'))

        # 遍历每一行
        scr_output(scr_5,'\n\n正在遍历工作表，请稍后…………\n\n')
        for row in range(biaotou_row+1, worksheet.max_row+1):
            # 变量赋值
            name_value = None
            birth_value = None
            first_value = None
            positive_value = None
            object_value = None
            ready_value = None
            become_value = None

            worksheet_row = str(row)
            # 姓名列存在，进行判断
            if name_col:
                name_value = worksheet[name_col + worksheet_row].value
                if name_value==None: # 如果表格单元格里缺失值，得到的是等于None，type为NoneType类型
                    worksheet[name_col + worksheet_row].fill = color
                    scr_output(scr_5,'\n单元格{}  错误信息：{}  {}'.format(worksheet[name_col + worksheet_row].coordinate,name_value, '姓名信息缺值'))
                else:
                    if not name_legal(name_value):
                        worksheet[name_col + worksheet_row].fill = color
                        scr_output(scr_5,'\n单元格{}  错误信息：{}  {}'.format(worksheet[name_col + worksheet_row].coordinate,name_value, '姓名不合法'))
            # else:
                # scr_output(scr_5,'警告：{}'.format('本表没有姓名列,本次不能检测'))
                # break

            # 所有日期：出现点、逗号的隔离机制，出现缺少位数的隔离机制（控制大于19000000小于20500000）,出现年月日的隔离机制，月份和日期出现0或者超出12或31
            # 出生年月列存在，进行判断
            if birth_col or (not birth_col and birth2_col):
                try:
                    birth_value = worksheet[birth_col + worksheet_row].value
                except:
                    birth_value = worksheet[birth2_col + worksheet_row].value[6:14]
                if birth_col and birth_value == None:
                    worksheet[birth_col + worksheet_row].fill = color
                    scr_output(scr_5,'\n单元格{}  错误信息：{}的{}  {}'.format(worksheet[birth_col + worksheet_row].coordinate,name_value,birth_value, '出生年月信息缺值'))
                else:
                    if not data_legal(birth_value):
                        worksheet[birth_col + worksheet_row].fill = color
                        scr_output(scr_5,'\n单元格{}  错误信息：{}的{}  {}'.format(worksheet[birth_col + worksheet_row].coordinate,name_value,birth_value, '出生年月日不合法'))
                        birth_value = None
            # 检测身份证号码
            if birth2_col:
                try:
                    birth2_value = worksheet[birth2_col + worksheet_row].value
                except: pass
                if not birth2_col and birth2_value == None:
                    worksheet[birth2_col + worksheet_row].fill = color
                    scr_output(scr_5,'\n单元格{}  错误信息：{}的{}  {}'.format(worksheet[birth2_col + worksheet_row].coordinate,name_value,birth2_value, '身份证号信息缺值'))
                else:
                    if not sfz_ID_legal(birth2_value):
                        worksheet[birth2_col + worksheet_row].fill = color
                        scr_output(scr_5,'\n单元格{}  错误信息：{}的{}  {}'.format(worksheet[birth2_col + worksheet_row].coordinate,name_value,birth2_value, '身份证号不合法'))
                        birth2_value = None
            # 首次递交入党申请书时间列存在，进行判断
            if first_col:
                first_value = worksheet[first_col + worksheet_row].value
                if first_value == None:
                    worksheet[first_col + worksheet_row].fill = color
                    scr_output(scr_5,'\n单元格{}  错误信息：{}的{}  {}'.format(worksheet[first_col + worksheet_row].coordinate,name_value,first_value, '首次递交入党申请书时间缺值'))
                else:
                    if not data_legal(first_value):
                        worksheet[first_col + worksheet_row].fill = color
                        scr_output(scr_5,'\n单元格{}  错误信息：{}的{}  {}'.format(worksheet[first_col + worksheet_row].coordinate,name_value,first_value, '首次递交入党申请书时间不合法'))
                        first_value = None

            # 第一关 年满十八
            if birth_value and first_value:
                pass1_1 = """if {}:""".format(gongshi1) # int(first_value) - int(birth_value) - 180000 < 0
                pass1_2 = """
                    worksheet[birth_col + worksheet_row].fill = color2
                    worksheet[first_col + worksheet_row].fill = color2
                    scr_output(scr_5,'\\n单元格{}  错误信息：{}出生日期{} 首次递交入党申请书{}  {}'.format(worksheet[first_col + worksheet_row].coordinate,name_value,birth_value,first_value, '未满十八'))
                """
                exec(pass1_1+pass1_2) # exec() 把字符串转换成正常的句子，即：去掉引号

            # 确认为入党积极分子时间列存在，进行判断
            if positive_col:
                positive_value = worksheet[positive_col + worksheet_row].value
                if positive_value == None:
                    worksheet[positive_col + worksheet_row].fill = color
                    scr_output(scr_5,'\n单元格{}  错误信息：{}的{}  {}'.format(worksheet[positive_col + worksheet_row].coordinate, name_value, positive_value,'确认为入党积极分子时间缺值'))
                else:
                    if not data_legal(positive_value):
                        worksheet[positive_col + worksheet_row].fill = color
                        scr_output(scr_5,'\n单元格{}  错误信息：{}的{}  {}'.format(worksheet[positive_col + worksheet_row].coordinate, name_value, positive_value,'确认为入党积极分子时间不合法'))
                        positive_value = None

            # 第二关 时间正向
            if first_value and positive_value:
                # pass2_1 = """if {}:""".format(gongshi2) # int(positive_value) - int(first_value) -15 <= 0
                # pass2_2 = """
                #     worksheet[first_col + worksheet_row].fill = color2
                #     worksheet[positive_col + worksheet_row].fill = color2
                #     scr_output(scr_5,'\\n单元格{}  错误信息：{}首次递交入党申请书{} 积极分子{}  {}'.format(worksheet[positive_col + worksheet_row].coordinate,name_value,first_value,positive_value, '两次时间有误'))
                # """
                # exec(pass2_1+pass2_2)

                if is_half_year(str(first_value), str(positive_value),setting=int(gongshi2)) == False:
                    worksheet[first_col + worksheet_row].fill = color2
                    worksheet[positive_col + worksheet_row].fill = color2
                    scr_output(scr_5,'\\n单元格{}  错误信息：{}首次递交入党申请书{} 积极分子{}  {}'.format(worksheet[positive_col + worksheet_row].coordinate,name_value,first_value,positive_value, '两次时间有误'))


            # 列为发展对象时间列存在，进行判断
            if object_col:
                object_value = worksheet[object_col + worksheet_row].value
                if object_value == None:
                    worksheet[object_col + worksheet_row].fill = color
                    scr_output(scr_5,'\n单元格{}  错误信息：{}的{}  {}'.format(worksheet[object_col + worksheet_row].coordinate, name_value, object_value,'列为发展对象时间缺值'))
                else:
                    if not data_legal(object_value):
                        worksheet[object_col + worksheet_row].fill = color
                        scr_output(scr_5,'\n单元格{}  错误信息：{}的{}  {}'.format(worksheet[object_col + worksheet_row].coordinate, name_value, object_value,'列为发展对象时间不合法'))
                        object_value = None

            # 第三关 积极分子满一年
            if positive_value and object_value:
                pass3_1 = """if {}:""".format(gongshi3) # int(object_value) - int(positive_value) - 10000 <= 0
                pass3_2 = """
                    worksheet[positive_col + worksheet_row].fill = color2
                    worksheet[object_col + worksheet_row].fill = color2
                    scr_output(scr_5,'\\n单元格{}  错误信息：{}积极分子{} 发展对象{}  {}'.format(worksheet[object_col + worksheet_row].coordinate,name_value,positive_value,object_value, '两次时间未满一年或有误'))
                """
                exec(pass3_1+pass3_2)

            # 发展为预备党员时间列存在，进行判断
            if ready_col:
                ready_value = worksheet[ready_col + worksheet_row].value
                if ready_value == None:
                    worksheet[ready_col + worksheet_row].fill = color
                    scr_output(scr_5,'\n单元格{}  错误信息：{}的{}  {}'.format(worksheet[ready_col + worksheet_row].coordinate, name_value, ready_value,'发展为预备党员时间缺值'))
                else:
                    if not data_legal(ready_value):
                        worksheet[ready_col + worksheet_row].fill = color
                        scr_output(scr_5,'\n单元格{}  错误信息：{}的{}  {}'.format(worksheet[ready_col + worksheet_row].coordinate, name_value, ready_value,'发展为预备党员时间不合法'))
                        ready_value = None

            # 第四关 时间线正向
            if object_value and ready_value:
                pass4_1 = """if {}:""".format(gongshi4) # int(ready_value) - int(object_value) <= 0
                pass4_2 = """
                    worksheet[object_col + worksheet_row].fill = color2
                    worksheet[ready_col + worksheet_row].fill = color2
                    scr_output(scr_5,'\\n单元格{}  错误信息：{}列为发展对象{} 发展为预备党员{}  {}'.format(worksheet[ready_col + worksheet_row].coordinate,name_value,object_value,ready_value, '两次时间有误'))
                """
                exec(pass4_1+pass4_2)

            # 预备党员转正时间列存在，进行判断
            if become_col:
                become_value = worksheet[become_col + worksheet_row].value
                if become_value == None:
                    worksheet[become_col + worksheet_row].fill = color
                    scr_output(scr_5,'\n单元格{}  错误信息：{}的{}  {}'.format(worksheet[become_col + worksheet_row].coordinate, name_value, become_value,'预备党员转正时间缺值'))
                else:
                    if not data_legal(become_value):
                        worksheet[become_col + worksheet_row].fill = color
                        scr_output(scr_5,'\n单元格{}  错误信息：{}的{}  {}'.format(worksheet[become_col + worksheet_row].coordinate, name_value, become_value,'预备党员转正时间不合法'))
                        become_value = None

            # 第五关 转正满一年
            if ready_value and become_value:
                pass5_1 = """if {}:""".format(gongshi5) # int(become_value) - int(ready_value) >= 0
                pass5_2 = """
                    worksheet[ready_col + worksheet_row].fill = color2
                    worksheet[become_col + worksheet_row].fill = color2
                    scr_output(scr_5,'\\n单元格{}  错误信息：{}发展为预备党员{} 转正时间{}  {}'.format(worksheet[become_col + worksheet_row].coordinate,name_value,ready_value,become_value, '两次未满一年或有误'))
                """
                exec(pass5_1+pass5_2)

        # 补充检测结业证号
        zhengshu_bianhao_value_list= []  # 存放所有的证书编号
        # print(zhengshu_bianhao_col_set) #  证书编号的列 {'O', 'N', 'P'}
        if len(list(zhengshu_bianhao_col_set)) == 0:
            pass
        elif len(list(zhengshu_bianhao_col_set)) ==1:
            for row in range(biaotou_row + 1, worksheet.max_row + 1):
                zhengshu_bianhao_value = worksheet[str(list(zhengshu_bianhao_col_set)[0]) + str(row)].value
                if zhengshu_bianhao_value == None:
                    worksheet[str(list(zhengshu_bianhao_col_set)[0]) + str(row)].fill = color
                    scr_output(scr_5,'\n单元格{}  错误信息：{}  {}'.format(worksheet[str(list(zhengshu_bianhao_col_set)[0]) + str(row)].coordinate,zhengshu_bianhao_value, '证书编号信息缺值'))
                else:
                    zhengshu_bianhao_value_list.append(str(zhengshu_bianhao_value))  # 写入当前证书编号
                    if not zhengshu_bianhao_legal(zhengshu_bianhao_value):
                        worksheet[str(list(zhengshu_bianhao_col_set)[0]) + str(row)].fill = color
                        scr_output(scr_5,'\n单元格{}  错误信息：{}  {}'.format(worksheet[str(list(zhengshu_bianhao_col_set)[0]) + str(row)].coordinate,zhengshu_bianhao_value, '证书编号不合法'))
        else:
            for s in list(zhengshu_bianhao_col_set):
                for row in range(biaotou_row + 1, worksheet.max_row + 1):
                    zhengshu_bianhao_value = worksheet[str(s) + str(row)].value
                    if zhengshu_bianhao_value == None:
                        worksheet[str(s) + str(row)].fill = color
                        scr_output(scr_5, '\n单元格{}  错误信息：{}  {}'.format(worksheet[str(s) + str(row)].coordinate,zhengshu_bianhao_value, '证书编号信息缺值'))
                    else:
                        zhengshu_bianhao_value_list.append(str(zhengshu_bianhao_value))
                        if not zhengshu_bianhao_legal(zhengshu_bianhao_value):
                            worksheet[str(s) + str(row)].fill = color
                            scr_output(scr_5,'\n单元格{}  错误信息：{}  {}'.format(worksheet[str(s) + str(row)].coordinate,zhengshu_bianhao_value, '证书编号不合法'))
        if zhengshu_bianhao_value_list != []:  # 查证书编号重复值操作
            zhengshu_bianhao_value_repitition_list = []
            dict = Counter(zhengshu_bianhao_value_list)
            # print(dict)  # Counter({'1': 2, '2': 1, '3': 1, '4': 1})
            a = sorted(dict.items(), key=lambda item: item[1], reverse=True)
            # print(a)  # [('2', 3), ('3', 3), ('1', 2), ('4', 1), ('5', 1), ('10', 1)]
            for i in range(len(a)):
                if a[i][1] >= 2:
                    zhengshu_bianhao_value_repitition_list.append(a[i][0])
            if  zhengshu_bianhao_value_repitition_list != []:
                if len(list(zhengshu_bianhao_col_set)) == 1:
                    for row in range(biaotou_row + 1, worksheet.max_row + 1):
                        zhengshu_bianhao_value = worksheet[str(list(zhengshu_bianhao_col_set)[0]) + str(row)].value
                        if zhengshu_bianhao_value != None and (str(zhengshu_bianhao_value) in zhengshu_bianhao_value_repitition_list):
                            worksheet[str(list(zhengshu_bianhao_col_set)[0]) +  str(row)].fill = color2
                            scr_output(scr_5, '\n单元格{}  错误信息：{}  {}'.format(worksheet[str(list(zhengshu_bianhao_col_set)[0]) + str(row)].coordinate,zhengshu_bianhao_value, '证书编号出现重复值'))
                else:
                    for s in list(zhengshu_bianhao_col_set):
                        for row in range(biaotou_row + 1, worksheet.max_row + 1):
                            zhengshu_bianhao_value = worksheet[str(s) + str(row)].value
                            if zhengshu_bianhao_value != None and (str(zhengshu_bianhao_value) in zhengshu_bianhao_value_repitition_list):
                                worksheet[str(s) + str(row)].fill = color2
                                scr_output(scr_5, '\n单元格{}  错误信息：{}  {}'.format(worksheet[str(s) + str(row)].coordinate,zhengshu_bianhao_value, '证书编号出现重复值'))
        try:
            main5_save(workbook,filename,path)
        except Exception as error:
            error = str(error)
            print('错误提示', error)
            scr_output(scr_5, '\n生成请示文件 失败！\n\n\n本次错误信息：{}\n\n'.format(error))
            scr_output(scr_5, '\n--------文件没有成功保存--------\n\n\n\n\n\n\n')
            messagebox.showinfo('错误提示', '生成标注文件 失败！\n错误信息：\n{}'.format(error))
        return workbook,filename,path
'''#################################################################################################################'''
# 输出窗口的实时输出
def scr_output(scr,information):
    scr.config(state=NORMAL)  # 开启可写入模式
    scr.mark_set('insert', 'end') # 将光标移动到最后一行
    if information:  # 如果是event事件触发的，那么删除一个换行符
        # information = information[:-1]
        scr.insert('insert', '\n' + str(information))
        scr.update()  # 插入后及时的更新
        scr.config(state=DISABLED)  # 关闭可写入模式
        scr.see(END)  # 使得聊天记录text默认显示底端
# 点击“选择文件夹”按钮调用该功能
def select_files(scr,pathin):
    path = askdirectory(title='选择文件夹')
    pathin.set(path.strip()) # strip() 用于移除字符串头尾指定的字符（默认为空格或者换行符）或字符序列，只能是移除字符串开头和结尾部分
    scr_output(scr,'---------------------------------------------------------------------------------------\n' \
                    '---------------------------------------------------------------------------------------\n\n\n' \
                    '本次选择的文件夹路径是：\n{}\n\n'.format(pathin.get()))
# 点击“选择文件”按钮调用该功能
def select_file(scr,pathin):
    path = askopenfilename(title='选择文件')
    pathin.set(path.strip()) # strip() 用于移除字符串头尾指定的字符（默认为空格或者换行符）或字符序列，只能是移除字符串开头和结尾部分
    scr_output(scr,'---------------------------------------------------------------------------------------\n' \
                    '---------------------------------------------------------------------------------------\n\n\n' \
                    '本次选择的文件路径是：\n{}\n\n'.format(pathin.get()))
    if scr == scr_5: # 说明是时间管理的选择文件
        name = path.split('/')[-1]
        nowtime = datetime.datetime.now().strftime("%Y_%m_%d")
        pathin2_5.set(name.split('.xlsx')[0] + '（审核标注） ' + nowtime)
# 获取参数/期数
def get_canshu(number_chosen):
    # print('number2_0 系数值',number2_0.get())
    # print('number_row_0 表头选项',number_row_0.get())
    # print('number_col_0 表头选项',number_col_0.get())
    # if number_row_0.get()==1:
    #     print('表头横向')
    # if number_col_0.get()==1:
    #     print('表头纵向')
    #
    # print('number2_1 小组数',number2_1.get())
    # print('number_row_1 按支部分',number_row_1.get())
    # print('number_col_1 按核定数分',number_col_1.get())
    # if number_row_1.get()==1:
    #     print('按支部分')
    # if number_col_1.get()==1:
    #     print('按核定数分')
    # # 各个期数进行同步
    list_number_chosen = [number_chosen_0,number_chosen_1,number_chosen_2,number_chosen_3,number_chosen_4]
    for i in list_number_chosen:
        i.set(number_chosen)
    pathin2_0.set('经济管理与法学学院分党校第{}期各支部入党积极分子合并系数表'.format(number_chosen_0.get()))
    pathin4_0.set('经济管理与法学学院分党校第{}期各支部入党积极分子名额分配表'.format(number_chosen_0.get()))
    pathin6_0.set('经济管理与法学学院分党校第{}期入党积极分子培训班学员花名册'.format(number_chosen_0.get()))
    pathin2_1.set('经济管理与法学学院分党校第{}期入党积极分子培训班分组名单'.format(number_chosen_1.get()))
    pathin2_2.set('经济管理与法学学院分党校第{}期入党积极分子结业考试座位表'.format(number_chosen_2.get()))
    pathin2_3.set('经济管理与法学学院分党校第{}期入党积极分子培训班座位表'.format(number_chosen_3.get()))
    pathin2_4.set('经济管理与法学学院分党校第{}期入党积极分子培训班小组成员通讯录及考勤表'.format(number_chosen_4.get()))
    # window.mainloop()
'''#################################################################################################################'''
# 初始化案例文件
# 随机获取日期函数
def getBirthday(year_min, year_max):
    # 随机生成年月日
    year = random.randint(year_min, year_max)
    month = random.randint(1, 12)
    # 判断每个月有多少天随机生成日
    if year % 4 == 0:
        if month in (1, 3, 5, 7, 8, 10, 12):
            day = random.randint(1, 31)
        elif month in (4, 6, 9, 11):
            day = random.randint(1, 30)
        else:
            day = random.randint(1, 29)
    else:
        if month in (1, 3, 5, 7, 8, 10, 12):
            day = random.randint(1, 31)
        elif month in (4, 6, 9, 11):
            day = random.randint(1, 30)
        else:
            day = random.randint(1, 28)
    # 小于10的月份前面加0
    if month < 10:
        month = '0' + str(month)
    if day < 10:
        day = '0' + str(day)
    birthday = str(year) + str(month) + str(day)
    return birthday
# 生成 参考 各支部学员花名册 参考文件
def case_file_book():
    from faker import Faker
    fake = Faker(locale='zh_CN')

    zhibu_list = ['电物支部', '工信支部', '会计一支部', '会计二支部', '国贸支部', '经济支部', '研一支部', '研二支部', '法学支部', '人营支部']
    zhuanye_list = [["电商", "物流"], ["工管", "信管", "企管"], ["会计", "ACCA"], ["会计", "ACCA"], ["国贸"], ["经济"],
                    ["工商管理", "公共管理", "应用经济学"], ["会计", "MBA", "法律（非法学）"], ["法学"], ["人管", "营销"]]
    # 遍历写入各个支部的学员册
    for zb in range(len(zhibu_list)):
        # 新建表
        workbook = openpyxl.load_workbook('mould\model\模板0 学员花名册.xlsx')  # 打开模板表
        worksheet = workbook.worksheets[0]

        # 删除红行
        worksheet.delete_rows(4)

        # 该支部学员册文件的人数上线，随机
        count = random.randint(10, 30)
        # print(count)
        # 写入学员册里面的人
        for row in range(4, count):
            # 学号  姓名	性别	出生年月	民族	籍贯	所属院系	年级	专业班次	现任学生干部职务 首次递交入党申请书时间  确认为入党积极分子时间  是否为团员  备注
            # 随机学号
            worksheet.cell(row, 1, fake.random_int(min=20150000000, max=20210000000))

            # 随机名字和性别
            if random.randint(0, 1) == 1:
                worksheet.cell(row, 2, fake.name_female())
                worksheet.cell(row, 3, '男')
            else:
                worksheet.cell(row, 2, fake.name_male())
                worksheet.cell(row, 3, '女')

            # 随机出生日期
            worksheet.cell(row, 4, getBirthday(year_min=1996, year_max=2003))

            # 随机民族
            mz = random.randint(1, 10)
            if mz < 10:
                worksheet.cell(row, 5, '汉族')
            if mz == 10:
                mz_list = ["汉族", "蒙古族", "回族", "藏族", "维吾尔族", "苗族", "彝族", "壮族", "布依族", "朝鲜族", "满族", "侗族", "瑶族", "白族",
                           "土家族", "哈尼族", "哈萨克族", "傣族", "黎族", "僳僳族",
                           "佤族", "畲族", "高山族", "拉祜族", "水族", "东乡族", "纳西族", "景颇族", "柯尔克孜族", "土族", "达斡尔族", "仫佬族", "羌族",
                           "布朗族", "撒拉族", "毛南族", "仡佬族", "锡伯族",
                           "阿昌族", "普米族", "塔吉克族", "怒族", "乌孜别克族", "俄罗斯族", "鄂温克族", "德昂族", "保安族", "裕固族", "京族", "塔塔尔族",
                           "独龙族", "鄂伦春族", "赫哲族", "门巴族", "珞巴族", "基诺族"]
                worksheet.cell(row, 5, random.choice(mz_list))

            # 随机籍贯
            # 这里懒得核对省份城市了，后人可以完善一下
            provinces = ["北京", "上海", "天津", "重庆", "内蒙古", "山西", "河北", "吉林", "江苏", "辽宁", "黑龙江", "安徽", "山东", "浙江", "江西",
                         "福建", "湖南", "湖北",
                         "河南", "广东", "广西", "贵州", "海南", "四川", "云南", "陕西", "甘肃省", "宁夏", "青海", "新疆", "西藏", "台湾", "香港",
                         "澳门"]
            cities = ["哈尔滨", "长春", "沈阳", "呼和浩特", "石家庄", "乌鲁木齐", "兰州", "西宁", "西安", "银川", "郑州", "济南", "太原", "合肥", "武汉",
                      "长沙", "南京", "成都", "贵阳", "昆明", "南宁", "拉萨",
                      "杭州", "南昌", "广州", "福州", "台北", "海口", "郴州", "宁乡", "怀化", "太原", "辛集", "邯郸", "沈阳", "娄底", "兴城", "北镇",
                      "阜新", "哈尔滨", "衡阳", "湘西", "张家界", "常德",
                      "六安", "巢湖", "马鞍山", "永安", "宁德", "嘉禾", "荆门", "潜江", "大冶", "宜都", "佛山", "深圳", "潮州", "惠州", "汕尾", "东莞",
                      "梧州", "湘潭", "长沙", "株洲", "益阳"]
            worksheet.cell(row, 6, random.choice(provinces) + random.choice(cities))

            # yuanxi_list = ["机械工程学院","电气工程学院","核科学技术学院","资源环境与安全工程学院","计算机学院","土木工程学院","建筑学院","松霖设计艺术学院",
            #                "化学化工学院","数理学院","衡阳医学院","药学院","公共卫生学院","护理学院","经济管理与法学学院","语言文学学院","马克思主义学院",
            #                "体育学院","国际学院","船山学院","创新创业学院","继续教育学院"]
            '''# 这里提供本校的一些院系专业供后人开发'''
            # zhuanye_dict = {"机械工程学院": ("机械设计制造及其自动化", "材料成型及控制工程", "过程装备与控制工程", "测控技术与仪器", "能源与动力工程", "车辆工程"),
            #                 "电气工程学院": ("电子信息工程", "电子信息科学与技术", "电气工程及其自动化", "自动化", "通信工程", "生物医学工程"),
            #                 "计算机学院": ("软件工程", "物联网工程", "网络工程", "数字媒体技术", "医学信息工程"),
            #                 "土木工程学院": ("土木工程", "建筑环境与能源应用工程", "给排水科学与工程", "建筑电气与智能化", "道路桥梁与渡河工程"),
            #                 "化学化工学院": ("化学工程与工艺", "制药工程", "高分子材料与工程", "无机非金属材料工程"),
            #                 "核科学技术学院": ("核工程与核技术", "辐射防护与核安全", "核化工与核燃料工程", "核物理"),
            #                 "资源环境与安全工程学院": ("矿物资源工程", "矿物加工工程", "资源勘查工程", "城市地下空间工程", "安全工程", "环境工程", "环保设备工程"),
            #                 "松霖设计艺术学院": ("工业设计", "建筑学", "风景园林", "城乡规划", "视觉传达设计", "环境设计", "产品设计", "数字媒体艺术"),
            #                 "数理学院": ("信息与计算科学"),
            #                 "衡阳医学院": ("临床医学", "医学检验技术", "医学影像学", "口腔医学", "麻醉学", "儿科学", "生物技术"),
            #                 "药学院": ("药学", "药物制剂", "预防医学", "卫生检验与检疫"),
            #                 "护理学院": ("护理学"),
            #                 "经济管理与法学学院": (
            #                 "工商管理", "人力资源管理", "市场营销", "会计学", "电子商务", "物流工程", "国际经济与贸易", "经济学", "法学", "信息管理与信息系统"),
            #                 "语言文学学院": ("英语", "翻译", "日语", "汉语言文学")}
            '''# 随机院系及专业'''
            # yuanxi = random.choice(list(zhuanye_dict)) # 随机院系
            yuanxi = '经济管理与法学学院'
            zhuanye = random.choice(zhuanye_list[zb])

            worksheet.cell(row, 7, yuanxi)

            # 随机年级
            nianji_list = [x for x in range(15, 22)]
            nianji = str(random.choice(nianji_list))
            worksheet.cell(row, 8, '20{}级'.format(nianji))

            # 随机专业班次
            banji = '{}{}{}班'.format(zhuanye, nianji, random.choice([x for x in range(1, 16)]))
            worksheet.cell(row, 9, banji)

            # 随机现任学生干部职务
            zhiwu_list = ['宣传委员', '生活委员', '班长', '团支书', '副班长', '组织委员', '文娱委员', '科协委员', '心理委员']
            zw = random.randint(0, 1)
            if zw == 0:
                worksheet.cell(row, 10, '无')
            if zw == 1:
                worksheet.cell(row, 10, banji + random.choice(zhiwu_list))

            # 首次递交入党申请书时间
            scrd = getBirthday(year_min=2019, year_max=2020)
            worksheet.cell(row, 11, scrd)

            # 确认为入党积极分子时间
            jjfz = getBirthday(year_min=int(scrd[0:4]) + 1, year_max=int(scrd[:4]) + 1)
            worksheet.cell(row, 12, jjfz)

            # 是否团员
            worksheet.cell(row, 13, '是')

            # QQ
            worksheet.cell(row, 14, str(fake.phone_number())[:9])

            # 推荐支部
            worksheet.cell(row, 15, str(zhibu_list[zb]))

            # 备注
            bz = random.randint(1, 10)
            beizhu = ['抗疫志愿者', '优秀团干部']
            if bz == 10:
                worksheet.cell(row, 16, random.choice(beizhu))
            else:
                pass
        # 获取四个区域
        max_row = worksheet.max_row  # 获得最大行数
        # max_column = worksheet.max_column  # 获得最大列数
        min_row = worksheet.min_row
        # min_column = worksheet.min_column
        # 给区域设置设置框线
        for row in tuple(worksheet[min_row:max_row]):
            for cell in row:
                cell.border = my_border('thin', 'thin', 'thin', 'thin')
                # 设置单元格对齐方式 Alignment(horizontal=水平对齐模式,vertical=垂直对齐模式,text_rotation=旋转角度,wrap_text=是否自动换行)
                cell.alignment = Alignment(horizontal='center', vertical='center', text_rotation=0, wrap_text=True)

        if os.path.exists('参考 各支部学员花名册') is False:
            print("文件夹不存在")
            os.mkdir("参考 各支部学员花名册")

        workbook.save('参考 各支部学员花名册/{} 入党积极分子名册（2021上）.xlsx'.format(zhibu_list[zb]))
# 生成 参考 各支部递交入党申请书人数 参考文件
def case_file_count():
    from openpyxl.utils import get_column_letter
    # # 字典不好搞遍历？？
    # zhibu_dict = {'电物支部':("电商","物流"),'工信支部':("工管","信管","企管"),'会计一支部':("会计","ACCA"),'法学支部':("法学"),
    #                 '会计二支部':("会计","ACCA"),'国贸支部':("国贸"),'经济支部':("经济"),'人营支部':("人管","营销"),
    #               '研一支部':("工商管理","公共管理","应用经济学"),'研二支部':("会计","MBA","法律（非法学）")}
    zhibu_list = ['电物支部', '工信支部', '会计一支部', '会计二支部', '国贸支部', '经济支部', '研一支部', '研二支部', '法学支部', '人营支部']
    zhuanye_list = [["电商", "物流"], ["工管", "信管", "企管"], ["会计", "ACCA"], ["会计", "ACCA"], ["国贸"], ["经济"],
                    ["工商管理", "公共管理", "应用经济学"], ["会计", "MBA", "法律（非法学）"], ["法学"], ["人管", "营销"]]

    # 遍历写入每个支部文件
    for zb in range(len(zhibu_list)):
        # 新建表
        workbook = openpyxl.Workbook()
        worksheet = workbook.worksheets[0]

        # 写入表头
        worksheet.cell(1, 1, '支部')
        worksheet.cell(1, 2, '班级')
        worksheet.cell(1, 3, '递交入党申请书人数')

        print('正在遍历支部是：{}'.format(zhibu_list[zb]))
        row = 2  # 从第二行开始写
        # 遍历写入专业
        zhuanye = zhuanye_list[zb]
        nianji = ['17', '18', '19', '20']
        for bj in range(len(zhuanye)):
            print('正在遍历专业是：{}'.format(zhuanye[bj]))
            # 随机写入班级   zhuanye_list[zb][bj]==每个专业
            class_names = []
            for i in range(random.randint(4, 12)):  # 设置人数概率
                class_n = zhuanye[bj] + random.choice(nianji) + str(random.randint(1, 5)) + '班'
                if class_n not in class_names:
                    worksheet.cell(row, 2, class_n)  # 写入班级全名
                    worksheet.cell(row, 3, random.randint(0, 12))  # 写入随机的递交入党申请书人数
                    row = row + 1
                class_names.append(class_n)

        # print('本次写到了行数：{}'.format(row))

        worksheet.cell(2, 1, zhibu_list[zb])  # 写入支部名字
        worksheet.merge_cells(start_row=2, start_column=1, end_row=row - 1, end_column=1)  # 合并单元格

        # 获取四个区域
        max_row = worksheet.max_row  # 获得最大行数
        max_column = worksheet.max_column  # 获得最大列数
        min_row = worksheet.min_row
        min_column = worksheet.min_column

        # 给区域设置设置框线
        for r in tuple(worksheet[min_row:max_row]):
            for cell in r:
                cell.border = my_border('thin', 'thin', 'thin', 'thin')
                # 设置单元格对齐方式 Alignment(horizontal=水平对齐模式,vertical=垂直对齐模式,text_rotation=旋转角度,wrap_text=是否自动换行)
                cell.alignment = Alignment(horizontal='center', vertical='center', text_rotation=0, wrap_text=True)

        # 区域自动调整列宽
        column_widths = []  # 定义用来获取当前列最大宽度的空列表
        for i, col in enumerate(
                worksheet.iter_cols(min_col=min_column, max_col=max_column, min_row=min_row, max_row=max_row)):
            for cell in col:
                value = cell.value
                if value is not None:
                    if isinstance(value, str) is False:
                        value = str(value)
                    try:
                        column_widths[i] = max(column_widths[i], len(value))
                    except IndexError:
                        column_widths.append(len(value))
        # print('column_widths', column_widths)  # 得到该列最大的一个单元格的宽度（字符串数量）
        for i, width in enumerate(column_widths):
            col_name = get_column_letter(min_column + i)  # 获取行字母表头
            value = column_widths[i] * 2  # 设置列宽为最大长度比例
            worksheet.column_dimensions[col_name].width = value

        if os.path.exists('参考 各支部递交入党申请书人数') is False:
            print("文件夹不存在")
            os.mkdir("参考 各支部递交入党申请书人数")

        workbook.save('参考 各支部递交入党申请书人数/{} 递交入党申请书人数（2021上）.xlsx'.format(zhibu_list[zb]))
'''#################################################################################################################'''
# 请示文件cookie的模板的识别（后续开发做准备）
def qingshi_model_cookie(cookie,yeardu,pici,year_up,year,month,day,party_name,party_num,first_people,people_num,people_sheet):
    if cookie == '100':  # 发展对象的请示
        a = "关于建议接收{}年度经济管理与法学学院".format(yeardu)
        b = "{}发展对象的请示".format(pici)
        c = "尊敬的院党委："
        d = "经过{}等{}个学生党支部委员会充分研究讨论，确认{}等{}名同志为{}年{}半年发展对象人选，建议学院党委将{}等{}名同志列为中共党员发展" \
            "对象，名单如下（{}）：".format(party_name, party_num, first_people, people_num, year, year_up, first_people,
                                       people_num,'排名以首字母为序' if peoplename.get() == 1 else '排名以班级为序')
        e = "请批示。"
        f = "经济管理与法学学院学生党建工作委员会"
        g = "{}年{}月{}日".format(year, month, day)
    if cookie == '010':  # 预备党员的请示
        a = "关于建议接收{}年度经济管理与法学学院".format(yeardu)
        b = "{}预备党员的请示".format(pici)
        c = "尊敬的院党委："
        d = "经过{}等{}个学生党支部召开支部大会充分讨论，认为{}等{}名同志符合预备党员的条件。现拟提请学院党委接受" \
            "{}等{}名同志为中共预备党员，名单如下（{}）：".format(party_name, party_num, first_people, people_num, first_people,
                                                    people_num,'排名以首字母为序' if peoplename.get() == 1 else '排名以班级为序')
        e = "请批示。"
        f = "经济管理与法学学院学生党建工作委员会"
        g = "{}年{}月{}日".format(year, month, day)
    if cookie == '001':  # 预备党员转正的请示
        a = "关于建议接收{}年度经济管理与法学学院".format(yeardu)
        b = "{}预备党员转正的请示".format(pici)
        c = "尊敬的院党委："
        d = "经过{}等{}个学生党支部召开支部大会充分讨论，确认{}等{}名同志为{}年{}半年预备党员转正人选，建议学院党委将" \
            "{}等{}名同志列为中共党员，名单如下（{}）：".format(party_name, party_num, first_people, people_num, year, year_up,
                                                   first_people, people_num,'排名以首字母为序' if peoplename.get() == 1 else '排名以班级为序' )
        e = "请批示。"
        f = "经济管理与法学学院学生党建工作委员会"
        g = "{}年{}月{}日".format(year, month, day)
    return a,b,c,d,e,f,g
# 请示文件的写入
def write_qingshi(cookie,yeardu,pici,year_up,year,month,day,party_name,party_num,first_people,people_num,people_sheet):
    print(peoplename.get())
    people_sheet = (sorted(scr_sheet6.get(1.0, 'end').split(), key=lambda keys: [pinyin(i, style=Style.TONE3) for i in keys]) if peoplename.get() == 1 else scr_sheet6.get(1.0, 'end').split())
    first_people = people_sheet[0]
    try:
        if type(people_sheet) is str: people_sheet = people_sheet.split()
        if cookie == '000':
            messagebox.showinfo('错误提示', '未选中请示的类型，请检查！')
            return
        if people_num != len(people_sheet):
            scr_output(scr_6, '\n生成请示文件 失败！\n错误信息：同志人数{}与人名数量{}不匹配，请检查！\n'.format(people_num,len(people_sheet)))
            messagebox.showinfo('错误提示', '生成请示文件 失败！\n错误信息：同志人数{}与人名数量{}不匹配，请检查！'.format(people_num,len(people_sheet)))
            return
        a,b,c,d,e,f,g = qingshi_model_cookie(cookie, yeardu, pici, year_up, year, month, day, party_name, party_num,first_people, people_num,people_sheet) # 执行下面注释代码的函数
        # if cookie=='100': # 发展对象的请示
        #     a = "关于建议接收{}年度经济管理与法学学院".format(yeardu)
        #     b = "{}发展对象的请示".format(pici)
        #     c = "尊敬的院党委："
        #     d = "经过{}等{}个学生党支部委员会充分研究讨论，确认{}等{}名同志为{}年{}半年发展对象人选，建议学院党委将{}等{}名同志列为中共党员发展" \
        #         "对象，名单如下（排名以班级为序）：".format(party_name,party_num,first_people,people_num,year,year_up,first_people,people_num)
        #     e = "请批示。"
        #     f = "经济管理与法学学院学生党建工作委员会"
        #     g = "{}年{}月{}日".format(year,month,day)
        # if cookie=='010': # 预备党员的请示
        #     a = "关于建议接收{}年度经济管理与法学学院".format(yeardu)
        #     b = "{}预备党员的请示".format(pici)
        #     c = "尊敬的院党委："
        #     d = "经过{}等{}个学生党支部召开支部大会充分讨论，认为{}等{}名同志符合预备党员的条件。现拟提请学院党委接受" \
        #         "{}等{}名同志为中共预备党员，名单如下（排名以班级为序）：" .format(party_name,party_num,first_people,people_num,first_people,people_num,)
        #     e = "请批示。"
        #     f = "经济管理与法学学院学生党建工作委员会"
        #     g = "{}年{}月{}日".format(year,month,day)
        # if cookie=='001': # 预备党员转正的请示
        #     a = "关于建议接收{}年度经济管理与法学学院".format(yeardu)
        #     b = "{}预备党员转正的请示".format(pici)
        #     c = "尊敬的院党委："
        #     d = "经过{}等{}个学生党支部召开支部大会充分讨论，确认{}等{}名同志为{}年{}半年预备党员转正人选，建议学院党委将" \
        #         "{}等{}名同志列为中共党员，名单如下（排名以班级为序）：" .format(party_name,party_num,first_people,people_num,year,year_up,first_people,people_num,)
        #     e = "请批示。"
        #     f = "经济管理与法学学院学生党建工作委员会"
        #     g = "{}年{}月{}日".format(year,month,day)
        doc = Document()
        # 判断人数，来设置表格
        if 0 <= people_num<= 64: # 四号字体
            doc.styles['Normal'].paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE  # 1.5倍行距
            doc.styles['Normal'].font.size = Pt(14)
            col_width = [2.43,1.9]
            row_height = 1
        if 64 < people_num<= 104: # 小四字体
            doc.styles['Normal'].paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE # 1.5倍行距
            doc.styles['Normal'].font.size = Pt(12)
            col_width = [2.15,1.8]
            row_height = 0.9
        if 104 < people_num <= 120:  # 小四字体
            doc.styles['Normal'].paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE  # 1.5倍行距
            doc.styles['Normal'].font.size = Pt(12)
            col_width = [2.15, 1.8]
            row_height = 0.8
        if 120 < people_num<= 168: # 小四字体
            doc.styles['Normal'].paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST # 最小倍倍行距
            doc.styles['Normal'].font.size = Pt(12)
            col_width = [1.98,1.8]
            row_height = 0.55
        if 168 < people_num<= 184: # 五号字体
            doc.styles['Normal'].paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST # 最小倍倍行距
            doc.styles['Normal'].font.size = Pt(10.5)
            col_width = [1.98,1.8]
            row_height = 0.55
        if 184 < people_num:
            doc.styles['Normal'].font.size = Pt(10)
            col_width = [1.98, 1.8]
            row_height = 0.55
            print('人数太多（大于184），请自行调整word中存在的格式问题。')
            scr_output(scr_6, '\n\n人数太多（大于184），请自行调整word中存在的格式问题。')
        # 标题样式
        doc.styles['Header'].font.name = 'Times New Roman'  # 设置英文字体
        doc.styles['Header']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')  # 设置中文字体
        doc.styles['Header'].font.bold = True  # 加粗
        doc.styles['Header'].font.size = Pt(16)
        doc.styles['Header'].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 居中对齐
        doc.styles['Header'].paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE # 1.5倍行距
        doc.styles['Header'].paragraph_format.space_before = Pt(0)  # 段前
        doc.styles['Header'].paragraph_format.space_after = Pt(0)  # 段后
        # 普通正文央视
        doc.styles['Footer'].font.name = 'Times New Roman'  # 设置英文字体
        doc.styles['Footer']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')  # 设置中文字体
        doc.styles['Footer'].font.size = Pt(14)
        doc.styles['Footer'].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY # 两端对齐
        doc.styles['Footer'].paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE # 1.5倍行距
        doc.styles['Footer'].paragraph_format.space_before = Pt(0)  # 段前
        doc.styles['Footer'].paragraph_format.space_after = Pt(0)  # 段后
        doc.styles['Footer'].paragraph_format.first_line_indent = doc.styles['Footer'].font.size * 2  #首行缩进2字符 1 英寸=2.54 厘米
        # 表格样式
        doc.styles['Normal'].font.name = 'Times New Roman'  # 设置英文字体
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')  # 设置中文字体
        # doc.styles['Normal'].font.size = Pt(12)
        doc.styles['Normal'].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.DISTRIBUTE # 分散对齐
        # doc.styles['Normal'].paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST # 最小倍倍行距
        doc.styles['Normal'].paragraph_format.space_before = Pt(0)  # 段前
        doc.styles['Normal'].paragraph_format.space_after = Pt(0)  # 段后
        doc.styles['Normal'].paragraph_format.first_line_indent = Inches(0)  #首行缩进2字符 1 英寸=2.54 厘米

        # 标题 两段
        doc.add_paragraph(a,style='Header')
        doc.add_paragraph(b,style='Header')
        # 称呼 一段 （首不设两字符）
        doc.add_paragraph(c,style='Footer').paragraph_format.first_line_indent=Inches(0) # 1 英寸=2.54 厘米
        # 正文
        doc.add_paragraph(d,style='Footer')

        table = doc.add_table(people_num//8 if people_num%8 == 0 else people_num//8+1, 8)
        table.autofit = True   # if is True 按窗口大小自动调整
        count = 0

        for row in range(len(table.rows)):
            table.rows[row].height = Cm(row_height)  # 调整行高
            for col in range(len(table.columns)):
                # print(行, 列)  # 可以查看表格输出结果
                table.cell(row, col).text = people_sheet[count]    # 写入人名
                # table.cell(行, 列).width = doc.styles['Normal'].font.size * len(people_sheet[count])
                # table.cell(行, 列).height = doc.styles['Normal'].font.size * 5
                table.cell(row, col).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER  # 上下居中（中部居中对齐）
                # table.cell(行, 列).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 水平居中（中部居中对齐）
                count += 1
                if count == people_num:
                    break
            if count == people_num:
                break
        # 修正列宽
        for col in range(len(table.columns)):
            maxlist = []
            for r in range(len(table.rows)):
                try:
                    maxlist.append(len(people_sheet[8*r + col]))
                    # print(people_sheet[8*r + col])
                except:pass
            if maxlist != []: maxnum = max(maxlist) # 每一列的最大值
            else: maxnum = 3  # 每一列的最大值
            table.cell(len(table.rows)-1, col).width = Cm( col_width[0] if maxnum==4 else col_width[1] ) # 调整列宽 2字:1.3 3字:1.8 4字:2.1
            # 要在最后一行设置列宽度，因为设置前面的，后面一行出现空格，前面设置的宽度就不生效了

        table.alignment = WD_TABLE_ALIGNMENT.CENTER  # 设置整个表格为居中对齐
        # table.autofit = True
        # 结束语
        doc.add_paragraph(e,style='Footer')
        doc.add_paragraph("",style='Footer')
        # 落款和时间
        doc.add_paragraph(f,style='Footer').alignment = WD_ALIGN_PARAGRAPH.RIGHT  # 右对齐
        doc.add_paragraph(g,style='Footer').alignment = WD_ALIGN_PARAGRAPH.RIGHT


        doc.save("关于建议接收{}等{}名同志为{}年度经济管理与法学学院".format(first_people,people_num,yeardu) + b + '.docx')
        messagebox.showinfo('小提示', '生成请示文件 成功！请注意检查word文件格式！')
        scr_output(scr_6, '\n\n生成请示文件 成功！请注意检查word文件格式！\n')

    except Exception as error:
        error = str(error)
        print('错误提示', error)
        scr_output(scr_6, '\n生成请示文件 失败！\n\n\n本次错误信息：{}\n\n'.format(error))
        scr_output(scr_6, '\n--------文件没有成功保存--------\n\n\n\n\n\n\n')
        messagebox.showinfo('错误提示', '生成请示文件 失败！\n错误信息：\n{}'.format(error))
# 请示管理 自动检测姓名列 更新多个变量值
def auto_qingshi_read():
    if scr_sheet6.get(1.0, 'end').split() != []:
        messagebox.showinfo('小提示', '已经识别到文本中已有人名，请勿重复生成，请检查！'+'\n'
                            +'如需重新生成，请记得Ctrl+A清除不需要的人名，以防出错！'+'\n'
                            +'注意：本提示只是温馨提示，是不会停止继续执行自动检测的')
    # print(pathin_6.get())
    if pathin_6.get() != '':
        path = pathin_6.get()

        xls_files = [x for x in os.listdir(path) if os.path.splitext(x)[1] == '.xls']
        if xls_files != []:  # 说明有xls文件
            xls_to_xlsx(path=path, sole=False)  # 给路径，让其自己转换成xlsx的
            scr_output(scr_6, '\n\n检测到有{}个xls格式文件，已经自动在原路径转换生成可读取的xlsx文件类型：\n'.format(len(xls_files)))

        xlsx_files = [x for x in os.listdir(path) if os.path.splitext(x)[1] == '.xlsx']  # 罗列当前目录内所有xlsx文件
        scr_output(scr_6, '\n\n需要提取名字的表格{}个'.format(len(xlsx_files)))
        scr_output(scr_6, '\n\n需要提取名字的表格有：\n{}'.format(xlsx_files))
        print('需要提取', len(xlsx_files), '个表格')
        print('提取表格有：\n', xlsx_files)  # 本目录下的xlsx文件名字列表
        list_names = []
        for p in xlsx_files:
            r, c = None, None
            xlsx_file = path + '/' + p
            workbook = openpyxl.load_workbook(filename=xlsx_file)
            worksheet = workbook.worksheets[0]
            # 获取名字信息
            for row in tuple(worksheet[1:3]):
                for cell in row:
                    # print(cell.value)
                    if cell.value == ('姓名' or '名字' or '名称'):
                        r = cell.row
                        c = cell.column
                        break
            if r != None and c != None:
                # print(r, c)
                # print(worksheet[c])
                list_name = list(cell.value for cell in [col for col in worksheet.columns][c - 1])[r:]
                for i in list_name:
                    list_names.append(i)
                list_names.append('\n') # 遍历完每个支部加换个行
                scr_output(scr_6, '\n\n提取出来的名单：\n{}'.format(list_name))
                print('\n\n提取出来的名单：\n{}'.format(list_name))
            else:
                print('找不到名字，请手动输入！')
                scr_output(scr_6, '\n找不到名字，请手动输入！\n')
        # print(list_names)
        scr_sheet6.insert('insert', ' '.join(i for i in list_names)) # 插入名字
        scr_sheet6.update()  # 插入后及时的更新
        scr_sheet6.see(END)  # 使得聊天记录text默认显示底端
    else:
        print('路径为空！')
        scr_output(scr_6, '\n路径为空！\n')
    people_sheet = scr_sheet6.get(1.0, 'end').split()
    if people_sheet != []:
        number11_6.set(people_sheet[0])
        number12_6.set(len(people_sheet))
    else: print('名单为空，请检查！')
# 请示模板修改，未完善
def qingshi_model_alter():
    messagebox.showinfo('小提示', '本版本只支持查看请示模板，暂不支持修改')
    cookie = str(number_6_1.get()) + str(number_6_2.get()) + str(number_6_3.get())
    if cookie == '000':
        messagebox.showinfo('错误提示', '未选中请示的类型，请检查！')
        return
    if cookie == '100':  # 发展对象的请示
        a = "关于建议接收{}年度经济管理与法学学院"
        b = "{}发展对象的请示"
        c = "尊敬的院党委："
        d = "经过{}等{}个学生党支部委员会充分研究讨论，确认{}等{}名同志为{}年{}半年发展对象人选，建议学院党委将{}等{}名同志列为中共党员发展" \
            "对象，"
        d2 = "名单如下（{}）：".format('排名以首字母为序' if peoplename.get() == 1 else '排名以班级为序')
        e = "{sheet}\n请批示。"
        f = "经济管理与法学学院学生党建工作委员会"
        g = "{}年{}月{}日"
    if cookie == '010':  # 预备党员的请示
        a = "关于建议接收{}年度经济管理与法学学院"
        b = "{}预备党员的请示"
        c = "尊敬的院党委："
        d = "经过{}等{}个学生党支部召开支部大会充分讨论，认为{}等{}名同志符合预备党员的条件。现拟提请学院党委接受{}等{}名同志为中共预备党员，"
        d2 = "名单如下（{}）：".format('排名以首字母为序' if peoplename.get() == 1 else '排名以班级为序')
        e = "{sheet}\n请批示。"
        f = "经济管理与法学学院学生党建工作委员会"
        g = "{}年{}月{}日"
    if cookie == '001':  # 预备党员转正的请示
        a = "关于建议接收{}年度经济管理与法学学院"
        b = "{}预备党员转正的请示"
        c = "尊敬的院党委："
        d = "经过{}等{}个学生党支部召开支部大会充分讨论，确认{}等{}名同志为{}年{}半年预备党员转正人选，建议学院党委将{}等{}名同志列为中共党员，"
        d2 = "名单如下（{}）：".format('排名以首字母为序' if peoplename.get() == 1 else '排名以班级为序')
        e = "{sheet}\n请批示。"
        f = "经济管理与法学学院学生党建工作委员会"
        g = "{}年{}月{}日"
    list_qingshi_model = [a,b,c,d,d2,e,f,g]
    def qingshi_model_save():
        scr_output(scr_6,'\n{}\n请示模板保存失败！，本版本模板不支持修改！\n'.format(qingshi_model_var.get()))
        qingshi_model.destroy()

    def qingshi_model_default():
        scr_output(scr_6,'\n模板已经是默认！\n')

    qingshi_model = Toplevel(window)
    qingshi_model.geometry("500x290+700+270")
    try:
        qingshi_model.iconbitmap('mould\data\logo.ico')
    except:pass
    # 窗口置顶
    qingshi_model.attributes("-topmost", 1)  # 1==True 处于顶层
    # 禁止窗口的拉伸
    qingshi_model.resizable(0, 0)
    # 窗口的标题
    qingshi_model.title("内置-{}请示模板-修改窗口".format(b[2:6]))

    # 定义变量
    qingshi_model_var= StringVar()
    scr_qingshi_model = scrolledtext.ScrolledText(qingshi_model, wrap=WORD)
    scr_qingshi_model.place(x=10, y=10, width=480,height=245)
    scr_qingshi_model.config(state=DISABLED)  # 关闭可写入模式
    for i in list_qingshi_model:
        scr_output(scr_qingshi_model, str(i) + '\n')

    button_qingshi_model = ttk.Button(qingshi_model, text="保存参数", command=qingshi_model_save)
    button_qingshi_model.place(x=250, y=260)

    button_qingshi_model = ttk.Button(qingshi_model, text="恢复默认", command=qingshi_model_default)
    button_qingshi_model.place(x=120, y=260)

    # 显示窗口(消息循环)
    qingshi_model.mainloop()
'''#################################################################################################################'''
# 批复文件cookie的模板的识别（后续开发做准备）
def pifu_model_cookie(cookie, yeardu, pici, year_up, qs_year,qs_month,qs_day,qingshi_name, year, month, day,party_name, party_num, first_people, people_num,people_sheet):
    if cookie == '100':  # 发展对象的批复
        a = "关于同意接收{}年度经济管理与法学学院".format(yeardu)
        b = "{}发展对象的批复".format(pici)
        c = "中共南华大学经济管理与法学学院"
        d = "{}等{}个学生党支部：".format(party_name,party_num)
        e = "收到了贵支部{}年{}月{}日“{}”，且公示无异议。" \
            "认为你们按照党员标准对入党积极分子进行了有效的培养和教育。根据《中国共产党发展党员工作细则》的要求，" \
            "院党委于{}年{}月{}日召开党委会，认真讨论和审核材料，现将{}等" \
            "{}名同志列为中共党员发展对象，名单如下（{}）：" \
            "".format(qs_year,qs_month,qs_day,qingshi_name,year,month,day,first_people,people_num,'排名以首字母为序' if peoplename.get() == 1 else '排名以班级为序')
        f = "望你们继续加强对发展对象的培养和考察。"
        g = "特此批复。"
        h = "党委书记签名：_______________"
        i = "中共南华大学经济管理与法学学院委员会（盖章）"
        j = "{}年{}月{}日".format(year, month, day)
    if cookie == '010':  # 预备党员的批复
        a = "关于同意接收{}年度经济管理与法学学院".format(yeardu)
        b = "{}预备党员的批复".format(pici)
        c = "中共南华大学经济管理与法学学院"
        d = "{}等{}个学生党支部：".format(party_name,party_num)
        e = "收到了贵支部{}年{}月{}日“{}”，且公示无异议。" \
            "认为你们按照党员标准对发展对象进行了有效的培养和教育。根据《中国共产党发展党员工作细则》的要求，" \
            "院党委于{}年{}月{}日召开党委会，认真讨论和审核材料，现确定" \
            "{}等{}名同志为中共预备党员，名单如下（{}）：" \
            "".format(qs_year,qs_month,qs_day,qingshi_name,year,month,day,first_people,people_num,'排名以首字母为序' if peoplename.get() == 1 else '排名以班级为序')
        f = "望你们继续加强对预备党员的培养和考察。"
        g = "特此批复。"
        h = "党委书记签名：_______________"
        i = "中共南华大学经济管理与法学学院委员会（盖章）"
        j = "{}年{}月{}日".format(year, month, day)
    if cookie == '001':  # 预备党员转正的批复
        a = "关于同意接收{}年度经济管理与法学学院".format(yeardu)
        b = "{}预备党员转正的批复".format(pici)
        c = "中共南华大学经济管理与法学学院"
        d = "{}等{}个学生党支部：".format(party_name,party_num)
        e = "{}等{}名同志向党支部提出了转为正式党员的书面申请。" \
            "院党委于{}年{}月{}日召开党委会，讨论通过{}等{}名同志预备党员转为正式党员的决议。" \
            "{}等{}同志从预备期满之日起成为中国共产党正式党员，党龄从即日算起。名单如下（{}）：" \
            "".format(first_people,people_num,year,month,day,first_people,people_num,first_people,people_num,'排名以首字母为序' if peoplename.get() == 1 else '排名以班级为序')
        f = None
        g = "特此批复。"
        h = "党委书记签名：_______________"
        i = "中共南华大学经济管理与法学学院委员会（盖章）"
        j = "{}年{}月{}日".format(year, month, day)
    return a, b, c, d, e, f, g, h, i, j
# 批复文件的写入
def write_pifu(cookie, yeardu, pici, year_up, qs_year,qs_month,qs_day,qingshi_name, year, month, day, party_name, party_num, first_people, people_num,people_sheet):
    people_sheet = (sorted(scr_sheet7.get(1.0, 'end').split(), key=lambda keys: [pinyin(i, style=Style.TONE3) for i in keys]) if peoplename.get() == 1 else scr_sheet7.get(1.0, 'end').split())
    first_people = people_sheet[0]
    try:
        if type(people_sheet) is str: people_sheet = people_sheet.split()
        if cookie == '000':
            messagebox.showinfo('错误提示', '未选中批复的类型，请检查！')
            return
        if people_num != len(people_sheet):
            scr_output(scr_7, '\n生成批复文件 失败！\n错误信息：同志人数{}与人名数量{}不匹配，请检查！\n'.format(people_num, len(people_sheet)))
            messagebox.showinfo('错误提示', '生成批复文件 失败！\n错误信息：同志人数{}与人名数量{}不匹配，请检查！'.format(people_num, len(people_sheet)))
            return
        a, b, c, d, e, f, g, h, i, j = pifu_model_cookie(cookie, yeardu, pici, year_up, qs_year,qs_month,qs_day,qingshi_name,
                                                year, month, day, party_name, party_num,first_people, people_num, people_sheet)
        doc = Document()
        # 判断人数，来设置表格
        if 0 <= people_num <= 64:  # 四号字体
            doc.styles['Normal'].paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE  # 1.5倍行距
            doc.styles['Normal'].font.size = Pt(14)
            col_width = [2.43, 1.9]
            row_height = 1
        if 64 < people_num <= 80:  # 小四字体
            doc.styles['Normal'].paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE  # 1.5倍行距
            doc.styles['Normal'].font.size = Pt(12)
            col_width = [2.15, 1.8]
            row_height = 0.9
        if 80 < people_num <= 88:  # 小四字体
            doc.styles['Normal'].paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE  # 1.5倍行距
            doc.styles['Normal'].font.size = Pt(12)
            col_width = [2.15, 1.8]
            row_height = 0.8
        if 88 < people_num <= 120:  # 小四字体
            doc.styles['Normal'].paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST  # 最小倍倍行距
            doc.styles['Normal'].font.size = Pt(12)
            col_width = [1.98, 1.8]
            row_height = 0.55
        if 120 < people_num <= 136:  # 五号字体
            doc.styles['Normal'].paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST  # 最小倍倍行距
            doc.styles['Normal'].font.size = Pt(10.5)
            col_width = [1.98, 1.8]
            row_height = 0.55
        if 136 < people_num:
            doc.styles['Normal'].font.size = Pt(10)
            col_width = [1.98, 1.8]
            row_height = 0.55
            print('人数太多（大于184），请自行调整word中存在的格式问题。')
            scr_output(scr_7, '人数太多（大于184），请自行调整word中存在的格式问题。')
        # 标题样式
        doc.styles['Header'].font.name = 'Times New Roman'  # 设置英文字体
        doc.styles['Header']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')  # 设置中文字体
        doc.styles['Header'].font.bold = True  # 加粗
        doc.styles['Header'].font.size = Pt(16)
        doc.styles['Header'].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 居中对齐
        doc.styles['Header'].paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE  # 1.5倍行距
        doc.styles['Header'].paragraph_format.space_before = Pt(0)  # 段前
        doc.styles['Header'].paragraph_format.space_after = Pt(0)  # 段后
        # 普通正文央视
        doc.styles['Footer'].font.name = 'Times New Roman'  # 设置英文字体
        doc.styles['Footer']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')  # 设置中文字体
        doc.styles['Footer'].font.size = Pt(14)
        doc.styles['Footer'].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY  # 两端对齐
        doc.styles['Footer'].paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE  # 1.5倍行距
        doc.styles['Footer'].paragraph_format.space_before = Pt(0)  # 段前
        doc.styles['Footer'].paragraph_format.space_after = Pt(0)  # 段后
        doc.styles['Footer'].paragraph_format.first_line_indent = doc.styles[
                                                                      'Footer'].font.size * 2  # 首行缩进2字符 1 英寸=2.54 厘米
        # 表格样式
        doc.styles['Normal'].font.name = 'Times New Roman'  # 设置英文字体
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')  # 设置中文字体
        # doc.styles['Normal'].font.size = Pt(12)
        doc.styles['Normal'].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.DISTRIBUTE  # 分散对齐
        # doc.styles['Normal'].paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST # 最小倍倍行距
        doc.styles['Normal'].paragraph_format.space_before = Pt(0)  # 段前
        doc.styles['Normal'].paragraph_format.space_after = Pt(0)  # 段后
        doc.styles['Normal'].paragraph_format.first_line_indent = Inches(0)  # 首行缩进2字符 1 英寸=2.54 厘米

        # 标题 两段
        doc.add_paragraph(a, style='Header')
        doc.add_paragraph(b, style='Header')
        # 称呼两段（首不设两字符）
        doc.add_paragraph(c, style='Footer').paragraph_format.first_line_indent = Inches(0)  # 1 英寸=2.54 厘米
        doc.add_paragraph(d, style='Footer').paragraph_format.first_line_indent = Inches(0)  # 1 英寸=2.54 厘米
        # 正文
        doc.add_paragraph(e, style='Footer')

        table = doc.add_table(people_num // 8 if people_num % 8 == 0 else people_num // 8 + 1, 8)
        table.autofit = True  # if is True 按窗口大小自动调整
        count = 0

        for row in range(len(table.rows)):
            table.rows[row].height = Cm(row_height)  # 调整行高
            for col in range(len(table.columns)):
                # print(行, 列)  # 可以查看表格输出结果
                table.cell(row, col).text = people_sheet[count]  # 写入人名
                # table.cell(行, 列).width = doc.styles['Normal'].font.size * len(people_sheet[count])
                # table.cell(行, 列).height = doc.styles['Normal'].font.size * 5
                table.cell(row, col).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER  # 上下居中（中部居中对齐）
                # table.cell(行, 列).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 水平居中（中部居中对齐）
                count += 1
                if count == people_num:
                    break
            if count == people_num:
                break
        # 修正列宽
        for col in range(len(table.columns)):
            maxlist = []
            for r in range(len(table.rows)):
                try:
                    maxlist.append(len(people_sheet[8 * r + col]))
                    # print(people_sheet[8*r + col])
                except:
                    pass
            if maxlist != []:
                maxnum = max(maxlist)  # 每一列的最大值
            else:
                maxnum = 3  # 每一列的最大值
            table.cell(len(table.rows) - 1, col).width = Cm(
                col_width[0] if maxnum == 4 else col_width[1])  # 调整列宽 2字:1.3 3字:1.8 4字:2.1
            # 要在最后一行设置列宽度，因为设置前面的，后面一行出现空格，前面设置的宽度就不生效了

        table.alignment = WD_TABLE_ALIGNMENT.CENTER  # 设置整个表格为居中对齐
        # table.autofit = True
        # 结束语
        if f != None:
            doc.add_paragraph(f, style='Footer')
        doc.add_paragraph(g, style='Footer')
        doc.add_paragraph("", style='Footer')
        # 落款和时间
        doc.add_paragraph(h, style='Footer').alignment = WD_ALIGN_PARAGRAPH.RIGHT  # 右对齐
        doc.add_paragraph(i, style='Footer').alignment = WD_ALIGN_PARAGRAPH.RIGHT
        doc.add_paragraph(j, style='Footer').alignment = WD_ALIGN_PARAGRAPH.RIGHT

        doc.save("关于同意接收{}等{}名同志为{}年度经济管理与法学学院".format(first_people, people_num, yeardu) + b + '.docx')
        messagebox.showinfo('小提示', '生成批复文件 成功！请注意检查word文件格式！')
        scr_output(scr_7, '\n\n生成批复文件 成功！请注意检查word文件格式！\n')

    except Exception as error:
        error = str(error)
        print('错误提示', error)
        scr_output(scr_7, '\n生成批复文件 失败！\n\n\n本次错误信息：{}\n\n'.format(error))
        scr_output(scr_7, '\n--------文件没有成功保存--------\n\n\n\n\n\n\n')
        messagebox.showinfo('错误提示', '生成批复文件 失败！\n错误信息：\n{}'.format(error))
# 批复管理 自动检测姓名列 更新多个变量值
def auto_pifu_read():
    if scr_sheet7.get(1.0, 'end').split() != []:
        messagebox.showinfo('小提示', '已经识别到文本中已有人名，请勿重复生成，请检查！' + '\n'
                            + '如需重新生成，请记得Ctrl+A清除不需要的人名，以防出错！' + '\n'
                            + '注意：本提示只是温馨提示，是不会停止继续执行自动检测的')
    # print(pathin_7.get())
    # 如果路径不为空，写入名单
    if pathin_7.get() != '':
        path = pathin_7.get()

        xls_files = [x for x in os.listdir(path) if os.path.splitext(x)[1] == '.xls']
        if xls_files != []:  # 说明有xls文件
            xls_to_xlsx(path=path, sole=False)  # 给路径，让其自己转换成xlsx的
            scr_output(scr_7, '\n\n检测到有{}个xls格式文件，已经自动在原路径转换生成可读取的xlsx文件类型：\n'.format(len(xls_files)))

        xlsx_files = [x for x in os.listdir(path) if os.path.splitext(x)[1] == '.xlsx']  # 罗列当前目录内所有xlsx文件
        scr_output(scr_7, '\n\n需要提取名字的表格{}个'.format(len(xlsx_files)))
        scr_output(scr_7, '\n\n需要提取名字的表格有：\n{}'.format(xlsx_files))
        print('需要提取', len(xlsx_files), '个表格')
        print('提取表格有：\n', xlsx_files)  # 本目录下的xlsx文件名字列表
        list_names = []
        for p in xlsx_files:
            r, c = None, None
            xlsx_file = path + '/' + p
            workbook = openpyxl.load_workbook(filename=xlsx_file)
            worksheet = workbook.worksheets[0]
            # 获取名字信息
            for row in tuple(worksheet[1:3]):
                for cell in row:
                    # print(cell.value)
                    if cell.value == ('姓名' or '名字' or '名称'):
                        r = cell.row
                        c = cell.column
                        break
            if r != None and c != None:
                # print(r, c)
                # print(worksheet[c])
                list_name = list(cell.value for cell in [col for col in worksheet.columns][c - 1])[r:]
                for i in list_name:
                    list_names.append(i)
                list_names.append('\n')  # 遍历完每个支部加换个行
                scr_output(scr_7, '\n\n提取出来的名单：\n{}'.format(list_name))
                print('\n\n提取出来的名单：\n{}'.format(list_name))
            else:
                print('找不到名字，请手动输入！')
                scr_output(scr_7, '\n找不到名字，请手动输入！\n')
        # print(list_names)
        scr_sheet7.insert('insert', ' '.join(i for i in list_names))  # 插入名字
        scr_sheet7.update()  # 插入后及时的更新
        scr_sheet7.see(END)  # 使得聊天记录text默认显示底端
    else:
        print('路径为空！')
        scr_output(scr_7, '\n路径为空！\n')
    # 获取名单
    people_sheet = scr_sheet7.get(1.0, 'end').split()
    if people_sheet != []:
        number11_7.set(people_sheet[0])
        number12_7.set(len(people_sheet))
        cookie = str(number_7_1.get()) + str(number_7_2.get()) + str(number_7_3.get())
        if cookie == '100':  # 发展对象的批复
            number16_7.set('关于建议接收{}等{}名同志为{}年度经济管理与法学学院{}发展对象的请示'.format(
                people_sheet[0],len(people_sheet),number1_7.get(),number2_7.get()))
        if cookie == '010':  # 预备党员的批复
            number16_7.set('关于建议接收{}等{}名同志为{}年度经济管理与法学学院{}预备党员的请示'.format(
                people_sheet[0],len(people_sheet),number1_7.get(),number2_7.get()))
        if cookie == '001':  # 预备党员转正的批复
            number16_7.set('关于建议接收{}等{}名同志为{}年度经济管理与法学学院{}预备党员转正的请示'.format(
                people_sheet[0],len(people_sheet),number1_7.get(),number2_7.get()))
        if cookie == '000':
            scr_output(scr_7, '\n没有选择需要生成哪种批复类型！\n')

        number13_7.set(number5_6.get()) # 更新支部请示时间
        number14_7.set(number6_6.get())
        number15_7.set(number7_6.get())

    else:
        print('名单为空，请检查！')
        scr_output(scr_7, '\n名单为空，请检查！\n')
# 批复模板修改，未完善
def pifu_model_alter():
    messagebox.showinfo('小提示', '本版本只支持查看批复模板，暂不支持修改')
    cookie = str(number_7_1.get()) + str(number_7_2.get()) + str(number_7_3.get())
    if cookie == '000':
        messagebox.showinfo('错误提示', '未选中批复的类型，请检查！')
        return
    if cookie == '100':  # 发展对象的批复
        a = "关于同意接收{}年度经济管理与法学学院"
        b = "{}发展对象的批复"
        c = "中共南华大学经济管理与法学学院"
        d = "{}等{}个学生党支部："
        e = "收到了贵支部{}年{}月{}日“{}”，且公示无异议。" \
            "认为你们按照党员标准对入党积极分子进行了有效的培养和教育。根据《中国共产党发展党员工作细则》的要求，" \
            "院党委于{}年{}月{}日召开党委会，认真讨论和审核材料，现将{}等" \
            "{}名同志列为中共党员发展对象，"
        e2 = "名单如下（{}）：".format('排名以首字母为序' if peoplename.get() == 1 else '排名以班级为序')
        f = "望你们继续加强对发展对象的培养和考察。"
        g = "特此批复。"
        h = "党委书记签名：_______________"
        i = "中共南华大学经济管理与法学学院委员会（盖章）"
        j = "{}年{}月{}日"
    if cookie == '010':  # 预备党员的批复
        a = "关于同意接收{}年度经济管理与法学学院"
        b = "{}预备党员的批复"
        c = "中共南华大学经济管理与法学学院"
        d = "{}等{}个学生党支部："
        e = "收到了贵支部{}年{}月{}日“{}”，且公示无异议。" \
            "认为你们按照党员标准对发展对象进行了有效的培养和教育。根据《中国共产党发展党员工作细则》的要求，" \
            "院党委于{}年{}月{}日召开党委会，认真讨论和审核材料，现确定" \
            "{}等{}名同志为中共预备党员，"
        e2 = "名单如下（{}）：".format('排名以首字母为序' if peoplename.get() == 1 else '排名以班级为序')
        f = "望你们继续加强对预备党员的培养和考察。"
        g = "特此批复。"
        h = "党委书记签名：_______________"
        i = "中共南华大学经济管理与法学学院委员会（盖章）"
        j = "{}年{}月{}日"
    if cookie == '001':  # 预备党员转正的批复
        a = "关于同意接收{}年度经济管理与法学学院"
        b = "{}预备党员转正的批复"
        c = "中共南华大学经济管理与法学学院"
        d = "{}等{}个学生党支部："
        e = "{}等{}名同志向党支部提出了转为正式党员的书面申请。" \
            "院党委于{}年{}月{}日召开党委会，讨论通过{}等{}名同志预备党员转为正式党员的决议。" \
            "{}等{}同志从预备期满之日起成为中国共产党正式党员，党龄从即日算起。"
        e2 = "名单如下（{}）：".format('排名以首字母为序' if peoplename.get() == 1 else '排名以班级为序')
        f = ''
        g = "特此批复。"
        h = "党委书记签名：_______________"
        i = "中共南华大学经济管理与法学学院委员会（盖章）"
        j = "{}年{}月{}日"
    list_pifu_model = [a, b, c, d, e, e2, f, g, h, i]

    def pifu_model_save():
        scr_output(scr_7, '\n{}\n批复模板保存失败！，本版本模板不支持修改！\n'.format(pifu_model_var.get()))
        pifu_model.destroy()

    def pifu_model_default():
        scr_output(scr_7, '\n模板已经是默认！\n')

    pifu_model = Toplevel(window)
    pifu_model.geometry("500x290+700+270")
    try:
        pifu_model.iconbitmap('mould\data\logo.ico')
    except:pass
    # 窗口置顶
    pifu_model.attributes("-topmost", 1)  # 1==True 处于顶层
    # 禁止窗口的拉伸
    pifu_model.resizable(0, 0)
    # 窗口的标题
    pifu_model.title("内置-{}批复模板-修改窗口".format(b[2:6]))

    # 定义变量
    pifu_model_var = StringVar()
    scr_pifu_model = scrolledtext.ScrolledText(pifu_model, wrap=WORD)
    scr_pifu_model.place(x=10, y=10, width=480, height=245)
    scr_pifu_model.config(state=DISABLED)  # 关闭可写入模式
    for i in list_pifu_model:
        scr_output(scr_pifu_model, str(i) + '\n')

    button_pifu_model = ttk.Button(pifu_model, text="保存参数", command=pifu_model_save)
    button_pifu_model.place(x=250, y=260)

    button_pifu_model = ttk.Button(pifu_model, text="恢复默认", command=pifu_model_default)
    button_pifu_model.place(x=120, y=260)

    # 显示窗口(消息循环)
    pifu_model.mainloop()
'''#################################################################################################################'''
# 备案报告cookie的模板的识别（后续开发做准备）
def beian_model_cookie(cookie,yeardu,pici,year_up, year,month,day, dw_year,dw_month,dw_day, first_people,people_num,people_sheet):
    if cookie == '10':  # 预备党员的备案报告
        a = "经济管理与法学学院预备党员报组织部备案报告"
        b = "校党委组织部："
        c = "学院党委于{}年{}月{}日召开党委会，现确定{}等{}名同志为中共党员预备党员，" \
            "名单如下（{}）：".format(dw_year,dw_month,dw_day,first_people,people_num,'排名以首字母为序' if peoplename.get() == 1 else '排名以班级为序')
        d = "学院将继续加强对预备党员的培养和考察。"
        e = "特此报告。"
        f = "中共南华大学经济管理与法学学院委员会（盖章）"
        g = "{}年{}月{}日".format(year, month, day)
    if cookie == '01':  # 预备党员转正的备案报告
        a = "经济管理与法学学院转正党员报组织部备案报告"
        b = "校党委组织部："
        c = "学院党委于{}年{}月{}日召开党委会，现确定{}等{}名同志为中共党员，" \
            "名单如下（{}）：".format(dw_year,dw_month,dw_day,first_people,people_num,'排名以首字母为序' if peoplename.get() == 1 else '排名以班级为序')
        d = "学院将继续加强对党员的培养和考察。"
        e = "特此报告。"
        f = "中共南华大学经济管理与法学学院委员会（盖章）"
        g = "{}年{}月{}日".format(year, month, day)

    return a,b,c,d,e,f,g
# 备案报告的写入
def write_beian(cookie,yeardu,pici,year_up, year,month,day, dw_year,dw_month,dw_day, first_people,people_num,people_sheet):
    people_sheet = (sorted(scr_sheet8.get(1.0, 'end').split(),key=lambda keys: [pinyin(i, style=Style.TONE3) for i in keys]) if peoplename.get() == 1 else scr_sheet8.get(1.0,'end').split())
    first_people = people_sheet[0]
    try:
        if type(people_sheet) is str: people_sheet = people_sheet.split()
        if cookie == '00':
            messagebox.showinfo('错误提示', '未选中备案报告的类型，请检查！')
            return
        if people_num != len(people_sheet):
            scr_output(scr_8, '\n生成备案报告 失败！\n错误信息：同志人数{}与人名数量{}不匹配，请检查！\n'.format(people_num,len(people_sheet)))
            messagebox.showinfo('错误提示', '生成备案报告 失败！\n错误信息：同志人数{}与人名数量{}不匹配，请检查！'.format(people_num,len(people_sheet)))
            return
        a,b,c,d,e,f,g = beian_model_cookie(cookie, yeardu, pici, year_up, year, month, day, dw_year,dw_month,dw_day, first_people, people_num,people_sheet)
        doc = Document()
        # 判断人数，来设置表格
        if 0 <= people_num<= 64: # 四号字体
            doc.styles['Normal'].paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE  # 1.5倍行距
            doc.styles['Normal'].font.size = Pt(14)
            col_width = [2.43,1.9]
            row_height = 1
        if 64 < people_num<= 104: # 小四字体
            doc.styles['Normal'].paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE # 1.5倍行距
            doc.styles['Normal'].font.size = Pt(12)
            col_width = [2.15,1.8]
            row_height = 0.9
        if 104 < people_num <= 120:  # 小四字体
            doc.styles['Normal'].paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE  # 1.5倍行距
            doc.styles['Normal'].font.size = Pt(12)
            col_width = [2.15, 1.8]
            row_height = 0.8
        if 120 < people_num<= 168: # 小四字体
            doc.styles['Normal'].paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST # 最小倍倍行距
            doc.styles['Normal'].font.size = Pt(12)
            col_width = [1.98,1.8]
            row_height = 0.55
        if 168 < people_num<= 184: # 五号字体
            doc.styles['Normal'].paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST # 最小倍倍行距
            doc.styles['Normal'].font.size = Pt(10.5)
            col_width = [1.98,1.8]
            row_height = 0.55
        if 184 < people_num:
            doc.styles['Normal'].font.size = Pt(10)
            col_width = [1.98, 1.8]
            row_height = 0.55
            print('人数太多（大于184），请自行调整word中存在的格式问题。')
            scr_output(scr_8, '人数太多（大于184），请自行调整word中存在的格式问题。')
        # 标题样式
        doc.styles['Header'].font.name = 'Times New Roman'  # 设置英文字体
        doc.styles['Header']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')  # 设置中文字体
        doc.styles['Header'].font.bold = True  # 加粗
        doc.styles['Header'].font.size = Pt(16)
        doc.styles['Header'].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 居中对齐
        doc.styles['Header'].paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE # 1.5倍行距
        doc.styles['Header'].paragraph_format.space_before = Pt(0)  # 段前
        doc.styles['Header'].paragraph_format.space_after = Pt(0)  # 段后
        # 普通正文央视
        doc.styles['Footer'].font.name = 'Times New Roman'  # 设置英文字体
        doc.styles['Footer']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')  # 设置中文字体
        doc.styles['Footer'].font.size = Pt(14)
        doc.styles['Footer'].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY # 两端对齐
        doc.styles['Footer'].paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE # 1.5倍行距
        doc.styles['Footer'].paragraph_format.space_before = Pt(0)  # 段前
        doc.styles['Footer'].paragraph_format.space_after = Pt(0)  # 段后
        doc.styles['Footer'].paragraph_format.first_line_indent = doc.styles['Footer'].font.size * 2  #首行缩进2字符 1 英寸=2.54 厘米
        # 表格样式
        doc.styles['Normal'].font.name = 'Times New Roman'  # 设置英文字体
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')  # 设置中文字体
        # doc.styles['Normal'].font.size = Pt(12)
        doc.styles['Normal'].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.DISTRIBUTE # 分散对齐
        # doc.styles['Normal'].paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST # 最小倍倍行距
        doc.styles['Normal'].paragraph_format.space_before = Pt(0)  # 段前
        doc.styles['Normal'].paragraph_format.space_after = Pt(0)  # 段后
        doc.styles['Normal'].paragraph_format.first_line_indent = Inches(0)  #首行缩进2字符 1 英寸=2.54 厘米

        # 标题 一段
        doc.add_paragraph(a,style='Header')
        # 称呼 一段 （首不设两字符）
        doc.add_paragraph(b,style='Footer').paragraph_format.first_line_indent=Inches(0) # 1 英寸=2.54 厘米
        # 正文
        doc.add_paragraph(c,style='Footer')

        table = doc.add_table(people_num//8 if people_num%8 == 0 else people_num//8+1, 8)
        table.autofit = True   # if is True 按窗口大小自动调整
        count = 0

        for row in range(len(table.rows)):
            table.rows[row].height = Cm(row_height)  # 调整行高
            for col in range(len(table.columns)):
                # print(行, 列)  # 可以查看表格输出结果
                table.cell(row, col).text = people_sheet[count]    # 写入人名
                # table.cell(行, 列).width = doc.styles['Normal'].font.size * len(people_sheet[count])
                # table.cell(行, 列).height = doc.styles['Normal'].font.size * 5
                table.cell(row, col).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER  # 上下居中（中部居中对齐）
                # table.cell(行, 列).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 水平居中（中部居中对齐）
                count += 1
                if count == people_num:
                    break
            if count == people_num:
                break
        # 修正列宽
        for col in range(len(table.columns)):
            maxlist = []
            for r in range(len(table.rows)):
                try:
                    maxlist.append(len(people_sheet[8*r + col]))
                    # print(people_sheet[8*r + col])
                except:pass
            if maxlist != []: maxnum = max(maxlist) # 每一列的最大值
            else: maxnum = 3  # 每一列的最大值
            table.cell(len(table.rows)-1, col).width = Cm( col_width[0] if maxnum==4 else col_width[1] ) # 调整列宽 2字:1.3 3字:1.8 4字:2.1
            # 要在最后一行设置列宽度，因为设置前面的，后面一行出现空格，前面设置的宽度就不生效了

        table.alignment = WD_TABLE_ALIGNMENT.CENTER  # 设置整个表格为居中对齐
        # table.autofit = True
        # 结束语
        doc.add_paragraph(d,style='Footer')
        doc.add_paragraph(e,style='Footer')
        doc.add_paragraph("",style='Footer')
        # 落款和时间
        doc.add_paragraph(f,style='Footer').alignment = WD_ALIGN_PARAGRAPH.RIGHT  # 右对齐
        doc.add_paragraph(g,style='Footer').alignment = WD_ALIGN_PARAGRAPH.RIGHT

        doc.save("经济管理与法学学院{}年度{}{}报组织部备案报告".format(yeardu,pici,a[9:13]) + '.docx')
        messagebox.showinfo('小提示', '生成备案报告 成功！请注意检查word文件格式！')
        scr_output(scr_8, '\n\n生成备案报告 成功！请注意检查word文件格式！\n')

    except Exception as error:
        error = str(error)
        print('错误提示', error)
        scr_output(scr_8, '\n生成备案报告 失败！\n\n\n本次错误信息：{}\n\n'.format(error))
        scr_output(scr_8, '\n--------文件没有成功保存--------\n\n\n\n\n\n\n')
        messagebox.showinfo('错误提示', '生成备案报告 失败！\n错误信息：\n{}'.format(error))
# 备案报告 自动检测姓名列 更新多个变量值
def auto_beian_read():
    if scr_sheet8.get(1.0, 'end').split() != []:
        messagebox.showinfo('小提示', '已经识别到文本中已有人名，请勿重复生成，请检查！'+'\n'
                            +'如需重新生成，请记得Ctrl+A清除不需要的人名，以防出错！'+'\n'
                            +'注意：本提示只是温馨提示，是不会停止继续执行自动检测的')
    # print(pathin_8.get())
    if pathin_8.get() != '':
        path = pathin_8.get()

        xls_files = [x for x in os.listdir(path) if os.path.splitext(x)[1] == '.xls']
        if xls_files != []:  # 说明有xls文件
            xls_to_xlsx(path=path, sole=False)  # 给路径，让其自己转换成xlsx的
            scr_output(scr_8, '\n\n检测到有{}个xls格式文件，已经自动在原路径转换生成可读取的xlsx文件类型：\n'.format(len(xls_files)))

        xlsx_files = [x for x in os.listdir(path) if os.path.splitext(x)[1] == '.xlsx']  # 罗列当前目录内所有xlsx文件
        scr_output(scr_8, '\n\n需要提取名字的表格{}个'.format(len(xlsx_files)))
        scr_output(scr_8, '\n\n需要提取名字的表格有：\n{}'.format(xlsx_files))
        print('需要提取', len(xlsx_files), '个表格')
        print('提取表格有：\n', xlsx_files)  # 本目录下的xlsx文件名字列表
        list_names = []
        for p in xlsx_files:
            r, c = None, None
            xlsx_file = path + '/' + p
            workbook = openpyxl.load_workbook(filename=xlsx_file)
            worksheet = workbook.worksheets[0]
            # 获取名字信息
            for row in tuple(worksheet[1:3]):
                for cell in row:
                    # print(cell.value)
                    if cell.value == ('姓名' or '名字' or '名称'):
                        r = cell.row
                        c = cell.column
                        break
            if r != None and c != None:
                # print(r, c)
                # print(worksheet[c])
                list_name = list(cell.value for cell in [col for col in worksheet.columns][c - 1])[r:]
                for i in list_name:
                    list_names.append(i)
                list_names.append('\n') # 遍历完每个支部加换个行
                scr_output(scr_8, '\n\n提取出来的名单：\n{}'.format(list_name))
                print('\n\n提取出来的名单：\n{}'.format(list_name))
            else:
                print('找不到名字，请手动输入！')
                scr_output(scr_8, '\n找不到名字，请手动输入！\n')
        # print(list_names)
        scr_sheet8.insert('insert', ' '.join(i for i in list_names)) # 插入名字
        scr_sheet8.update()  # 插入后及时的更新
        scr_sheet8.see(END)  # 使得聊天记录text默认显示底端
    else:
        print('路径为空！')
        scr_output(scr_8, '\n路径为空！\n')
    people_sheet = scr_sheet8.get(1.0, 'end').split()
    if people_sheet != []:
        number13_8.set(people_sheet[0])
        number14_8.set(len(people_sheet))
        number_chosen10_8.set(number_chosen5_8.get())
        number_chosen11_8.set(number_chosen6_8.get())
        number_chosen12_8.set(number_chosen7_8.get())
    else: print('名单为空，请检查！')
# 备案报告模板修改，未完善
def beian_model_alter():
    messagebox.showinfo('小提示', '本版本只支持查看备案报告模板，暂不支持修改')
    cookie = str(number_8_1.get()) + str(number_8_2.get())
    if cookie == '00':
        messagebox.showinfo('错误提示', '未选中备案报告的类型，请检查！')
        return
    if cookie == '10':  # 预备党员的备案报告
        a = "经济管理与法学学院预备党员报组织部备案报告"
        b = "校党委组织部："
        c = "学院党委于{}年{}月{}日召开党委会，现确定{}等{}名同志为中共党员预备党员，"
        c2 = "名单如下（{}）：".format('排名以首字母为序' if peoplename.get() == 1 else '排名以班级为序')
        d = "学院将继续加强对预备党员的培养和考察。"
        e = "特此报告。"
        f = "中共南华大学经济管理与法学学院委员会（盖章）"
        g = "{}年{}月{}日"
    if cookie == '01':  # 预备党员转正的备案报告
        a = "经济管理与法学学院转正党员报组织部备案报告"
        b = "校党委组织部："
        c = "学院党委于{}年{}月{}日召开党委会，现确定{}等{}名同志为中共党员，"
        c2 = "名单如下（{}）：".format('排名以首字母为序' if peoplename.get() == 1 else '排名以班级为序')
        d = "学院将继续加强对党员的培养和考察。"
        e = "特此报告。"
        f = "中共南华大学经济管理与法学学院委员会（盖章）"
        g = "{}年{}月{}日"
    list_beian_model = [a,b,c,c2,d,e,f,g]
    def beian_model_save():
        scr_output(scr_8,'\n{}\n备案报告模板保存失败！，本版本模板不支持修改！\n'.format(beian_model_var.get()))
        beian_model.destroy()

    def beian_model_default():
        scr_output(scr_8,'\n模板已经是默认！\n')

    beian_model = Toplevel(window)
    beian_model.geometry("500x290+700+270")
    try:
        beian_model.iconbitmap('mould\data\logo.ico')
    except:pass
    # 窗口置顶
    beian_model.attributes("-topmost", 1)  # 1==True 处于顶层
    # 禁止窗口的拉伸
    beian_model.resizable(0, 0)
    # 窗口的标题
    beian_model.title("内置-{}备案报告模板-修改窗口".format(b[2:6]))

    # 定义变量
    beian_model_var= StringVar()
    scr_beian_model = scrolledtext.ScrolledText(beian_model, wrap=WORD)
    scr_beian_model.place(x=10, y=10, width=480,height=245)
    scr_beian_model.config(state=DISABLED)  # 关闭可写入模式
    for i in list_beian_model:
        scr_output(scr_beian_model, str(i) + '\n')

    button_beian_model = ttk.Button(beian_model, text="保存参数", command=beian_model_save)
    button_beian_model.place(x=250, y=260)

    button_beian_model = ttk.Button(beian_model, text="恢复默认", command=beian_model_default)
    button_beian_model.place(x=120, y=260)

    # 显示窗口(消息循环)
    beian_model.mainloop()
'''#################################################################################################################'''
# 生成试卷
def generate_test_paper(diffcult,year,qishu,danxuan_num,panduan_num,duoxuan_num,tiankong_num,jianda_num,lunsu_num):
    '''
    经济管理与法学学院{2021}年第{八}期入党积极分子培训班
    结业考试
    说明：本试卷共五大题，39小题，满分100分。考试时长：100分钟。
    一、单选题（共20小题；每小题1分，满分20分，每小题只有一个选项符合题意，请把正确答案填入下列表格中）
    题号	1	2	3	4	5	6	7	8	9	10	总分
    得分
    题号	11	12	13	14	15	16	17	18	19	20
    得分

    二、判断题（共10小题；每小题1分，满分10分；正确的打“√”，错误的打“×”，，请把正确答案填入下列表格中）
    题号	21	22	23	24	25	26	27	28	29	30	总分
    得分

    三、填空题（共5题；每空1分，满分15分）

    四、简答题（共4小题；满分25分）

    五、论述开放题（共1小题；满分30分）

    正文          宋体小四      Normal
    标题1         宋体小三      Heading 1
    标题2         黑体小四      Heading 2
    说明字体       宋体五号     Title
    '''
    try:
        filepath = 'mould\model\题库.xlsx'
        os.rename("mould\data\orique_data", filepath)
        workbook = openpyxl.load_workbook(filepath)
        os.rename(filepath, "mould\data\orique_data")
    except:
        try:
            filepath = "mould\data\orique_data" +'.xlsx'
            os.rename("mould\data\orique_data",filepath)
            workbook = openpyxl.load_workbook(filepath)
            os.rename(filepath,"mould\data\orique_data")
        except Exception as error:
            messagebox.showinfo('错误提示', '尝试打开题库失败！\n错误信息：\n{}'.format(error))
    # print(filepath)
    try:
        sum_num = int(danxuan_num)+int(panduan_num)+int(duoxuan_num)+int(tiankong_num)+int(jianda_num)+int(lunsu_num)
        # 有时候我们希望读取到公式计算出来的结果，可以使用load_workbook()中的data_only属性, data_only=True
        # workbook = openpyxl.load_workbook(filepath)
        danxuan_sheet = workbook.worksheets[0]
        panduan_sheet = workbook.worksheets[1]
        duoxuan_sheet = workbook.worksheets[2]
        tiankong_sheet = workbook.worksheets[3]
        jianda_sheet = workbook.worksheets[4]
        lunsu_sheet = workbook.worksheets[5]

        # 打开文档
        test_doc = Document('mould\model\模板5 试卷.docx')
        answer_doc = Document()

        # for i in test_doc.styles:
        #     if i.type == WD_STYLE_TYPE.PARAGRAPH:
        #         print(i.name)

        a = '经济管理与法学学院{}年第{}期入党积极分子培训班'.format(year, qishu)
        b = '结业考试'
        c = '说明：本试卷共五大题，{}小题，满分100分。考试时长：100分钟。'.format(sum_num)
        test_doc.add_paragraph(a, style='Heading 1')  # 文章标题
        test_doc.add_paragraph(b, style='Heading 1')  # 文章标题
        test_doc.add_paragraph(c, style='Title')  # 说明
        # test_doc.styles['Normal'].paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE # 1.5倍行距
        answer_doc.styles['Normal'].font.name = 'Times New Roman'  # 设置英文字体
        answer_doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')  # 设置中文字体
        questions_num = 1  # 题号

        d = '一、单选题（共{}小题；每小题1分，满分{}分，每小题只有一个选项符合题意，请把正确答案填入下列表格中）'.format(danxuan_num, danxuan_num)
        test_doc.add_paragraph(d, style='Heading 2')  # 标题
        answer_doc.add_paragraph(d)  # 答案写入标题
        # 增加一个表格 没有成功
        # test_doc.add_table(rows=danxuan_num//10*2 , cols=12)
        # for c in test_doc.tables[0].columns[0].cells:
        #     c.width = Cm(1.5)
        # for c in test_doc.tables[0].columns[len(test_doc.tables[0].columns)-1].cells:
        #     c.width = Cm(1.5)
        # for c in range(1,len(test_doc.tables[0].rows)+1):
        #     if c%2 == 0:
        #         test_doc.tables[0].cell(c-1,0).text = '答案'
        #         test_doc.tables[0].cell(c-1,0).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        #         test_doc.tables[0].cell(c-1,0).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        #
        #     else:
        #         test_doc.tables[0].cell(c-1,0).text = '题号'
        #         test_doc.tables[0].cell(c-1,0).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        #         test_doc.tables[0].cell(c-1,0).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        #
        # test_doc.tables[0].cell(0,len(test_doc.tables[0].columns)-1).text = '总分'
        # test_doc.tables[0].cell(0,len(test_doc.tables[0].columns)-1).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        # test_doc.tables[0].cell(0,len(test_doc.tables[0].columns)-1).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        # tihao_num = 1
        # for r in range(1,len(test_doc.tables[0].rows)+1):
        #     if r%2 != 0:
        #         for c in range(2, len(test_doc.tables[0].columns)):
        #             test_doc.tables[0].cell(r-1,c-1).text = str(tihao_num)
        #             test_doc.tables[0].cell(r-1,c-1).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        #             test_doc.tables[0].cell(r-1,c-1).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        #             tihao_num += 1
        danxuan_set = set()
        answer_danxuan_para = ''
        # print(danxuan_sheet.max_row - 1)
        while (len(danxuan_set) < int(danxuan_num)):
            i = random.randint(2, danxuan_sheet.max_row)
            while danxuan_sheet.cell(row=i, column=4).value == None:
                i = random.randint(2, danxuan_sheet.max_row)
            if i not in danxuan_set:
                danxuan_set.add(i)
                # print(i)
                # print(questions_num)
                # print(danxuan_sheet.cell(row=i, column=4).value)
                danxuan_para = str(questions_num) + '、' + danxuan_sheet.cell(row=i, column=4).value  # 问题题目
                print(danxuan_para)
                test_doc.add_paragraph(danxuan_para)  # 写入题目
                # run = danxuan_paragraph.add_run(danxuan_para).bold = True # 单句加粗
                A = str(danxuan_sheet.cell(row=i, column=5).value)
                B = str(danxuan_sheet.cell(row=i, column=6).value)
                C = str(danxuan_sheet.cell(row=i, column=7).value)
                D = str(danxuan_sheet.cell(row=i, column=8).value)
                # print(A,C,B,D,len(A)+len(B)+len(C)+len(D))
                if (len(A) + len(B) + len(C) + len(D)) <= 20:  # 1行
                    danxuan_temp = '   A、%s   B、%s   C、%s   D、%s\n' % (A, B, C, D)  # 选项
                elif (len(A) + len(B)) <= 26 and (len(C) + len(D)) <= 26:  # 2行
                    danxuan_temp = '   A、%s   B、%s\n   C、%s   D、%s\n' % (A, B, C, D)  # 选项
                elif (len(A) + len(B)) > 26 or (len(C) + len(D)) > 26:  # 4行
                    danxuan_temp = '   A、%s\n   B、%s\n   C、%s\n   D、%s\n' % (A, B, C, D)  # 选项 四行
                test_doc.add_paragraph(danxuan_temp)  # 写入选项
                answer_danxuan_para += str(questions_num) + '、' + danxuan_sheet.cell(row=i,
                                                                                     column=9).value + '        '  # 答案
                if len(danxuan_set) % 5 == 0:
                    answer_doc.add_paragraph(answer_danxuan_para)  # 写入答案
                    answer_danxuan_para = ''
                questions_num += 1
            else:
                pass

        e = '二、判断题（共{}小题；每小题1分，满分{}分；正确的打“√”，错误的打“×”，，请把正确答案填入下列表格中）'.format(panduan_num, panduan_num)
        test_doc.add_paragraph(e, style='Heading 2')  # 标题
        answer_doc.add_paragraph(e)  # 答案写入标题
        panduan_set = set()
        answer_panduan_para = ''
        while (len(panduan_set) < int(panduan_num)):
            i = random.randint(2, panduan_sheet.max_row)
            while panduan_sheet.cell(row=i, column=4).value == None:
                i = random.randint(2, panduan_sheet.max_row)
            if i not in panduan_set:
                panduan_set.add(i)
                panduan_para = str(questions_num) + '、' + panduan_sheet.cell(row=i, column=4).value  # 问题
                test_doc.add_paragraph(panduan_para)
                answer_panduan_para += str(questions_num) + '、' + panduan_sheet.cell(row=i, column=5).value + '        '
                if len(panduan_set) % 5 == 0:
                    answer_doc.add_paragraph(answer_panduan_para)  # 写入答案
                    answer_panduan_para = ''
                questions_num += 1
            else:
                pass

        f = '三、填空题（共{}题；每空1分，满分{}分）'.format(tiankong_num, 10)
        test_doc.add_paragraph(f, style='Heading 2')  # 标题
        answer_doc.add_paragraph(f)  # 答案写入标题
        tiankong_set = set()
        answer_tiankong_para = ''
        while (len(tiankong_set) < int(tiankong_num)):
            i = random.randint(2, tiankong_sheet.max_row)
            while tiankong_sheet.cell(row=i, column=4).value == None:
                i = random.randint(2, tiankong_sheet.max_row)
            if i not in tiankong_set:
                tiankong_set.add(i)
                tiankong_para = str(questions_num) + '、' + tiankong_sheet.cell(row=i, column=4).value  # 问题
                test_doc.add_paragraph(tiankong_para)
                answer_tiankong_para += str(questions_num) + '、' + tiankong_sheet.cell(row=i, column=5).value + '\n'
                questions_num += 1
            else:
                pass
        answer_doc.add_paragraph(answer_tiankong_para)  # 写入答案

        g = '四、简答题（共{}小题；满分{}分）'.format(jianda_num, 25)
        test_doc.add_paragraph(g, style='Heading 2')  # 标题
        answer_doc.add_paragraph(g)  # 答案写入标题
        jianda_set = set()
        answer_jianda_para = ''
        while (len(jianda_set) < int(jianda_num)):
            i = random.randint(2, jianda_sheet.max_row)
            while jianda_sheet.cell(row=i, column=4).value == None:
                i = random.randint(2, jianda_sheet.max_row)
            if i not in jianda_set:
                jianda_set.add(i)
                jianda_para = str(questions_num) + '、' + jianda_sheet.cell(row=i, column=4).value + '\n\n\n\n\n\n'  # 问题
                test_doc.add_paragraph(jianda_para)
                answer_jianda_para += str(questions_num) + '、' + jianda_sheet.cell(row=i, column=5).value + '\n\n'
                questions_num += 1
            else:
                pass
        answer_doc.add_paragraph(answer_jianda_para)  # 写入答案

        h = '五、论述开放题（共{}小题；满分{}分）'.format(lunsu_num, 30)
        test_doc.add_paragraph(h, style='Heading 2')  # 标题
        answer_doc.add_paragraph(h)  # 答案写入标题
        lunsu_set = set()
        answer_lunsu_para = ''
        while (len(lunsu_set) < int(lunsu_num)):
            i = random.randint(2, lunsu_sheet.max_row)
            while lunsu_sheet.cell(row=i, column=4).value == None:
                i = random.randint(2, lunsu_sheet.max_row)
            if i not in lunsu_set:
                lunsu_set.add(i)
                lunsu_para = str(questions_num) + '、' + lunsu_sheet.cell(row=i, column=4).value  # 问题
                test_doc.add_paragraph(lunsu_para)
                answer_lunsu_para += str(questions_num) + '、' + lunsu_sheet.cell(row=i, column=5).value
                questions_num += 1
            else:
                pass
        answer_doc.add_paragraph(answer_lunsu_para)  # 写入答案
        # 保存文件
        test_doc.save(a+b+'试卷'+'.docx')
        answer_doc.save(a+b+'答案''.docx')
        scr_output(scr_2, '\n--------文件成功保存--------\n\n\n\n\n\n\n')
        messagebox.showinfo('小提示', '生成积极分子结业试卷 成功！\n')
    except Exception as error:
        scr_output(scr_2, '\n生成积极分子结业试卷 失败！\n\n\n本次错误信息：{}\n\n'.format(error))
        scr_output(scr_2, '\n--------文件没有成功保存--------\n\n\n\n\n\n\n')
        messagebox.showinfo('错误提示', '生成积极分子结业试卷 失败！\n错误信息：\n{}'.format(error))
'''#################################################################################################################'''
# 支部管理 生成各支部的请示和批复
# 支部请示文件cookie的模板的识别（后续开发做准备）
def zhibu_qingshi_model_cookie(cookie,party_name,year,month,day,zd_year,zd_month,zd_day,first_people,people_num,people_sheet,
                               yeardu,year_up,dy_sum,dy_true,dy_wait,dy_in,dy_true_in):
    if cookie == '100':  # 发展对象的请示
        a = "关于建议将{}等{}人列为中共党员发展对象的请示".format(first_people, people_num)
        b = "尊敬的学院党委："
        c = "{}等{}人，自    年  月至    年  月期间先后递交入党申请书，经各支部支委会讨论研究，" \
            "于    年  月至    年  月期间先后确定为入党积极分子并参加党校培训，学习优秀，获得结业证书。".format(first_people,people_num)
        d = "该{}人自递交入党申请书以来，以实际行动向党组织靠拢，以党员标准严格要求自己。政治上，认真学习党的理论，" \
            "坚决拥护党的领导、方针和政策，与党中央保持一致，入党动机端正；思想上，树立正确的人生观、价值观和世界观，" \
            "坚定共产主义信念，热爱祖国，热爱人民，严格要求自己，做到身未入党思想先入党；工作上，充分发挥了不怕苦、" \
            "不怕累、乐于奉献的精神，有强烈的责任心和集体荣誉感，起到了先锋模范作用；学习上，态度端正，成绩优良，" \
            "既牢固掌握本专业的基础知识和技能，又广泛学习其他学科的知识；在生活上，勤俭节约，诚实守信；作风上，" \
            "求真务实，言行一致，廉洁自律；纪律上，自觉遵守校纪校规，无任何违法违纪违规情况。经较长时间的培养和教育，" \
            "该{}人进步明显，对党的认识深刻，各方面表现突出，党员和群众对{}人评价良好，已基本符合入党条件。" \
            "".format(people_num,people_num,people_num)
        e = "鉴于以上表现，经支委会讨论研究，确认{}等{}人为{}年{}党员发展对象人选，建议院党委将其列为中共党员发展对象。" \
            "名单如下（{}）：".format(first_people,people_num,yeardu,year_up,'排名以首字母为序' if peoplename.get() == 1 else '排名以班级为序')
        f = "请批示。"
        g = "党支部书记签字：______________"
        h = "中共南华大学经济管理与法学学院"
        i = "{}".format(party_name)
        j = "{}年{}月{}日".format(year, month, day)
    if cookie == '010':  # 预备党员的请示
        a = "关于建议接收{}等{}名同志为中共预备党员的请示".format(first_people, people_num)
        b = "尊敬的学院党委："
        c = "{}等{}人向党组织提出了入党申请，该{}人主动接受党的入党积极分子、发展对象阶段的培养和考察，" \
            "在各方面都能以共产党员的标准严格要求自己，群众反映良好。党支部对该{}人进行了严格考察，认真审核材料。" \
            "于{}年{}月{}日召开支部大会讨论，认为{}等{}名同志符合党员的条件，提请学院党委接收其为预备党员，" \
            "名单如下（{}）：".format(first_people, people_num,people_num,people_num,
                                    zd_year,zd_month,zd_day,first_people, people_num,'排名以首字母为序' if peoplename.get() == 1 else '排名以班级为序')
        d = None
        e = None
        f = "请批示。"
        g = "党支部书记签字：______________"
        h = "中共南华大学经济管理与法学学院"
        i = "{}".format(party_name)
        j = "{}年{}月{}日".format(year, month, day)
    if cookie == '001':  # 预备党员转正的请示
        a = "关于支部同意{}等{}名同志转为正式党员呈报院党委审批的请示".format(first_people, people_num)
        b = "尊敬的学院党委："
        c = "经济管理与法学学院{}于{}年{}月{}日召开支部大会，讨论{}等共{}名同志的入党转正申请。" \
            "大会认为{}等共{}名同志在预备期间，能以党员标准严格要求自己，重视政治理论学习，学习刻苦，成绩优异，" \
            "积极参加学校的各项活动，有较强的组织纪律观念，发挥了一个共产党员应有的作用。 " \
            "".format(party_name,zd_year,zd_month,zd_day,first_people, people_num,first_people,people_num)
        d = "本支部共有党员{}名，其中正式党员{}名，预备党员{}名。到会党员{}名，其中正式党员{}名，" \
            "有表决权的到会人数超过应到会人数的半数，符合人数要求。经无记名表决，" \
            "{}名正式党员一致同意{}共{}名同志按期转为正式党员。名单如下（{}）：" \
            "".format(dy_sum,dy_true,dy_wait,dy_in,dy_true_in,dy_true_in,first_people, people_num,'排名以首字母为序' if peoplename.get() == 1 else '排名以班级为序')
        e = None
        f = None
        g = "党支部书记签字：______________"
        h = "中共南华大学经济管理与法学学院"
        i = "{}".format(party_name)
        j = "{}年{}月{}日".format(year, month, day)
    return a,b,c,d,e,f,g,h,i,j
# 支部请示文件的写入
def write_zhibu_qingshi(cookie,party_name,year,month,day,zd_year,zd_month,zd_day,first_people,people_num,people_sheet,
                               yeardu,year_up,dy_sum,dy_true,dy_wait,dy_in,dy_true_in):
    people_sheet = (sorted(scr_sheet1_11.get(1.0, 'end').split(),key=lambda keys: [pinyin(i, style=Style.TONE3) for i in keys]) if peoplename.get() == 1 else scr_sheet1_11.get(1.0,'end').split())
    first_people = people_sheet[0]
    try:
        if type(people_sheet) is str: people_sheet = people_sheet.split()
        if cookie == '000':
            messagebox.showinfo('错误提示', '未选中支部请示的类型，请检查！')
            return
        if people_num != len(people_sheet):
            scr_output(scr1_11, '\n生成支部请示文件 失败！\n错误信息：支部同志人数{}与支部人名数量{}不匹配，请检查！\n'.format(people_num,len(people_sheet)))
            messagebox.showinfo('错误提示', '生成支部请示文件 失败！\n错误信息：同志人数{}与人名数量{}不匹配，请检查！'.format(people_num,len(people_sheet)))
            return
        a,b,c,d,e,f,g,h,i,j = zhibu_qingshi_model_cookie(cookie,party_name,year,month,day,zd_year,zd_month,zd_day,first_people,people_num,people_sheet,
                               yeardu,year_up,dy_sum,dy_true,dy_wait,dy_in,dy_true_in)
        doc = Document()
        # 判断人数，来设置表格
        if 0 <= people_num<= 64: # 小四字体
            doc.styles['Normal'].paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE  # 1.5倍行距
            doc.styles['Normal'].font.size = Pt(12)
            col_width = [2.43,1.9]
            row_height = 1
        if 64 < people_num<= 104: # 小四字体
            doc.styles['Normal'].paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE # 1.5倍行距
            doc.styles['Normal'].font.size = Pt(12)
            col_width = [2.15,1.8]
            row_height = 0.9
        if 104 < people_num <= 120:  # 小四字体
            doc.styles['Normal'].paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE  # 1.5倍行距
            doc.styles['Normal'].font.size = Pt(12)
            col_width = [2.15, 1.8]
            row_height = 0.8
        if 120 < people_num<= 168: # 小四字体
            doc.styles['Normal'].paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST # 最小倍倍行距
            doc.styles['Normal'].font.size = Pt(12)
            col_width = [1.98,1.8]
            row_height = 0.55
        if 168 < people_num<= 184: # 五号字体
            doc.styles['Normal'].paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST # 最小倍倍行距
            doc.styles['Normal'].font.size = Pt(10.5)
            col_width = [1.98,1.8]
            row_height = 0.55
        if 184 < people_num:
            doc.styles['Normal'].font.size = Pt(10)
            col_width = [1.98, 1.8]
            row_height = 0.55
            print('支部人数太多（大于184），请自行调整word中存在的格式问题。')
            scr_output(scr1_11, '支部人数太多（大于184），请自行调整word中存在的格式问题。')
        # 标题样式
        doc.styles['Header'].font.name = 'Times New Roman'  # 设置英文字体
        doc.styles['Header']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')  # 设置中文字体
        doc.styles['Header'].font.bold = True  # 加粗
        doc.styles['Header'].font.size = Pt(14)
        doc.styles['Header'].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 居中对齐
        doc.styles['Header'].paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE # 1.5倍行距
        doc.styles['Header'].paragraph_format.space_before = Pt(0)  # 段前
        doc.styles['Header'].paragraph_format.space_after = Pt(0)  # 段后
        # 普通正文央视
        doc.styles['Footer'].font.name = 'Times New Roman'  # 设置英文字体
        doc.styles['Footer']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')  # 设置中文字体
        doc.styles['Footer'].font.size = Pt(12)
        doc.styles['Footer'].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY # 两端对齐
        doc.styles['Footer'].paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE # 1.5倍行距
        doc.styles['Footer'].paragraph_format.space_before = Pt(0)  # 段前
        doc.styles['Footer'].paragraph_format.space_after = Pt(0)  # 段后
        doc.styles['Footer'].paragraph_format.first_line_indent = doc.styles['Footer'].font.size * 2  #首行缩进2字符 1 英寸=2.54 厘米
        # 表格样式
        doc.styles['Normal'].font.name = 'Times New Roman'  # 设置英文字体
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')  # 设置中文字体
        # doc.styles['Normal'].font.size = Pt(12)
        doc.styles['Normal'].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.DISTRIBUTE # 分散对齐
        # doc.styles['Normal'].paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST # 最小倍倍行距
        doc.styles['Normal'].paragraph_format.space_before = Pt(0)  # 段前
        doc.styles['Normal'].paragraph_format.space_after = Pt(0)  # 段后
        doc.styles['Normal'].paragraph_format.first_line_indent = Inches(0)  #首行缩进2字符 1 英寸=2.54 厘米

        # 标题 一段
        doc.add_paragraph(a,style='Header')
        # 称呼 一段 （首不设两字符）
        doc.add_paragraph(b,style='Footer').paragraph_format.first_line_indent=Inches(0) # 1 英寸=2.54 厘米
        # 正文
        doc.add_paragraph(c,style='Footer')
        if d != None:
            doc.add_paragraph(d,style='Footer')
        if e != None:
            doc.add_paragraph(e,style='Footer')

        table = doc.add_table(people_num//8 if people_num%8 == 0 else people_num//8+1, 8)
        table.autofit = True   # if is True 按窗口大小自动调整
        count = 0

        for row in range(len(table.rows)):
            table.rows[row].height = Cm(row_height)  # 调整行高
            for col in range(len(table.columns)):
                # print(行, 列)  # 可以查看表格输出结果
                table.cell(row, col).text = people_sheet[count]    # 写入人名
                # table.cell(行, 列).width = doc.styles['Normal'].font.size * len(people_sheet[count])
                # table.cell(行, 列).height = doc.styles['Normal'].font.size * 5
                table.cell(row, col).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER  # 上下居中（中部居中对齐）
                # table.cell(行, 列).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 水平居中（中部居中对齐）
                count += 1
                if count == people_num:
                    break
            if count == people_num:
                break
        # 修正列宽
        for col in range(len(table.columns)):
            maxlist = []
            for r in range(len(table.rows)):
                try:
                    maxlist.append(len(people_sheet[8*r + col]))
                    # print(people_sheet[8*r + col])
                except:pass
            if maxlist != []: maxnum = max(maxlist) # 每一列的最大值
            else: maxnum = 3  # 每一列的最大值
            table.cell(len(table.rows)-1, col).width = Cm( col_width[0] if maxnum==4 else col_width[1] ) # 调整列宽 2字:1.3 3字:1.8 4字:2.1
            # 要在最后一行设置列宽度，因为设置前面的，后面一行出现空格，前面设置的宽度就不生效了

        table.alignment = WD_TABLE_ALIGNMENT.CENTER  # 设置整个表格为居中对齐
        # table.autofit = True
        # 结束语
        if f != None:
            doc.add_paragraph(f,style='Footer')
        # 落款二段和时间一段
        doc.add_paragraph(g,style='Footer').alignment = WD_ALIGN_PARAGRAPH.RIGHT  # 右对齐
        doc.add_paragraph(h,style='Footer').alignment = WD_ALIGN_PARAGRAPH.RIGHT  # 右对齐
        doc.add_paragraph(i,style='Footer').alignment = WD_ALIGN_PARAGRAPH.RIGHT
        doc.add_paragraph(j,style='Footer').alignment = WD_ALIGN_PARAGRAPH.RIGHT

        index = None
        for i in range(len(zhibu_allname)):
            if zhibu_allname[i] == party_name:
                index = i
        doc.save('{} {}.docx'.format(party_name if index == None else zhibu_list[index],a))
        messagebox.showinfo('小提示', '生成支部请示文件 成功！请注意检查word文件格式！')
        scr_output(scr1_11, '\n\n生成支部请示文件 成功！请注意检查word文件格式！\n')

    except Exception as error:
        error = str(error)
        print('错误提示', error)
        scr_output(scr1_11, '\n生成支部请示文件 失败！\n\n\n本次错误信息：{}\n\n'.format(error))
        scr_output(scr1_11, '\n--------文件没有成功保存--------\n\n\n\n\n\n\n')
        messagebox.showinfo('错误提示', '生支部成请示文件 失败！\n错误信息：\n{}'.format(error))
# 支部请示管理 自动检测姓名列 更新多个变量值
def auto_zhibu_qingshi_read():
    if scr_sheet1_11.get(1.0, 'end').split() != []:
        messagebox.showinfo('小提示', '已经识别到文本中已有人名，请勿重复生成，请检查！'+'\n'
                            +'如需重新生成，请记得Ctrl+A清除不需要的人名，以防出错！'+'\n'
                            +'注意：本提示只是温馨提示，是不会停止继续执行自动检测的')
    # print(pathin_6.get())
    if pathin1_11.get() != '':
        path = pathin1_11.get()
        if os.path.splitext(path)[1] == '.xls':  # 说明是xls文件
            xls_to_xlsx(path=path, sole=True)  # 给路径，让其自己转换成xlsx的
            scr_output(scr1_11, '\n\n检测到本文件是xls格式文件，已经自动在原路径转换生成可读取的xlsx文件类型：\n')
            path = os.path.splitext(path)[1] + '.xlsx'
            pathin1_11.set(path)
        scr_output(scr1_11, '\n\n用于提取名单的表格为：\n{}'.format(path))
        r, c = None, None
        workbook = openpyxl.load_workbook(filename=path)
        worksheet = workbook.worksheets[0]
        # 获取名字信息
        for row in tuple(worksheet[1:3]):
            for cell in row:
                # print(cell.value)
                if cell.value == ('姓名' or '名字' or '名称'):
                    r = cell.row
                    c = cell.column
                    break
        if r != None and c != None:
            # print(r, c)
            # print(worksheet[c])
            list_name = list(cell.value for cell in [col for col in worksheet.columns][c - 1])[r:]
            scr_output(scr1_11, '\n\n提取出来的名单：\n{}'.format(list_name))
            # print('\n\n提取出来的名单：\n{}'.format(list_name))
        else:
            # print('找不到名字，请手动输入！')
            scr_output(scr1_11, '\n找不到名字，请手动输入！\n')
            return
        # print(list_names)
        scr_sheet1_11.insert('insert', ' '.join(i for i in list_name)) # 插入名字
        scr_sheet1_11.update()  # 插入后及时的更新
        scr_sheet1_11.see(END)  # 使得聊天记录text默认显示底端
    else:
        print('路径为空！')
        scr_output(scr1_11, '\n路径为空！\n')
    people_sheet = scr_sheet1_11.get(1.0, 'end').split()
    if people_sheet != []:
        number5_11.set(people_sheet[0])
        number6_11.set(len(people_sheet))
    else: print('名单为空，请检查！')
# 支部请示模板修改，未完善
def zhibu_qingshi_model_alter():
    messagebox.showinfo('小提示', '本版本只支持查看支部请示模板，暂不支持修改')
    cookie = str(number_11_1.get()) + str(number_11_2.get()) + str(number_11_3.get())
    if cookie == '000':
        messagebox.showinfo('错误提示', '未选中支部请示的类型，请检查！')
        return
    if cookie == '100':  # 发展对象的请示
        name = '发展对象'
        a = "关于建议将{}等{}人列为中共党员发展对象的请示"
        b = "尊敬的学院党委："
        c = "{}等{}人，自{}年{}月至{}年{}月期间先后递交入党申请书，经各支部支委会讨论研究，" \
            "于{}年{}月至{}年{}月期间先后确定为入党积极分子并参加党校培训，学习优秀，获得结业证书。"
        c2 = "该{}人自递交入党申请书以来，以实际行动向党组织靠拢，以党员标准严格要求自己。政治上，认真学习党的理论，" \
            "坚决拥护党的领导、方针和政策，与党中央保持一致，入党动机端正；思想上，树立正确的人生观、价值观和世界观，" \
            "坚定共产主义信念，热爱祖国，热爱人民，严格要求自己，做到身未入党思想先入党；工作上，充分发挥了不怕苦、" \
            "不怕累、乐于奉献的精神，有强烈的责任心和集体荣誉感，起到了先锋模范作用；学习上，态度端正，成绩优良，" \
            "既牢固掌握本专业的基础知识和技能，又广泛学习其他学科的知识；在生活上，勤俭节约，诚实守信；作风上，" \
            "求真务实，言行一致，廉洁自律；纪律上，自觉遵守校纪校规，无任何违法违纪违规情况。经较长时间的培养和教育，" \
            "该{}人进步明显，对党的认识深刻，各方面表现突出，党员和群众对{}人评价良好，已基本符合入党条件。"
        d = "鉴于以上表现，经支委会讨论研究，确认{}等{}人为{}年{}党员发展对象人选，建议院党委将其列为中共党员发展对象。"
        e = "名单如下（{}）：".format('排名以首字母为序' if peoplename.get() == 1 else '排名以班级为序')
        f = "请批示。"
        g = "党支部书记签字：______________"
        h = "中共南华大学经济管理与法学学院"
        i = "{支部全称}"
        j = "{}年{}月{}日"
    if cookie == '010':  # 预备党员的请示
        name = '预备党员'
        a = "关于建议接收{}等{}名同志为中共预备党员的请示"
        b = "尊敬的学院党委："
        c = "{}等{}人向党组织提出了入党申请，该{}人主动接受党的入党积极分子、发展对象阶段的培养和考察，" \
            "在各方面都能以共产党员的标准严格要求自己，群众反映良好。党支部对该{}人进行了严格考察，认真审核材料。" \
            "于{}年{}月{}日召开支部大会讨论，认为{}等{}名同志符合党员的条件，提请学院党委接收其为预备党员，"
        c2 = "名单如下（{}）：".format('排名以首字母为序' if peoplename.get() == 1 else '排名以班级为序')
        d = ''
        e = ''
        f = "请批示。"
        g = "党支部书记签字：______________"
        h = "中共南华大学经济管理与法学学院"
        i = "{支部全称}"
        j = "{}年{}月{}日"
    if cookie == '001':  # 预备党员转正的请示
        name = '预备党员转正'
        a = "支部同意{}等{}同志转为正式党员呈报院党委审批的请示"
        b = "尊敬的学院党委："
        c = "经济管理与法学学院{}于{}年{}月{}日召开支部大会，讨论{}等共{}名同志的入党转正申请。" \
            "大会认为{}等共{}名同志在预备期间，能以党员标准严格要求自己，重视政治理论学习，学习刻苦，成绩优异，" \
            "积极参加学校的各项活动，有较强的组织纪律观念，发挥了一个共产党员应有的作用。 "
        d = "本支部共有党员{}名，其中正式党员{}名，预备党员{}名。到会党员{}名，其中正式党员{}名，" \
            "有表决权的到会人数超过应到会人数的半数，符合人数要求。经无记名表决，" \
            "{}名正式党员一致同意{}共{}名同志按期转为正式党员。"
        e = ''
        f = ''
        g = "党支部书记签字：______________"
        h = "中共南华大学经济管理与法学学院"
        i = "{支部全称}"
        j = "{}年{}月{}日"

    list_zhibu_qingshi_model = [a,b,c,c2, d,e,f,g,h,i,j]
    def zhibu_qingshi_model_save():
        scr_output(scr1_11,'\n{}\n支部请示模板保存失败！，本版本模板不支持修改！\n')
        zhibu_qingshi_model.destroy()

    def zhibu_qingshi_model_default():
        scr_output(scr1_11,'\n模板已经是默认！\n')

    zhibu_qingshi_model = Toplevel(window)
    zhibu_qingshi_model.geometry("500x290+700+270")
    try:
        zhibu_qingshi_model.iconbitmap('mould\data\logo.ico')
    except:pass
    # 窗口置顶
    zhibu_qingshi_model.attributes("-topmost", 1)  # 1==True 处于顶层
    # 禁止窗口的拉伸
    zhibu_qingshi_model.resizable(0, 0)
    # 窗口的标题
    zhibu_qingshi_model.title("内置-{}-支部请示模板-修改窗口".format(name))

    # 定义变量
    qingshi_model_var= StringVar()
    scr_zhibu_qingshi_model = scrolledtext.ScrolledText(zhibu_qingshi_model, wrap=WORD)
    scr_zhibu_qingshi_model.place(x=10, y=10, width=480,height=245)
    scr_zhibu_qingshi_model.config(state=DISABLED)  # 关闭可写入模式
    for i in list_zhibu_qingshi_model:
        scr_output(scr_zhibu_qingshi_model, str(i) + '\n')

    button_zhibu_qingshi_model = ttk.Button(zhibu_qingshi_model, text="保存参数", command=zhibu_qingshi_model_save)
    button_zhibu_qingshi_model.place(x=250, y=260)

    button_zhibu_qingshi_model = ttk.Button(zhibu_qingshi_model, text="恢复默认", command=zhibu_qingshi_model_default)
    button_zhibu_qingshi_model.place(x=120, y=260)

    # 显示窗口(消息循环)
    zhibu_qingshi_model.mainloop()
'''#################################################################################################################'''
# 支部批复文件cookie的模板的识别（后续开发做准备）
def zhibu_pifu_model_cookie(cookie, party_name,qs_year,qs_month,qs_day,year, month, day,first_people, people_num,people_sheet):
    if cookie == '100':  # 发展对象的批复
        a = "关于同意将{}等{}人列为中共党员发展对象的批复".format(first_people,people_num)
        b = ""
        c = "中共南华大学经济管理与法学学院{}：".format(party_name)
        d = "收到了贵支部{}年{}月{}日“关于建议将{}等{}人列为中共党员发展对象的请示”，且公示无异议。" \
            "认为你们按照党员标准对入党积极分子进行了有效的培养和教育。根据《中国共产党发展党员工作细则》的要求，" \
            "院党委于{}年{}月{}日召开党委会，认真讨论和审核材料，现将{}等" \
            "{}名同志列为中共党员发展对象，名单如下（{}）：" \
            "".format(qs_year,qs_month,qs_day,first_people,people_num,year,month,day,first_people,people_num,'排名以首字母为序' if peoplename.get() == 1 else '排名以班级为序')
        e = "望你们继续加强对发展对象的培养和考察。"
        f = "特此批复。"
        g = "党委书记签名：_______________"
        h = "中共南华大学经济管理与法学学院委员会（盖章）"
        i = "{}年{}月{}日".format(year, month, day)
    if cookie == '010':  # 预备党员的批复
        a = "关于同意接收{}等{}名同志为中共预备党员的批复".format(first_people,people_num)
        b = ""
        c = "中共南华大学经济管理与法学学院{}：".format(party_name)
        d = "收到了贵支部{}年{}月{}日“关于建议将{}等{}人列为中共预备党员的请示”，且公示无异议。" \
            "认为你们按照党员标准对发展对象进行了有效的培养和教育。根据《中国共产党发展党员工作细则》的要求，" \
            "院党委于{}年{}月{}日召开党委会，认真讨论和审核材料，现确定" \
            "{}等{}名同志为中共预备党员，名单如下（{}）：" \
            "".format(qs_year,qs_month,qs_day,first_people,people_num,year,month,day,first_people,people_num,'排名以首字母为序' if peoplename.get() == 1 else '排名以班级为序')
        e = "望你们继续加强对预备党员的培养和考察。"
        f = "特此批复。"
        g = "党委书记签名：_______________"
        h = "中共南华大学经济管理与法学学院委员会（盖章）"
        i = "{}年{}月{}日".format(year, month, day)
    if cookie == '001':  # 预备党员转正的批复
        a = "学院党委对支部同意{}等{}名同志转为正式党员决议的审批意见".format(first_people,people_num)
        b = ""
        c = "中共南华大学经济管理与法学学院{}：".format(party_name)
        d = "{}等{}名同志向党支部提出了转为正式党员的书面申请。学院党委在{}年{}月{}日召开党委会，" \
            "讨论通过你支部关于{}等{}名同志预备党员转为正式党员的决议。" \
            "{}等{}名同志从预备期满之日起成为中国共产党正式党员，党龄从即日算起。名单如下" \
            "（{}）：".format(first_people,people_num,year,month,day,first_people,people_num,first_people,people_num,'排名以首字母为序' if peoplename.get() == 1 else '排名以班级为序')
        e = None
        f = None
        g = "党委书记签名：_______________"
        h = "中共南华大学经济管理与法学学院委员会（盖章）"
        i = "{}年{}月{}日".format(year, month, day)
    return a, b, c, d, e, f, g, h, i
# 支部批复文件的写入
def write_zhibu_pifu(cookie, party_name,qs_year,qs_month,qs_day,year, month, day,first_people, people_num,people_sheet):
    people_sheet = (sorted(scr_sheet2_11.get(1.0, 'end').split(),key=lambda keys: [pinyin(i, style=Style.TONE3) for i in keys]) if peoplename.get() == 1 else scr_sheet2_11.get(1.0,'end').split())
    first_people = people_sheet[0]
    try:
        if type(people_sheet) is str: people_sheet = people_sheet.split()
        if cookie == '000':
            messagebox.showinfo('错误提示', '未选中支部批复的类型，请检查！')
            return
        if people_num != len(people_sheet):
            scr_output(scr2_11, '\n生成支部批复文件 失败！\n错误信息：同志人数{}与人名数量{}不匹配，请检查！\n'.format(people_num, len(people_sheet)))
            messagebox.showinfo('错误提示', '生成支部批复文件 失败！\n错误信息：同志人数{}与人名数量{}不匹配，请检查！'.format(people_num, len(people_sheet)))
            return
        a, b, c, d, e, f, g, h, i = zhibu_pifu_model_cookie(cookie, party_name,qs_year,qs_month,qs_day,year, month, day,
                                                      first_people, people_num,people_sheet)
        doc = Document()
        # 判断人数，来设置表格
        if 0 <= people_num <= 64:  # 小四号字体
            doc.styles['Normal'].paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE  # 1.5倍行距
            doc.styles['Normal'].font.size = Pt(12)
            col_width = [2.43, 1.9]
            row_height = 1
        if 64 < people_num <= 80:  # 小四字体
            doc.styles['Normal'].paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE  # 1.5倍行距
            doc.styles['Normal'].font.size = Pt(12)
            col_width = [2.15, 1.8]
            row_height = 0.9
        if 80 < people_num <= 88:  # 小四字体
            doc.styles['Normal'].paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE  # 1.5倍行距
            doc.styles['Normal'].font.size = Pt(12)
            col_width = [2.15, 1.8]
            row_height = 0.8
        if 88 < people_num <= 120:  # 小四字体
            doc.styles['Normal'].paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST  # 最小倍倍行距
            doc.styles['Normal'].font.size = Pt(12)
            col_width = [1.98, 1.8]
            row_height = 0.55
        if 120 < people_num <= 136:  # 五号字体
            doc.styles['Normal'].paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST  # 最小倍倍行距
            doc.styles['Normal'].font.size = Pt(10.5)
            col_width = [1.98, 1.8]
            row_height = 0.55
        if 136 < people_num:
            doc.styles['Normal'].font.size = Pt(10)
            col_width = [1.98, 1.8]
            row_height = 0.55
            print('支部人数太多（大于184），请自行调整word中存在的格式问题。')
            scr_output(scr2_11, '支部人数太多（大于184），请自行调整word中存在的格式问题。')
        # 标题样式
        doc.styles['Header'].font.name = 'Times New Roman'  # 设置英文字体
        doc.styles['Header']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')  # 设置中文字体
        doc.styles['Header'].font.bold = True  # 加粗
        doc.styles['Header'].font.size = Pt(14)
        doc.styles['Header'].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 居中对齐
        doc.styles['Header'].paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE  # 1.5倍行距
        doc.styles['Header'].paragraph_format.space_before = Pt(0)  # 段前
        doc.styles['Header'].paragraph_format.space_after = Pt(0)  # 段后
        # 普通正文央视
        doc.styles['Footer'].font.name = 'Times New Roman'  # 设置英文字体
        doc.styles['Footer']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')  # 设置中文字体
        doc.styles['Footer'].font.size = Pt(12)
        doc.styles['Footer'].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY  # 两端对齐
        doc.styles['Footer'].paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE  # 1.5倍行距
        doc.styles['Footer'].paragraph_format.space_before = Pt(0)  # 段前
        doc.styles['Footer'].paragraph_format.space_after = Pt(0)  # 段后
        doc.styles['Footer'].paragraph_format.first_line_indent = doc.styles[
                                                                      'Footer'].font.size * 2  # 首行缩进2字符 1 英寸=2.54 厘米
        # 表格样式
        doc.styles['Normal'].font.name = 'Times New Roman'  # 设置英文字体
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')  # 设置中文字体
        # doc.styles['Normal'].font.size = Pt(12)
        doc.styles['Normal'].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.DISTRIBUTE  # 分散对齐
        # doc.styles['Normal'].paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST # 最小倍倍行距
        doc.styles['Normal'].paragraph_format.space_before = Pt(0)  # 段前
        doc.styles['Normal'].paragraph_format.space_after = Pt(0)  # 段后
        doc.styles['Normal'].paragraph_format.first_line_indent = Inches(0)  # 首行缩进2字符 1 英寸=2.54 厘米

        # 标题 一段
        doc.add_paragraph(a, style='Header')
        # 称呼两段（首不设两字符）
        doc.add_paragraph(b, style='Footer').paragraph_format.first_line_indent = Inches(0)  # 1 英寸=2.54 厘米
        doc.add_paragraph(c, style='Footer').paragraph_format.first_line_indent = Inches(0)  # 1 英寸=2.54 厘米
        # 正文
        doc.add_paragraph(d, style='Footer')

        table = doc.add_table(people_num // 8 if people_num % 8 == 0 else people_num // 8 + 1, 8)
        table.autofit = True  # if is True 按窗口大小自动调整
        count = 0

        for row in range(len(table.rows)):
            table.rows[row].height = Cm(row_height)  # 调整行高
            for col in range(len(table.columns)):
                # print(行, 列)  # 可以查看表格输出结果
                table.cell(row, col).text = people_sheet[count]  # 写入人名
                # table.cell(行, 列).width = doc.styles['Normal'].font.size * len(people_sheet[count])
                # table.cell(行, 列).height = doc.styles['Normal'].font.size * 5
                table.cell(row, col).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER  # 上下居中（中部居中对齐）
                # table.cell(行, 列).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 水平居中（中部居中对齐）
                count += 1
                if count == people_num:
                    break
            if count == people_num:
                break
        # 修正列宽
        for col in range(len(table.columns)):
            maxlist = []
            for r in range(len(table.rows)):
                try:
                    maxlist.append(len(people_sheet[8 * r + col]))
                    # print(people_sheet[8*r + col])
                except:
                    pass
            if maxlist != []:
                maxnum = max(maxlist)  # 每一列的最大值
            else:
                maxnum = 3  # 每一列的最大值
            table.cell(len(table.rows) - 1, col).width = Cm(
                col_width[0] if maxnum == 4 else col_width[1])  # 调整列宽 2字:1.3 3字:1.8 4字:2.1
            # 要在最后一行设置列宽度，因为设置前面的，后面一行出现空格，前面设置的宽度就不生效了

        table.alignment = WD_TABLE_ALIGNMENT.CENTER  # 设置整个表格为居中对齐
        # table.autofit = True
        # 结束语
        if e != None:
            doc.add_paragraph(e, style='Footer')
        if f != None:
            doc.add_paragraph(f, style='Footer')
        doc.add_paragraph("", style='Footer')
        # 落款和时间
        doc.add_paragraph(g, style='Footer').alignment = WD_ALIGN_PARAGRAPH.RIGHT  # 右对齐
        doc.add_paragraph(h, style='Footer').alignment = WD_ALIGN_PARAGRAPH.RIGHT  # 右对齐
        doc.add_paragraph(i, style='Footer').alignment = WD_ALIGN_PARAGRAPH.RIGHT

        index = None
        for i in range(len(zhibu_allname)):
            if zhibu_allname[i] == party_name:
                index = i
        doc.save('{} {}.docx'.format(party_name if index == None else zhibu_list[index], a))

        messagebox.showinfo('小提示', '生成批复文件 成功！请注意检查word文件格式！')
        scr_output(scr_7, '\n\n生成批复文件 成功！请注意检查word文件格式！\n')

    except Exception as error:
        error = str(error)
        print('错误提示', error)
        scr_output(scr2_11, '\n生成批复文件 失败！\n\n\n本次错误信息：{}\n\n'.format(error))
        scr_output(scr2_11, '\n--------文件没有成功保存--------\n\n\n\n\n\n\n')
        messagebox.showinfo('错误提示', '生成批复文件 失败！\n错误信息：\n{}'.format(error))
# 支部批复管理 自动检测姓名列 更新多个变量值
def auto_zhibu_pifu_read():
    if scr_sheet2_11.get(1.0, 'end').split() != []:
        messagebox.showinfo('小提示', '已经识别到文本中已有人名，请勿重复生成，请检查！' + '\n'
                            + '如需重新生成，请记得Ctrl+A清除不需要的人名，以防出错！' + '\n'
                            + '注意：本提示只是温馨提示，是不会停止继续执行自动检测的')
    # print(pathin_7.get())
    # 如果路径不为空，写入名单
    if pathin2_11.get() != '':
        path = pathin2_11.get()
        if os.path.splitext(path)[1] == '.xls':  # 说明是xls文件
            xls_to_xlsx(path=path, sole=True)  # 给路径，让其自己转换成xlsx的
            scr_output(scr2_11, '\n\n检测到本文件是xls格式文件，已经自动在原路径转换生成可读取的xlsx文件类型：\n')
            path = os.path.splitext(path)[1] + '.xlsx'
            pathin2_11.set(path)
        scr_output(scr2_11, '\n\n用于提取名单的表格为：\n{}'.format(path))
        r, c = None, None
        workbook = openpyxl.load_workbook(filename=path)
        worksheet = workbook.worksheets[0]
        # 获取名字信息
        for row in tuple(worksheet[1:3]):
            for cell in row:
                # print(cell.value)
                if cell.value == ('姓名' or '名字' or '名称'):
                    r = cell.row
                    c = cell.column
                    break
        if r != None and c != None:
            # print(r, c)
            # print(worksheet[c])
            list_name = list(cell.value for cell in [col for col in worksheet.columns][c - 1])[r:]
            scr_output(scr2_11, '\n\n提取出来的名单：\n{}'.format(list_name))
            # print('\n\n提取出来的名单：\n{}'.format(list_name))
        else:
            # print('找不到名字，请手动输入！')
            scr_output(scr2_11, '\n找不到名字，请手动输入！\n')
            return
        scr_sheet2_11.insert('insert', ' '.join(i for i in list_name))  # 插入名字
        scr_sheet2_11.update()  # 插入后及时的更新
        scr_sheet2_11.see(END)  # 使得聊天记录text默认显示底端
    else:
        print('路径为空！')
        scr_output(scr2_11, '\n路径为空！\n')
    # 获取名单
    people_sheet = scr_sheet2_11.get(1.0, 'end').split()
    if people_sheet != []:
        number13_11.set(people_sheet[0])
        number14_11.set(len(people_sheet))

        number15_11.set(number2_11.get()) # 更新支部请示时间
        number16_11.set(number3_11.get())
        number17_11.set(number4_11.get())

    else:
        print('名单为空，请检查！')
        scr_output(scr2_11, '\n名单为空，请检查！\n')
# 支部批复模板修改，未完善
def zhibu_pifu_model_alter():
    messagebox.showinfo('小提示', '本版本只支持查看批复模板，暂不支持修改')
    cookie = str(number_11_4.get()) + str(number_11_5.get()) + str(number_11_6.get())
    if cookie == '000':
        messagebox.showinfo('错误提示', '未选中批复的类型，请检查！')
        return
    if cookie == '100':  # 发展对象的批复
        name = '发展对象'
        a = "关于同意将{}等{}人列为中共党员发展对象的批复"
        b = ""
        c = "中共南华大学经济管理与法学学院{支部全称}："
        d = "收到了贵支部{}年{}月{}日“关于建议将{}等{}人列为中共党员发展对象的请示”，且公示无异议。" \
            "认为你们按照党员标准对入党积极分子进行了有效的培养和教育。根据《中国共产党发展党员工作细则》的要求，" \
            "院党委于{}年{}月{}日召开党委会，认真讨论和审核材料，现将{}等" \
            "{}名同志列为中共党员发展对象，"
        d2 = "名单如下（{}）：".format('排名以首字母为序' if peoplename.get() == 1 else '排名以班级为序')
        e = "望你们继续加强对发展对象的培养和考察。"
        f = "特此批复。"
        g = "党委书记签名：_______________"
        h = "中共南华大学经济管理与法学学院委员会（盖章）"
        i = "{}年{}月{}日"
    if cookie == '010':  # 预备党员的批复
        name = '预备党员'
        a = "关于同意接收{}等{}名同志为中共预备党员的批复"
        b = ""
        c = "中共南华大学经济管理与法学学院{支部全称}："
        d = "收到了贵支部{}年{}月{}日“关于建议将{}等{}人列为中共预备党员的请示”，且公示无异议。" \
            "认为你们按照党员标准对发展对象进行了有效的培养和教育。根据《中国共产党发展党员工作细则》的要求，" \
            "院党委于{}年{}月{}日召开党委会，认真讨论和审核材料，现确定" \
            "{}等{}名同志为中共预备党员，"
        d2 = "名单如下（{}）：".format('排名以首字母为序' if peoplename.get() == 1 else '排名以班级为序')
        e = "望你们继续加强对预备党员的培养和考察。"
        f = "特此批复。"
        g = "党委书记签名：_______________"
        h = "中共南华大学经济管理与法学学院委员会（盖章）"
        i = "{}年{}月{}日"
    if cookie == '001':  # 预备党员转正的批复
        name = '预备党员转正'
        a = "学院党委对支部同意{}等{}名同志转为正式党员决议的审批意见"
        b = ""
        c = "中共南华大学经济管理与法学学院{支部全称}："
        d = "{}等{}名同志向党支部提出了转为正式党员的书面申请。学院党委在{}年{}月{}日召开党委会，" \
            "讨论通过你支部关于{}等{}名同志预备党员转为正式党员的决议。" \
            "{}等{}名同志从预备期满之日起成为中国共产党正式党员，党龄从即日算起。"
        d2 = "名单如下（{}）：".format('排名以首字母为序' if peoplename.get() == 1 else '排名以班级为序')
        e = ''
        f = ''
        g = "党委书记签名：_______________"
        h = "中共南华大学经济管理与法学学院委员会（盖章）"
        i = "{}年{}月{}日"
    list_pifu_model = [a, b, c, d, d2, e, f, g, h, i]

    def zhibu_pifu_model_save():
        scr_output(scr2_11, '\n{}\n批复模板保存失败！，本版本模板不支持修改！\n'.format(zhibu_pifu_model_var.get()))
        zhibu_pifu_model.destroy()

    def zhibu_pifu_model_default():
        scr_output(scr2_11, '\n模板已经是默认！\n')

    zhibu_pifu_model = Toplevel(window)
    zhibu_pifu_model.geometry("500x290+700+270")
    try:
        zhibu_pifu_model.iconbitmap('mould\data\logo.ico')
    except:pass
    # 窗口置顶
    zhibu_pifu_model.attributes("-topmost", 1)  # 1==True 处于顶层
    # 禁止窗口的拉伸
    zhibu_pifu_model.resizable(0, 0)
    # 窗口的标题
    zhibu_pifu_model.title("内置-{}-支部批复模板-修改窗口".format(name))

    # 定义变量
    zhibu_pifu_model_var = StringVar()
    scr_zhibu_pifu_model = scrolledtext.ScrolledText(zhibu_pifu_model, wrap=WORD)
    scr_zhibu_pifu_model.place(x=10, y=10, width=480, height=245)
    scr_zhibu_pifu_model.config(state=DISABLED)  # 关闭可写入模式
    for i in list_pifu_model:
        scr_output(scr_zhibu_pifu_model, str(i) + '\n')

    button_zhibu_pifu_model = ttk.Button(zhibu_pifu_model, text="保存参数", command=zhibu_pifu_model_save)
    button_zhibu_pifu_model.place(x=250, y=260)

    button_zhibu_pifu_model = ttk.Button(zhibu_pifu_model, text="恢复默认", command=zhibu_pifu_model_default)
    button_zhibu_pifu_model.place(x=120, y=260)

    # 显示窗口(消息循环)
    zhibu_pifu_model.mainloop()
'''#################################################################################################################'''
# 支部积极分子备案报告模板修改，未完善
def zhibu_beian_model_alter():
    messagebox.showinfo('小提示', '本版本只支持查看积极分子备案报告模板，暂不支持修改')

    name = '积极分子'
    a = "关于推荐{}等{}人为入党积极分子的备案报告"
    b = ""
    c = "中共南华大学经济管理与法学学院委员会"
    d = "{}等{}人向党组织递交了入党申请书，现由正式党员和群团组织推荐，" \
        "经支部考察，该5人在学习、工作、生活中，思想进步，积极向上，作风正派，表现优秀，有较好群众基础，" \
        "支委会研究确定将其列为入党积极分子培养（名单附后）。"
    e = "特此呈报学院党委备案。"
    f = ""
    g = "{支部名称}"
    h = "书记签名（加盖支部公章）："
    i = "{}年{}月{}日"
    j = "附备案人员名单："
    k = "{备案名单}"

    list_beian_model = [a, b, c, d, e, f, g, h, i, j, k]

    def zhibu_beian_model_save():
        scr_output(scr3_11, '\n{}\n积极分子备案报告模板保存失败！，本版本模板不支持修改！\n'.format(zhibu_beian_model_var.get()))
        zhibu_beian_model.destroy()

    def zhibu_beian_model_default():
        scr_output(scr3_11, '\n模板已经是默认！\n')

    zhibu_beian_model = Toplevel(window)
    zhibu_beian_model.geometry("500x290+700+270")
    try:
        zhibu_beian_model.iconbitmap('mould\data\logo.ico')
    except:
        pass
    # 窗口置顶
    zhibu_beian_model.attributes("-topmost", 1)  # 1==True 处于顶层
    # 禁止窗口的拉伸
    zhibu_beian_model.resizable(0, 0)
    # 窗口的标题
    zhibu_beian_model.title("内置-{}-支部积极分子备案报告模板-修改窗口".format(name))

    # 定义变量
    zhibu_beian_model_var = StringVar()
    scr_zhibu_beian_model = scrolledtext.ScrolledText(zhibu_beian_model, wrap=WORD)
    scr_zhibu_beian_model.place(x=10, y=10, width=480, height=245)
    scr_zhibu_beian_model.config(state=DISABLED)  # 关闭可写入模式
    for i in list_beian_model:
        scr_output(scr_zhibu_beian_model, str(i) + '\n')

    button_zhibu_beian_model = ttk.Button(zhibu_beian_model, text="保存参数", command=zhibu_beian_model_save)
    button_zhibu_beian_model.place(x=250, y=260)

    button_zhibu_beian_model = ttk.Button(zhibu_beian_model, text="恢复默认", command=zhibu_beian_model_default)
    button_zhibu_beian_model.place(x=120, y=260)

    # 显示窗口(消息循环)
    zhibu_beian_model.mainloop()
# 支部积极分子备案报告管理 自动检测姓名列 更新多个变量值
def auto_zhibu_beian_read():
    # if scr_sheet3_11.get(1.0, 'end').split() != []:
    #     messagebox.showinfo('小提示', '已经识别到文本中已有人名，请勿重复生成，请检查！' + '\n'
    #                         + '如需重新生成，请记得Ctrl+A清除不需要的人名，以防出错！' + '\n'
    #                         + '注意：本提示只是温馨提示，是不会停止继续执行自动检测的')
    # # print(pathin_7.get())
    # 如果路径不为空，写入名单
    if pathin3_11.get() != '':
        from openpyxl.utils import get_column_letter
        # path = pathin3_11.get()
        # print(path)
        # if os.path.splitext(path)[1] == '.xls':  # 说明是xls文件
        #     xls_to_xlsx(path=path, sole=True)  # 给路径，让其自己转换成xlsx的
        #     scr_output(scr3_11, '\n\n检测到本文件是xls格式文件，已经自动在原路径转换生成可读取的xlsx文件类型：\n')
        #     path = os.path.splitext(path)[1] + '.xlsx'
        #     pathin3_11.set(path)
        # scr_output(scr3_11, '\n\n用于提取名单的表格为：\n{}'.format(path))
        # r, c = None, None
        # workbook = openpyxl.load_workbook(filename=path)
        # worksheet = workbook.worksheets[0]
        # # 获取名字信息
        # for row in tuple(worksheet[1:3]):
        #     for cell in row:
        #         # print(cell.value)
        #         if cell.value == ('姓名' or '名字' or '名称'):
        #             r = cell.row
        #             c = cell.column
        #             break
        # if r != None and c != None:
        #     # print(r, c)
        #     # print(worksheet[c])
        #     list_name = list(cell.value for cell in [col for col in worksheet.columns][c - 1])[r:]
        #     scr_output(scr3_11, '\n\n提取出来的名单：\n{}'.format(list_name))
        #     # print('\n\n提取出来的名单：\n{}'.format(list_name))
        # else:
        #     # print('找不到名字，请手动输入！')
        #     scr_output(scr3_11, '\n找不到名字，请手动输入！\n')
        #     return
        # scr_sheet3_11.insert('insert', ' '.join(i for i in list_name))  # 插入名字
        # scr_sheet3_11.update()  # 插入后及时的更新
        # scr_sheet3_11.see(END)  # 使得聊天记录text默认显示底端
        filepath = pathin3_11.get()
        if os.path.splitext(filepath)[1] == '.xls':  # 说明是xls文件
            xls_to_xlsx(path=filepath, sole=True)  # 给路径，让其自己转换成xlsx的
            scr_output(scr3_11, '\n\n检测到本文件是xls格式文件，已经自动在原路径转换生成可读取的xlsx文件类型：\n')
            path = os.path.splitext(filepath)[0] + '.xlsx'
        workbook = openpyxl.load_workbook(filepath)
        new_workbook = openpyxl.Workbook()
        worksheet = workbook.worksheets[0]
        new_worksheet = new_workbook.worksheets[0]

        biaotou_r = '3'  # 默认第三行表头
        for row in tuple(worksheet[1:3]):
            for cell in row:
                # print(cell.value)
                if cell.value == ('姓名' or '名字' or '名称'):
                    biaotou_r = cell.row

        a, b, c, d, e, f, g = 0, 0, 0, 0, 0, 0, 0
        one_cell = worksheet[str(biaotou_r)]  # 获取第1行的数据
        for i in one_cell:
            # print(i.value)
            # print(i.column)
            if i.value != None:
                if '姓名' in str(i.value):
                    a = i.column
                if '出生' in str(i.value):
                    b = i.column
                if '申请书' in str(i.value):
                    c = i.column
                if '专业' in str(i.value) or "班级" in str(i.value):
                    d = i.column
                if '支部' in str(i.value):
                    e = i.column
                if '备注' in str(i.value):
                    f = i.column
                if '党员' in str(i.value):
                    g = i.column
        if a != 0 and b != 0 and c != 0 and d != 0 and e != 0 and f != 0 and g != 0:
            scr_output(scr3_11, '已经找到全部表头，导入的名单没有问题！')
        else:  # if a == 0 or b == 0 or c == 0 or d == 0 or e == 0 or f == 0:
            # print("没有找到全部表头，请检查导入的名单文件！")
            messagebox.showinfo("提示", "没有找到全部表头，部分信息可能有缺失，请检查导入的名单文件！！！")
            scr_output(scr3_11, '\n没有找到全部表头，请检查导入的名单文件！\n')
            scr_output(scr3_11, '\n\n缺失信息有：{} {} {} {} {} {} {}\n\n'.format('姓名' if a==0 else '',
                                                    '出生日期' if b==0 else '', '申请书日期' if c==0 else '',
                                                    '推荐两名党员' if g==0 else '', '推荐群团组织' if d==0 else '',
                                                    '推荐党支部/党小组' if e==0 else '', '备注' if f==0 else ''))

        print(a, b, c, d, e, f, g)
        # 写入表头
        biaotou = ['序号', '姓名', '出生日期', '申请书日期', '推荐两名党员', '推荐群团组织', '推荐党支部/党小组', '备注']
        for bt in range(len(biaotou)):  # 添加数据
            new_worksheet.cell(row=1, column=bt + 1, value=str(biaotou[bt]))

        r_count = 2
        for r in range(worksheet.max_row):
            if worksheet.cell(int(biaotou_r) + 1, a).value == None:
                continue
            # 写序号
            new_worksheet.cell(row=r_count + r, column=1, value=str(str(r_count - 1 + r) if str(worksheet.cell(int(biaotou_r) + 1 + r, a).value) != 'None' else str('')))
            # 写入姓名
            if a != 0:
                new_worksheet.cell(row=r_count + r, column=2, value=str(str(worksheet.cell(int(biaotou_r) + 1 + r, a).value) if str(worksheet.cell(int(biaotou_r) + 1 + r, a).value) != 'None' else str('')))
                # print(str(worksheet.cell(int(biaotou_r) + 1 +r, a).value))
            # 写入出生日期
            if b != 0:
                date = str(str(worksheet.cell(int(biaotou_r) + 1 + r, b).value) if str(worksheet.cell(int(biaotou_r) + 1 + r, b).value) != 'None' else str(''))
                if len(date) == 6 and str(date).isdigit() == True:
                    new_worksheet.cell(row=r_count + r, column=3, value='{}年{}月{}日'.format(date[0:2],date[2:4],date[4:6]))
                else:
                    new_worksheet.cell(row=r_count + r, column=3, value=str(date) if str(date) != 'None' else str(''))
            # 写入申请书日期
            if c != 0:
                date2 = str(str(worksheet.cell(int(biaotou_r) + 1 + r, c).value) if str(worksheet.cell(int(biaotou_r) + 1 + r, c).value) != 'None' else str(''))
                if len(date2) == 6 and str(date2).isdigit() == True:
                    new_worksheet.cell(row=r_count + r, column=4, value='{}年{}月{}日'.format(date2[0:2],date2[2:4],date2[4:6]))
                else:
                    new_worksheet.cell(row=r_count + r, column=4, value=str(date2) if str(date2) != 'None' else str(''))
            # # 写入推荐两名党员
            if g != 0:
                g_v = str(worksheet.cell(int(biaotou_r) + 1 + r, g).value)
                if g_v != None and str(g_v) != 'None' :
                    new_worksheet.cell(row=r_count + r, column=5, value=str(g_v) if str(g_v) != 'None' else str(''))
            # 写入推荐群团组织
            if d != 0:
                new_worksheet.cell(row=r_count + r, column=6, value=str(str(worksheet.cell(int(biaotou_r) + 1 + r, d).value) + '团支部' if str(worksheet.cell(int(biaotou_r) + 1 + r, d).value) != 'None' else str('')))
            # 写入推荐党支部/党小组
            if e != 0:
                # print(worksheet.cell(int(biaotou_r) + 1 + r, e).value)
                zhibu_value = worksheet.cell(int(biaotou_r) + 1 + r, e).value
                if zhibu_value != None:
                    index = None
                    for i in range(len(zhibu_list)):
                        if zhibu_list[i] == worksheet.cell(int(biaotou_r) + 1 + r, e).value:
                            index = i
                            zhibu_value = zhibu_allname[i]
                    zhibu_value = '经济管理与法学学院' + zhibu_value
                    new_worksheet.cell(row=r_count + r, column=7, value=str(zhibu_value))  # 全称
                else:
                    new_worksheet.cell(row=r_count + r, column=7, value=str(''))  # 全称
            # 写入备注
            if f != 0:
                new_worksheet.cell(row=r_count + r, column=8, value=str(worksheet.cell(int(biaotou_r) + 1 + r, f).value if str(worksheet.cell(int(biaotou_r) + 1 + r, f).value) != 'None' else str('')))
        # print( [i for i in new_worksheet[new_worksheet.min_row:new_worksheet.max_row]])

        # 删除空行（涉及删除行后，行数的索引值发生变化，所以复杂了点）
        zero_row_list = []
        for row in range(1, new_worksheet.max_row + 1):
            if new_worksheet.cell(row, 1).value == None or  new_worksheet.cell(row,2).value == '':
                zero_row_list.append(row)
        delete_row = 0
        # print('\n\n开始删除{}的空行…………\n'.format(xlsx_files[n]))
        scr_output(scr3_11, '\n开始删除{}的空行/空列和张三示例行…………\n'.format("文件"))
        for j in zero_row_list:
            # scr_output(scr_0,'\n原本该空值在第{}行,正在删除该空值目前所在的第{}行'.format(j , j -delete_row)
            # print('原本该空值在第{}行,正在删除该空值目前所在的第{}行'.format(j , j -delete_row))
            new_worksheet.delete_rows(j - delete_row, 1)
            delete_row = delete_row + 1

        # 删除空行后，序号不一致，重新写
        xuhao = 1
        for row in range(2, new_worksheet.max_row + 1):
            new_worksheet.cell(row=row, column=1, value=str(xuhao))
            xuhao += 1

        scr_output(scr3_11, '\n开始给区域设置设置框线…………\n')
        scr_output(scr3_11, '\n开始居中对齐…………\n')
        # 给区域设置设置框线
        # print(new_worksheet[new_worksheet.min_row:new_worksheet.max_row])
        for row in tuple(new_worksheet[new_worksheet.min_row:new_worksheet.max_row]):
            if new_worksheet.min_row == new_worksheet.max_row: # 有时候只有一行
                row.border = my_border('thin', 'thin', 'thin', 'thin')
                # 设置单元格对齐方式 Alignment(horizontal=水平对齐模式,vertical=垂直对齐模式,text_rotation=旋转角度,wrap_text=是否自动换行)
                alignment = Alignment(horizontal='center', vertical='center', text_rotation=0, wrap_text=True)
                row.alignment = alignment
            else:
                for cell in row:
                    cell.border = my_border('thin', 'thin', 'thin', 'thin')
                    # 设置单元格对齐方式 Alignment(horizontal=水平对齐模式,vertical=垂直对齐模式,text_rotation=旋转角度,wrap_text=是否自动换行)
                    alignment = Alignment(horizontal='center', vertical='center', text_rotation=0, wrap_text=True)
                    cell.alignment = alignment

        scr_output(scr3_11, '\n开始区域自动调整列宽…………\n')
        # 区域自动调整列宽]
        column_widths = []  # 定义用来获取当前列最大宽度的空列表
        for i, col in enumerate(new_worksheet.iter_cols(min_col=new_worksheet.min_column, max_col=new_worksheet.max_column,
                                                    min_row=new_worksheet.min_row)):
            for cell in col:
                value = cell.value
                if value is not None:
                    if isinstance(value, str) is False:
                        value = str(value)
                    try:
                        column_widths[i] = max(column_widths[i], len(value))
                    except IndexError:
                        column_widths.append(len(value))
        # print('column_widths', column_widths)  # 得到该列最大的一个单元格的宽度（字符串数量）
        for i, width in enumerate(column_widths):
            col_name = get_column_letter(new_worksheet.min_column + i)  # 获取行字母表头
            value = column_widths[i] * 2  # 设置列宽为最大长度比例
            new_worksheet.column_dimensions[col_name].width = value

        scr_output(scr3_11, '\n开始生成积极分子备案报告名单信息表…………\n')

        party_name = number1_11.get()
        index = None
        for i in range(len(zhibu_allname)):
            if zhibu_allname[i] == party_name:
                index = i
        new_workbook.save("{} 积极分子备案报告名单.xlsx".format(party_name if index == None else zhibu_list[index]))
        messagebox.showinfo("提示","积极分子备案报告名单.xlsx，保存成功！请对表中信息进行核对和调整，然后再执行生成备案报告的操作")

        # 获取名单
        if new_worksheet.min_row != new_worksheet.max_row:
            number27_11.set(new_worksheet.cell(2, 2).value)
            number28_11.set(int(new_worksheet.max_row) - 1)
            for r in range(new_worksheet.max_row):
                for c in range(new_worksheet.max_column):
                    v = new_worksheet.cell(r + 1, c + 1).value
                    scr_sheet3_11.insert('insert', str(v if v != None else '空') + ' ')  # 插入名字
            scr_sheet3_11.update()  # 插入后及时的更新
            scr_sheet3_11.see(END)  # 使得聊天记录text默认显示底端

        else:
            print('名单为空，请检查！')
            scr_output(scr3_11, '\n名单为空，请检查！\n')

    else:
        print('路径为空！')
        scr_output(scr3_11, '\n路径为空！\n')

# python-docx库设置表格单元格的边框
def set_cell_border(cell: _Cell, **kwargs):
    from docx.oxml import OxmlElement
    """
    Set cell`s border
    Usage:
    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
        bottom={"sz": 12, "color": "#00FF00", "val": "single"},
        left={"sz": 24, "val": "dashed", "shadow": "true"},
        right={"sz": 12, "val": "dashed"},
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # check for tag existnace, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    # list over all available tags
    for edge in ('left', 'top', 'right', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)

            # check for tag existnace, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))
# 支部积极分子备案报告文件的写入
def write_zhibu_beian(party_name,year,month,day,first_people,people_num,new_workbook):
    index = None
    for i in range(len(zhibu_allname)):
        if zhibu_allname[i] == party_name:
            index = i
    party_name = party_name if index == None else zhibu_list[index]
    try:
        workbook = openpyxl.load_workbook(party_name + ' ' + new_workbook)
    except:
        messagebox.showinfo('小提示', '导入文件“{}”失败，找不到积极分子备案报告名单！请检文件是否存在或被误删！'.format(party_name + ' ' + new_workbook))
        scr_output(scr3_11, '\n\n导入文件“{}”失败，找不到积极分子备案报告名单！请检文件是否存在或被误删！\n'.format(party_name + ' ' + new_workbook))

    worksheet = workbook.worksheets[0]
    people_sheet = worksheet.max_row-1
    # try:
    if people_num != people_sheet:
        scr_output(scr3_11, '\n生成支部积极分子备案报告文件 失败！\n错误信息：积极分子人数{}与人名数量{}不匹配，请检查！\n'.format(people_num, people_sheet))
        messagebox.showinfo('错误提示', '生成支部积极分子备案报告文件 失败！\n错误信息：积极分子人数{}与人名数量{}不匹配，请检查！'.format(people_num, people_sheet))
        return
    name = '积极分子'
    a = "关于推荐{}等{}人为入党积极分子的备案报告".format(first_people,people_num)
    b = ""
    c = "中共南华大学经济管理与法学学院委员会"
    d = "{}等{}人向党组织递交了入党申请书，现由正式党员和群团组织推荐，" \
        "经支部考察，该{}人在学习、工作、生活中，思想进步，积极向上，作风正派，表现优秀，有较好群众基础，" \
        "支委会研究确定将其列为入党积极分子培养（名单附后）。".format(first_people,people_num,people_num)
    e = "特此呈报学院党委备案。"
    f = ""
    g = "经济管理与法学学院{}".format(party_name)
    h = "书记签名（加盖支部公章）：_______________"
    i = "{}年{}月{}日".format(year,month,day)
    j = "附备案人员名单："
    # k = "{}".format("")

    doc = Document()
    # 判断人数，来设置表格
    # 小四号字体
    doc.styles['Normal'].paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE  # 1.5倍行距
    doc.styles['Normal'].font.size = Pt(12)
    col_width = [2.43, 1.9]
    row_height = 1

    # 标题样式
    doc.styles['Header'].font.name = 'Times New Roman'  # 设置英文字体
    doc.styles['Header']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')  # 设置中文字体
    doc.styles['Header'].font.bold = True  # 加粗
    doc.styles['Header'].font.size = Pt(14)
    doc.styles['Header'].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 居中对齐
    doc.styles['Header'].paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE  # 1.5倍行距
    doc.styles['Header'].paragraph_format.space_before = Pt(0)  # 段前
    doc.styles['Header'].paragraph_format.space_after = Pt(0)  # 段后
    # 普通正文央视
    doc.styles['Footer'].font.name = 'Times New Roman'  # 设置英文字体
    doc.styles['Footer']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')  # 设置中文字体
    doc.styles['Footer'].font.size = Pt(12)
    doc.styles['Footer'].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY  # 两端对齐
    doc.styles['Footer'].paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE  # 1.5倍行距
    doc.styles['Footer'].paragraph_format.space_before = Pt(0)  # 段前
    doc.styles['Footer'].paragraph_format.space_after = Pt(0)  # 段后
    doc.styles['Footer'].paragraph_format.first_line_indent = doc.styles['Footer'].font.size * 2  # 首行缩进2字符 1 英寸=2.54 厘米
    # 表格样式
    doc.styles['Normal'].font.name = 'Times New Roman'  # 设置英文字体
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')  # 设置中文字体
    # doc.styles['Normal'].font.size = Pt(12)
    doc.styles['Normal'].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.DISTRIBUTE  # 分散对齐
    # doc.styles['Normal'].paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST # 最小倍倍行距
    doc.styles['Normal'].paragraph_format.space_before = Pt(0)  # 段前
    doc.styles['Normal'].paragraph_format.space_after = Pt(0)  # 段后
    doc.styles['Normal'].paragraph_format.first_line_indent = Inches(0)  # 首行缩进2字符 1 英寸=2.54 厘米

    # 标题 一段
    doc.add_paragraph(a, style='Header')
    # 称呼两段（首不设两字符）
    doc.add_paragraph(b, style='Footer').paragraph_format.first_line_indent = Inches(0)  # 1 英寸=2.54 厘米
    doc.add_paragraph(c, style='Footer').paragraph_format.first_line_indent = Inches(0)  # 1 英寸=2.54 厘米
    # 正文
    doc.add_paragraph(d, style='Footer')
    doc.add_paragraph(e, style='Footer') # "特此呈报学院党委备案。"
    doc.add_paragraph(f, style='Footer')

    # 落款和时间
    doc.add_paragraph(g, style='Footer').alignment = WD_ALIGN_PARAGRAPH.RIGHT  # 右对齐
    doc.add_paragraph(h, style='Footer').alignment = WD_ALIGN_PARAGRAPH.RIGHT  # 右对齐 # "书记签名（加盖支部公章）：_______________"
    doc.add_paragraph(i, style='Footer').alignment = WD_ALIGN_PARAGRAPH.RIGHT  # "{}年{}月{}日".format(year,month,day)

    doc.add_paragraph(j, style='Footer') # "附备案人员名单："

    table = doc.add_table(worksheet.max_row,worksheet.max_column)
    table.autofit = True  # if is True 按窗口大小自动调整

    # 写入表格
    for row in range(len(table.rows)):
        table.rows[row].height = Cm(row_height)  # 调整行高
        for col in range(len(table.columns)):
            # print(行, 列)  # 可以查看表格输出结果
            set_cell_border(cell=table.cell(row, col),
                            top={"sz": 12, "val": "single", "color": "#000000", "space": "0"},
                            bottom={"sz": 12, "val": "single", "color": "#000000", "space": "0"},
                            left={"sz": 12, "val": "single", "color": "#000000", "space": "0"},
                            right={"sz": 12, "val": "single", "color": "#000000", "space": "0"},
                            )  # 设置边框
            if str(worksheet.cell(row+1, col+1).value) != 'None':
                table.cell(row, col).text = worksheet.cell(row+1, col+1).value  # 写入
                print(worksheet.cell(row+1, col+1).value)
                table.cell(row, col).width = doc.styles['Normal'].font.size * 1
                table.cell(row, col).height = doc.styles['Normal'].font.size * 5
                table.cell(row, col).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER  # 上下居中（中部居中对齐）
                table.cell(row, col).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 水平居中（中部居中对齐）

    # 修正列宽
    for col in range(len(table.columns)):
        maxlist = []
        for r in range(len(table.rows)):
            try:
                maxlist.append(len(str(table.cell(r,col).value)))
                # print(people_sheet[8*r + col])
            except:
                pass
        if maxlist != []:
            maxnum = max(maxlist)  # 每一列的最大值
        else:
            maxnum = 3  # 每一列的最大值
        table.cell(len(table.rows) - 1, col).width = Cm(
            col_width[0] if maxnum == 4 else col_width[1])  # 调整列宽 2字:1.3 3字:1.8 4字:2.1
        # 要在最后一行设置列宽度，因为设置前面的，后面一行出现空格，前面设置的宽度就不生效了

    table.alignment = WD_TABLE_ALIGNMENT.CENTER  # 设置整个表格为居中对齐
    table.autofit = True

    # # 结束语
    # if e != None:
    #     doc.add_paragraph(e, style='Footer')
    # if f != None:
    #     doc.add_paragraph(f, style='Footer')
    # doc.add_paragraph("", style='Footer')

    index = None
    for i in range(len(zhibu_allname)):
        if zhibu_allname[i] == party_name:
            index = i
    doc.save('{} {}.docx'.format(party_name if index == None else zhibu_list[index], a))

    messagebox.showinfo('小提示', '生成积极分子备案报告文件 成功！请注意检查word文件格式！')
    scr_output(scr3_11, '\n\n生成积极分子备案报告文件 成功！请注意检查word文件格式！\n')

    # except Exception as error:
    #     error = str(error)
    #     print('错误提示', error)
    #     scr_output(scr3_11, '\n生成积极分子备案报告文件 失败！\n\n\n本次错误信息：{}\n\n'.format(error))
    #     scr_output(scr3_11, '\n--------文件没有成功保存--------\n\n\n\n\n\n\n')
    #     messagebox.showinfo('错误提示', '生成积极分子备案报告文件 失败！\n错误信息：\n{}'.format(error))
'''#################################################################################################################'''
# # 将pdf转换成word（.docx）格式
# def pdf2word(path,out_path,sole=True):
#     path = eval(repr(path).replace(r'\\\\', r'/'))  # 把 \\ 替换成 /  不然会报错  一根是转义的\
#     path = eval(repr(path).replace('/', r'\\'))  # 把 / 替换成 \  不然会报错，
#     if sole==True:
#         # os.path.splitext(path)把文件名字分成两部分，名字和后缀
#         if os.path.splitext(path)[1] != '.pdf':
#             print('给定文件不是pdf文件')
#             return
#         pdf_file = path
#         word_file = os.path.splitext(path)[0] + '.docx'
#         cv = Converter(pdf_file)# 也支持相对路径
#         try:
#             cv.convert(word_file)
#         except Exception as error:
#             error = str(error)
#             print('错误提示', error)
#             if error == 'No parsed pages. Please parse page first.':
#                 print('错误是：pdf2docx.converter.ConversionException: No parsed pages. Please parse page first.')
#                 print('用word转pdf的PDF文件，再回来pdf转word会报这个错误！')
#         cv.close()
#     if sole==False:
#         # 判断有没有文件
#         if os.listdir(path) == []:
#             print("文件夹为空，请检查！")
#             return
#         for file in os.listdir(path): #  # os.listdir(path) 罗列文件夹下面的所有文件
#             if os.path.splitext(file)[1] != '.pdf':
#                 continue
#             file_name = os.path.splitext(file)[0]
#             pdf_file = path + '\\' + file
#             word_file = path + '\\' + file_name + '.docx'
#             cv = Converter(pdf_file)
#             try:
#                 cv.convert(word_file)
#             except Exception as error:
#                 error = str(error)
#                 print('错误提示', error)
#                 if error == 'No parsed pages. Please parse page first.':
#                     print('错误是：pdf2docx.converter.ConversionException: No parsed pages. Please parse page first.')
#                     print('用word转pdf的PDF文件，再回来pdf转word会报这个错误！')
#             cv.close()
# # 将doc和docx文件转换成pdf格式
# def word2pdf(path,out_path,sole=True):
#     import win32com.client as win32
#     path = eval(repr(path).replace(r'\\\\', r'/'))  # 把 \\ 替换成 /  不然会报错  一根是转义的\
#     path = eval(repr(path).replace('/', r'\\'))  # 把 / 替换成 \  不然会报错，
#     # 注意：word文件路径和生成pdf文件路径一定要使用绝对路径
#     # word = win32.Dispatch('Word.Application')
#     word = win32.gencache.EnsureDispatch('Word.Application')
#     if sole==True:
#         if (os.path.splitext(path)[1] != '.doc') and (os.path.splitext(path)[1] != '.docx'):
#             print('给定文件不是.doc或.docx文件')
#             return
#         doc = word.Documents.Open(path)
#         pdf_file = os.path.splitext(path)[0] + ".pdf"  # 生成pdf文件路径名称
#         doc.SaveAs(pdf_file, FileFormat=17)
#         print("文件{}完成.docx到.pdf的转换！".format(path))
#         doc.Close()
#         word.Quit()
#     if sole==False:
#         for dirpath, dirnames, filenames in os.walk(path): # path是文件夹地址
#             # dirpath是文件夹路径，dirnames为空，filenames是文件夹下面所有的文件名字
#             # 判断有没有文件
#             if filenames==[]:
#                 print("文件夹为空，请检查！")
#                 return
#             # 判断是不是含有.doc或者.docx文件
#             elif ".doc" or ".docx" in filenames:
#                 for file in filenames:
#                     if file.lower().endswith(".docx"):
#                         pdf_file = file.replace(".docx", ".pdf")
#                         word_file =(dirpath + '/'+ file)
#                         pdf_file =(dirpath + '/' + pdf_file)
#                         doc = word.Documents.Open(word_file)
#                         doc.SaveAs(pdf_file, FileFormat = 17)
#                         print("文件{}完成.docx到.pdf的转换！".format(word_file))
#                         doc.Close()
#                     elif file.lower().endswith(".doc"):
#                         pdf_file = file.replace(".doc", ".pdf")
#                         word_file =(dirpath +'\\' + file)
#                         pdf_file =(dirpath +'\\' + pdf_file)
#                         doc = word.Documents.Open(word_file)
#                         doc.SaveAs(pdf_file, FileFormat = 17)
#                         print("文件{}完成.doc到.pdf的转换！".format(word_file))
#                         doc.Close()
#         word.Quit()
def pdf2word_or_word2pdf(path,out_path,what_to_what,sole=True):
    from pdf2docx import Converter
    import win32com.client as win32
    print(path,out_path,what_to_what,sole)
    if what_to_what =='00' or sole == '00':
        messagebox.showinfo('小提示', '没有勾选是哪个转换模式')
        return
    if what_to_what =='10': what_to_what = 'pdf2word'
    if what_to_what =='01': what_to_what = 'word2pdf'
    if sole =='10': sole = True
    if sole =='01': sole = False
    path = eval(repr(path).replace(r'\\\\', r'/'))  # 把 \\ 替换成 /  不然会报错  一根是转义的\
    path = eval(repr(path).replace('/', r'\\'))  # 把 / 替换成 \  不然会报错，
    out_path = eval(repr(out_path).replace(r'\\\\', r'/'))  # 把 \\ 替换成 /  不然会报错  一根是转义的\
    out_path = eval(repr(out_path).replace('/', r'\\'))  # 把 / 替换成 \  不然会报错，
    try:
        scr_output(scr_10, '---------------------------------------------------------------------------------------'
                           '\n\n{}\n温馨提示：\n正在尝试开始转换，请稍后。\n文件较大或文件夹内文件较多时候，转换较慢，请耐心等候~\n\n'.format(what_to_what))
        if what_to_what == 'pdf2word':
            # 获取文件上级目录路径的方法
            # path = '/Users/caowei/数据集/Stable/EN-2300-103-RevA.txt'
            # p_path = os.path.abspath(os.path.join(path, "..")) # C:\Users\caowei\数据集\Stable
            # saveDir = os.path.dirname(path) # /Users/caowei/数据集/Stable
            if sole == True:
                # os.path.splitext(path)把文件名字分成两部分，名字和后缀
                if os.path.splitext(path)[1] != '.pdf':
                    # print('给定文件不是pdf文件')
                    scr_output(scr_10, '{}错误提示:\n给定文件不是pdf文件\n'.format(what_to_what))
                    messagebox.showinfo('{}错误提示'.format(what_to_what), '\n给定文件不是pdf文件\n')
                    return
                pdf_file = path
                word_file = out_path + '\\' + os.path.basename(path) + '.docx'
                cv = Converter(pdf_file)  # 也支持相对路径
                try:
                    cv.convert(word_file)
                    scr_output(scr_10, "\n文件'{}'完成.pdf到.docx的转换！\n保存路径：{}".format(os.path.basename(path), word_file))
                    messagebox.showinfo('小提示', '{}转换成功！'.format(what_to_what))
                except Exception as error:
                    # error = str(error)
                    # print('错误提示', error)
                    scr_output(scr_10, '\n{}错误提示：\n错误是：{}\n'.format(what_to_what,error))
                    if error == 'No parsed pages. Please parse page first.':
                        # print('错误是：pdf2docx.converter.ConversionException: No parsed pages. Please parse page first.')
                        # print('用word转pdf的PDF文件，再回来pdf转word会报这个错误！')
                        scr_output(scr_10,'错误提示：\n错误是：pdf2docx.converter.ConversionException: No parsed pages. Please parse page first.\n' )
                        scr_output(scr_10, '\n用word转pdf的PDF文件，再回来pdf转word会报这个错误！\n')
                cv.close()
            if sole == False:
                    # 判断有没有文件
                    if os.listdir(path) == []:
                        print("\n文件夹为空，请检查！\n")
                        scr_output(scr_10, '{}错误提示：\n文件夹为空，请检查！\n'.format(what_to_what))
                        messagebox.showinfo('{}错误提示'.format(what_to_what), "\n文件夹为空，请检查！\n")
                        return
                    for file in os.listdir(path):  # # os.listdir(path) 罗列文件夹下面的所有文件
                        if os.path.splitext(file)[1] != '.pdf':
                            continue
                        file_name = os.path.splitext(file)[0]
                        pdf_file = path + '\\' + file
                        word_file = out_path + '\\' + file_name + '.docx'
                        cv = Converter(pdf_file)
                        try:
                            cv.convert(word_file)
                        except Exception as error:
                            # error = str(error)
                            # print('错误提示', error)
                            if error == 'No parsed pages. Please parse page first.':
                                # print('错误是：pdf2docx.converter.ConversionException: No parsed pages. Please parse page first.')
                                # print('用word转pdf的PDF文件，再回来pdf转word会报这个错误！')
                                scr_output(scr_10, '错误提示：\n错误是：pdf2docx.converter.ConversionException: No parsed pages. Please parse page first.\n')
                                scr_output(scr_10, '\n用word转pdf的PDF文件，再回来pdf转word会报这个错误！\n')
                        scr_output(scr_10, "\n文件'{}'完成.pdf到.docx的转换！\n保存路径：{}".format(file_name,word_file))
                    cv.close()
                    messagebox.showinfo('小提示', '{}转换成功！'.format(what_to_what))
        if what_to_what =='word2pdf':
            # 注意：word文件路径和生成pdf文件路径一定要使用绝对路径
            # word = win32.Dispatch('Word.Application')
            word = win32.gencache.EnsureDispatch('Word.Application')
            if sole == True:
                if (os.path.splitext(path)[1] != '.doc') and (os.path.splitext(path)[1] != '.docx'):
                    # print('给定文件不是.doc或.docx文件')
                    scr_output(scr_10, '{}错误提示：\n给定文件不是.doc或.docx文件\n'.format(what_to_what))
                    messagebox.showinfo('{}错误提示'.format(what_to_what), '\n给定文件不是.doc或.docx文件\n')
                    return
                doc = word.Documents.Open(path)
                pdf_file = os.path.splitext(path)[0] + ".pdf"  # 生成pdf文件路径名称
                doc.SaveAs(pdf_file, FileFormat=17)
                doc.Close()
                word.Quit()
                # print("文件{}完成.docx到.pdf的转换！".format(path))
                scr_output(scr_10, "\n文件'{}'完成.docx到.pdf的转换！\n保存路径：{}".format(os.path.dirname(path), pdf_file))
                messagebox.showinfo('小提示', '{}转换成功！'.format(what_to_what))
            if sole == False:
                for dirpath, dirnames, filenames in os.walk(path):  # path是文件夹地址
                    # dirpath是文件夹路径，dirnames为空，filenames是文件夹下面所有的文件名字
                    # 判断有没有文件
                    if filenames == []:
                        # print("文件夹为空，请检查！")
                        scr_output(scr_10, '{}错误提示：\n文件夹为空，请检查！\n'.format(what_to_what))
                        messagebox.showinfo('{}错误提示'.format(what_to_what), "\n文件夹为空，请检查！\n")
                        return
                    # 判断是不是含有.doc或者.docx文件
                    elif ".doc" or ".docx" in filenames:
                        for file in filenames:
                            if file.lower().endswith(".docx"):
                                pdf_file = file.replace(".docx", ".pdf")
                                word_file = (path + '/' + file)
                                pdf_file = (out_path + '/' + pdf_file)
                                doc = word.Documents.Open(word_file)
                                doc.SaveAs(pdf_file, FileFormat=17)
                                # print("文件{}完成.docx到.pdf的转换！".format(word_file))
                                scr_output(scr_10, "\n文件'{}'完成.docx到.pdf的转换！\n保存路径：{}".format(file,pdf_file))
                                doc.Close()
                            elif file.lower().endswith(".doc"):
                                pdf_file = file.replace(".doc", ".pdf")
                                word_file = (dirpath + '\\' + file)
                                pdf_file = (out_path + '\\' + pdf_file)
                                doc = word.Documents.Open(word_file)
                                doc.SaveAs(pdf_file, FileFormat=17)
                                # print("文件{}完成.doc到.pdf的转换！".format(word_file))
                                scr_output(scr_10, "\n文件'{}'完成.doc到.pdf的转换！\n".format(file))
                                doc.Close()
                word.Quit()
                messagebox.showinfo('小提示', '{}转换成功！'.format(what_to_what))
    except Exception as error:
        error = str(error)
        messagebox.showinfo('{}错误提示'.format(what_to_what), error)
        scr_output(scr_10,'\n尝试{}失败！\n错误提示：\n'.format(what_to_what, error))
        # print('错误提示', error)
'''#################################################################################################################'''
def getText(wordname):
    texts = []
    d = Document(wordname)
    for para in d.paragraphs:
        texts.append(para.text)
    return texts
def is_Chinese(word):
    for ch in word:
        if '\u4e00' <= ch <= '\u9fff':
            return True
    return False
def msplit(s, seperators=',|\.|\?|，|。|？|！'):
    import re
    # 这些符号来断句
    return re.split(seperators, s)
def readDocx(docfile):
    # print('*' * 80)
    # print('文件', docfile, '加载中……')
    t1 = datetime.datetime.now()
    paras = getText(docfile)
    segs = []
    for p in paras:
        temp = []
        for s in msplit(p):
            if len(s) > 2:
                temp.append(s.replace(' ', ""))
        if len(temp) > 0:
            segs.append(temp)
    t2 = datetime.datetime.now()
    # print('加载完成，用时: ', t2 - t1)
    showInfo(segs, docfile)
    return segs
def showInfo(doc, filename='filename'):
    chars = 0
    segs = 0
    for p in doc:
        for s in p:
            segs = segs + 1
            chars = chars + len(s)
    # print('段落数: {0:>8d} 个。'.format(len(doc)))
    # print('短句数: {0:>8d} 句。'.format(segs))
    # print('字符数: {0:>8d} 个。'.format(chars))
# from win32com import client
# def doc2docx(file_path):
#     file_path = os.path.abspath(file_path)
#     print(file_path)
#     word = client.Dispatch("Word.Application")  # 打开word应用程序
#     doc = word.Documents.Open(file_path)  # 打开word文件
#
#     a = os.path.split(file_path)  # 分离路径和文件
#     b = os.path.splitext(a[-1])[0]  # 拿到文件名
#     c = "{}\\{}.docx".format(a, b)
#     doc.SaveAs(c, 12)  # 另存为后缀为".docx"的文件，其中参数12或16指docx文件
#     doc.Close()  # 关闭原来word文件
#     word.Quit()
#     return c
def all_files_path(rootDir, filepaths):
    for root, dirs, files in os.walk(rootDir):  # 分别代表根目录、文件夹、文件
        for file in files:  # 遍历文件
            file_path = os.path.join(root, file)  # 获取文件绝对路径
            if os.path.splitext(file_path)[1] == '.docx':
                filepaths.append(file_path)  # 将文件路径添加进列表
            # if os.path.splitext(file_path)[1] == '.doc':   # doc忽略，问题如何解决……？？？？？？？？？？
            #     c = doc2docx(file_path)
            #     filepaths.append(c)
        # for dir in dirs:  # 遍历目录下的子目录  # 翻倍重复了  是个Bug
        #     dir_path = os.path.join(root, dir)  # 获取子目录路径
        #     all_files_path(dir_path,filepaths)  # 递归调用
    return filepaths
def compareParagraph(doc1, i, doc2, j, count_set, ratio_set, filepath, min_segment=5):
    # c = str('\n')
    # print("正在compareParagraph……")
    """
    功能为比较两个段落的相似度，返回结果为两个段落中相同字符的长度与较短段落长度的比值。
    :param p1: 行
    :param p2: 列
    :param min_segment = 5: 最小段的长度
    """
    p1 = doc1[i]
    p2 = doc2[j]
    len1 = sum([len(s) for s in p1])
    len2 = sum([len(s) for s in p2])
    if len1 < 10 or len2 < 10:
        return []

    listc = []
    for s1 in p1:
        if len(s1) < min_segment:
            continue
        for s2 in p2:
            if len(s2) < min_segment:
                continue
            if s2 in s1:
                listc.append(s2)
            elif s1 in s2:
                listc.append(s1)

    # 取两个字符串的最短的一个进行比值计算
    count = sum([len(s) for s in listc])
    ratio = float(count) / min(len1, len2)
    if count > 50 and ratio > 0.5:  # 超高重复率才显示窗口
        scr_output(scr_12, '\n' + ' 发现超高重复率内容 '.center(80, '*'))
        scr_output(scr_12, '\n与{}段落 存在相似\n'.format(filepath))
        scr_output(scr_12, '文件1第{0:0>4d}段内容：{1}'.format(i + 1, p1))
        scr_output(scr_12, '文件2第{0:0>4d}段内容：{1}'.format(j + 1, p2))
        scr_output(scr_12, '\n相同内容：{}'.format(listc))
        scr_output(scr_12, '相同字符比：{1:.2f}%\n相同字符数： {0}\n'.format(count, ratio * 100))
        scr_output(scr_12, ' 结束 '.center(80, '*') + '\n')

    if count > count_set and ratio > ratio_set:
        c = '\n' + ' 发现相同内容 '.center(80, '*')
        c += '\n与{}段落 存在相似\n'.format(filepath)
        c += '文件1第{0:0>4d}段内容：{1}'.format(i + 1, p1)
        c += '文件2第{0:0>4d}段内容：{1}'.format(j + 1, p2)
        c += '\n\n相同内容：{}'.format(listc)
        c += '相同字符比：{1:.2f}%\n相同字符数： {0}\n'.format(count, ratio * 100)
        c += ' 结束 '.center(90, '*') + '\n'
        return c
    else:
        return None
def cha_chong_shujuku(path, count_set, ratio_set, data, type, sole_many):
    try:
        if sole_many != '10' and sole_many != '01':
            messagebox.showinfo('小提示', '请选择文件位置')

        # if data == None and number_12_2.get() == 0:
        try:
            dirpath = 'mould\\chachong' # if data[1] == None and data[0] == '全部材料'
        except:
            messagebox.showinfo('小提示', 'mould\\chachong，文件夹不存在')
            return
        if data[1] != None and  data[0] == '1':  # 是否新建数据库
            scr_output(scr_12, '\n已选择的查重数据库路径：{}\n'.format(data[1]))
            dirpath = data[1]
        else:
            if type[0] != '全部材料':
                scr_output(scr_12, '\n已选择的查重数据库路径：mould\\chachong\\{}\n'.format(type[0]))
                dirpath = 'mould\\chachong\\{}'.format(type[0])

        filepaths = []  # 初始化文件列表用来
        filepaths = all_files_path(dirpath,filepaths)
        # print(filepaths)
        if len(filepaths) == 0:
            messagebox.showinfo('小提示', '数据库没有找到文件')
            return
        scr_output(scr_12, ' ...初始化数据库... '.center(80, '*') + '\n数据库文件数量为：{}篇\n'.format(len(filepaths)))
        t1 = datetime.datetime.now()
        if sole_many == '10':  # 单个文件
            if os.path.splitext(path)[1] != '.docx':
                messagebox.showinfo('小提示', '请选择docx后缀的word格式')
                return
            c = ' 开始比对... '.center(80, '*') + '\n数据库文件数量为：{}篇\n'.format(len(filepaths))
            for fi in range(len(filepaths)): # 遍历数据库每一个文件地址
                if len(filepaths) >= 50 :
                    if fi != 0 and fi % 50 == 0:
                        scr_output(scr_12, '\n已经完成比对{}篇完成，累计用时: {}'.format(fi,datetime.datetime.now() - t1))
                    if fi % round(len(filepaths)*0.1,0) == 0 and len(filepaths) >= 1000: # 十分之一汇报一次
                        scr_output(scr_12, '处理进行中，已处理文章数量 {0:>4d} (总数 {1:0>4d} ） '.format(fi, len(filepaths)))
                        c = c + '处理进行中，已处理文章数量 {0:>4d} (总数 {1:0>4d} ） '.format(fi, len(filepaths))
                doc1 = readDocx(path)
                doc2 = readDocx(filepaths[fi])
                for i in range(len(doc1)):
                    for j in range(len(doc2)):
                        cs = compareParagraph(doc1, i, doc2, j, count_set, ratio_set, filepaths[fi])
                        if isinstance(cs, str) == True: c  = c + cs   # 判断是否是字符串型
            t2 = datetime.datetime.now()
            scr_output(scr_12, '\n比对完成，对比数据库数据量：{}篇，总用时: {}\n详细报告请移步查重报告文件\n'.format(len(filepaths), t2 - t1))
            c += '\n比对完成，对比数据库数据量：{}篇，总用时: {}\n'.format(len(filepaths), t2 - t1)
            # content = scr_12.get('1.0', 'end-1c')
            with open('{} 查重报告.txt'.format(os.path.splitext(path)[0]), 'w') as f:
                f.write(c)
        if sole_many == '01': # 文件夹
            fileps = []  # 初始化文件列表用来
            all_files_path(path,fileps)
            for fp in fileps:
                content2 = '\n\n' + ' 开始比对... '.center(80, '*') + '\n文章：{}\n'.format(fp)
                scr_output(scr_12, '\n文章：{}\n'.format(fp) + ' 开始比对... '.center(80, '*'))
                doc3 = readDocx(fp)
                for fi in range(len(filepaths)):  # 遍历数据库每一个文件地址
                    if len(filepaths) >= 50:
                        if fi != 0 and fi % 50 == 0:
                            scr_output(scr_12, '\n已经完成比对{}篇完成，累计用时: {}'.format(fi, datetime.datetime.now() - t1))
                        if fi % round(len(filepaths) / 10, 0) == 0 and len(filepaths) >= 1000:  # 十分之一汇报一次
                            scr_output(scr_12, '处理进行中，已处理文章数量 {0:>4d} (总数 {1:0>4d} ） '.format(fi, len(filepaths)))
                            content2 += '处理进行中，已处理文章数量 {0:>4d} (总数 {1:0>4d} ） '.format(fi, len(filepaths))
                    doc4 = readDocx(filepaths[fi])
                    for i in range(len(doc3)):
                        for j in range(len(doc4)):
                            contents2 = compareParagraph(doc3, i, doc4, j, count_set, ratio_set, filepaths[fi])
                            if isinstance(contents2, str) == True: content2 += contents2  # 判断是否是字符串型
                t2 = datetime.datetime.now()
                scr_output(scr_12, '\n比对完成，对比数据库数据量：{}篇，总用时: {}\n详细报告请移步查重报告文件\n'.format(len(filepaths), t2 - t1))
                content2 += '\n比对完成，对比数据库数据量：{}篇，总用时: {}\n'.format(len(filepaths),t2 - t1)
                t1 = datetime.datetime.now()
                # content = scr_12.get('1.0', 'end-1c')
                # if not os.path.exists(os.path.splitext(fp)[0]): os.mkdir(os.path.splitext(fp)[0])  # 如果不存在，创建目录
                with open('{} 查重报告.txt'.format(os.path.splitext(fp)[0]), 'w') as f:
                    f.write(content2)
    except Exception as error:
        error = str(error)
        messagebox.showinfo('{}错误提示'.format('查重'), error)
        scr_output(scr_10,'\n尝试{}失败！\n错误提示：\n'.format('查重', error))
        # print('错误提示', error)

def cha_chong_tongpici(dirpath, count_set, ratio_set):
    if dirpath == '':
        messagebox.showinfo('小提示', '请选择文件夹位置')
    filepaths = []  # 初始化文件列表用来
    filepaths = all_files_path(dirpath,filepaths)
    # print(filepaths)
    if len(filepaths) <= 1:
        messagebox.showinfo('小提示', '数据库文件夹没有找到文件或只有一个文件')
        return
    scr_output(scr_12, ' ...初始化数据库... '.center(80, '*') + '\n数据库文件数量为：{}篇\n'.format(len(filepaths)))
    t1 = datetime.datetime.now()
    fileps = []  # 初始化文件列表用来
    all_files_path(dirpath,fileps)
    for fp in fileps:
        content2 = '\n\n' + ' 开始比对... '.center(80, '*') + '\n文章：{}\n'.format(fp)
        scr_output(scr_12, '\n文章：{}\n'.format(fp) + ' 开始比对... '.center(80, '*'))
        doc3 = readDocx(fp)
        for fi in range(len(filepaths)):  # 遍历数据库每一个文件地址
            if fp != filepaths[fi]: # 不是同一个文件
                if len(filepaths) >= 50:
                    if fi != 0 and fi % 50 == 0:
                        scr_output(scr_12, '\n已经完成比对{}篇完成，累计用时: {}'.format(fi, datetime.datetime.now() - t1))
                    if fi % round(len(filepaths) / 10, 0) == 0 and len(filepaths) >= 1000:  # 十分之一汇报一次
                        scr_output(scr_12, '处理进行中，已处理文章数量 {0:>4d} (总数 {1:0>4d} ） '.format(fi, len(filepaths)))
                        content2 += '处理进行中，已处理文章数量 {0:>4d} (总数 {1:0>4d} ） '.format(fi, len(filepaths))
                doc4 = readDocx(filepaths[fi])
                for i in range(len(doc3)):
                    for j in range(len(doc4)):
                        contents2 = compareParagraph(doc3, i, doc4, j, count_set, ratio_set, filepaths[fi])
                        if isinstance(contents2, str) == True: content2 += contents2  # 判断是否是字符串型
        t2 = datetime.datetime.now()
        scr_output(scr_12, '\n比对完成，对比数据库数据量：{}篇，总用时: {}\n详细报告请移步查重报告文件\n'.format(len(filepaths)-1, t2 - t1))
        content2 += '\n比对完成，对比数据库数据量：{}篇，总用时: {}\n'.format(len(filepaths)-1,t2 - t1)
        t1 = datetime.datetime.now()
        # content = scr_12.get('1.0', 'end-1c')
        # if not os.path.exists(os.path.splitext(fp)[0]): os.mkdir(os.path.splitext(fp)[0])  # 如果不存在，创建目录
        with open('{} 查重报告.txt'.format(os.path.splitext(fp)[0]), 'w') as f:
            f.write(content2)
'''#################################################################################################################'''
# 帮助系列
def open_abspath(): #打开执行文件夹目录
    path = os.path.abspath('.')
    # print(path)
    os.system("explorer.exe %s" % path)
    # os.getcwd()#获得当前工作目录
    # os.path.abspath('.')#获得当前工作目录
    # os.path.abspath('..')#获得当前工作目录的父目录
    # os.path.abspath(os.curdir)#获得当前工作目录
def version_():
    import webbrowser
    if messagebox.askokcancel('小提示','请点击‘确定’跳转到指定网址，请牢记云盘访问密码是：0911 \n如果不能成功打开，请自行访问蓝奏云链接：https://wwu.lanzoui.com/b02i9hztg '
                                    '\n。如果还是打不开，请联系作者获取最新链接'):
        url = 'https://wwu.lanzoui.com/b02i9hztg'
        webbrowser.open(url = url, new = 0)
def about_():  # 关于软件
    import webbrowser
    messagebox.showinfo('小提示','正在尝试打开软件说明文档，可能稍有延迟，请等待加载\n') # PS: 如果系统提示文件打开方式，请选择用记事本或其他文本编辑工具打开
    # url = 'https://github.com/HongWu-122/Party-building-decision-support-system' # 源码地址
    try:
        webbrowser.open(url='mould\README_zh.html', new=0)
    except:
        messagebox.showinfo('错误提示', '打开失败！原因：{}\n'.format(error))

def help_():  # 打开帮助
    messagebox.showinfo('小提示','正在尝试打开帮助和更新文档，可能稍有延迟，请等待加载\n') # PS: 如果系统提示文件打开方式，请选择用记事本或其他文本编辑工具打开
    try:
        os.system("start " + 'mould\README_zh.md')
    except Exception as error:
        messagebox.showinfo('错误提示','打开失败！原因：{}\n'.format(error))

def video_():
    import webbrowser
    if messagebox.askokcancel("小提示", "应广大用户的要求，本软件为许多电脑小白们配套了视频教程帮助更快的入门，如果不能成功打开，请自行访问：https://www.bilibili.com/video/BV1m44y1H75s/\n请点击‘确定’跳转到指定网址"):
        url = 'https://www.bilibili.com/video/BV1m44y1H75s/'
        # url = 'https://www.bilibili.com/video/BV1Wf4y1H7zd?p=1&share_medium=android&share_plat=android&share_session_id=54db8f2e-9f2e-4068-a2ac-39b17f10c149&share_source=COPY&share_tag=s_i&timestamp=1630294478&unique_k=qgq4kr'
        # urlv1 = https://b23.tv/qgq4kr
        # urlv2 = https://www.bilibili.com/video/BV1m44y1H75s/
        webbrowser.open(url = url, new = 0)
def case_():
    if messagebox.askokcancel("小提示", "为了方便广大用户更好上手、更容易熟悉本程序操作流程，\n本软件编写了初始化案例程序，可以为您生成初始化的案例文件，\n点击‘确定’开始生成"):
        case_file_book()
        case_file_count()
def author_():
    messagebox.showinfo('小提示','感谢您不离不弃一如既往地使用本软件，如果您在使用过程中发现任何BUG，请联系我~\n也可以通过菜单栏编辑功能联系我\n'
                               '如果您也想要开发一款软件或工具的需求，可以联系我告诉我需求和想法。\n当然如果您懂得如何使用python，并且有意愿加入本团队共同开发本软件的，也可以联系我哦~\nQQ：1228815090')
def thank_():
    messagebox.showinfo('致谢','致谢党建组织部的每一位同学，特别要感谢会计213班李晶晶同学，感谢他们不求回报、默默付出，以及在本软件开发、运行、维护的过程中所给予的帮助、建议和鼓励！\n')
'''#################################################################################################################'''
# 菜单栏的有趣玩意
def change_background_image():
    pass
#   from PIL import Image as imm
#   from PIL import ImageTk as imk
#     global a,window,bg_im_label,background_image_list
#
#     im_label = imk.PhotoImage(imm.open(r'{}'.format(random.choice(background_image_list))).resize((580, 320)))
#     bg_im_label = Label(tab0, image=im_label)  # , textvariable=var , compound=CENTER , text='hello'
#     bg_im_label.place(x=680, y=70)
#     window.mainloop()
def open_background_music():  # 开启音乐
    try:
        import pygame
        pygame.mixer.init()
        # print('v.get()',v.get())
        if v.get()==1:  # 打勾状态
            file_path = r'mould\media'
            file_lists = os.listdir(file_path)
            # os.listdir()用于返回指定的文件夹下包含的文件或文件夹名字的列表，这个列表按字母顺序排序。
            # os.walk可以用于遍历指定文件下所有的子目录、非目录子文件。
            a = 0

            file_list = []
            for i in file_lists:
                file_list.append('mould/media/' + file_lists[a])
                a = a + 1
            # print(random.choice(file_list))
            b = random.choice(file_list)  # 在列表里随机取出一个数返回  需要定义一个变量储存 不然random出来的file_list每次都会变化
            messagebox.showinfo('小提示', '正在播放"{}"'.format(b))
            track = pygame.mixer.music.load( b )   ###########################################
            pygame.mixer.music.play()
        else:
            pygame.mixer.music.stop()
    except:
        messagebox.showinfo('小提示', 'mould/media/文件夹没有音乐或者其他原因，无法播放')
def change_home_style():
    global window,window_style
    from ttkbootstrap import Style as ttkStyle
    theme_list = ['orange_theme','cerculean','darkly'] # ,'simplex','morph','journal','solar','superhero','vapor','cyborg'
    # style = ttkStyle(theme='cerculean')
    new_window_style = random.choice(theme_list)
    if window_style == None:
        window_style = new_window_style
    else:
        while new_window_style == window_style:
            new_window_style = random.choice(theme_list)
    window_style = new_window_style
    style = ttkStyle(theme= window_style)
    window = style.master
    window.mainloop()
#自定义的线程函数类
def thread_it(func, *args):
    import threading
    '''将函数放入线程中执行'''
    # 创建线程
    t = threading.Thread(target=func, args=args)
    # 守护线程
    t.setDaemon(True)
    # 启动线程
    t.start()
def xi_sayings():
    def find():
        try:
            import linecache  # 获取一行数据
            txt = open(r'mould/data/sayings.txt', 'rb')
            data = txt.read().decode('utf-8')  # python3一定要加上这句不然会编码报错！
            txt.close()

            # 获取txt的总行数！
            n = data.count('\n')
            print("总行数", n)
            # 选取随机的数
            i = random.randint(1, (n + 1))
            print("本次使用的行数", i)
            ###得到对应的i行的数据
            line = linecache.getline(r'mould/data/sayings.txt', i)
            print(line)
            return line
        except:
            return "找不到语录文件：mould/data/sayings.txt！"

    def say(line):
        try:
            import win32com.client as win32
            # 1.创建一个播报对象
            speaker = win32.Dispatch("SAPI.SpVoice")
            # 2.通过这个播报器对象，直接播放对应的语音字符串就可以
            speaker.Speak(line)
        except:
            messagebox.showinfo('小提示', '很遗憾，语音播报模块不可用')

    line = find()
    window_xijp = Toplevel()
    window_xijp.title('学习党性语录')
    # 设置背景颜色
    # window_xijp['background'] = 'gray20'
    # 禁止窗口的拉伸
    window_xijp.resizable(0, 0)  # == (Flase,Flase)
    # 窗口置顶
    window_xijp.attributes("-topmost", 1)  # 1==True 处于顶层
    window_xijp.wm_attributes("-alpha", 0.7)  # 透明度(0.0~1.0
    try:
        window_xijp.iconbitmap('mould\data\logo.ico')
    except:
        pass
    b = len(line) * 19 + 44
    print("len:  %s" % b)
    if b >= 1500:
        b = len(line) / 2 * 19 + 44
        # root.geometry("%dx1+0+70" % b)
        window_xijp.geometry("%dx140+0+112" % b)  # x乘以y 加上坐标 x要等于b的字符串长度，y一行取80，两行取120
        print("新:  %d" % b)
        l = int(len(line) / 2)
        line1 = line[:l]
        line2 = line[l:]
        a1 = Label(window_xijp, text=line1, font=("华文中宋", 14, "normal"),compound=CENTER,wraplength=102) # # 自动换行  wraplength 是以像素为单位进行换行的
        a1.place(x=20, y=20)
        a2 = Label(window_xijp, text=line2, font=("华文中宋", 14, "normal"))
        a2.place(x=20, y=60)
        button = Button(window_xijp, text="点击播放语录", width="10", command=lambda: thread_it(say(line)))
        button.place(x=20, y=100)
        label = Label(window_xijp, text='To YOU', font=('华文中宋', 13)).place(x=int(int("%d" % b) - 93), y=110)
        # print(int(int("%d"% b)-90))
    else:
        # root.geometry("%dx1+0+20" % b)
        window_xijp.geometry("%dx97+0+112" % b)
        # 为了区别root和tl，我们向tl中添加了一个Label
        a = Label(window_xijp, text=line, font=("华文中宋", 14, "normal", ))
        a.place(x=20, y=20)
        button = Button(window_xijp, text="点击播放语录", width="10", command=lambda: thread_it(say(line)))
        button.place(x=20, y=60)
        label = Label(window_xijp, text='To YOU', font=('华文中宋', 13)).place(x=int(int("%d" % b) - 93), y=60)
        # print(int(int("%d" % b) - 90))
    window_xijp.mainloop()
def my_heart():
    from PIL import Image as imm
    from PIL import ImageTk as imk
    def winning():
        print('恭喜你，猜中了')
        xingxing_path = 'mould/data/xingxing'
        # if not os.path.exists(xingxing_path):  # 判断是否存在文件
        f = open(xingxing_path,'a+',encoding='utf-8')
        f.write('*')
        f.close()
        messagebox.showinfo('小提示', '恭喜你，猜中了啊，获得星星+1。星星可以用于联系作者，是消耗品哦。并会在小黑实验室荣誉墙上展示，重启软件后生效。')


    def encourage():
        xi_sayings()
        # try:
        #     import linecache
        #     获取一行数据
        #     txt = open(r'mould/data/sayings.txt', 'rb')
        #     data = txt.read().decode('utf-8')  # python3一定要加上这句不然会编码报错！
        #     txt.close()
        #     # 获取txt的总行数！
        #     n = data.count('\n')
        #     print("总行数", n)
        #     # 选取随机的数
        #     i = random.randint(1, (n + 1))
        #     print("本次使用的行数", i)
        #     ###得到对应的i行的数据
        #     line = linecache.getline(r'mould/data/sayings.txt', i)
        #     print(line)
        #     # messagebox.showinfo('小提示', '接近了，再接再厉')
        #     messagebox.showinfo('小提示', line)
        # except:
        #     messagebox.showinfo('小提示', "找不到语录文件：mould/data/sayings.txt！")

    def tip():
        messagebox.showinfo('小提示', '恭喜你，获得位置提示，位置在第{}个(从左往右，从上到下依次数)'.format(e))

    # messagebox.showinfo('小提示', '正在紧张制作中')
    win_heart = Toplevel()
    win_heart.geometry('601x481')
    win_heart.title("吾心年年望雪待月，漫倚吟矶不得忘")
    win_heart.attributes("-topmost", 1)  # 1==True 处于顶层
    try:
        im_label = imk.PhotoImage(
            imm.open(r'{}'.format('mould/data/my_heart.jpg')).resize((600, 480)))  # resize 按size大小调整图像.resize((600,480))
        label = Label(win_heart, image=im_label)
        label.pack()

        im = imm.open("mould/data/my_heart.jpg", "r").resize((600, 480))
        # im = ImageEnhance.Contrast(im).enhance(0.9) # 亮度
        # im = ImageEnhance.Color(im).enhance(0.9) # 饱和度
    except:pass
    # 照片剪切坐标
    x1 = 0
    y1 = 0
    x2 = 120
    y2 = 40
    # 按钮坐标
    a = 0
    b = 0
    l = 0  # 根据列表对变量名字命名，列表切片
    d = 0  # 判断是否有中奖概率
    e = None
    im_button = []
    button = []
    for k in range(1, 61):
        im_button.append('im_button{}'.format(k))
        button.append('button{}'.format(k))
    print(button)
    for j in range(1, 13):  # 竖
        for i in range(1, 6):  # 横
            box = (x1, y1, x2, y2)
            try:
                region = im.crop(box)  # 现成切割
                im_button[l] = imk.PhotoImage(region)  # 直接调用
                # region.save('jpg/{}_{}.jpg'.format(j, i))  # 保存本地
                # im_button[l] = ImageTk.PhotoImage(Image.open(r'Media_files\jpg\{}_{}.jpg'.format(j, i)).resize((120, 40)))  # 本地打开方法
                button[l] = Button(win_heart, image=im_button[l])  # , state='active'
            except:
                button[l] = Button(win_heart)  # , state='active'
            c = random.randint(1, 60)
            print(c, l)
            if c != l and c % 3 == 0:
                button[l].configure(command=encourage)
                print('c 鼓励', l + 1)
            if c == l:
                d = d + 1
                button[l].configure(command=winning)
                e = l + 1
                print('位置在', l + 1)
            if c != l and c in [13, 17, 23, 29, 47, 37, 57] and e != None:
                print('恭喜你，获得提示，位置在{}'.format(e))
                button[l].configure(command=tip)

            button[l].place(x=a, y=b, width=120, height=40)
            x1 = x1 + 120
            x2 = x2 + 120
            a = a + 120
            l = l + 1
        x1 = 0
        y1 = y1 + 40
        x2 = 120
        y2 = y2 + 40
        a = 0
        b = b + 40

    if d == 0:
        print('本次中奖概率为零')
        messagebox.showinfo('小提示', '本次中奖概率为零')
    else:
        d = (d / 60) * 100
        print('本次中奖概率为{:.2f}%'.format(d))
        messagebox.showinfo('小提示', '本次中奖概率为{:.2f}%'.format(d))

    win_heart.mainloop()
def sending():
    import yagmail  # 邮件模块
    global xingxing_Y
    if  xingxing_Y == True:
        a = text.get('1.0',END)
        # print('全文:\n',a)
        content = '全文:\n{}'.format(a)
        print(content)
        # print(type(content))        #  <class 'str'>
        try:
            import win32com.client as win32
            # 1.创建一个播报对象
            speaker = win32.Dispatch("SAPI.SpVoice")
        except:
            print('语音播报模块不可用')
            messagebox.showinfo(title='飞鸽传书', message="很遗憾，语音播报模块不可用，但不影响发送功能")
        try:
            print('正在运行发送邮件模块')
            messagebox.showinfo(title='飞鸽传书', message="点击确定，开始运行发送邮件模块，期间软件可能会有点卡，请耐心等待！")

            # 连接服务器
            # 用户名、授权码、服务器地址
            yag_server = yagmail.SMTP(user='h13902688308@163.com', password='NCXGETOKOEQRQYVK', host='smtp.163.com')
            '''接着，通过 send() 函数，将邮件发送出去'''
            # 发送对象列表
            email_to = ['1228815090@qq.com', ]
            email_title = '党建使用者私信'
            email_content = str(platform.uname()) + '\n' + content
            # 附件列表
            # email_attachments = ['./attachments/report.png', ]
            # 发送邮件
            yag_server.send(email_to, email_title, email_content)  # , email_attachments
            '''邮件发送完毕之后，关闭连接即可'''
            # 关闭连接
            yag_server.close()

            time.sleep(1)
            try:
                # import win32com.client as win32
                # # 1.创建一个播报对象
                # speaker = win32.Dispatch("SAPI.SpVoice")
                speaker.Speak("发送成功")
            except:
                pass
            xingxing_Y = False
            messagebox.showinfo(title='飞鸽传书', message="感谢来信，发送成功！")
        except:
            try:
                speaker = win32.Dispatch("SAPI.SpVoice")
                speaker.Speak("发送失败")
            except:
                pass
            messagebox.showinfo(title='飞鸽传书', message="哎呀，小鸽子发送失败了，请检查下网络连接，或者其他问题功能不可用")
    else:
        messagebox.showinfo(title='飞鸽传书', message="哎呀，小鸽子发送失败了，星星不够！")
def private_chat():
    global xingxing_Y
    xingxing_Y = False
    xingxing_path = 'mould/data/xingxing'
    if not os.path.exists(xingxing_path):
        messagebox.showinfo(title='飞鸽传书', message="很遗憾，你还没有玩过游戏，没有星星。（PS：通过小游戏可以拿到星星，通过消耗一颗星星可以向作者发送消息哦，此功能需要联网！）")
        return
    else:
        try:
            file = open(xingxing_path, mode='r', encoding='utf-8')
            linexx = file.readline()
            if len(linexx) >= 1:
                new_linexx = linexx[:-1]
                new_file = open(xingxing_path, mode='w+', encoding='utf-8')
                new_file.write(new_linexx)
                new_file.close()
                xingxing_Y = True
            else:
                messagebox.showinfo(title='飞鸽传书', message="很遗憾，你的星星不够！（PS：通过小游戏可以拿到星星，通过消耗一颗星星可以向作者发送消息哦，此功能需要联网！）")
                return
        except:
            messagebox.showinfo(title='飞鸽传书', message="很遗憾，出问题了！请联系作者")
            return

    def help():
        messagebox.showinfo('飞鸽小提示','点击在这个窗口文本框里，输入内容，最后点击发送。\n跟电脑自带的文本文档编辑器差不多操作，就是多了个发送按钮\n显示界面不是很好，条件简陋了点，和记事本差不多，应该是不限字数的。\n写完就点击发送.注意在此之前别关闭窗口哦，否则内容就没了哦。')

    def myopen():   # 打开本地
        global filename
        filename = askopenfilename(defaultextension='.txt')
        if filename == '':
            filename = None
        else:
            chat_Toplevel.title('飞鸽传书  ' + os.path.basename(filename))
            text.delete(1.0, END)
            f = open(filename, 'r')
            text.insert(1.0, f.read())
            f.close()

    def new():   # 新建文件
        global chat_Toplevel, filename, text
        chat_Toplevel.title('纸笺')
        filename = None
        text.delete(1.0, END)

    def save():   # 保存
        global filename
        try:
            f = open(filename, 'w')
            msg = text.get(1.0, 'end')
            f.write(msg)
            f.close()
        except:
            saveas()

    def saveas():   # 另存为
        f = asksaveasfilename(initialfile='纸笺.txt', defaultextension='.txt')
        global filename
        filename = f
        fh = open(f, 'w')
        msg = text.get(1.0, END)
        fh.write(msg)
        fh.close()
        chat_Toplevel.title('飞鸽传书  ' + os.path.basename(f))

    def cut(): # 剪切
        global text
        text.event_generate('<<Cut>>')
    def copy(): # 复制
        global text
        text.event_generate('<<Copy>>')
    def paste(): # 粘贴
        global text
        text.event_generate('<<Paste>>')
    def select_all(): # 全选
        global text
        # text.event_generate('sel', '1.0', 'end')
        # text.event_generate(" SEL('1.0', 'end')")
        # text.selection('1.0', 'end')
        text.event_generate("<<SelectAll>>")

    def find():   # 查找
        global chat_Toplevel
        t = Toplevel(chat_Toplevel)
        t.title('查找')
        # 设置窗口大小
        t.geometry('260x60+200+250')
        t.transient(chat_Toplevel)
        Label(t, text='查找:').grid(row=0, column=0, sticky='e')
        v = StringVar()
        e = Entry(t, width=20, textvariable=v)
        e.grid(row=0, column=1, padx=2, pady=2, sticky='we')
        e.focus_set()
        c = IntVar()
        Checkbutton(t, text='不区分大小写', variable=c).grid(row=1, column=1, sticky='e')
        Button(t, text='查找所有', command=lambda: search(v.get(), c.get(), text, t, e)).grid(row=0, column=2,
                                                                                             sticky='e' + 'w', padx=2,
                                                                                             pady=2)

        def close_search():
            text.tag_remove('match', '1.0', END)
            t.destroy()

        t.protocol('WM_DELETE_WINDOW', close_search)


    def search(needle, cssnstv, text, t, e):    # 搜索
        text.tag_remove('match', '1.0', END)
        count = 0
        if needle:
            pos = '1.0'
            while True:
                pos = text.search(needle, pos, nocase=cssnstv, stopindex=END)
                if not pos: break
                lastpos = pos + str(len(needle))
                text.tag_add('match', pos, lastpos)
                count += 1
                pos = lastpos
            text.tag_config('match', foreground='red', background='yellow')
            e.focus_set()
            t.title(str(count) + '个被匹配')

    def font():    # 调字体
        def change_a(self):    #  不加self 会报错 TypeError: change_a() takes 0 positional arguments but 1 was given
            global a,text,chat_Toplevel
            print('v.get',v.get())
            text['font'] = ('微软雅黑', '{}'.format(v.get()))
            text.pack(expand=YES, side=LEFT, fill=BOTH)

        global chat_Toplevel
        t = Toplevel(chat_Toplevel)
        t.title('字体')
        # 设置窗口大小
        t.geometry('260x60+200+280')
        v = StringVar()
        Scale(t,
              from_=10,  # 设置最大值
              to=30,  # 设置最小值
              resolution=1,  # 设置步距值
              orient=HORIZONTAL,  # 设置水平方向
              variable=v,  # 绑定变量
              command=change_a  # 设置回调函数
              ).pack()


    global chat_Toplevel,text,a
    chat_Toplevel = Toplevel()
    try:
        chat_Toplevel.iconbitmap('mould\data\logo.ico')
    except:
        pass
    chat_Toplevel.geometry('500x500+270+270')
    chat_Toplevel.title('联系作者（请在文末留下您的联系方式）')
    # chat_Toplevel.attributes("-topmost", 1)  # 1==True 处于顶层
    # 创建滚动条和文本框
    scroll = Scrollbar(chat_Toplevel)                         # 创建滚动条
    text = Text(chat_Toplevel, font=('微软雅黑', '15'))    # 创建文本框
    # 设置文本框初始内容
    text.insert('insert', '请在此处开始编辑文字，在文末记得留下您的联系方式：（这小戳文字可以删除）\n')

    # 将滚动条和文本框分别填充
    scroll.pack(side=RIGHT, fill=Y)     # side指定Scrollbar为居右；fill指定填充满整个剩余区域     # side是滚动条放置的位置，上下左右。fill是将滚动条沿着y轴填充
    text.pack(expand=YES,side=LEFT,fill=BOTH)         # 将文本框填充进窗口的左侧         # expand=YES 支持扩张yes   fill=BOTH 填充XY
    # 将滚动条与文本框互相关联
    scroll.config(command=text.yview)        # 指定Scrollbar的command的回调函数是Listbar的yview     # 将文本框关联到滚动条上，滚动条滑动，文本框跟随滑动
    text.config(yscrollcommand=scroll.set)   # 将滚动条关联到文本框

    # 发送按钮
    # button = Button(chat_Toplevel,text='发送',command=sending)
    # button.place(x=220,y=460,width=60,height=30)

    # 菜单栏
    menubar=Menu(chat_Toplevel,tearoff = False)

    filemenu = Menu(menubar,tearoff = False)
    filemenu.add_command(label='新建纸条', command=new)       # , accelerator='Ctrl+N'
    filemenu.add_command(label='打开', command=myopen)    # , accelerator='Ctrl+O'
    filemenu.add_command(label='保存', command=save)      # , accelerator='Ctrl+S'
    filemenu.add_command(label='另存为', command=saveas)  # , accelerator='Ctrl+Shift+S'
    menubar.add_cascade(label=' 文件 ', menu=filemenu)

    editmenu = Menu(menubar,tearoff = False)
    editmenu.add_command(label='剪切', accelerator='Ctrl+X', command=cut)
    editmenu.add_command(label='复制', accelerator='Ctrl+C', command=copy)
    editmenu.add_command(label='粘贴', accelerator='Ctrl+V', command=paste)
    editmenu.add_command(label='全选', accelerator='Ctrl+A', command=select_all)
    menubar.add_cascade(label=' 编辑 ', menu=editmenu)

    menubar.add_command(label=' 字体 ', command=font)
    menubar.add_command(label=' 查找 ', command=find)
    menubar.add_command(label=' 帮助 ',command=help)
    menubar.add_command(label=' 点击发送 ',command=sending)

    def popup(event):
        menubar.post(event.x_chat_Toplevel,event.y_chat_Toplevel)
    #绑定鼠标右键
    text.bind('<Button-3>',popup)

    # 放在菜单栏
    chat_Toplevel.config(menu=menubar)

    chat_Toplevel.mainloop()
def zuzhibu():
    from PIL import Image as imm
    from PIL import ImageTk as imk
    zu_zhi_bu = Toplevel(window)
    zu_zhi_bu.geometry("615x600+200+200")
    try:
        zu_zhi_bu.iconbitmap('mould\data\logo.ico')
    except:
        pass
    # 窗口置顶
    # zu_zhi_bu.attributes("-topmost", 1)  # 1==True 处于顶层
    # 禁止窗口的拉伸
    zu_zhi_bu.resizable(0, 0)
    # 窗口的标题
    zu_zhi_bu.title("欢迎来到组织部")

    # 定义变量
    zu_zhi_bu_var = StringVar()

    one_zuzhibu = '''本部门旨在统一协调学生党建、党支部工作，严格把关党员发展工作，培育组织能力强、熟悉党员材料，发展程序的高素质党员。\n
    组织工作部主要职能为协助书记及副书记协调党建委员会及学院各支部的日常工作。具体工作包括党员发展工作、入党积极分子培训、发展对象培训、预备党员培训和党员培训名单的收集审核，
    以及每期入党积极分子、发展对象培训、预备党员转正培训和党员培训的证书制作、党费收缴工作、各部门活动审批工作以及其他日常性工作。
    '''
    two_zuzhibu = '''
第一届学生党建工作委员会组织工作部
组织工作部
部长：
刘佳
副部长：
齐双悦
王灿

第二届学生党建工作委员会组织工作部
部长：
修梓洋 
副部长：
工管181班 赵慧
经济183班 江南
干事：
电商192班 谭周豪
法学182班 易也博
国贸181班 彭思雅
国贸192班 左玉冰
商类198班 黄永健
会计182班 蒋茜
会计185班 陈婉溶
经济181班 杨明嘉
营销182班 贺娜

第三届学生党建工作委员会组织工作部
部长：
工管181班 赵慧 
部长助理：
工企192班 黄永健
温琳艳
张祯婧
干事：
会计ACCA201班 王翊成
商类207班 何泓霖
电商192班 田梦凡
会计ACCA201班 严雨婷
商类213班 邓琴
会计192班 周国燕
法学191班 宋辰
法学201班 刘星雨

第四届学生党建工作委员会组织工作部
部长助理：
黄永健 
副部长：
会计204班 邓琴
经济201班 刘聪颖 
副部长助理：
会计ACCA201班 王翊成
法学201班 刘星雨
干事：
会计202班 赵碧
商类2105班 彭琼卉
会计ACCA201班 赵妍珂  
物流211班 张许兰 
国贸212班 石琳 
经济212班 陈晨
国贸211班 王静
经济202班 唐明敏
国贸212班 孔洁滢
商类2106班 李晶晶
'''

    scr_zzb = scrolledtext.ScrolledText(zu_zhi_bu, wrap=WORD)
    scr_zzb.place(x=10, y=10, width=310, height=250)
    scr_output(scr=scr_zzb, information=one_zuzhibu)
    scr_zzb.config(state=DISABLED)  # 关闭可写入模式

    scr_zzb2 = scrolledtext.ScrolledText(zu_zhi_bu, wrap=WORD)
    scr_zzb2.place(x=330, y=10, width=285, height=250)
    scr_output(scr=scr_zzb2, information=two_zuzhibu)
    scr_zzb2.config(state=DISABLED)  # 关闭可写入模式

    background_image_list = []
    for filename in os.listdir(r'mould/picture'):
        if filename[-4:] == ".jpg":
            background_image_list.append(str(r'mould/picture/' + filename))
    # print(background_image_list)
    im_label = imk.PhotoImage(imm.open(r'{}'.format(random.choice(background_image_list))).resize((580, 320)))
    bg_im_label = Label(zu_zhi_bu, image=im_label)  # , textvariable=var , compound=CENTER , text='hello'
    bg_im_label.place(x=10, y=260)

    # 显示窗口(消息循环)
    zu_zhi_bu.mainloop()
'''#################################################################################################################'''
# 统计分析系列
class VerticalScrolledFrame(Frame):
    """A pure Tkinter scrollable frame that actually works!
    * Use the 'interior' attribute to place widgets inside the scrollable frame
    * Construct and pack/place/grid normally
    * This frame only allows vertical scrolling

    """
    def __init__(self, parent, *args, **kw):
        Frame.__init__(self, parent, *args, **kw)

        # create a canvas object and a vertical scrollbar for scrolling it
        vscrollbar = Scrollbar(self, orient=VERTICAL)
        vscrollbar.pack(fill=Y, side=RIGHT, expand=FALSE)
        canvas = Canvas(self, bd=0, highlightthickness=0,
                        yscrollcommand=vscrollbar.set)
        canvas.pack(side=LEFT, fill=BOTH, expand=TRUE)
        vscrollbar.config(command=canvas.yview)

        # reset the view
        canvas.xview_moveto(0)
        canvas.yview_moveto(0)

        # create a frame inside the canvas which will be scrolled with it
        self.interior = interior = Frame(canvas)
        interior_id = canvas.create_window(0, 0, window=interior,
                                           anchor=NW)

        # track changes to the canvas and frame width and sync them,
        # also updating the scrollbar
        def _configure_interior(event):
            # update the scrollbars to match the size of the inner frame
            size = (interior.winfo_reqwidth(), interior.winfo_reqheight())
            # print("size",size)
            canvas.config(scrollregion="0 0 %s %s" % size)
            if interior.winfo_reqwidth() != canvas.winfo_width():
                # update the canvas's width to fit the inner frame
                canvas.config(width=interior.winfo_reqwidth())
        interior.bind('<Configure>', _configure_interior)

        def _configure_canvas(event):
            if interior.winfo_reqwidth() != canvas.winfo_width():
                # update the inner frame's width to fill the canvas
                canvas.itemconfigure(interior_id, width=canvas.winfo_width())
        canvas.bind('<Configure>', _configure_canvas)
def statistics_Toplevel():
    from tkintertable import TableCanvas, TableModel
    statistics_Toplevel = Toplevel()
    screenwidth = statistics_Toplevel.winfo_screenwidth()
    screenheight = statistics_Toplevel.winfo_screenheight()
    width = screenwidth-100  # 1300
    height = screenheight-100  # 830
    statistics_Toplevel.geometry('%dx%d+%d+%d' % (width, height, (screenwidth - width) / 2, (screenheight - height - 50) / 2))  # 设置主窗口大小
    try:
        statistics_Toplevel.iconbitmap('mould\data\logo.ico')
    except:pass
    # 测试 统计分析
    mighty18_1 = ttk.LabelFrame(statistics_Toplevel, text='') # ,style='Frame1.TFrame'
    mighty18_1.place(x=0,y=0,width=width,height=400) # 64
    # table = TableCanvas(mighty18_1)
    # table.show()
    model = TableModel()
    # table = TableCanvas(mighty18_1, model=model)
    table = TableCanvas(mighty18_1, model,
                cellwidth=120, cellbackgr='#e3f698',
                thefont=('Arial',12),rowheight=24, rowheaderwidth=30,
                rowselectedcolor='yellow', editable=True,rows=100, cols=100,)
    try:
        table.importCSV('mould/model/statistics.csv')
    except:pass
    table.show()
    button = Button(master=statistics_Toplevel, text="保存", command= lambda :save_table(table,"statistics_test.txt") )
    # 按钮放在下边
    button.place(x=170, y=415)
    statistics_Toplevel.mainloop()
def save_table(table,filename):
    table.save(filename)
    messagebox.showinfo("提示","文件：{}，保存成功！（该功能正在测试，导出文件有问题，请注意！）".format(filename))
def keshihua():
    import numpy as np
    try:
        import matplotlib.pyplot as plt
        # 支持中文
        plt.rcParams['font.sans-serif'] = ['SimHei']  # 用来正常显示中文标签
        plt.rcParams['axes.unicode_minus'] = False  # 用来正常显示负号
        from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg, NavigationToolbar2Tk
        from matplotlib.backend_bases import key_press_handler
        from matplotlib.figure import Figure
    except:
        messagebox.showinfo("提示", "缺少matplotlib的配置文件mpl-data，统计功能不可用！")
        return
    win = tkinter.Toplevel()
    win.title("党员发展情况可视化")
    # f = Figure(figsize=(5, 4), dpi=100)
    # a = f.add_subplot(111)  # 添加子图:1行1列第1个
    f = plt.figure()

    year = ['2019年', '2020年', '2021年', '2022年']
    # jj = np.array([672, 150, 441, 364])
    # fz = np.array([195, 178, 428, 159])
    # yb = np.array([138, 167, 406, 153])
    # dy = np.array([147, 138, 134, 360])
    jj = np.array([672, 400, 441, 364])
    fz = np.array([195, 178, 428, 159])
    yb = np.array([138, 167, 406, 153])
    dy = np.array([147, 138, 134, 360])
    # 此处的 _ 下划线表示将循环取到的值放弃，只得到[0,1,2,3,4]
    # ind = [x for x, _ in enumerate(year)]
    #绘制堆叠图
    plt.bar(year, jj, width=0.5, label='积极分子', color='gold') # , bottom=
    plt.bar(year, fz, width=0.5, label='发展对象', color='orange')
    plt.bar(year, yb, width=0.5, label='预备党员', color='#CD853F')
    plt.bar(year, dy, width=0.5, label='转正党员', color='red')

    #设置坐标轴
    # plt.xticks(ind, year)
    plt.ylabel("人数")
    plt.xlabel("年份")
    plt.legend(loc="upper right")
    plt.title("党员发展堆叠图")
    # plt.show()


    # 将绘制的图形显示到tkinter:创建属于root的canvas画布,并将图f置于画布上
    canvas = FigureCanvasTkAgg(f, master=win)
    canvas.draw()  # 注意show方法已经过时了,这里改用draw
    plt.draw()

    canvas.get_tk_widget().pack(side=tkinter.TOP,  # 上对齐
                                fill=tkinter.BOTH,  # 填充方式
                                expand=tkinter.YES)  # 随窗口大小调整而调整

    # matplotlib的导航工具栏显示上来(默认是不会显示它的)
    toolbar = NavigationToolbar2Tk(canvas, win)
    toolbar.update()
    canvas._tkcanvas.pack(side=tkinter.TOP,  # get_tk_widget()得到的就是_tkcanvas
                          fill=tkinter.BOTH,
                          expand=tkinter.YES)
def keshihua2():
    import numpy as np
    try:
        import matplotlib.pyplot as plt
        # 支持中文
        plt.rcParams['font.sans-serif'] = ['SimHei']  # 用来正常显示中文标签
        plt.rcParams['axes.unicode_minus'] = False  # 用来正常显示负号
        from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg, NavigationToolbar2Tk
        from matplotlib.backend_bases import key_press_handler
        from matplotlib.figure import Figure
    except:
        messagebox.showinfo("提示", "缺少matplotlib的配置文件mpl-data，统计功能不可用！")
        return
    win = tkinter.Toplevel()
    win.title("数据统计可视化")
    # f = Figure(figsize=(5, 4), dpi=100)
    # a = f.add_subplot(111)  # 添加子图:1行1列第1个
    f = plt.figure()

    labels = ['积极分子', '发展对象','预备党员', '转正党员']
    a = [672, 195, 138, 147]
    b = [400, 178, 167, 138]
    c = [441, 428, 406, 134]
    d = [364, 159, 153, 360]
    year = ['2019年', '2020年', '2021年', '2022年']

    x = np.arange(len(labels))  # 标签位置
    width = 0.1  # 柱状图的宽度，可以根据自己的需求和审美来改

    # fig, ax = plt.subplots()
    rects1 = plt.bar(x - width * 2, a, width, label=labels[0])
    rects2 = plt.bar(x - width + 0.01, b, width, label=labels[1])
    rects3 = plt.bar(x + 0.02, c, width, label=labels[2])
    rects4 = plt.bar(x + width + 0.03, d, width, label=labels[3])

    # 为y轴、标题和x轴等添加一些文本。
    plt.ylabel('人数', fontsize=16)
    plt.xlabel('年份', fontsize=16)
    plt.title('近四年党员发展人数柱状图')
    plt.xticks(x,year)
    # plt.xticklabels(labels)
    plt.legend()

    def autolabel(rects):
        """在*rects*中的每个柱状条上方附加一个文本标签，显示其高度"""
        for rect in rects:
            height = rect.get_height()
            plt.annotate('{}'.format(height),
                        xy=(rect.get_x() + rect.get_width() / 2, height),
                        xytext=(0, 3),  # 3点垂直偏移
                        textcoords="offset points",
                        ha='center', va='bottom')

    autolabel(rects1)
    autolabel(rects2)
    autolabel(rects3)
    autolabel(rects4)

    # fig.tight_layout()


    # 将绘制的图形显示到tkinter:创建属于root的canvas画布,并将图f置于画布上
    canvas = FigureCanvasTkAgg(f, master=win)
    canvas.draw()  # 注意show方法已经过时了,这里改用draw
    plt.draw()

    canvas.get_tk_widget().pack(side=tkinter.TOP,  # 上对齐
                                fill=tkinter.BOTH,  # 填充方式
                                expand=tkinter.YES)  # 随窗口大小调整而调整

    # matplotlib的导航工具栏显示上来(默认是不会显示它的)
    toolbar = NavigationToolbar2Tk(canvas, win)
    toolbar.update()
    canvas._tkcanvas.pack(side=tkinter.TOP,  # get_tk_widget()得到的就是_tkcanvas
                          fill=tkinter.BOTH,
                          expand=tkinter.YES)
def statistics_vision(filename_i):
    from tkintertable import TableCanvas, TableModel
    # table = TableCanvas(mighty18_1)
    # table.show()
    model = TableModel()
    # table = TableCanvas(mighty18_1, model=model)
    table = TableCanvas(mighty18_2, model,
                        cellwidth=120, cellbackgr='#e3f698',
                        thefont=('Arial', 12), rowheight=24, rowheaderwidth=30,
                        rowselectedcolor='yellow', editable=True, rows=100, cols=100, )
    filename = 'mould//csv//' + filename_i + ".csv"
    # try:
    table.importCSV(filename) # 只支持gbk的编码文件导入
    table.show()
    print(filename)
    # except:
    #     messagebox.showinfo("提示","文件：{}，导入错误！".format(filename))
def excel2csv(excel_file):
    import pandas as pd
    # 打开excel文件
    workbook = openpyxl.load_workbook(excel_file)
    # 获取所有sheet名字
    sheet_names = workbook.sheetnames
    for worksheet_name in sheet_names:
        # 遍历每个sheet并用Pandas读取
        data_xls = pd.read_excel(excel_file, worksheet_name, index_col=None)
        # 获取excel当前目录
        dir_path = os.path.abspath(os.path.dirname(excel_file))
        # 转换成csv并保存到excel所在目录下的csv文件夹中
        csv_path = 'mould\\csv\\'  # dir_path +
        if not os.path.exists(csv_path):
            os.mkdir(csv_path)
        data_xls.to_csv(csv_path + worksheet_name + '.csv', index=None, encoding='gbk') # utf-8-sig
    print("生成csv成功！")
def zairu():
    filename = "mould/model/模板7 请示批复一览表.xlsx"
    try:
        excel2csv(filename)
    except:
        messagebox.showinfo("提示", "文件：{}，载入失败，请检查文件位置是否存在！".format(filename))
'''#################################################################################################################'''
# 登录界面
def index_Toplevel():
    index = tkinter.Toplevel()  # 创建主窗口
    # index.attributes('-alpha',1)  #窗口背景透明化
    index.title('智慧党建-登录界面')  # 设置主窗口标题
    index.iconbitmap('mould\data\logo.ico')
    # # 窗口置顶
    # index.attributes("-topmost", 1)  # 1==True 处于顶层
    screenwidth = index.winfo_screenwidth()
    screenheight = index.winfo_screenheight()
    width = 500  # 1300
    height = 400  # 830
    index.geometry('%dx%d+%d+%d' % (width, height, (screenwidth - width) / 2, (screenheight - height - 100) / 2))  # 设置主窗口大小

    # 下面两行代码的作用是固定窗口大小，不可拉动调节
    index.maxsize(500, 300)
    index.minsize(500, 300)

    # **************************************************
    #                 添加组件

    # 加载图片
    canvas = tkinter.Canvas(index, width=500, height=300, bg=None)
    try:
        image_file = tkinter.PhotoImage(file="mould/data/login.gif")
        image = canvas.create_image(250, 0, anchor='n', image=image_file)
    except:pass
    canvas.pack()

    # 账号与密码文字标签
    account_lable = tkinter.Label(index, text='账号', bg='SandyBrown', fg='white', font=('Arial', 12), width=5,
                                  height=1) # lightskyblue  OrangeRed  SandyBrown Moccasin
    account_lable.place(relx=0.29, rely=0.4)
    pasw_lable = tkinter.Label(index, text='密码', bg='SandyBrown', fg='white', font=('Arial', 12), width=5, height=1)
    pasw_lable.place(relx=0.29, rely=0.5)

    # 账号与密码输入框
    account = tkinter.Entry(index, width=20, highlightthickness=1, highlightcolor='Coral',
                            relief='groove')  # 账号输入框
    account.place(relx=0.4, rely=0.4)  # 添加进主页面,relx和rely意思是与父元件的相对位置
    password = tkinter.Entry(index, show='*', highlightthickness=1, highlightcolor='Coral',
                             relief='groove')  # 密码输入框
    password.place(relx=0.4, rely=0.5)  # 添加进主页面

    user = {"邓琴": "hongwu1120",
            "李晶晶": "hongwu1120",
            "彭琼卉": "hongwu1120",
            "王静": "hongwu1120",
            "黄永健": "hongwu1120",
            "super_admin": "hongwu1120",
            }  # 定义一个字典来存储用户的信息(key :账号 , value：密码)

    # 登录按钮处理函数
    def login():
        ac = account.get()
        ps = password.get()
        if (ac == "" or ps == ""):
            messagebox.showinfo("用户登录", "请完整填写信息！！")  # messagebox的方法
        elif user.get(ac) != ps:
            account.delete(0, 'end')  # 清空文本框的内容
            password.delete(0, 'end')  # 清空文本框的内容
            messagebox.showinfo("用户登录", "账号或者密码有误！")  # messagebox的方法
        else:
            account.delete(0, 'end')  # 清空文本框的内容
            password.delete(0, 'end')  # 清空文本框的内容
            messagebox.showinfo("用户登录", "登录成功！即将进入菜单界面....")  # messagebox的方法
            index.destroy()

            # os.system("python demo2.py")
            tabControl.add(tab0_4, text=' 积极分子培训 ')  # Add the tab 二级菜单

            LICENCE_path = os.getcwd() + "\\mould\\data\\" + 'LICENCE'  # os.getcwd()返回当前文件所在的目录绝对路径
            file = open(LICENCE_path, mode='w', encoding='utf-8')
            file.write(str(platform.uname())+ '\n' + 'True')
            file.close()
            print('创建文件成功')

            window.mainloop()

            return True

    def reguest():
        print("注册")
        # messagebox.showinfo("用户注册", "暂未开放注册")  # messagebox的方法
        if messagebox.askokcancel("用户注册", "暂未开放注册，点击确定可以直接进入！"):
            index.destroy()

        # ac = account.get()
        # ps = password.get()
        # if (ac == "" or ps == ""):
        #     messagebox.showinfo("用户登录", "请完整填写信息！！")  # messagebox的方法
        # elif ac in user:
        #     account.delete(0, 'end')  # 清空文本框的内容
        #     password.delete(0, 'end')  # 清空文本框的内容
        #     messagebox.showinfo("用户注册", "账号已存在！")  # messagebox的方法
        # else:
        #     user[ac] = ps
        #     account.delete(0, 'end')  # 清空文本框的内容
        #     password.delete(0, 'end')  # 清空文本框的内容
        #     messagebox.showinfo("用户注册", "注册成功！")  # messagebox的方法

    # 登录与注册按钮
    loginBtn = tkinter.Button(index, text='登录', font=('宋体', 12), width=4, height=1, command=login, relief='solid',
                              bd=0.5, bg='Cornsilk') # lightcyan
    loginBtn.place(relx=0.41, rely=0.63)
    loginBtn = tkinter.Button(index, text='注册', font=('宋体', 12), width=4, height=1, bd=0.5, command=reguest,
                              relief='solid', bg='Cornsilk')
    loginBtn.place(relx=0.56, rely=0.63)

    # **************************************************
    index.mainloop()  # 使窗口不断刷新，应该放在代码最后一句
    window.mainloop()
'''#################################################################################################################'''
# 定义初始化
# 公式
list_gongshi = ['int(first_value) - int(birth_value) - 180000 <= 0',
                # 'int(positive_value) - int(first_value) -180 <= 0',
                180,
                'int(object_value) - int(positive_value) - 10000 <= 0',
                'int(ready_value) - int(object_value) <= 0',
                'int(become_value) - int(ready_value) - 10000 > 0']
list_gongshi_simple = ['18','0','1','0','1',   '0','180','0','0','0',    '包括当天','包括当天','包括当天','包括当天','不包括当天']
# 支部全称
zhibu_allname = ["学生第一党支部","学生第二党支部","学生第三党支部","学生第四党支部","学生第五党支部","学生第六党支部","学生第七党支部",
                 "学生第八党支部",'研究生第一党支部', '研究生第二党支部','研究生第三党支部',
                 '电子商务与物流工程专业学生党支部',  '工商管理与信管专业学生党支部', '会计学专业第一学生党支部',
                 '会计学专业第二学生党支部', '国际经济与贸易专业学生党支部', '经济学专业学生党支部', '法学专业学生党支部',
                '人力资源管理与市场营销专业学生党支部']
zhibu_list = ['第一支部', '第二支部', '第三支部', '第四支部', '第五支部', '第六支部', '第七支部',
              '第八支部', '研一支部', '研二支部', '研三支部',
                '电物支部', '工信支部', '会计一支部',
                '会计二支部', '国贸支部', '经济支部' '法学支部', '人营支部']

now = datetime.datetime.now()
window_style = None

'''# -1 主窗口的控件 ############################################################'''

# style = ttkStyle(theme='cerculean')
# style = ttkStyle(theme='simplex')
# window = style.master

# messagebox.showinfo('TO YOU', message='毛主席万岁')

# 分区
tabControl = ttk.Notebook(window)       # Create Tab Control 总菜单

tab0_4 =  ttk.Frame(tabControl)
# tabControl.add(tab0_4, text=' 积极分子培训 ')    # Add the tab 二级菜单

LICENCE_path = os.getcwd() + "\\mould\\data\\" + 'LICENCE'  # os.getcwd()返回当前文件所在的目录绝对路径
try:
    if os.path.exists(LICENCE_path):  # 判断是否存在文件
        file = open(LICENCE_path, encoding='utf-8')
        line1 = file.readline()
        line2 = file.readline()
        if  line1 == str(platform.uname())+"\n" and line2 == 'True':
            tabControl.add(tab0_4, text=' 积极分子培训 ')  # Add the tab 二级菜单
except:
    pass

tabControl_jjfz = ttk.Notebook(tab0_4)       # Create Tab Control 二级菜单
tab0 = ttk.Frame(tabControl_jjfz)            # Create a tab
tabControl_jjfz.add(tab0, text=' 学员册 ')    # Add the tab
tab1 = ttk.Frame(tabControl_jjfz)
tabControl_jjfz.add(tab1, text=' 分组名单 ')
tab3 = ttk.Frame(tabControl_jjfz)
tabControl_jjfz.add(tab3, text=' 座位考勤 ')
# tab4 合并至 tab3，tab4不再沿用
tab2 = ttk.Frame(tabControl_jjfz)
tabControl_jjfz.add(tab2, text=' 培训考试 ')

# tab14 =  ttk.Frame(tabControl)
# tabControl.add(tab14, text=' 发展对象培训 ')    # Add the tab 二级菜单
# tab15 =  ttk.Frame(tabControl)
# tabControl.add(tab15, text=' 预备党员培训 ')    # Add the tab 二级菜单
# tab16 =  ttk.Frame(tabControl)
# tabControl.add(tab16, text=' 党员培训 ')    # Add the tab 二级菜单

tab6_8 =  ttk.Frame(tabControl)
tabControl.add(tab6_8, text=' 请示批复备案 ')    # Add the tab 二级菜单

tabControl_qspf = ttk.Notebook(tab6_8)       # Create Tab Control 二级菜单
tab6 = ttk.Frame(tabControl_qspf)
tabControl_qspf.add(tab6, text=' 请示管理 ')
tab7 = ttk.Frame(tabControl_qspf)
tabControl_qspf.add(tab7, text=' 批复管理 ')
tab8 = ttk.Frame(tabControl_qspf)
tabControl_qspf.add(tab8, text=' 备案报告 ')

tab11 = ttk.Frame(tabControl)
tabControl.add(tab11, text=' 支部管理 ')

tabControl_zbgl = ttk.Notebook(tab11)       # Create Tab Control 二级菜单
tab11_c = ttk.Frame(tabControl_zbgl)
tabControl_zbgl.add(tab11_c, text=' 两推一备案 ')
tab11_a = ttk.Frame(tabControl_zbgl)
tabControl_zbgl.add(tab11_a, text=' 请示管理 ')
tab11_b = ttk.Frame(tabControl_zbgl)
tabControl_zbgl.add(tab11_b, text=' 批复管理 ')

tab5 = ttk.Frame(tabControl)
tabControl.add(tab5, text=' 时间管理 ')

tab17 = ttk.Frame(tabControl)
tabControl.add(tab17, text=' 党员管理 ')

tab12 = ttk.Frame(tabControl)
tabControl.add(tab12, text=' 智慧查重 ')

tab18 = ttk.Frame(tabControl)
tabControl.add(tab18, text=' 统计分析 ')

tab10 = ttk.Frame(tabControl)
tabControl.add(tab10, text=' 通用功能 ')

tab9 = ttk.Frame(tabControl)
tabControl.add(tab9, text=' 帮助 ')

tab13 = ttk.Frame(tabControl)
tabControl.add(tab13, text=' 小黑实验室 ')


tabControl.pack(expand=1, fill="both")  # Pack to make visible
tabControl_jjfz.pack(expand=1, fill="both")  # Pack to make visible
tabControl_qspf.pack(expand=1, fill="both")  # Pack to make visible
tabControl_zbgl.pack(expand=1, fill="both")  # Pack to make visible



def demo():
    messagebox.showinfo('提示','未开发，有什么好的开发建议和需求功能，请联系我！')
    pass



menubar = Menu(window)

submenu1 = Menu(menubar, tearoff=0)
submenu1.add_command(label=' 登录 ', command=index_Toplevel)
submenu1.add_command(label=' 打开 ', command=demo)
submenu1.add_command(label=' 返回 ', command=demo)
submenu1.add_command(label=' 关闭 ', command=demo)
menubar.add_cascade(label=' 菜单 ', menu=submenu1)   # 将menubar的menu属性指定为submenu1，即submenu1为menubar的下拉菜单

submenu2 = Menu(menubar,tearoff=0)
submenu2.add_command(label='青葱岁月（组织部）', command=zuzhibu)
submenu2.add_command(label='学习语录（多学习）', command=xi_sayings)
submenu2.add_command(label='我心向党（小游戏）', command=my_heart)
submenu2.add_command(label='为民服务（联系我）', command=private_chat)
submenu2.add_command(label='努力工作 ', command=demo)
submenu2.add_command(label='奋斗终生', command=demo)
menubar.add_cascade(label=' 编辑 ',menu=submenu2)


submenu3 = Menu(menubar, tearoff=0)
submenu3.add_command(label='更换主页风格',command=change_home_style)
# submenu3.add_command(label='更换背景图片',command=change_background_image)
v = IntVar()
submenu3.add_checkbutton(label='打开背景音乐',command=open_background_music,variable = v)
menubar.add_cascade(label=' 视图 ',menu=submenu3)

submenu4 = Menu(menubar, tearoff=0)
submenu4.add_command(label='打开目录',command=open_abspath)
submenu4.add_command(label='查看帮助',command=help_)
submenu4.add_command(label='检查版本',command=version_)

submenu4.add_command(label='视频教程',command=video_)
# submenu4.add_separator()    # 将各个菜单项使用分隔符隔开
submenu4.add_command(label='合作反馈',command=author_)
# submenu4.add_separator()
submenu4.add_command(label='关于软件',command=about_)
# submenu4.add_command(label='部门介绍',command=author_)
submenu4.add_command(label='致谢',command=thank_)
menubar.add_cascade(label='关于',menu=submenu4)

window['menu'] = menubar    # == window.config(menu=menubar)  == window.configure(menu=menubar)

print('')
# # 背景
# canvas = Canvas(window, width=580, height=320, bd=0, highlightthickness=0) # 使用canvas,其中width,height参数为画布的宽和高，bd(borderwidth)为 文本框边框宽度
# imgpath = 'mould/picture/1.gif'
# img = imm.open(imgpath)
# photo = imk.PhotoImage(img)
#
# canvas.create_image(0, 0, image=photo)
# canvas.place(x=0,y=0)
# canvas.create_window(0,0, width=100, height=20)

# 测试 主题颜色
# s = ttk.Style()
# s.configure('Frame1.TFrame', background='#fcdfc5')


'''# 18    统计分析 ############################################################'''

# 创建一个按钮,并把上面那个函数绑定过来
button = Button(master=tab18, text="载入表格", command=zairu)
button.place(x = 140, y = 415) # 按钮放在下边

button = Button(master=tab18, text="可视化", command=keshihua)
button.place(x = 240, y = 415)

button = Button(master=tab18, text="数据分析", command=keshihua2)
button.place(x = 340, y = 415)

button = Button(master=tab18, text="全屏显示", command=statistics_Toplevel)
button.place(x = 440, y = 415)

button = Button(master=tab18, text="保存", command=lambda: save_table(table, "测试_保存的统计表格.txt"))
button.place(x = 540, y = 415)

# 选择学期滚动条
mighty18_2 = VerticalScrolledFrame(tab18)
mighty18_2.place(x = 10, y = 40, height=400)
label = Label(master=tab18, text="请选择学期：")
label.place(x = 10, y = 10)
buttons = []

try:
    workbook = openpyxl.load_workbook("mould/model/模板7 请示批复一览表.xlsx")
    statistics_time = workbook.sheetnames
    # print(workbook.sheetnames)
except:
    statistics_time = ['2018下', '2019上', '2019下', '2020上', '2020下', '2021上', '2021中', '2021下', '2022上', '2022下', '2023上', '2023下', '2024上', '2024下']

for i in statistics_time:
    buttons.append(Button(master=mighty18_2.interior, text = i , command = lambda arg=i:statistics_vision(arg))) # command = lambda : statistics_vision(i)
    # 上述代码实现，导入多个按钮的动态传参数效果
    buttons[-1].pack()
# print(buttons)

# 表格展示区
mighty18_2 = ttk.LabelFrame(tab18, text='') # ,style='Frame1.TFrame'
mighty18_2.place(x=120,y=10,width=650-120-10,height=400) # 64
# # table = TableCanvas(mighty18_1)
# # table.show()
# model = TableModel()
# # table = TableCanvas(mighty18_1, model=model)
# table = TableCanvas(mighty18_2, model,
#             cellwidth=120, cellbackgr='#e3f698',
#             thefont=('Arial',12),rowheight=24, rowheaderwidth=30,
#             rowselectedcolor='yellow', editable=True,rows=100, cols=100,)

'''# 14  15   16  17  党员管理  待完善分区的控件 ############################################################'''
# # 布局Frame
# mightys = ttk.LabelFrame(tab14, text='')
# mightys.place(x=0,y=0,width=1400,height=600)
# # 标签
# labels = ttk.Label(mightys, text="本功能正在开发，敬请期待！\n全新版本，正在紧张制作！\n请加入本软件答疑交流QQ群\n群号：866367580\n开发不易，敬请谅解！", font=("黑体", 25))
# labels.place(x=10,y=10)
#
# mightys = ttk.LabelFrame(tab15, text='')
# mightys.place(x=0,y=0,width=1400,height=600)
# # 标签
# labels = ttk.Label(mightys, text="本功能正在开发，敬请期待！\n全新版本，正在紧张制作！\n请加入本软件答疑交流QQ群\n（群号：866367580）\n开发不易，敬请谅解！", font=("黑体", 25))
# labels.place(x=10,y=10)
#
# mightys = ttk.LabelFrame(tab16, text='')
# mightys.place(x=0,y=0,width=1400,height=600)
# # 标签
# labels = ttk.Label(mightys, text="本功能正在开发，敬请期待！\n全新版本，正在紧张制作！\n请加入本软件答疑交流QQ群\n（群号：866367580）\n开发不易，敬请谅解！", font=("黑体", 25))
# labels.place(x=10,y=10)

mightys = ttk.LabelFrame(tab17, text='')
mightys.place(x=0,y=0,width=650,height=690)
# 标签
labels = ttk.Label(mightys, text="本功能正在开发，敬请期待！\n全新版本，正在紧张制作！\n有什么更好的开发建议，请联系作者！\n请加入本软件答疑交流QQ群\n（群号：866367580）\n开发不易，敬请谅解！", font=("黑体", 25))
labels.place(x=10,y=10)



















'''# 0 学员册分区的控件 ############################################################'''

# 布局Frame
mighty0 = ttk.LabelFrame(tab0, text='递交入党申请书人数名额分配') # ,style='Frame1.TFrame'
mighty0.place(x=10,y=10,width=630,height=260)

# 标签
label5_0 = ttk.Label(mighty0, text="选择本次培训期数：")
label5_0.place(x=10,y=10)
label3_0 = ttk.Label(mighty0, text="系数值：")
label3_0.place(x=220,y=10)
label1_0 = ttk.Label(mighty0, text="各支部系数文件夹：")
label1_0.place(x=10,y=40)
label2_0 = ttk.Label(mighty0, text="合并系数表文件名：")
label2_0.place(x=10,y=70)
label6_0 = ttk.Label(mighty0, text="合并后的文件路径：")
label6_0.place(x=10,y=110)
label7_0 = ttk.Label(mighty0, text="名额分配表文件名：")
label7_0.place(x=10,y=140)
label1_0 = ttk.Label(mighty0, text="各支部学员文件夹：")
label1_0.place(x=10,y=180)
label2_0 = ttk.Label(mighty0, text="合并学员册文件名：")
label2_0.place(x=10,y=210)
label4_0 = ttk.Label(tab0, text="输出窗口：")
label4_0.place(x=10,y=270)

# 输入期数
number_0 = StringVar()
number_chosen_0 = ttk.Combobox(mighty0, width=12, textvariable=number_0, state='readonly')
number_chosen_0['values'] = ('一','二','三','四','五','六','七','八','九','十','十一','十二','十三','十四','十五','十六','十七','十八','十九','二十')
number_chosen_0.place(x=130,y=10,width=70)
number_chosen_0.current(10)   # 设置初始显示值，值为元组['values']的下标
number_chosen_0.config(state='readonly')  # 设为只读模式

# 输入框
pathin1_0 = StringVar()  # 定义变量
entry_pathin1_0 = ttk.Entry(mighty0, textvariable=pathin1_0)  # 输入框    # entry不能和grid连写，否则会报错
entry_pathin1_0.place(x=130,y=40,width=380)
createToolTip(entry_pathin1_0, '这里不需要输入')   # Add Tooltip

pathin2_0 = StringVar()  # 定义变量
entry_pathin2_0 = ttk.Entry(mighty0, textvariable=pathin2_0)  # 输入框    # entry不能和grid连写，否则会报错
entry_pathin2_0.place(x=130,y=70,width=380)
createToolTip(entry_pathin2_0, '这里可以选择需要输入生成的文件名称')   # Add Tooltip
pathin2_0.set('经济管理与法学学院分党校第{}期各支部入党积极分子合并系数表'.format(number_chosen_0.get()))

# 输入框
pathin3_0 = StringVar()  # 定义变量
entry_pathin3_0 = ttk.Entry(mighty0, textvariable=pathin3_0)  # 输入框    # entry不能和grid连写，否则会报错
entry_pathin3_0.place(x=130,y=110,width=380)
createToolTip(entry_pathin3_0, '这里不需要输入')   # Add Tooltip

pathin4_0 = StringVar()  # 定义变量
entry_pathin4_0 = ttk.Entry(mighty0, textvariable=pathin4_0)  # 输入框    # entry不能和grid连写，否则会报错
entry_pathin4_0.place(x=130,y=140,width=380)
createToolTip(entry_pathin4_0, '这里可以选择需要输入生成的文件名称')   # Add Tooltip
pathin4_0.set('经济管理与法学学院分党校第{}期各支部入党积极分子名额分配表'.format(number_chosen_0.get()))

# 输入框
pathin5_0 = StringVar()  # 定义变量
entry_pathin5_0 = ttk.Entry(mighty0, textvariable=pathin5_0)  # 输入框    # entry不能和grid连写，否则会报错
entry_pathin5_0.place(x=130,y=180,width=380)
createToolTip(entry_pathin5_0, '这里不需要输入')   # Add Tooltip

pathin6_0 = StringVar()  # 定义变量
entry_pathin6_0 = ttk.Entry(mighty0, textvariable=pathin6_0)  # 输入框    # entry不能和grid连写，否则会报错
entry_pathin6_0.place(x=130,y=210,width=380)
createToolTip(entry_pathin6_0, '这里可以选择需要输入生成的文件名称')   # Add Tooltip
pathin6_0.set('经济管理与法学学院分党校第{}期入党积极分子培训班学员花名册'.format(number_chosen_0.get()))

# 系数值数字选择框
number2_0 =StringVar()
spin = Spinbox(mighty0,
        from_=0,  # 设置最小值
        to=1,  # 设置最大值
        increment=0.05,  # 设置增量值为5，这个与Scale的resolution意思相同
        textvariable = number2_0
        )
number2_0.set(0.75)
spin.place(x=270,y=10,width=70)

number_row_0 = IntVar()
check1 = Checkbutton(mighty0, text="表头横行", variable=number_row_0)
check1.state(['selected'])
number_row_0.set(1)
check1.place(x=360,y=10)

number_col_0 = IntVar()
check2 = Checkbutton(mighty0, text="表头纵行", variable=number_col_0)
check2.state(['disabled'])  # Clears (turns off) the checkbutton.
check2.place(x=440,y=10)

# GUI Callback function
def checkCallback(*ignoredArgs):
    # only enable one checkbutton
    if number_row_0.get():
        check2.configure(state='disabled')
    else:
        check2.configure(state='normal')
    if number_col_0.get():
        check1.configure(state='disabled')
    else:
        check1.configure(state='normal')
# trace the state of the two checkbuttons  #？？
number_col_0.trace('w', lambda unused0, unused1, unused2: checkCallback())
number_row_0.trace('w', lambda unused0, unused1, unused2: checkCallback())

number_0_0 = IntVar()
check3 = Checkbutton(tab0, text="自定义学员册模板", variable=number_0_0)
check3.state(['active'])  # Clears (turns off) the checkbutton.
check3.place(x=400,y=270)
# active, disabled, focus, pressed, selected, background,readonly, alternate, invalid

# 文本框
scr_0 = scrolledtext.ScrolledText(tab0, width=88, height=25, wrap=WORD)
scr_0.place(x=7,y=295)
scr_0.config(state=DISABLED)  # 关闭可写入模式



button3_1 = ttk.Button(mighty0, text="获取参数", command=lambda : get_canshu(number_chosen_0.get()))
button3_1.place(x=530,y=10)

button1_0 = ttk.Button(mighty0, text="选择文件夹", command=lambda : select_files(scr_0, pathin1_0))
button1_0.place(x=530,y=40)
button2_0 = ttk.Button(mighty0, text="开始合并", command=lambda : main0_1(path=pathin1_0.get(), filename=pathin2_0.get()))
button2_0.place(x=530,y=70)

button4_0 = ttk.Button(mighty0, text="选择文件", command=lambda : select_file(scr_0, pathin3_0))
button4_0.place(x=530,y=110)
button5_0 = ttk.Button(mighty0, text="开始生成", command=lambda : main0_2(path=pathin3_0.get(), filename=pathin4_0.get(), xishu=number2_0.get()))
button5_0.place(x=530,y=140)

button6_0 = ttk.Button(mighty0, text="选择文件夹", command=lambda : select_files(scr_0,pathin5_0))
button6_0.place(x=530,y=180)
button7_0 = ttk.Button(mighty0, text="开始合并", command=lambda : main0_3(path=pathin5_0.get(), filename=pathin6_0.get(), qishu=number_chosen_0.get()))
button7_0.place(x=530,y=210)








'''# 1 分组名单分区的控件 ############################################################'''
# 布局Frame
mighty1 = ttk.LabelFrame(tab1, text='生成分组名单')
mighty1.place(x=10,y=10,width=630,height=130)
mighty2 = ttk.LabelFrame(tab1, text='生成分组名单2')
mighty2.place(x=10,y=140,width=630,height=55)

# 标签
label1_1 = ttk.Label(mighty1, text="选择本次培训期数：")
label1_1.place(x=10,y=20)
label5_1 = ttk.Label(mighty1, text="小组数：")
label5_1.place(x=220,y=20)
label2_1 = ttk.Label(mighty1, text="学员册文件路径：")
label2_1.place(x=10,y=50)
label3_1 = ttk.Label(mighty1, text="生成分组名单文件名：")
label3_1.place(x=10,y=80)
label4_1 = ttk.Label(tab1, text="输出窗口：")
label4_1.place(x=10,y=200)
label5_1 = ttk.Label(mighty2, text="生成分组名单2文件名：")
label5_1.place(x=5,y=4)

# 输入期数
number_1 = StringVar()
number_chosen_1 = ttk.Combobox(mighty1, width=12, textvariable=number_1, state='readonly')
number_chosen_1['values'] = ('一','二','三','四','五','六','七','八','九','十','十一','十二','十三','十四','十五','十六','十七','十八','十九','二十')
number_chosen_1.place(x=130,y=20,width=70)
number_chosen_1.current(10)   # 设置初始显示值，值为元组['values']的下标
number_chosen_1.config(state='readonly')  # 设为只读模式

# 小组数
number2_1 = StringVar()
number2_chosen_1 = ttk.Combobox(mighty1, width=12, textvariable=number2_1, state='readonly')
number2_chosen_1['values'] = ('一','二','三','四','五','六','七','八','九','十','十一','十二','十三','十四','十五','十六','十七','十八','十九','二十')
number2_chosen_1.place(x=270,y=20,width=70)
number2_chosen_1.current(9)   # 设置初始显示值，值为元组['values']的下标
number2_chosen_1.config(state='readonly')  # 设为只读模式

# 输入框
pathin_1 = StringVar()  # 定义变量
entry_pathin_1 = ttk.Entry(mighty1, textvariable=pathin_1)  # 输入框    # entry不能和grid连写，否则会报错
entry_pathin_1.place(x=130,y=50,width=380)
createToolTip(entry_pathin_1, '这里不需要输入')   # Add Tooltip

pathin2_1 = StringVar()  # 定义变量
entry_pathin2_1 = ttk.Entry(mighty1, textvariable=pathin2_1)  # 输入框    # entry不能和grid连写，否则会报错
entry_pathin2_1.place(x=130,y=80,width=380)
createToolTip(entry_pathin2_1, '这里可以选择需要输入生成的文件名称')   # Add Tooltip
pathin2_1.set('经济管理与法学学院分党校第{}期入党积极分子培训班分组名单'.format(number_chosen_1.get()))

pathin4_1 = StringVar()  # 定义变量
entry_pathin4_1 = ttk.Entry(mighty2, textvariable=pathin4_1)  # 输入框    # entry不能和grid连写，否则会报错
entry_pathin4_1.place(x=130,y=2,width=380)
createToolTip(entry_pathin4_1, '这里可以选择需要输入生成的文件名称')   # Add Tooltip
pathin4_1.set('经济管理与法学学院分党校第{}期入党积极分子培训班分组名单2'.format(number_chosen_1.get()))

number_row_1 = IntVar()
check1_1 = Checkbutton(mighty1, text="按支部分", variable=number_row_1)
check1_1.state(['selected'])
number_row_1.set(1)
check1_1.place(x=360,y=20)

number_col_1 = IntVar()
check2_1 = Checkbutton(mighty1, text="按核定数", variable=number_col_1)
check2_1.state(['disabled'])  # Clears (turns off) the checkbutton.
check2_1.place(x=440,y=20)
# GUI Callback function
def checkCallback_1(*ignoredArgs):
    # only enable one checkbutton
    if number_row_1.get():
        check2_1.configure(state='disabled')
    else:
        check2_1.configure(state='normal')
    if number_col_1.get():
        check1_1.configure(state='disabled')
    else:
        check1_1.configure(state='normal')
# trace the state of the two checkbuttons  #？？
number_col_1.trace('w', lambda unused0, unused1, unused2: checkCallback_1())
number_row_1.trace('w', lambda unused0, unused1, unused2: checkCallback_1())

# 文本框
scr_1 = scrolledtext.ScrolledText(tab1, width=88, height=30, wrap=WORD)
scr_1.place(x=7,y=225)
scr_1.config(state=DISABLED)  # 关闭可写入模式

# 按钮
button3_1 = ttk.Button(mighty1, text="获取参数", command=lambda : get_canshu(number_chosen_1.get()))
button3_1.place(x=530,y=20)

button1_1 = ttk.Button(mighty1, text="选择学员册", command=lambda : select_file(scr_1,pathin_1))
button1_1.place(x=530,y=50)

button2_1 = ttk.Button(mighty1, text="开始生成", command=lambda : main1(path=pathin_1.get(),
                                                                    filename=pathin2_1.get(), qishu=number_chosen_1.get()))
button2_1.place(x=530,y=80)

# 生成分组名单2
button3_1 = ttk.Button(mighty2, text="开始生成", command=lambda : fenzu_main2(workbook_name=pathin2_1.get(),
                                                                    output_workbook_name=pathin4_1.get(), qishu=number_chosen_1.get()))
button3_1.place(x=530,y=0)




'''# 2 培训考试 考试表分区的控件 ############################################################'''
# 布局Frame
mighty2_1 = ttk.LabelFrame(tab2, text='生成考试座位表')
mighty2_1.place(x=10,y=10,width=630,height=110)

# 标签
label1_2 = ttk.Label(mighty2_1, text="选择本次培训期数：")
label1_2.place(x=10,y=0)
label2_2 = ttk.Label(mighty2_1, text="学员册文件路径：")
label2_2.place(x=10,y=30)
label3_2 = ttk.Label(mighty2_1, text="生成考试表文件名：")
label3_2.place(x=10,y=60)
label4_2 = ttk.Label(tab2, text="输出窗口：")
label4_2.place(x=10,y=230)

# 输入期数
number_2 = StringVar()
number_chosen_2 = ttk.Combobox(mighty2_1, width=12, textvariable=number_2, state='readonly')
number_chosen_2['values'] = ('一','二','三','四','五','六','七','八','九','十','十一','十二','十三','十四','十五','十六','十七','十八','十九','二十')
number_chosen_2.place(x=130,y=0)
number_chosen_2.current(10)   # 设置初始显示值，值为元组['values']的下标
number_chosen_2.config(state='readonly')  # 设为只读模式

# 输入框
pathin_2 = StringVar()  # 定义变量
entry_pathin_2 = ttk.Entry(mighty2_1, textvariable=pathin_2)  # 输入框    # entry不能和grid连写，否则会报错
entry_pathin_2.place(x=130,y=30,width=380)
createToolTip(entry_pathin_2, '这里不需要输入')   # Add Tooltip

pathin2_2 = StringVar()  # 定义变量
entry_pathin2_2 = ttk.Entry(mighty2_1, textvariable=pathin2_2)  # 输入框    # entry不能和grid连写，否则会报错
entry_pathin2_2.place(x=130,y=60,width=380)
createToolTip(entry_pathin2_2, '这里可以选择需要输入生成的文件名称')   # Add Tooltip
pathin2_2.set('经济管理与法学学院分党校第{}期入党积极分子培训班考试座位表'.format(number_chosen_2.get()))

# 文本框
scr_2 = scrolledtext.ScrolledText(tab2, width=88, height=30, wrap=WORD)
scr_2.place(x=7,y=255)
scr_2.config(state=DISABLED)  # 关闭可写入模式

# 按钮
button3_2 = ttk.Button(mighty2_1, text="获取期数", command=lambda : get_canshu(number_chosen_2.get()))
button3_2.place(x=530,y=0)

button1_2 = ttk.Button(mighty2_1, text="选择学员册", command=lambda : select_file(scr_2,pathin_2))
button1_2.place(x=530,y=30)

button2_2 = ttk.Button(mighty2_1, text="开始生成", command=lambda : main2(path=pathin_2.get(), filename=pathin2_2.get(), qishu=number_chosen_2.get()))
button2_2.place(x=530,y=60)

# 布局Frame
mighty2_2 = ttk.LabelFrame(tab2, text='生成考试试卷及答案')
mighty2_2.place(x=10,y=120,width=630,height=110)

# 第一行
label5_2 = ttk.Label(mighty2_2, text="选择试卷难易等级：")
label5_2.place(x=10,y=0)
number5_2 = StringVar()
number_chosen5_2 = ttk.Combobox(mighty2_2, width=8, textvariable=number5_2, state='readonly')
number_chosen5_2['values'] = ('难','中','易','极易') ################################# 以下部分只能可读，不设选项，未开发
number_chosen5_2.place(x=130,y=0)
number_chosen5_2.current(3)   # 设置初始显示值，值为元组['values']的下标 #################################
number_chosen5_2.config(state='readonly')  # 设为只读模式

label13_2 = ttk.Label(mighty2_2, text="选择试卷培训年度：")
label13_2.place(x=10,y=30)
number13_2 = StringVar()
number_chosen13_2 = ttk.Combobox(mighty2_2, width=8, textvariable=number13_2, state='readonly')
number_chosen13_2['values'] = tuple(2015+i for i in range(20))
number_chosen13_2.place(x=130,y=30)
number_chosen13_2.current(7)   # 设置初始显示值，值为元组['values']的下标
number_chosen13_2.config(state='readonly')  # 设为只读模式

label14_2 = ttk.Label(mighty2_2, text="选择试卷培训期数：")
label14_2.place(x=10,y=60)
number14_2 = StringVar()
number_chosen14_2 = ttk.Combobox(mighty2_2, width=8, textvariable=number14_2, state='readonly')
number_chosen14_2['values'] = ('一','二','三','四','五','六','七','八','九','十','十一','十二','十三','十四','十五','十六','十七','十八','十九','二十')
number_chosen14_2.place(x=130,y=60)
number_chosen14_2.current(10)   # 设置初始显示值，值为元组['values']的下标
number_chosen14_2.config(state='readonly')  # 设为只读模式

# 第二行
label6_2 = ttk.Label(mighty2_2, text="选择试卷题型数量：")
label6_2.place(x=245,y=0)

label7_2 = ttk.Label(mighty2_2, text="单选题")
label7_2.place(x=360,y=0)
number7_2 = StringVar()
number_chosen7_2 = ttk.Combobox(mighty2_2, width=6, textvariable=number7_2, state='readonly')
number_chosen7_2['values'] = ['共'+str(i)+'道' for i in range(1,41)] ##################################
number_chosen7_2.place(x=410,y=0)
number_chosen7_2.current(19)   # 设置初始显示值，值为元组['values']的下标 #################################
number_chosen7_2.config(state='readonly')  # 设为只读模式

label8_2 = ttk.Label(mighty2_2, text="判断题")
label8_2.place(x=360,y=30)
number8_2 = StringVar()
number_chosen8_2 = ttk.Combobox(mighty2_2, width=6, textvariable=number8_2, state='readonly')
number_chosen8_2['values'] = ['共'+str(i)+'道' for i in range(1,41)] #################################
number_chosen8_2.place(x=410,y=30)
number_chosen8_2.current(9)   # 设置初始显示值，值为元组['values']的下标 #################################
number_chosen8_2.config(state='readonly')  # 设为只读模式

label9_2 = ttk.Label(mighty2_2, text="多选题")
label9_2.place(x=490,y=0)
number9_2 = StringVar()
number_chosen9_2 = ttk.Combobox(mighty2_2, width=6, textvariable=number9_2, state='readonly')
number_chosen9_2['values'] = ['共'+str(i)+'道' for i in range(41)] #################################
number_chosen9_2.place(x=540,y=0)
number_chosen9_2.current(0)   # 设置初始显示值，值为元组['values']的下标 #################################
number_chosen9_2.config(state='readonly')  # 设为只读模式

label10_2 = ttk.Label(mighty2_2, text="填空题")
label10_2.place(x=490,y=30)
number10_2 = StringVar()
number_chosen10_2 = ttk.Combobox(mighty2_2, width=6, textvariable=number10_2, state='readonly')
number_chosen10_2['values'] = ['共'+str(i)+'道' for i in range(1,41)] #################################
number_chosen10_2.place(x=540,y=30)
number_chosen10_2.current(4)   # 设置初始显示值，值为元组['values']的下标 #################################
number_chosen10_2.config(state='readonly')  # 设为只读模式

label11_2 = ttk.Label(mighty2_2, text="简答题")
label11_2.place(x=360,y=60)
number11_2 = StringVar()
number_chosen11_2 = ttk.Combobox(mighty2_2, width=6, textvariable=number11_2, state='readonly')
number_chosen11_2['values'] = ['共'+str(i)+'道' for i in range(1,41)] #################################
number_chosen11_2.place(x=410,y=60)
number_chosen11_2.current(2)   # 设置初始显示值，值为元组['values']的下标 #################################
number_chosen11_2.config(state='readonly')  # 设为只读模式

label12_2 = ttk.Label(mighty2_2, text="论述题")
label12_2.place(x=490,y=60)
number12_2 = StringVar()
number_chosen12_2 = ttk.Combobox(mighty2_2, width=6, textvariable=number12_2, state='readonly')
number_chosen12_2['values'] = ['共'+str(i)+'道' for i in range(1,41)] #################################
number_chosen12_2.place(x=540,y=60)
number_chosen12_2.current(0)   # 设置初始显示值，值为元组['values']的下标 #################################
number_chosen12_2.config(state='readonly')  # 设为只读模式

button4_2 = ttk.Button( mighty2_2, text="开始合成试卷",
         command=lambda : generate_test_paper(diffcult=number5_2.get(),
         year=number13_2.get(),qishu=number14_2.get(),
         danxuan_num=(number7_2.get().strip('共')).strip('道'),panduan_num=(number8_2.get().strip('共')).strip('道'),
         duoxuan_num=(number9_2.get().strip('共')).strip('道'),tiankong_num=(number10_2.get().strip('共')).strip('道'),
         jianda_num=(number11_2.get().strip('共')).strip('道'),lunsu_num=(number12_2.get().strip('共')).strip('道')))
button4_2.place(x=245,y=58,width=100)

def auto_default_topic():
    number5_2.set('极易')
    number7_2.set('共20道')
    number8_2.set('共10道')
    number9_2.set('共0道')
    number10_2.set('共5道')
    number11_2.set('共3道')
    number12_2.set('共1道')

button4_2 = ttk.Button( mighty2_2, text="恢复默认题型", command=auto_default_topic)
button4_2.place(x=245,y=28,width=100)













'''# 3 座位考勤 座位表分区的控件 ############################################################'''

# 布局Frame
mighty3 = ttk.LabelFrame(tab3, text='生成学员座位表')
mighty3.place(x=10,y=10,width=630,height=110)

# 标签
label1_3 = ttk.Label(mighty3, text="选择本次培训期数：")
label1_3.place(x=10,y=0)
label2_3 = ttk.Label(mighty3, text="分组名单文件路径：")
label2_3.place(x=10,y=30)
label3_3 = ttk.Label(mighty3, text="生成座位表文件名：")
label3_3.place(x=10,y=60)
label4_3 = ttk.Label(tab3, text="输出窗口：")
label4_3.place(x=10,y=230)

# 输入期数
number_3 = StringVar()
number_chosen_3 = ttk.Combobox(mighty3, width=12, textvariable=number_3, state='readonly')
number_chosen_3['values'] = ('一','二','三','四','五','六','七','八','九','十','十一','十二','十三','十四','十五','十六','十七','十八','十九','二十')
number_chosen_3.place(x=130,y=0)
number_chosen_3.current(10)   # 设置初始显示值，值为元组['values']的下标
number_chosen_3.config(state='readonly')  # 设为只读模式

# 输入框
pathin_3 = StringVar()  # 定义变量
entry_pathin_3 = ttk.Entry(mighty3, textvariable=pathin_3)  # 输入框    # entry不能和grid连写，否则会报错
entry_pathin_3.place(x=130,y=30,width=380)
createToolTip(entry_pathin_3, '这里不需要输入')   # Add Tooltip

pathin2_3 = StringVar()  # 定义变量
entry_pathin2_3 = ttk.Entry(mighty3, textvariable=pathin2_3)  # 输入框    # entry不能和grid连写，否则会报错
entry_pathin2_3.place(x=130,y=60,width=380)
createToolTip(entry_pathin2_3, '这里可以选择需要输入生成的文件名称')   # Add Tooltip
pathin2_3.set('经济管理与法学学院分党校第{}期入党积极分子培训班座位表'.format(number_chosen_3.get()))

# 文本框
scr_3 = scrolledtext.ScrolledText(tab3, width=88, height=30, wrap=WORD)
scr_3.place(x=7,y=255)
scr_3.config(state=DISABLED)  # 关闭可写入模式

# 按钮
button3_3 = ttk.Button(mighty3, text="获取期数", command=lambda : get_canshu(number_chosen_3.get()))
button3_3.place(x=530,y=0)

button1_3 = ttk.Button(mighty3, text="选择分组名单", command=lambda : select_file(scr_3,pathin_3))
button1_3.place(x=530,y=30)

button2_3 = ttk.Button(mighty3, text="开始生成", command=lambda : main3(path=pathin_3.get(), filename=pathin2_3.get(), qishu=number_chosen_3.get()))
button2_3.place(x=530,y=60)





'''# 4 考勤表分区的控件  V2版本开始，合并至3座位表处############################################################'''
# 布局Frame
mighty4 = ttk.LabelFrame(tab3, text='生成小组考勤表')
mighty4.place(x=10,y=120,width=630,height=110)

# 标签
label1_4 = ttk.Label(mighty4, text="选择本次培训期数：")
label1_4.place(x=10,y=0)
label2_4 = ttk.Label(mighty4, text="分组名单文件路径：")
label2_4.place(x=10,y=30)
label3_4 = ttk.Label(mighty4, text="生成考勤表文件夹：")
label3_4.place(x=10,y=60)
# label4_4 = ttk.Label(tab3, text="输出窗口：")
# label4_4.place(x=10,y=200)

# 输入期数
number_4 = StringVar()
number_chosen_4 = ttk.Combobox(mighty4, width=12, textvariable=number_4, state='readonly')
number_chosen_4['values'] = ('一','二','三','四','五','六','七','八','九','十','十一','十二','十三','十四','十五','十六','十七','十八','十九','二十')
number_chosen_4.place(x=130,y=0)
number_chosen_4.current(10)   # 设置初始显示值，值为元组['values']的下标
number_chosen_4.config(state='readonly')  # 设为只读模式

# 输入框
pathin_4 = StringVar()  # 定义变量
entry_pathin_4 = ttk.Entry(mighty4, textvariable=pathin_4)  # 输入框    # entry不能和grid连写，否则会报错
entry_pathin_4.place(x=130,y=30,width=380)
createToolTip(entry_pathin_4, '这里不需要输入')   # Add Tooltip

pathin2_4 = StringVar()  # 定义变量
entry_pathin2_4 = ttk.Entry(mighty4, textvariable=pathin2_4)  # 输入框    # entry不能和grid连写，否则会报错
entry_pathin2_4.place(x=130,y=60,width=380)
createToolTip(entry_pathin2_4, '这里可以选择需要输入生成的文件名称')   # Add Tooltip
pathin2_4.set('经济管理与法学学院分党校第{}期入党积极分子培训班小组成员通讯录及考勤表'.format(number_chosen_4.get()))

# # 文本框 与scr_3共用
# scr_4 = scrolledtext.ScrolledText(tab3, width=88, height=30, wrap=WORD)
# scr_4.place(x=7,y=225)
# scr_4.config(state=DISABLED)  # 关闭可写入模式

# 按钮
button3_4 = ttk.Button(mighty4, text="获取期数", command=lambda : get_canshu(number_chosen_4.get()))
button3_4.place(x=530,y=0)

button1_4 = ttk.Button(mighty4, text="选择分组名单", command=lambda : select_file(scr_3,pathin_4))
button1_4.place(x=530,y=30)

button2_4 = ttk.Button(mighty4, text="开始生成", command=lambda : main4(path=pathin_4.get(), filename=pathin2_4.get(), qishu=number_chosen_4.get()))
button2_4.place(x=530,y=60)










'''# 5 时间管理分区的控件 ############################################################'''
# 布局Frame
mighty5 = ttk.LabelFrame(tab5, text='时间审核管理')
mighty5.place(x=10,y=10,width=630,height=210)
mighty1_5 = ttk.LabelFrame(mighty5, text='时间列')
mighty1_5.place(x=5,y=65,width=620,height=80)
# 标签
label2_5 = ttk.Label(mighty5, text="需审核的文件路径：")
label2_5.place(x=10,y=10)
label3_5 = ttk.Label(mighty5, text="生成审核标注文件：")
label3_5.place(x=10,y=150)
label4_5 = ttk.Label(tab5, text="输出窗口：")
label4_5.place(x=10,y=230)
label5_5 = ttk.Label(mighty5, text="选择表头行")
label5_5.place(x=60,y=40)
label6_5 = ttk.Label(mighty5, text="选择工作表")
label6_5.place(x=220,y=40)


# 表头行
# 检测列填入选项卡（自动填入，如果未打勾默认不能填写）
number9_5 = StringVar()
number_chosen9_5 = ttk.Combobox(mighty5, width=8, textvariable=number9_5)
number_chosen9_5['values'] = tuple('第{}行'.format(str(i)) for i in range(1,6))
number_chosen9_5.place(x=130,y=40)
number_chosen9_5.current(0)   # 设置初始显示值，值为元组['values']的下标
number_chosen9_5.config(state='readonly')  # 设为只读模式
createToolTip(number_chosen9_5, '这里展示自动检测后的的文件表头在哪一行，您也可以自行选择')   # Add Tooltip

# 工作表
# 检测列填入选项卡（自动填入，如果未打勾默认不能填写）
number10_5 = StringVar()
number_chosen10_5 = ttk.Combobox(mighty5, width=8, textvariable=number10_5)
number_chosen10_5['values'] = tuple('Sheet' + str(i) for i in range(1,21))
number_chosen10_5.place(x=290,y=40)
number_chosen10_5.current(0)   # 设置初始显示值，值为元组['values']的下标
number_chosen10_5.config(state='readonly')  # 设为只读模式
createToolTip(number_chosen10_5, '这里可以选择审核文件里面的哪一个工作表')   # Add Tooltip

def disabled(number_row_5,number_chosen_5):
    if number_row_5.get() == 0:  # 没有打勾
        number_chosen_5.config(state='disabled')
        number_chosen_5.set('')
    if number_row_5.get() == 1:
        number_chosen_5.config(state='normal')

# 时间列的六列

# 检测打勾选项卡
number_row1_5 = IntVar()
check1_5 = Checkbutton(mighty1_5, text="  姓名 ", variable=number_row1_5,command=lambda : disabled(number_row1_5,number_chosen1_5))
check1_5.state(['selected'])
number_row1_5.set(1)
check1_5.place(x=10,y=0)
# 检测列填入选项卡（自动填入，如果未打勾默认不能填写）
number1_5 = StringVar()
number_chosen1_5 = ttk.Combobox(mighty1_5, width=3, textvariable=number1_5 )
number_chosen1_5['values'] = ('A','B','C','D','E','F','G','H','I','J','K','L','M','N','O','P','Q','R','S','T','U','V','W','X','Y','Z')
number_chosen1_5.place(x=85,y=0)
# number_chosen1_5.current(7)   # 设置初始显示值，值为元组['values']的下标
# number_chosen1_5.config(state='readonly')  # 设为只读模式 state='readonly', # 设置状态 normal(可选可输入)、readonly(只可选)、 disabled


# 检测打勾选项卡
number_row2_5 = IntVar()
check2_5 = Checkbutton(mighty1_5, text="身份证号", variable=number_row2_5,command=lambda : disabled(number_row2_5,number_chosen2_5))
check2_5.state(['selected'])
number_row2_5.set(1)
check2_5.place(x=10,y=30)
# 检测列填入选项卡（自动填入，如果未打勾默认不能填写）
number2_5 = StringVar()
number_chosen2_5 = ttk.Combobox(mighty1_5, width=3, textvariable=number2_5)
number_chosen2_5['values'] = ('A','B','C','D','E','F','G','H','I','J','K','L','M','N','O','P','Q','R','S','T','U','V','W','X','Y','Z')
number_chosen2_5.place(x=85,y=30)
# number_chosen2_5.current(7)   # 设置初始显示值，值为元组['values']的下标
# number_chosen2_5.config(state='readonly')  # 设为只读模式


# 检测打勾选项卡
number_row3_5 = IntVar()
check3_5 = Checkbutton(mighty1_5, text="出生年月", variable=number_row3_5,command=lambda : disabled(number_row3_5,number_chosen3_5))
check3_5.state(['selected'])
number_row3_5.set(1)
check3_5.place(x=150,y=0)
# 检测列填入选项卡（自动填入，如果未打勾默认不能填写）
number3_5 = StringVar()
number_chosen3_5 = ttk.Combobox(mighty1_5, width=3, textvariable=number3_5)
number_chosen3_5['values'] = ('A','B','C','D','E','F','G','H','I','J','K','L','M','N','O','P','Q','R','S','T','U','V','W','X','Y','Z')
number_chosen3_5.place(x=225,y=0)
# number_chosen3_5.current(7)   # 设置初始显示值，值为元组['values']的下标
# number_chosen3_5.config(state='readonly')  # 设为只读模式


# 检测打勾选项卡
number_row4_5 = IntVar()
check4_5 = Checkbutton(mighty1_5, text="申请入党", variable=number_row4_5,command=lambda : disabled(number_row4_5,number_chosen4_5))
check4_5.state(['selected'])
number_row4_5.set(1)
check4_5.place(x=150,y=30)
# 检测列填入选项卡（自动填入，如果未打勾默认不能填写）
number4_5 = StringVar()
number_chosen4_5 = ttk.Combobox(mighty1_5, width=3, textvariable=number4_5)
number_chosen4_5['values'] = ('A','B','C','D','E','F','G','H','I','J','K','L','M','N','O','P','Q','R','S','T','U','V','W','X','Y','Z')
number_chosen4_5.place(x=225,y=30)
# number_chosen4_5.current(7)   # 设置初始显示值，值为元组['values']的下标
# number_chosen4_5.config(state='readonly')  # 设为只读模式


# 检测打勾选项卡
number_row5_5 = IntVar()
check5_5 = Checkbutton(mighty1_5, text="积极分子", variable=number_row5_5,command=lambda : disabled(number_row5_5,number_chosen5_5))
check5_5.state(['selected'])
number_row5_5.set(1)
check5_5.place(x=290,y=0)
# 检测列填入选项卡（自动填入，如果未打勾默认不能填写）
number5_5 = StringVar()
number_chosen5_5 = ttk.Combobox(mighty1_5, width=3, textvariable=number5_5)
number_chosen5_5['values'] = ('A','B','C','D','E','F','G','H','I','J','K','L','M','N','O','P','Q','R','S','T','U','V','W','X','Y','Z')
number_chosen5_5.place(x=365,y=0)
# number_chosen5_5.current(7)   # 设置初始显示值，值为元组['values']的下标
# number_chosen5_5.config(state='readonly')  # 设为只读模式


# 检测打勾选项卡
number_row6_5 = IntVar()
check6_5 = Checkbutton(mighty1_5, text="发展对象", variable=number_row6_5,command=lambda : disabled(number_row6_5,number_chosen6_5))
check6_5.state(['selected'])
number_row6_5.set(1)
check6_5.place(x=290,y=30)
# 检测列填入选项卡（自动填入，如果未打勾默认不能填写）
number6_5 = StringVar()
number_chosen6_5 = ttk.Combobox(mighty1_5, width=3, textvariable=number6_5)
number_chosen6_5['values'] = ('A','B','C','D','E','F','G','H','I','J','K','L','M','N','O','P','Q','R','S','T','U','V','W','X','Y','Z')
number_chosen6_5.place(x=365,y=30)
# number_chosen6_5.current(7)   # 设置初始显示值，值为元组['values']的下标
# number_chosen6_5.config(state='readonly')  # 设为只读模式


# 检测打勾选项卡
number_row7_5 = IntVar()
check7_5 = Checkbutton(mighty1_5, text="预备党员", variable=number_row7_5,command=lambda : disabled(number_row7_5,number_chosen7_5))
check7_5.state(['selected'])
number_row7_5.set(1)
check7_5.place(x=430,y=0)
# 检测列填入选项卡（自动填入，如果未打勾默认不能填写）
number7_5 = StringVar()
number_chosen7_5 = ttk.Combobox(mighty1_5, width=3, textvariable=number7_5)
number_chosen7_5['values'] = ('A','B','C','D','E','F','G','H','I','J','K','L','M','N','O','P','Q','R','S','T','U','V','W','X','Y','Z')
number_chosen7_5.place(x=505,y=0)
# number_chosen7_5.current(7)   # 设置初始显示值，值为元组['values']的下标
# number_chosen7_5.config(state='readonly')  # 设为只读模式


# 检测打勾选项卡
number_row8_5 = IntVar()
check8_5 = Checkbutton(mighty1_5, text="党员转正", variable=number_row8_5,command=lambda : disabled(number_row8_5,number_chosen8_5))
check8_5.state(['selected'])
number_row8_5.set(1)
check8_5.place(x=430,y=30)
# 检测列填入选项卡（自动填入，如果未打勾默认不能填写）
number8_5 = StringVar()
number_chosen8_5 = ttk.Combobox(mighty1_5, width=3, textvariable=number8_5)
number_chosen8_5['values'] = ('A','B','C','D','E','F','G','H','I','J','K','L','M','N','O','P','Q','R','S','T','U','V','W','X','Y','Z')
number_chosen8_5.place(x=505,y=30)
# number_chosen8_5.current(7)   # 设置初始显示值，值为元组['values']的下标
# number_chosen8_5.config(state='readonly')  # 设为只读模式





# 输入框
pathin_5 = StringVar()  # 定义变量
entry_pathin_5 = ttk.Entry(mighty5, textvariable=pathin_5)  # 输入框    # entry不能和grid连写，否则会报错
entry_pathin_5.place(x=130,y=10,width=380)
createToolTip(entry_pathin_5, '这里不需要输入')   # Add Tooltip

pathin2_5 = StringVar()  # 定义变量
entry_pathin_5 = ttk.Entry(mighty5, textvariable=pathin2_5)  # 输入框    # entry不能和grid连写，否则会报错
entry_pathin_5.place(x=130,y=150,width=380)
createToolTip(entry_pathin_5, '这里可以选择需要输入生成的标注文件名称')   # Add Tooltip
pathin2_5.set('时间审核标注文件')

# 文本框
scr_5 = scrolledtext.ScrolledText(tab5, width=88, height=29, wrap=WORD)
scr_5.place(x=7,y=255)
scr_5.config(state=DISABLED)  # 关闭可写入模式

# 按钮
button1_5 = ttk.Button(mighty1_5, text="自动\n检测", command=lambda:auto_time_management(pathin_5.get()))
button1_5.place(x=560,y=0, width=53, height=50)

button2_5 = ttk.Button(mighty5, text="选择文件", command=lambda : select_file(scr_5, pathin_5))
button2_5.place(x=530,y=10)

button3_5 = ttk.Button(mighty5, text="开始审核", command=lambda : main5(path=pathin_5.get(), filename=pathin2_5.get()))
button3_5.place(x=530,y=150)

button4_5 = ttk.Button(mighty5, text="编辑公式", command=gongshi)
button4_5.place(x=390,y=40)

button5_5 = ttk.Button(mighty5, text="一键多表", command=one_key_jiance)
button5_5.place(x=530,y=40)
createToolTip(button5_5, '这里同时在一个文件中一键审核多个sheet表，并生成审核表（功能测试阶段）')   # Add Tooltip



'''# 6 请示管理的控件 ############################################################'''
# 布局Frame
mighty6_1 = ttk.LabelFrame(tab6, text='总请示管理')
mighty6_1.place(x=10,y=30,width=630,height=220)

# 选择是哪一类型的请示
number_6_1 = IntVar()
check6_1 = Checkbutton(tab6, text="发展对象请示", variable=number_6_1)
check6_1.state(['active'])
# check6_1.state(['disabled'])
# number_6_1.set(1)# 默认不勾选
check6_1.place(x=120,y=10)
number_6_2 = IntVar()
check6_2 = Checkbutton(tab6, text="预备党员请示", variable=number_6_2)
check6_2.state(['active'])  # Clears (turns off) the checkbutton.
check6_2.place(x=270,y=10)
number_6_3 = IntVar()
check6_3 = Checkbutton(tab6, text="党员转正请示", variable=number_6_3)
check6_3.state(['active'])  # Clears (turns off) the checkbutton.
check6_3.place(x=420,y=10)
# GUI Callback function
def checkCallback_3(*ignoredArgs):
    # only enable one checkbutton
    if number_6_1.get():  # ==1
        check6_2.configure(state='disabled')
        check6_3.configure(state='disabled')
    else:
        check6_2.configure(state='normal')
        check6_3.configure(state='normal')
def checkCallback_4(*ignoredArgs):
    if number_6_2.get():
        check6_1.configure(state='disabled')
        check6_3.configure(state='disabled')
    else:
        check6_1.configure(state='normal')
        check6_3.configure(state='normal')
def checkCallback_5(*ignoredArgs):
    if number_6_3.get():
        check6_1.configure(state='disabled')
        check6_2.configure(state='disabled')
    else:
        check6_1.configure(state='normal')
        check6_2.configure(state='normal')
# trace the state of the two checkbuttons  #？？
number_6_1.trace('w', lambda unused0, unused1, unused2: checkCallback_3())
number_6_2.trace('w', lambda unused0, unused1, unused2: checkCallback_4())
number_6_3.trace('w', lambda unused0, unused1, unused2: checkCallback_5())

# 第一行标签
label1_6 = ttk.Label(mighty6_1, text="年度")
label1_6.place(x=0,y=10)
number1_6 = StringVar()
number_chosen1_6 = ttk.Combobox(mighty6_1, width=5, textvariable=number1_6)
number_chosen1_6['values'] = tuple(2015+i for i in range(20))
number_chosen1_6.place(x=30,y=10)
number_chosen1_6.current(now.year-2015)   # 设置初始显示值，值为元组['values']的下标
number_chosen1_6.config(state='readonly')  # 设为只读模式
createToolTip(number_chosen1_6, '这里展示请示年度时间，您需要自行选择')   # Add Tooltip

label2_6 = ttk.Label(mighty6_1, text="批次")
label2_6.place(x=100,y=10)
number2_6 = StringVar()
number_chosen2_6 = ttk.Combobox(mighty6_1, width=5, textvariable=number2_6)
number_chosen2_6['values'] = ('第一批','第二批','第三批','第四批','第五批')
number_chosen2_6.place(x=130,y=10)
number_chosen2_6.current(0)   # 设置初始显示值，值为元组['values']的下标
number_chosen2_6.config(state='readonly')  # 设为只读模式
createToolTip(number_chosen2_6, '这里展示请示批次时间，您需要自行选择')   # Add Tooltip

label3_6 = ttk.Label(mighty6_1, text="半年")
label3_6.place(x=200,y=10)
number3_6 = StringVar()
number_chosen3_6 = ttk.Combobox(mighty6_1, width=5, textvariable=number3_6)
number_chosen3_6['values'] = ('上','下')
number_chosen3_6.place(x=230,y=10)
number_chosen3_6.current(0)   # 设置初始显示值，值为元组['values']的下标
number_chosen3_6.config(state='readonly')  # 设为只读模式
createToolTip(number_chosen3_6, '这里展示请示上下半年时间，您需要自行选择')   # Add Tooltip

label4_6 = ttk.Label(mighty6_1, text="落款：")
label4_6.place(x=300,y=10)
# number4_6 = StringVar()

label5_6 = ttk.Label(mighty6_1, text="年")
label5_6.place(x=393,y=11)
number5_6 = StringVar()
number_chosen5_6 = ttk.Combobox(mighty6_1, width=4, textvariable=number5_6)
number_chosen5_6['values'] = tuple(2015+i for i in range(20))
number_chosen5_6.place(x=340,y=10)
number_chosen5_6.current(now.year-2015)   # 设置初始显示值，值为元组['values']的下标
number_chosen5_6.config(state='readonly')  # 设为只读模式
createToolTip(number_chosen5_6, '这里展示请示落款时间，您需要自行选择')   # Add Tooltip

label6_6 = ttk.Label(mighty6_1, text="月")
label6_6.place(x=447,y=11)
number6_6 = StringVar()
number_chosen6_6 = ttk.Combobox(mighty6_1, width=2, textvariable=number6_6)
number_chosen6_6['values'] = tuple(1+i for i in range(12))
number_chosen6_6.place(x=410,y=10)
number_chosen6_6.current(now.month-1)   # 设置初始显示值，值为元组['values']的下标
number_chosen6_6.config(state='readonly')  # 设为只读模式
createToolTip(number_chosen6_6, '这里展示请示落款时间，您需要自行选择')   # Add Tooltip

label7_6 = ttk.Label(mighty6_1, text="日")
label7_6.place(x=503,y=11)
number7_6 = StringVar()
number_chosen7_6 = ttk.Combobox(mighty6_1, width=2, textvariable=number7_6)
number_chosen7_6['values'] = tuple(1+i for i in range(31))
number_chosen7_6.place(x=465,y=10)
number_chosen7_6.current(now.day-1)   # 设置初始显示值，值为元组['values']的下标
number_chosen7_6.config(state='readonly')  # 设为只读模式
createToolTip(number_chosen7_6, '这里展示请示落款时间，您需要自行选择')   # Add Tooltip

button1_6 = ttk.Button(mighty6_1, text="模板修改", command = qingshi_model_alter) #
button1_6.place(x=530,y=10)

# 第二行标签
label8_6 = ttk.Label(mighty6_1, text="各支部请示送审表文件夹：")
label8_6.place(x=0,y=40)
pathin_6 = StringVar()  # 定义变量
entry_pathin_6 = ttk.Entry(mighty6_1, textvariable=pathin_6)  # 输入框    # entry不能和grid连写，否则会报错
entry_pathin_6.place(x=145,y=40,width=365)
createToolTip(entry_pathin_6, '这里不需要输入')   # Add Tooltip
button2_6 = ttk.Button(mighty6_1, text="选择文件夹", command = lambda : select_files(scr_6, pathin_6))
button2_6.place(x=530,y=40)

# 第三行标签
label9_6 = ttk.Label(mighty6_1, text="经")
label9_6.place(x=10,y=71)

number9_6 = StringVar()
number_chosen9_6 = ttk.Combobox(mighty6_1, width=16, textvariable=number9_6)
number_chosen9_6['values'] = zhibu_allname
number_chosen9_6.place(x=30,y=70)
number_chosen9_6.current(0)   # 设置初始显示值，值为元组['values']的下标
# number_chosen9_6.config(state='readonly')  # 设为只读模式
createToolTip(number_chosen9_6, '这里展示请示的首个支部名，您需要确认并自行选择')   # Add Tooltip

labe10_6 = ttk.Label(mighty6_1, text="等")
labe10_6.place(x=170,y=71)
number10_6 = StringVar()
number_chosen10_6 = ttk.Combobox(mighty6_1, width=4, textvariable=number10_6)
number_chosen10_6['values'] = [1+i for i in range(15)]
number_chosen10_6.place(x=190,y=70)
number_chosen10_6.current(10)   # 设置初始显示值，值为元组['values']的下标
number_chosen10_6.config(state='readonly')  # 设为只读模式
createToolTip(number_chosen10_6, '这里展示请示的支部数量，您需要确认并自行选择')   # Add Tooltip

labe11_6 = ttk.Label(mighty6_1, text="个支部，确认")
labe11_6.place(x=245,y=71)
number11_6 = StringVar()
number_chosen11_6 = ttk.Combobox(mighty6_1, width=7, textvariable=number11_6)
number_chosen11_6['values'] = ['张三李四']
number_chosen11_6.place(x=325,y=70)
number_chosen11_6.current(0)   # 设置初始显示值，值为元组['values']的下标
# number_chosen11_6.config(state='readonly')  # 设为只读模式
createToolTip(number_chosen11_6, '这里展示自动检测后的首名同志，您也可以自行选择')   # Add Tooltip

labe12_6 = ttk.Label(mighty6_1, text="等")
labe12_6.place(x=405,y=71)
number12_6 = StringVar()
number_chosen12_6 = ttk.Combobox(mighty6_1, width=4, textvariable=number12_6)
number_chosen12_6['values'] = tuple(1+i for i in range(200))
number_chosen12_6.place(x=425,y=70)
number_chosen12_6.current(99)   # 设置初始显示值，值为元组['values']的下标
# number_chosen12_6.config(state='readonly')  # 设为只读模式
createToolTip(number_chosen12_6, '这里展示自动检测后的同志数量，您也可以自行选择')   # Add Tooltip

labe11_6 = ttk.Label(mighty6_1, text="名同志")
labe11_6.place(x=480,y=71)

# 第四行
labe12_6 = ttk.Label(mighty6_1, text="具体名单：（人名间用空格隔开）")
labe12_6.place(x=5,y=100)
# 表格文本框
scr_sheet6 = scrolledtext.ScrolledText(mighty6_1, width=72, height=5, wrap=WORD)
scr_sheet6.place(x=5,y=125)

button3_6 = ttk.Button(mighty6_1, text="自动识别", command = auto_qingshi_read )  #
button3_6.place(x=530,y=70, height=50)


button4_6 = ttk.Button(mighty6_1, text="生成",
    command =lambda : write_qingshi(cookie = str(number_6_1.get()) + str(number_6_2.get()) + str(number_6_3.get()),
                                        yeardu = number1_6.get(),pici = number2_6.get(),year_up = number3_6.get(),
                                        year = number5_6.get(),month = number6_6.get(),day = number7_6.get(),
                                        party_name = number9_6.get(),party_num = int(number10_6.get()),
                                        first_people = number11_6.get(),people_num = int(number12_6.get()),
                                        people_sheet =  (sorted(scr_sheet6.get(1.0,'end').split(),key=lambda keys:[pinyin(i, style=Style.TONE3) for i in keys])
                                        if peoplename.get() == 1 else scr_sheet6.get(1.0,'end').split())
                                    )) #获取文本框第一行到全部的内容
button4_6.place(x=530,y=125, height=70)

# 文本框
label13_6 = ttk.Label(tab6, text="输出窗口：")
label13_6.place(x=7, y=255)
scr_6 = scrolledtext.ScrolledText(tab6, width=88, height=27, wrap=WORD)
scr_6.place(x=7,y=275)
scr_6.config(state=DISABLED)  # 关闭可写入模式










'''# 7 批复管理的控件 ############################################################'''
# 布局Frame
mighty7_1 = ttk.LabelFrame(tab7, text='总批复管理')
mighty7_1.place(x=10,y=30,width=630,height=250)

# 选择是哪一类型的批复
number_7_1 = IntVar()
check7_1 = Checkbutton(tab7, text="发展对象批复", variable=number_7_1)
check7_1.state(['active'])
# number_7_1.set(1)# 默认不勾选
check7_1.place(x=120,y=10)
number_7_2 = IntVar()
check7_2 = Checkbutton(tab7, text="预备党员批复", variable=number_7_2)
check7_2.state(['active'])  # Clears (turns off) the checkbutton.
check7_2.place(x=270,y=10)
number_7_3 = IntVar()
check7_3 = Checkbutton(tab7, text="党员转正批复", variable=number_7_3)
check7_3.state(['active'])  # Clears (turns off) the checkbutton.
check7_3.place(x=420,y=10)
# GUI Callback function
def checkCallback_6(*ignoredArgs):
    # only enable one checkbutton
    if number_7_1.get():  # ==1
        check7_2.configure(state='disabled')
        check7_3.configure(state='disabled')
    else:
        check7_2.configure(state='normal')
        check7_3.configure(state='normal')
def checkCallback_7(*ignoredArgs):
    if number_7_2.get():
        check7_1.configure(state='disabled')
        check7_3.configure(state='disabled')
    else:
        check7_1.configure(state='normal')
        check7_3.configure(state='normal')
def checkCallback_8(*ignoredArgs):
    if number_7_3.get():
        check7_1.configure(state='disabled')
        check7_2.configure(state='disabled')
    else:
        check7_1.configure(state='normal')
        check7_2.configure(state='normal')
# trace the state of the two checkbuttons  #？？
number_7_1.trace('w', lambda unused0, unused1, unused2: checkCallback_6())
number_7_2.trace('w', lambda unused0, unused1, unused2: checkCallback_7())
number_7_3.trace('w', lambda unused0, unused1, unused2: checkCallback_8())

# 第一行标签
label1_7 = ttk.Label(mighty7_1, text="年度")
label1_7.place(x=0,y=10)
number1_7 = StringVar()
number_chosen1_7 = ttk.Combobox(mighty7_1, width=5, textvariable=number1_7)
number_chosen1_7['values'] = tuple(2015+i for i in range(20))
number_chosen1_7.place(x=30,y=10)
number_chosen1_7.current(now.year-2015)   # 设置初始显示值，值为元组['values']的下标
number_chosen1_7.config(state='readonly')  # 设为只读模式
createToolTip(number_chosen1_7, '这里展示批复年度时间，您需要自行选择')   # Add Tooltip

label2_7 = ttk.Label(mighty7_1, text="批次")
label2_7.place(x=100,y=10)
number2_7 = StringVar()
number_chosen2_7 = ttk.Combobox(mighty7_1, width=5, textvariable=number2_7)
number_chosen2_7['values'] = ('第一批','第二批','第三批','第四批','第五批')
number_chosen2_7.place(x=130,y=10)
number_chosen2_7.current(0)   # 设置初始显示值，值为元组['values']的下标
number_chosen2_7.config(state='readonly')  # 设为只读模式
createToolTip(number_chosen2_7, '这里展示批复批次时间，您需要自行选择')   # Add Tooltip

label3_7 = ttk.Label(mighty7_1, text="半年")
label3_7.place(x=200,y=10)
number3_7 = StringVar()
number_chosen3_7 = ttk.Combobox(mighty7_1, width=5, textvariable=number3_7)
number_chosen3_7['values'] = ('上','下')
number_chosen3_7.place(x=230,y=10)
number_chosen3_7.current(0)   # 设置初始显示值，值为元组['values']的下标
number_chosen3_7.config(state='readonly')  # 设为只读模式
createToolTip(number_chosen3_7, '这里展示批复上下半年时间，您需要自行选择')   # Add Tooltip

label4_7 = ttk.Label(mighty7_1, text="落款：")
label4_7.place(x=300,y=10)
number4_7 = StringVar()

label5_7 = ttk.Label(mighty7_1, text="年")
label5_7.place(x=393,y=11)
number5_7 = StringVar()
number_chosen5_7 = ttk.Combobox(mighty7_1, width=4, textvariable=number5_7)
number_chosen5_7['values'] = tuple(2015+i for i in range(20))
number_chosen5_7.place(x=340,y=10)
number_chosen5_7.current(now.year-2015)   # 设置初始显示值，值为元组['values']的下标
number_chosen5_7.config(state='readonly')  # 设为只读模式
createToolTip(number_chosen5_7, '这里展示批复落款时间，您需要自行选择')   # Add Tooltip

label6_7 = ttk.Label(mighty7_1, text="月")
label6_7.place(x=447,y=11)
number6_7 = StringVar()
number_chosen6_7 = ttk.Combobox(mighty7_1, width=2, textvariable=number6_7)
number_chosen6_7['values'] = tuple(1+i for i in range(12))
number_chosen6_7.place(x=410,y=10)
number_chosen6_7.current(now.month-1)   # 设置初始显示值，值为元组['values']的下标
number_chosen6_7.config(state='readonly')  # 设为只读模式
createToolTip(number_chosen6_7, '这里展示批复落款时间，您需要自行选择')   # Add Tooltip

label7_7 = ttk.Label(mighty7_1, text="日")
label7_7.place(x=503,y=11)
number7_7 = StringVar()
number_chosen7_7 = ttk.Combobox(mighty7_1, width=2, textvariable=number7_7)
number_chosen7_7['values'] = tuple(1+i for i in range(31))
number_chosen7_7.place(x=465,y=10)
number_chosen7_7.current(now.day-1)   # 设置初始显示值，值为元组['values']的下标
number_chosen7_7.config(state='readonly')  # 设为只读模式
createToolTip(number_chosen7_7, '这里展示批复落款时间，您需要自行选择')   # Add Tooltip

button1_7 = ttk.Button(mighty7_1, text="模板修改", command = pifu_model_alter) #
button1_7.place(x=530,y=10)

# 第二行标签
label8_7 = ttk.Label(mighty7_1, text="各支部批复送审表文件夹：")
label8_7.place(x=0,y=40)
pathin_7 = StringVar()  # 定义变量
entry_pathin_7 = ttk.Entry(mighty7_1, textvariable=pathin_7)  # 输入框    # entry不能和grid连写，否则会报错
entry_pathin_7.place(x=145,y=40,width=365)
createToolTip(entry_pathin_7, '这里不需要输入')   # Add Tooltip
button2_7 = ttk.Button(mighty7_1, text="选择文件夹", command = lambda : select_files(scr_7, pathin_7))
button2_7.place(x=530,y=40)

# 第三行标签
label9_7 = ttk.Label(mighty7_1, text="收到")
label9_7.place(x=0,y=71)

number9_7 = StringVar()
number_chosen9_7 = ttk.Combobox(mighty7_1, width=17, textvariable=number9_7)
number_chosen9_7['values'] = zhibu_allname
number_chosen9_7.place(x=30,y=70)
number_chosen9_7.current(0)   # 设置初始显示值，值为元组['values']的下标
# number_chosen9_7.config(state='readonly')  # 设为只读模式
createToolTip(number_chosen9_7, '这里展示批复回复的首个支部名，您需要确认并自行选择')   # Add Tooltip

labe10_7 = ttk.Label(mighty7_1, text="等")
labe10_7.place(x=170,y=71)
number10_7 = StringVar()
number_chosen10_7 = ttk.Combobox(mighty7_1, width=4, textvariable=number10_7)
number_chosen10_7['values'] = [1+i for i in range(15)]
number_chosen10_7.place(x=190,y=70)
number_chosen10_7.current(10)   # 设置初始显示值，值为元组['values']的下标
number_chosen10_7.config(state='readonly')  # 设为只读模式
createToolTip(number_chosen10_7, '这里展示批复的支部数量，您需要确认并自行选择')   # Add Tooltip

labe11_7 = ttk.Label(mighty7_1, text="个支部，确认")
labe11_7.place(x=245,y=71)
number11_7 = StringVar()
number_chosen11_7 = ttk.Combobox(mighty7_1, width=7, textvariable=number11_7)
number_chosen11_7['values'] = ['张三李四']
number_chosen11_7.place(x=325,y=70)
number_chosen11_7.current(0)   # 设置初始显示值，值为元组['values']的下标
# number_chosen11_7.config(state='readonly')  # 设为只读模式
createToolTip(number_chosen11_7, '这里展示自动检测后的首名同志，您也可以自行选择')   # Add Tooltip

labe12_7 = ttk.Label(mighty7_1, text="等")
labe12_7.place(x=405,y=71)
number12_7 = StringVar()
number_chosen12_7 = ttk.Combobox(mighty7_1, width=4, textvariable=number12_7)
number_chosen12_7['values'] = tuple(1+i for i in range(200))
number_chosen12_7.place(x=425,y=70)
number_chosen12_7.current(99)   # 设置初始显示值，值为元组['values']的下标
# number_chosen12_7.config(state='readonly')  # 设为只读模式
createToolTip(number_chosen12_7, '这里展示自动检测后的同志数量，您也可以自行选择')   # Add Tooltip

labe11_7 = ttk.Label(mighty7_1, text="名同志")
labe11_7.place(x=480,y=71)

# 第四行
label16_7 = ttk.Label(mighty7_1, text="其请示名:")
label16_7.place(x=0,y=100)
number16_7 = StringVar()
entry_pathin2_7 = ttk.Entry(mighty7_1, textvariable=number16_7)  # 请示名字输入框
entry_pathin2_7.place(x=60,y=100,width=450)
createToolTip(entry_pathin2_7, '这里展示自动检测后的请示名字（党委会收到的支部请示），您也可以自行选择')   # Add Tooltip

label12_7 = ttk.Label(mighty7_1, text="支部请示落款：")
label12_7.place(x=255,y=130)

label13_7 = ttk.Label(mighty7_1, text="年")
label13_7.place(x=393,y=130)
number13_7 = StringVar()
number_chosen13_7 = ttk.Combobox(mighty7_1, width=4, textvariable=number13_7)
number_chosen13_7['values'] = tuple(2015+i for i in range(20))
number_chosen13_7.place(x=340,y=128)
number_chosen13_7.current(now.year-2015)   # 设置初始显示值，值为元组['values']的下标
number_chosen13_7.config(state='readonly')  # 设为只读模式
createToolTip(number_chosen13_7, '这里展示请示落款时间，您需要自行选择')   # Add Tooltip

label14_7 = ttk.Label(mighty7_1, text="月")
label14_7.place(x=447,y=130)
number14_7 = StringVar()
number_chosen14_7 = ttk.Combobox(mighty7_1, width=2, textvariable=number14_7)
number_chosen14_7['values'] = tuple(1+i for i in range(12))
number_chosen14_7.place(x=410,y=128)
number_chosen14_7.current(now.month-1)   # 设置初始显示值，值为元组['values']的下标
number_chosen14_7.config(state='readonly')  # 设为只读模式
createToolTip(number_chosen14_7, '这里展示请示落款时间，您需要自行选择')   # Add Tooltip

label15_7 = ttk.Label(mighty7_1, text="日")
label15_7.place(x=503,y=130)
number15_7 = StringVar()
number_chosen15_7 = ttk.Combobox(mighty7_1, width=2, textvariable=number15_7)
number_chosen15_7['values'] = tuple(1+i for i in range(31))
number_chosen15_7.place(x=465,y=128)
number_chosen15_7.current(now.day-1)   # 设置初始显示值，值为元组['values']的下标
number_chosen15_7.config(state='readonly')  # 设为只读模式
createToolTip(number_chosen15_7, '这里展示请示落款时间，您需要自行选择')   # Add Tooltip

# 第五行
labe17_7 = ttk.Label(mighty7_1, text="具体名单：（人名间用空格隔开）")
labe17_7.place(x=0,y=130)
# 表格文本框
scr_sheet7 = scrolledtext.ScrolledText(mighty7_1, width=72, height=5, wrap=WORD)
scr_sheet7.place(x=5,y=155)

button3_7 = ttk.Button(mighty7_1, text="自动识别", command = auto_pifu_read )  #
button3_7.place(x=530,y=70, height=80)

button4_7 = ttk.Button(mighty7_1, text="生成",
    command =lambda : write_pifu(cookie = str(number_7_1.get()) + str(number_7_2.get()) + str(number_7_3.get()),
                                        yeardu = number1_7.get(),pici = number2_7.get(),year_up = number3_7.get(),
                                        qs_year=number13_7.get(), qs_month=number14_7.get(), qs_day=number15_7.get(),
                                        qingshi_name=number16_7.get(),
                                        year = number5_7.get(),month = number7_7.get(),day = number7_7.get(),
                                        party_name = number9_7.get(),party_num = int(number10_7.get()),
                                        first_people = number11_7.get(),people_num = int(number12_7.get()),
                                 people_sheet=(sorted(scr_sheet7.get(1.0, 'end').split(),
                                                      key=lambda keys: [pinyin(i, style=Style.TONE3) for i in keys])
                                               if peoplename.get() == 1 else scr_sheet7.get(1.0, 'end').split())
                       ))  # 获取文本框第一行到全部的内容
button4_7.place(x=530,y=155, height=70)

# 文本框
label18_7 = ttk.Label(tab7, text="输出窗口：")
label18_7.place(x=7, y=285)
scr_7 = scrolledtext.ScrolledText(tab7, width=88, height=25, wrap=WORD)
scr_7.place(x=7,y=305)
scr_7.config(state=DISABLED)  # 关闭可写入模式








'''# 8 备案报告的控件 ############################################################'''
# 布局Frame
mighty8_1 = ttk.LabelFrame(tab8, text='总备案报告')
mighty8_1.place(x=10,y=30,width=630,height=220)

# 选择是哪一类型的备案报告
number_8_1 = IntVar()
check8_1 = Checkbutton(tab8, text="预备党员报组织部备案报告", variable=number_8_1)
check8_1.state(['active'])  # Clears (turns off) the checkbutton.
# number_8_1.set(1) # 默认不勾选
check8_1.place(x=130,y=10)
number_8_2 = IntVar()
check8_2 = Checkbutton(tab8, text="党员转正报组织部备案报告", variable=number_8_2)
check8_2.state(['active'])  # Clears (turns off) the checkbutton.
check8_2.place(x=320,y=10)
# GUI Callback function
def checkCallback_9(*ignoredArgs):
    if number_8_1.get():
        check8_2.configure(state='disabled')
    else:
        check8_2.configure(state='normal')
    if number_8_2.get():
        check8_1.configure(state='disabled')
    else:
        check8_1.configure(state='normal')
# trace the state of the two checkbuttons  #？？
number_8_1.trace('w', lambda unused0, unused1, unused2: checkCallback_9())
number_8_2.trace('w', lambda unused0, unused1, unused2: checkCallback_9())

# 第一行标签
label1_8 = ttk.Label(mighty8_1, text="年度")
label1_8.place(x=0,y=10)
number1_8 = StringVar()
number_chosen1_8 = ttk.Combobox(mighty8_1, width=5, textvariable=number1_8)
number_chosen1_8['values'] = tuple(2015+i for i in range(20))
number_chosen1_8.place(x=30,y=10)
number_chosen1_8.current(now.year-2015)   # 设置初始显示值，值为元组['values']的下标
number_chosen1_8.config(state='readonly')  # 设为只读模式
createToolTip(number_chosen1_8, '这里展示备案报告年度时间，您需要自行选择')   # Add Tooltip

label2_8 = ttk.Label(mighty8_1, text="批次")
label2_8.place(x=100,y=10)
number2_8 = StringVar()
number_chosen2_8 = ttk.Combobox(mighty8_1, width=5, textvariable=number2_8)
number_chosen2_8['values'] = ('第一批','第二批','第三批','第四批','第五批')
number_chosen2_8.place(x=130,y=10)
number_chosen2_8.current(0)   # 设置初始显示值，值为元组['values']的下标
number_chosen2_8.config(state='readonly')  # 设为只读模式
createToolTip(number_chosen2_8, '这里展示备案报告批次时间，您需要自行选择')   # Add Tooltip

label3_8 = ttk.Label(mighty8_1, text="半年")
label3_8.place(x=200,y=10)
number3_8 = StringVar()
number_chosen3_8 = ttk.Combobox(mighty8_1, width=5, textvariable=number3_8)
number_chosen3_8['values'] = ('上','下')
number_chosen3_8.place(x=230,y=10)
number_chosen3_8.current(0)   # 设置初始显示值，值为元组['values']的下标
number_chosen3_8.config(state='readonly')  # 设为只读模式
createToolTip(number_chosen3_8, '这里展示备案报告上下半年时间，您需要自行选择')   # Add Tooltip

label4_8 = ttk.Label(mighty8_1, text="落款：")
label4_8.place(x=300,y=10)
# number4_8 = StringVar()

label5_8 = ttk.Label(mighty8_1, text="年")
label5_8.place(x=393,y=11)
number5_8 = StringVar()
number_chosen5_8 = ttk.Combobox(mighty8_1, width=4, textvariable=number5_8)
number_chosen5_8['values'] = tuple(2015+i for i in range(20))
number_chosen5_8.place(x=340,y=10)
number_chosen5_8.current(now.year-2015)   # 设置初始显示值，值为元组['values']的下标
number_chosen5_8.config(state='readonly')  # 设为只读模式
createToolTip(number_chosen5_8, '这里展示备案报告落款时间，您需要自行选择')   # Add Tooltip

label6_8 = ttk.Label(mighty8_1, text="月")
label6_8.place(x=447,y=11)
number6_8 = StringVar()
number_chosen6_8 = ttk.Combobox(mighty8_1, width=2, textvariable=number6_8)
number_chosen6_8['values'] = tuple(1+i for i in range(12))
number_chosen6_8.place(x=410,y=10)
number_chosen6_8.current(now.month-1)   # 设置初始显示值，值为元组['values']的下标
number_chosen6_8.config(state='readonly')  # 设为只读模式
createToolTip(number_chosen6_8, '这里展示备案报告落款时间，您需要自行选择')   # Add Tooltip

label7_8 = ttk.Label(mighty8_1, text="日")
label7_8.place(x=513,y=11)
number7_8 = StringVar()
number_chosen7_8 = ttk.Combobox(mighty8_1, width=2, textvariable=number7_8)
number_chosen7_8['values'] = tuple(1+i for i in range(31))
number_chosen7_8.place(x=475,y=10)
number_chosen7_8.current(now.day-1)   # 设置初始显示值，值为元组['values']的下标
number_chosen7_8.config(state='readonly')  # 设为只读模式
createToolTip(number_chosen7_8, '这里展示备案报告落款时间，您需要自行选择')   # Add Tooltip

button1_8 = ttk.Button(mighty8_1, text="模板修改", command = beian_model_alter) #
button1_8.place(x=530,y=10)

# 第二行标签
label8_8 = ttk.Label(mighty8_1, text="各支部请示送审表文件夹：")
label8_8.place(x=0,y=40)
pathin_8 = StringVar()  # 定义变量
entry_pathin_8 = ttk.Entry(mighty8_1, textvariable=pathin_8)  # 输入框    # entry不能和grid连写，否则会报错
entry_pathin_8.place(x=145,y=40,width=365)
createToolTip(entry_pathin_8, '这里不需要输入')   # Add Tooltip
button2_8 = ttk.Button(mighty8_1, text="选择文件夹", command = lambda : select_files(scr_8, pathin_8))
button2_8.place(x=530,y=40)

# 第三行标签
label9_8 = ttk.Label(mighty8_1, text="党委会召开时间：")
label9_8.place(x=0,y=71)

label10_8 = ttk.Label(mighty8_1, text="年")
label10_8.place(x=153,y=71)
number10_8 = StringVar()
number10_8 = StringVar()
number_chosen10_8 = ttk.Combobox(mighty8_1, width=4, textvariable=number10_8)
number_chosen10_8['values'] = tuple(2015+i for i in range(20))
number_chosen10_8.place(x=100,y=70)
number_chosen10_8.current(now.year-2015)   # 设置初始显示值，值为元组['values']的下标
number_chosen10_8.config(state='readonly')  # 设为只读模式
createToolTip(number_chosen10_8, '这里展示备案报告的党委会召开时间（即批复时间），您需要确认并自行选择')   # Add Tooltip

labe11_8 = ttk.Label(mighty8_1, text="月")
labe11_8.place(x=207,y=71)
number11_8 = StringVar()
number_chosen11_8 = ttk.Combobox(mighty8_1, width=2, textvariable=number11_8)
number_chosen11_8['values'] = tuple(1+i for i in range(12))
number_chosen11_8.place(x=170,y=70)
number_chosen11_8.current(now.month-1)   # 设置初始显示值，值为元组['values']的下标
number_chosen11_8.config(state='readonly')  # 设为只读模式
createToolTip(number_chosen11_8, '这里展示备案报告的党委会召开时间，您需要确认并自行选择')   # Add Tooltip

label12_8 = ttk.Label(mighty8_1, text="日")
label12_8.place(x=263,y=71)
number12_8 = StringVar()
number_chosen12_8 = ttk.Combobox(mighty8_1, width=2, textvariable=number12_8)
number_chosen12_8['values'] = tuple(1+i for i in range(31))
number_chosen12_8.place(x=225,y=70)
number_chosen12_8.current(now.day-1)   # 设置初始显示值，值为元组['values']的下标
number_chosen12_8.config(state='readonly')  # 设为只读模式
createToolTip(number_chosen12_8, '这里展示备案报告的党委会召开时间，您需要自行选择')   # Add Tooltip

labe13_8 = ttk.Label(mighty8_1, text="-->确认")
labe13_8.place(x=275,y=71)
number13_8 = StringVar()
number_chosen13_8 = ttk.Combobox(mighty8_1, width=7, textvariable=number13_8)
number_chosen13_8['values'] = ['张三李四']
number_chosen13_8.place(x=325,y=70)
number_chosen13_8.current(0)   # 设置初始显示值，值为元组['values']的下标
# number_chosen11_8.config(state='readonly')  # 设为只读模式
createToolTip(number_chosen13_8, '这里展示自动检测后的首名同志，您也可以自行选择')   # Add Tooltip

labe14_8 = ttk.Label(mighty8_1, text="等")
labe14_8.place(x=405,y=71)
number14_8 = StringVar()
number_chosen14_8 = ttk.Combobox(mighty8_1, width=4, textvariable=number14_8)
number_chosen14_8['values'] = tuple(1+i for i in range(200))
number_chosen14_8.place(x=425,y=70)
number_chosen14_8.current(99)   # 设置初始显示值，值为元组['values']的下标
# number_chosen12_8.config(state='readonly')  # 设为只读模式
createToolTip(number_chosen14_8, '这里展示自动检测后的同志数量，您也可以自行选择')   # Add Tooltip

labe15_8 = ttk.Label(mighty8_1, text="名同志")
labe15_8.place(x=480,y=71)

# 第四行
labe16_8 = ttk.Label(mighty8_1, text="具体名单：（人名间用空格隔开）")
labe16_8.place(x=5,y=100)
# 表格文本框
scr_sheet8 = scrolledtext.ScrolledText(mighty8_1, width=72, height=5, wrap=WORD)
scr_sheet8.place(x=5,y=125)


button3_8 = ttk.Button(mighty8_1, text="自动识别", command = auto_beian_read )  #
button3_8.place(x=530,y=70, height=50)

button4_8 = ttk.Button(mighty8_1, text="生成",
    command =lambda : write_beian(cookie = str(number_8_1.get()) + str(number_8_2.get()),
                                        yeardu = number1_8.get(),pici = number2_8.get(),year_up = number3_8.get(),
                                        year = number5_8.get(),month = number6_8.get(),day = number7_8.get(),
                                        dw_year = number10_8.get(), dw_month = number11_8.get(), dw_day = number12_8.get(),
                                        first_people = number13_8.get(),people_num = int(number14_8.get()),
                                  people_sheet=(sorted(scr_sheet8.get(1.0, 'end').split(),
                                                       key=lambda keys: [pinyin(i, style=Style.TONE3) for i in keys])
                                                if peoplename.get() == 1 else scr_sheet8.get(1.0, 'end').split())
                       ))  # 获取文本框第一行到全部的内容
button4_8.place(x=530,y=125, height=70)

# 文本框
label18_8 = ttk.Label(tab8, text="输出窗口：")
label18_8.place(x=7, y=255)
scr_8 = scrolledtext.ScrolledText(tab8, width=88, height=27, wrap=WORD)
scr_8.place(x=7,y=275)
scr_8.config(state=DISABLED)  # 关闭可写入模式






'''# 11 支部管理的控件 ############################################################'''
number1_11 = StringVar()
def label_zhibu_guanli(tab):
    global number1_11,number_chosen1_11
    # 第一行标签
    label1_11 = ttk.Label(tab, text="支部全称：")
    label1_11.place(x=15, y=5)
    number_chosen1_11 = ttk.Combobox(tab, textvariable=number1_11)
    number_chosen1_11['values'] = zhibu_allname
    number_chosen1_11.place(x=75, y=5, width=563)
    # number_chosen2_11.config(state='readonly')  # 设为只读模式
    createToolTip(number_chosen1_11, '这里不需要输入')  # Add Tooltip




scr1_11 = scrolledtext.ScrolledText(tab11_a, width=88, height=24, wrap=WORD)
scr1_11.place(x=7, y=290)
scr1_11.config(state=DISABLED)  # 关闭可写入模式    scr_11.place(x=7, y=280)
scr1_11.config(state=DISABLED)  # 关闭可写入模式
scr2_11 = scrolledtext.ScrolledText(tab11_b, width=88, height=25, wrap=WORD)
scr2_11.place(x=7, y=280)
scr2_11.config(state=DISABLED)  # 关闭可写入模式    scr_11.place(x=7, y=280)
scr2_11.config(state=DISABLED)  # 关闭可写入模式
scr3_11 = scrolledtext.ScrolledText(tab11_c, width=88, height=25, wrap=WORD)
scr3_11.place(x=7, y=280)
scr3_11.config(state=DISABLED)  # 关闭可写入模式    scr_11.place(x=7, y=280)
scr3_11.config(state=DISABLED)  # 关闭可写入模式

label_zhibu_guanli(tab11_a)
label_zhibu_guanli(tab11_b)
label_zhibu_guanli(tab11_c)

# 文本框
label45_11 = ttk.Label(tab11_a, text="输出窗口：")
label45_11.place(x=7, y=265)
# label45_11.place_forget()  # 默认不勾选，即不隐藏
# 文本框
label46_11 = ttk.Label(tab11_b, text="输出窗口：")
label46_11.place(x=7, y=250)
# label46_11.place_forget()  # 默认不勾选，即不隐藏
# 文本框
label47_11 = ttk.Label(tab11_c, text="输出窗口：")
label47_11.place(x=7, y=250)
# label47_11.place_forget()  # 默认不勾选，即不隐藏

number_chosen1_11.current(0)  # 设置初始显示值，值为元组['values']的下标

# 第一块 布局Frame
mighty1_11 = ttk.LabelFrame(tab11_a, text='各支部请示')
mighty1_11.place(x=10,y=40,width=630,height=210)

# 第一行标签 左边标签
label29_11 = ttk.Label(tab11_a, text="年度")
label29_11.place(x=45, y=38)
createToolTip(label29_11, '（发展对象请示原句：鉴于以上表现，经支委会讨论研究，确认{}等{}人为{}年{}党员发展对象人选，）')
number18_11 = StringVar()
number_chosen18_11 = ttk.Combobox(tab11_a, width=5, textvariable=number18_11)
number_chosen18_11['values'] = tuple(2015 + i for i in range(20))
number_chosen18_11.place(x=75, y=37)
number_chosen18_11.current(now.year - 2015)  # 设置初始显示值，值为元组['values']的下标
number_chosen18_11.config(state='readonly')  # 设为只读模式
createToolTip(number_chosen18_11, '这里展示支部发展对象请示需要的年度时间，您需要自行选择')  # Add Tooltip
label30_11 = ttk.Label(tab11_a, text="上下半年")
label30_11.place(x=140, y=38)
createToolTip(label30_11, '（发展对象请示原句：鉴于以上表现，经支委会讨论研究，确认{}等{}人为{}年{}党员发展对象人选，）')
number19_11 = StringVar()
number_chosen19_11 = ttk.Combobox(tab11_a, width=5, textvariable=number19_11)
number_chosen19_11['values'] = ('上', '下')
number_chosen19_11.place(x=195, y=37)
number_chosen19_11.current(0)  # 设置初始显示值，值为元组['values']的下标
number_chosen19_11.config(state='readonly')  # 设为只读模式
createToolTip(number_chosen19_11, '这里展示支部发展对象请示需要的上下半年时间，您需要自行选择')  # Add Tooltip

# 第一行右边标签
label31_11 = ttk.Label(tab11_a, text="总党员数")
label31_11.place(x=20, y=38)
createToolTip(label31_11, '（党员转正请示原句：本支部共有党员{}名，其中正式党员{}名，预备党员{}名。到会党员{}名，其中正式党员{}名，）')
number20_11 = StringVar()
number_chosen20_11 = ttk.Combobox(tab11_a, width=5, textvariable=number20_11)
number_chosen20_11['values'] = tuple(1+i for i in range(100))
number_chosen20_11.place(x=75, y=37)
number_chosen20_11.current(11)   # 设置初始显示值，值为元组['values']的下标
# number_chosen20_11.config(state='readonly')  # 设为只读模式
createToolTip(number_chosen20_11, '这里展示支部党员转正请示需要的总党员数量（包括支部正式党员和预备党员），您需要自行选择')   # Add Tooltip

label32_11 = ttk.Label(tab11_a, text="正式党员")
label32_11.place(x=140, y=38)
createToolTip(label32_11, '（党员转正请示原句：本支部共有党员{}名，其中正式党员{}名，预备党员{}名。到会党员{}名，其中正式党员{}名，）')
number21_11 = StringVar()
number_chosen21_11 = ttk.Combobox(tab11_a, width=5, textvariable=number21_11)
number_chosen21_11['values'] = tuple(1+i for i in range(50))
number_chosen21_11.place(x=195, y=37)
number_chosen21_11.current(5)   # 设置初始显示值，值为元组['values']的下标
# number_chosen21_11.config(state='readonly')  # 设为只读模式
createToolTip(number_chosen21_11, '这里展示支部党员转正请示需要的支部正式党员数量，您需要自行选择')   # Add Tooltip

label33_11 = ttk.Label(tab11_a, text="预备党员")
createToolTip(label33_11, '（党员转正请示原句：本支部共有党员{}名，其中正式党员{}名，预备党员{}名。到会党员{}名，其中正式党员{}名，）')
label33_11.place(x=260, y=38)
number22_11 = StringVar()
number_chosen22_11 = ttk.Combobox(tab11_a, width=5, textvariable=number22_11)
number_chosen22_11['values'] = tuple(1+i for i in range(50))
number_chosen22_11.place(x=315, y=37)
number_chosen22_11.current(5)   # 设置初始显示值，值为元组['values']的下标
# number_chosen22_11.config(state='readonly')  # 设为只读模式
createToolTip(number_chosen22_11, '这里展示支部党员转正请示需要的支部预备党员数量，您需要自行选择')   # Add Tooltip

label34_11 = ttk.Label(tab11_a, text="到会党员")
createToolTip(label34_11, '（党员转正请示原句：本支部共有党员{}名，其中正式党员{}名，预备党员{}名。到会党员{}名，其中正式党员{}名，）')
label34_11.place(x=380, y=38)
number23_11 = StringVar()
number_chosen23_11 = ttk.Combobox(tab11_a, width=5, textvariable=number23_11)
number_chosen23_11['values'] = tuple(1+i for i in range(50))
number_chosen23_11.place(x=435, y=37)
number_chosen23_11.current(11)   # 设置初始显示值，值为元组['values']的下标
# number_chosen23_11.config(state='readonly')  # 设为只读模式
createToolTip(number_chosen23_11, '这里展示支部党员转正请示需要的到会党员数量，您需要自行选择')   # Add Tooltip

label35_11 = ttk.Label(tab11_a, text="到会正式党员")
createToolTip(label35_11, '（党员转正请示原句：本支部共有党员{}名，其中正式党员{}名，预备党员{}名。到会党员{}名，其中正式党员{}名，）')
label35_11.place(x=500, y=38)
number24_11 = StringVar()
number_chosen24_11 = ttk.Combobox(tab11_a, width=5, textvariable=number24_11)
number_chosen24_11['values'] = tuple(1+i for i in range(50))
number_chosen24_11.place(x=580, y=37)
number_chosen24_11.current(5)   # 设置初始显示值，值为元组['values']的下标
# number_chosen24_11.config(state='readonly')  # 设为只读模式
createToolTip(number_chosen24_11, '这里展示支部党员转正请示需要的支部到会的正式党员数量，您需要自行选择')   # Add Tooltip

# 暂时隐藏控件
label29_11.place_forget()
number_chosen18_11.place_forget()
label30_11.place_forget()
number_chosen19_11.place_forget()

label31_11.place_forget()
label32_11.place_forget()
label33_11.place_forget()
label34_11.place_forget()
label35_11.place_forget()
number_chosen20_11.place_forget()
number_chosen21_11.place_forget()
number_chosen22_11.place_forget()
number_chosen23_11.place_forget()
number_chosen24_11.place_forget()

# 选择是哪一类型的请示
number_11_1 = IntVar()
check11_1 = Checkbutton(mighty1_11, text="发展对象请示", variable=number_11_1)
check11_1.state(['active'])
# number_11_1.set(1) # 默认不勾选
# active, disabled, focus, pressed, selected, background,readonly, alternate, invalid
check11_1.place(x=100,y=0)
number_11_2 = IntVar()
check11_2 = Checkbutton(mighty1_11, text="预备党员请示", variable=number_11_2)
check11_2.state(['active'])  # Clears (turns off) the checkbutton.
check11_2.place(x=250,y=0)
number_11_3 = IntVar()
check11_3 = Checkbutton(mighty1_11, text="党员转正请示", variable=number_11_3)
check11_3.state(['active'])  # Clears (turns off) the checkbutton.
check11_3.place(x=400,y=0)


# GUI Callback function
def checkCallback_13(*ignoredArgs):
    global mighty1_11,label29_11,number_chosen18_11,label30_11,number_chosen19_11
    # only enable one checkbutton
    if number_11_1.get():  # ==1
        check11_2.configure(state='disabled')
        check11_3.configure(state='disabled')

        # UI向下移动，删除底下的一个“输出窗口”标签，拉出上面一行变量标签
        label29_11.place(x=75, y=38)
        number_chosen18_11.place(x=105, y=37)
        label30_11.place(x=175, y=38)
        number_chosen19_11.place(x=230, y=37)
        # label46_11.place_forget() # label46_11.place(x=7, y=270)
        mighty1_11.place(x=10, y=60, width=630, height=210)
        # mighty2_11.place(x=10, y=270, width=630, height=210)
    else:
        check11_2.configure(state='normal')
        check11_3.configure(state='normal')

        # UI向上移动，浮出底下的一个“输出窗口”标签，删除上面一行变量标签
        label29_11.place_forget()
        number_chosen18_11.place_forget()
        label30_11.place_forget()
        number_chosen19_11.place_forget()
        # label46_11.place(x=7, y=460)
        mighty1_11.place(x=10, y=40, width=630, height=210)
        # mighty2_11.place(x=10, y=250, width=630, height=210)
def checkCallback_14(*ignoredArgs):
    if number_11_2.get():
        check11_1.configure(state='disabled')
        check11_3.configure(state='disabled')
    else:
        check11_1.configure(state='normal')
        check11_3.configure(state='normal')
def checkCallback_15(*ignoredArgs):
    global mighty1_11,label31_11,label32_11,label33_11,label34_11,label35_11,number_chosen20_11,number_chosen21_11,number_chosen22_11,number_chosen23_11,number_chosen24_11
    if number_11_3.get():
        check11_1.configure(state='disabled')
        check11_2.configure(state='disabled')
        print("xxxxx")
        label31_11.place(x=20, y=38)
        number_chosen20_11.place(x=75, y=37)
        label32_11.place(x=140, y=38)
        number_chosen21_11.place(x=195, y=37)
        label33_11.place(x=260, y=38)
        number_chosen22_11.place(x=315, y=37)
        label34_11.place(x=380, y=38)
        number_chosen23_11.place(x=435, y=37)
        label35_11.place(x=500, y=38)
        number_chosen24_11.place(x=580, y=37)
        # label46_11.place_forget()
        mighty1_11.place(x=10, y=60, width=630, height=210)
        # mighty2_11.place(x=10, y=270, width=630, height=210)
    else:
        check11_1.configure(state='normal')
        check11_2.configure(state='normal')

        label31_11.place_forget()
        label32_11.place_forget()
        label33_11.place_forget()
        label34_11.place_forget()
        label35_11.place_forget()
        number_chosen20_11.place_forget()
        number_chosen21_11.place_forget()
        number_chosen22_11.place_forget()
        number_chosen23_11.place_forget()
        number_chosen24_11.place_forget()
        # label46_11.place(x=7, y=460)
        mighty1_11.place(x=10, y=40, width=630, height=210)
        # mighty2_11.place(x=10, y=250, width=630, height=210)

# trace the state of the two checkbuttons  #？？
number_11_1.trace('w', lambda unused0, unused1, unused2: checkCallback_13())
number_11_2.trace('w', lambda unused0, unused1, unused2: checkCallback_14())
number_11_3.trace('w', lambda unused0, unused1, unused2: checkCallback_15())

button1_11 = ttk.Button(mighty1_11, text="模板修改", command=zhibu_qingshi_model_alter)
button1_11.place(x=530, y=0)
# 第二行标签
label2_11 = ttk.Label(mighty1_11, text="名单文件导入：")
label2_11.place(x=10, y=30)
pathin1_11 = StringVar()  # 定义变量
entry_pathin1_11 = ttk.Entry(mighty1_11, textvariable=pathin1_11)  # 输入框    # entry不能和grid连写，否则会报错
entry_pathin1_11.place(x=100, y=30, width=420)
createToolTip(entry_pathin1_11, '这里不需要输入，您可以选择导入请示一览表表格文件')  # Add Tooltip
button2_11 = ttk.Button(mighty1_11, text="选择文件", command=lambda: select_file(scr1_11, pathin1_11))
button2_11.place(x=530, y=30)

# 第三行
label3_11 = ttk.Label(mighty1_11, text="请示落款时间：")
label3_11.place(x=10,y=60)

label4_11 = ttk.Label(mighty1_11, text="年")
label4_11.place(x=153,y=61)
number2_11 = StringVar()
number_chosen2_11 = ttk.Combobox(mighty1_11, width=4, textvariable=number2_11)
number_chosen2_11['values'] = tuple(2015+i for i in range(20))
number_chosen2_11.place(x=100,y=60)
number_chosen2_11.current(now.year-2015)   # 设置初始显示值，值为元组['values']的下标
number_chosen2_11.config(state='readonly')  # 设为只读模式
createToolTip(number_chosen2_11, '这里展示支部请示落款时间，您需要自行选择')   # Add Tooltip

label5_11 = ttk.Label(mighty1_11, text="月")
label5_11.place(x=207,y=61)
number3_11 = StringVar()
number_chosen3_11 = ttk.Combobox(mighty1_11, width=2, textvariable=number3_11)
number_chosen3_11['values'] = tuple(1+i for i in range(12))
number_chosen3_11.place(x=170,y=60)
number_chosen3_11.current(now.month-1)   # 设置初始显示值，值为元组['values']的下标
number_chosen3_11.config(state='readonly')  # 设为只读模式
createToolTip(number_chosen3_11, '这里展示支部请示落款时间，您需要自行选择')   # Add Tooltip

label6_11 = ttk.Label(mighty1_11, text="日")
label6_11.place(x=263,y=61)
number4_11 = StringVar()
number_chosen4_11 = ttk.Combobox(mighty1_11, width=2, textvariable=number4_11)
number_chosen4_11['values'] = tuple(1+i for i in range(31))
number_chosen4_11.place(x=225,y=60)
number_chosen4_11.current(now.day-1)   # 设置初始显示值，值为元组['values']的下标
number_chosen4_11.config(state='readonly')  # 设为只读模式
createToolTip(number_chosen4_11, '这里展示支部请示落款时间，您需要自行选择')   # Add Tooltip

label7_11 = ttk.Label(mighty1_11, text="确认")
label7_11.place(x=285, y=61)
number5_11 = StringVar()
number_chosen5_11 = ttk.Combobox(mighty1_11, width=7, textvariable=number5_11)
number_chosen5_11['values'] = ['张三李四']
number_chosen5_11.place(x=325, y=60)
number_chosen5_11.current(0)  # 设置初始显示值，值为元组['values']的下标
# number_chosen11_8.config(state='readonly')  # 设为只读模式
createToolTip(number_chosen5_11, '这里展示自动检测后的首名同志，您也可以自行选择')  # Add Tooltip

label8_11 = ttk.Label(mighty1_11, text="等")
label8_11.place(x=405, y=61)
number6_11 = StringVar()
number_chosen6_11 = ttk.Combobox(mighty1_11, width=4, textvariable=number6_11)
number_chosen6_11['values'] = tuple(1 + i for i in range(200))
number_chosen6_11.place(x=425, y=60)
number_chosen6_11.current(99)  # 设置初始显示值，值为元组['values']的下标
# number_chosen12_8.config(state='readonly')  # 设为只读模式
createToolTip(number_chosen6_11, '这里展示自动检测后的同志数量，您也可以自行选择')  # Add Tooltip

label9_11 = ttk.Label(mighty1_11, text="名同志")
label9_11.place(x=480, y=61)

# 第四行
label10_11 = ttk.Label(mighty1_11, text="具体名单：（人名间用空格隔开）")
label10_11.place(x=5, y=90)

label11_11 = ttk.Label(mighty1_11, text="支部大会时间：")
label11_11.place(x=250, y=90)

label12_11 = ttk.Label(mighty1_11, text="年")
label12_11.place(x=393, y=90)
number7_11 = StringVar()
number_chosen7_11 = ttk.Combobox(mighty1_11, width=4, textvariable=number7_11)
number_chosen7_11['values'] = tuple(2015 + i for i in range(20))
number_chosen7_11.place(x=340, y=88)
number_chosen7_11.current(now.year - 2015)  # 设置初始显示值，值为元组['values']的下标
number_chosen7_11.config(state='readonly')  # 设为只读模式
createToolTip(number_chosen7_11, '这里展示支部大会时间，您需要自行选择')  # Add Tooltip

label13_11 = ttk.Label(mighty1_11, text="月")
label13_11.place(x=447, y=90)
number8_11 = StringVar()
number_chosen8_11 = ttk.Combobox(mighty1_11, width=2, textvariable=number8_11)
number_chosen8_11['values'] = tuple(1 + i for i in range(12))
number_chosen8_11.place(x=410, y=88)
number_chosen8_11.current(now.month - 1)  # 设置初始显示值，值为元组['values']的下标
number_chosen8_11.config(state='readonly')  # 设为只读模式
createToolTip(number_chosen8_11, '这里展示支部大会时间，您需要自行选择')  # Add Tooltip

label14_11 = ttk.Label(mighty1_11, text="日")
label14_11.place(x=503, y=90)
number9_11 = StringVar()
number_chosen9_11 = ttk.Combobox(mighty1_11, width=2, textvariable=number9_11)
number_chosen9_11['values'] = tuple(1 + i for i in range(31))
number_chosen9_11.place(x=465, y=88)
number_chosen9_11.current(now.day - 1)  # 设置初始显示值，值为元组['values']的下标
number_chosen9_11.config(state='readonly')  # 设为只读模式
createToolTip(number_chosen9_11, '这里展示支部大会时间，您需要自行选择')  # Add Tooltip

# 表格文本框
scr_sheet1_11 = scrolledtext.ScrolledText(mighty1_11, width=72, height=5, wrap=WORD)
scr_sheet1_11.place(x=5, y=115)

button3_11 = ttk.Button(mighty1_11, text="自动识别", command=auto_zhibu_qingshi_read)  #
button3_11.place(x=530, y=60, height=50)
button4_11 = ttk.Button(mighty1_11, text="生成",
                        command=lambda: write_zhibu_qingshi(party_name=number1_11.get(),
                        cookie=str(number_11_1.get())+str(number_11_2.get())+str(number_11_3.get()),
                        year=number2_11.get(),month=number3_11.get(),day=number4_11.get(),
                        zd_year=number7_11.get(),zd_month=number8_11.get(),zd_day=number9_11.get(),
                        first_people=number5_11.get(),people_num=int(number6_11.get()),
                        people_sheet =  (sorted(scr_sheet1_11.get(1.0,'end').split(),key=lambda keys:[pinyin(i, style=Style.TONE3) for i in keys])
                        if peoplename.get() == 1 else scr_sheet1_11.get(1.0,'end').split()),
                        yeardu=number18_11.get(),
                        year_up=number19_11.get(),dy_sum=number20_11.get(),dy_true=number21_11.get(),
                        dy_wait=number22_11.get(),dy_in=number23_11.get(),dy_true_in=number24_11.get()))
button4_11.place(x=530, y=115, height=70)

# 第二块 布局Frame
mighty2_11 = ttk.LabelFrame(tab11_b, text='各支部批复')
# mighty2_11.place(x=10,y=250,width=630,height=210)
mighty2_11.place(x=10,y=40,width=630,height=210)

# 第二块 布局Frame
# 选择是哪一类型的请示
number_11_4 = IntVar()
check11_4 = Checkbutton(mighty2_11, text="发展对象批复", variable=number_11_4)
check11_4.state(['active'])
# number_11_4.set(1) #  # 默认不勾选 0
check11_4.place(x=100,y=0)
number_11_5 = IntVar()
check11_5 = Checkbutton(mighty2_11, text="预备党员批复", variable=number_11_5)
check11_5.state(['active'])  # Clears (turns off) the checkbutton.
check11_5.place(x=250,y=0)
number_11_6 = IntVar()
check11_6 = Checkbutton(mighty2_11, text="党员转正批复", variable=number_11_6)
check11_6.state(['active'])  # Clears (turns off) the checkbutton.
check11_6.place(x=400,y=0)
# GUI Callback function
def checkCallback_16(*ignoredArgs):
    # only enable one checkbutton
    if number_11_4.get():  # ==1
        check11_5.configure(state='disabled')
        check11_6.configure(state='disabled')
    else:
        check11_5.configure(state='normal')
        check11_6.configure(state='normal')
def checkCallback_17(*ignoredArgs):
    if number_11_5.get():
        check11_4.configure(state='disabled')
        check11_6.configure(state='disabled')
    else:
        check11_4.configure(state='normal')
        check11_6.configure(state='normal')
def checkCallback_18(*ignoredArgs):
    if number_11_6.get():
        check11_4.configure(state='disabled')
        check11_5.configure(state='disabled')
    else:
        check11_4.configure(state='normal')
        check11_5.configure(state='normal')
# trace the state of the two checkbuttons  #？？
number_11_4.trace('w', lambda unused0, unused1, unused2: checkCallback_16())
number_11_5.trace('w', lambda unused0, unused1, unused2: checkCallback_17())
number_11_6.trace('w', lambda unused0, unused1, unused2: checkCallback_18())

button5_11 = ttk.Button(mighty2_11, text="模板修改",command=zhibu_pifu_model_alter) # ,
button5_11.place(x=530, y=0)
# 第二行标签
label15_6 = ttk.Label(mighty2_11, text="名单文件导入：")
label15_6.place(x=10, y=30)
pathin2_11 = StringVar()  # 定义变量
entry_pathin2_11 = ttk.Entry(mighty2_11, textvariable=pathin2_11)  # 输入框    # entry不能和grid连写，否则会报错
entry_pathin2_11.place(x=100, y=30, width=420)
createToolTip(entry_pathin2_11, '这里不需要输入，您可以选择导入请示一览表表格文件')  # Add Tooltip
button6_11 = ttk.Button(mighty2_11, text="选择文件", command=lambda: select_file(scr2_11, pathin2_11))
button6_11.place(x=530, y=30)

# 第三行
label16_11 = ttk.Label(mighty2_11, text="批复落款时间：")
label16_11.place(x=10,y=60)

label17_11 = ttk.Label(mighty2_11, text="年")
label17_11.place(x=153,y=61)
number10_11 = StringVar()
number_chosen10_11 = ttk.Combobox(mighty2_11, width=4, textvariable=number10_11)
number_chosen10_11['values'] = tuple(2015+i for i in range(20))
number_chosen10_11.place(x=100,y=60)
number_chosen10_11.current(now.year-2015)   # 设置初始显示值，值为元组['values']的下标
number_chosen10_11.config(state='readonly')  # 设为只读模式
createToolTip(number_chosen10_11, '这里展示批复落款时间（院党委会时间），您需要自行选择')   # Add Tooltip

label18_11 = ttk.Label(mighty2_11, text="月")
label18_11.place(x=207,y=61)
number11_11 = StringVar()
number_chosen11_11 = ttk.Combobox(mighty2_11, width=2, textvariable=number11_11)
number_chosen11_11['values'] = tuple(1+i for i in range(12))
number_chosen11_11.place(x=170,y=60)
number_chosen11_11.current(now.month-1)   # 设置初始显示值，值为元组['values']的下标
number_chosen11_11.config(state='readonly')  # 设为只读模式
createToolTip(number_chosen11_11, '这里展示批复落款时间（院党委会时间），您需要自行选择')   # Add Tooltip

label19_11 = ttk.Label(mighty2_11, text="日")
label19_11.place(x=263,y=61)
number12_11 = StringVar()
number_chosen12_11 = ttk.Combobox(mighty2_11, width=2, textvariable=number12_11)
number_chosen12_11['values'] = tuple(1+i for i in range(31))
number_chosen12_11.place(x=225,y=60)
number_chosen12_11.current(now.day-1)   # 设置初始显示值，值为元组['values']的下标
number_chosen12_11.config(state='readonly')  # 设为只读模式
createToolTip(number_chosen12_11, '这里展示批复落款时间（院党委会时间），您需要自行选择')   # Add Tooltip

label20_11 = ttk.Label(mighty2_11, text="确认")
label20_11.place(x=285, y=61)
number13_11 = StringVar()
number_chosen13_11 = ttk.Combobox(mighty2_11, width=7, textvariable=number13_11)
number_chosen13_11['values'] = ['张三李四']
number_chosen13_11.place(x=325, y=60)
number_chosen13_11.current(0)  # 设置初始显示值，值为元组['values']的下标
# number_chosen11_8.config(state='readonly')  # 设为只读模式
createToolTip(number_chosen13_11, '这里展示自动检测后的首名同志，您也可以自行选择')  # Add Tooltip

label21_11 = ttk.Label(mighty2_11, text="等")
label21_11.place(x=405, y=61)
number14_11 = StringVar()
number_chosen14_11 = ttk.Combobox(mighty2_11, width=4, textvariable=number14_11)
number_chosen14_11['values'] = tuple(1 + i for i in range(200))
number_chosen14_11.place(x=425, y=60)
number_chosen14_11.current(99)  # 设置初始显示值，值为元组['values']的下标
# number_chosen12_8.config(state='readonly')  # 设为只读模式
createToolTip(number_chosen14_11, '这里展示自动检测后的同志数量，您也可以自行选择')  # Add Tooltip

label22_11 = ttk.Label(mighty2_11, text="名同志")
label22_11.place(x=480, y=61)

# 第四行
label23_11 = ttk.Label(mighty2_11, text="具体名单：（人名间用空格隔开）")
label23_11.place(x=5, y=90)

label24_11 = ttk.Label(mighty2_11, text="支部请示时间：")
label24_11.place(x=255, y=90)

label25_11 = ttk.Label(mighty2_11, text="年")
label25_11.place(x=393, y=90)
number15_11 = StringVar()
number_chosen15_11 = ttk.Combobox(mighty2_11, width=4, textvariable=number15_11)
number_chosen15_11['values'] = tuple(2015 + i for i in range(20))
number_chosen15_11.place(x=340, y=88)
number_chosen15_11.current(now.year - 2015)  # 设置初始显示值，值为元组['values']的下标
number_chosen15_11.config(state='readonly')  # 设为只读模式
createToolTip(number_chosen15_11, '这里展示支部请示时间，您需要自行选择')  # Add Tooltip

label26_11 = ttk.Label(mighty2_11, text="月")
label26_11.place(x=447, y=90)
number16_11 = StringVar()
number_chosen16_11 = ttk.Combobox(mighty2_11, width=2, textvariable=number16_11)
number_chosen16_11['values'] = tuple(1 + i for i in range(12))
number_chosen16_11.place(x=410, y=88)
number_chosen16_11.current(now.month - 1)  # 设置初始显示值，值为元组['values']的下标
number_chosen16_11.config(state='readonly')  # 设为只读模式
createToolTip(number_chosen16_11, '这里展示支部请示时间，您需要自行选择')  # Add Tooltip

label27_11 = ttk.Label(mighty2_11, text="日")
label27_11.place(x=503, y=90)
number17_11 = StringVar()
number_chosen17_11 = ttk.Combobox(mighty2_11, width=2, textvariable=number17_11)
number_chosen17_11['values'] = tuple(1 + i for i in range(31))
number_chosen17_11.place(x=465, y=88)
number_chosen17_11.current(now.day - 1)  # 设置初始显示值，值为元组['values']的下标
number_chosen17_11.config(state='readonly')  # 设为只读模式
createToolTip(number_chosen17_11, '这里展示支部请示时间，您需要自行选择')  # Add Tooltip

# 表格文本框
scr_sheet2_11 = scrolledtext.ScrolledText(mighty2_11, width=72, height=5, wrap=WORD)
scr_sheet2_11.place(x=5, y=115)

button7_11 = ttk.Button(mighty2_11, text="自动识别", command=auto_zhibu_pifu_read)  #
button7_11.place(x=530, y=60, height=50)
button8_11 = ttk.Button(mighty2_11,
                        text="生成",
                        command=lambda: write_zhibu_pifu(party_name=number1_11.get(),
                        cookie=str(number_11_4.get())+str(number_11_5.get())+str(number_11_6.get()),
                        qs_year=number15_11.get(),qs_month=number16_11.get(),qs_day=number17_11.get(),
                        year=number10_11.get(), month=number11_11.get(), day=number12_11.get(),
                        first_people=number13_11.get(), people_num=int(number14_11.get()),
                                                         people_sheet=(sorted(scr_sheet2_11.get(1.0, 'end').split(),
       key=lambda keys:[pinyin(i, style=Style.TONE3) for i in keys]) if peoplename.get() == 1 else scr_sheet2_11.get(1.0, 'end').split()
                                                                       )
                                                         )
                        )

button8_11.place(x=530, y=115, height=70)




# 第三块 布局Frame  两推一备案
mighty3_11 = ttk.LabelFrame(tab11_c, text='积极分子备案报告')
mighty3_11.place(x=10,y=40,width=630,height=210)

# 第二行标签
label36_11 = ttk.Label(mighty3_11, text="培训名单导入：")
label36_11.place(x=10, y=30)
pathin3_11 = StringVar()  # 定义变量
entry_pathin3_11 = ttk.Entry(mighty3_11, textvariable=pathin3_11)  # 输入框    # entry不能和grid连写，否则会报错
entry_pathin3_11.place(x=100, y=30, width=420)
createToolTip(entry_pathin3_11, '这里不需要输入，您可以选择导入积极分子表格文件')  # Add Tooltip

# 第三行
label37_11 = ttk.Label(mighty3_11, text="备案落款时间：")
label37_11.place(x=10,y=60)

label38_11 = ttk.Label(mighty3_11, text="年")
label38_11.place(x=153,y=61)
number25_11 = StringVar()
number_chosen25_11 = ttk.Combobox(mighty3_11, width=4, textvariable=number25_11)
number_chosen25_11['values'] = tuple(2015+i for i in range(20))
number_chosen25_11.place(x=100,y=60)
number_chosen25_11.current(now.year-2015)   # 设置初始显示值，值为元组['values']的下标
number_chosen25_11.config(state='readonly')  # 设为只读模式
createToolTip(number_chosen25_11, '这里展示备案落款时间，您需要自行选择')   # Add Tooltip

label39_11 = ttk.Label(mighty3_11, text="月")
label39_11.place(x=207,y=61)
number26_11 = StringVar()
number_chosen26_11 = ttk.Combobox(mighty3_11, width=2, textvariable=number26_11)
number_chosen26_11['values'] = tuple(1+i for i in range(12))
number_chosen26_11.place(x=170,y=60)
number_chosen26_11.current(now.month-1)   # 设置初始显示值，值为元组['values']的下标
number_chosen26_11.config(state='readonly')  # 设为只读模式
createToolTip(number_chosen26_11, '这里展示备案落款时间，您需要自行选择')   # Add Tooltip

label40_11 = ttk.Label(mighty3_11, text="日")
label40_11.place(x=263,y=61)
number29_11 = StringVar()
number_chosen27_11 = ttk.Combobox(mighty3_11, width=2, textvariable=number29_11)
number_chosen27_11['values'] = tuple(1+i for i in range(31))
number_chosen27_11.place(x=225,y=60)
number_chosen27_11.current(now.day-1)   # 设置初始显示值，值为元组['values']的下标
number_chosen27_11.config(state='readonly')  # 设为只读模式
createToolTip(number_chosen27_11, '这里展示备案落款时间，您需要自行选择')   # Add Tooltip

label41_11 = ttk.Label(mighty3_11, text="确认")
label41_11.place(x=285, y=61)
number27_11 = StringVar()
number_chosen28_11 = ttk.Combobox(mighty3_11, width=7, textvariable=number27_11)
number_chosen28_11['values'] = ['张三李四']
number_chosen28_11.place(x=325, y=60)
number_chosen28_11.current(0)  # 设置初始显示值，值为元组['values']的下标
# number_chosen11_8.config(state='readonly')  # 设为只读模式
createToolTip(number_chosen13_11, '这里展示自动检测后的首名积极分子，您也可以自行选择')  # Add Tooltip

label42_11 = ttk.Label(mighty3_11, text="等")
label42_11.place(x=405, y=61)
number28_11 = StringVar()
number_chosen29_11 = ttk.Combobox(mighty3_11, width=4, textvariable=number28_11)
number_chosen29_11['values'] = tuple(1 + i for i in range(200))
number_chosen29_11.place(x=425, y=60)
number_chosen29_11.current(99)  # 设置初始显示值，值为元组['values']的下标
# number_chosen12_8.config(state='readonly')  # 设为只读模式
createToolTip(number_chosen29_11, '这里展示自动检测后的积极分子数量，您也可以自行选择')  # Add Tooltip

label43_11 = ttk.Label(mighty3_11, text="名同志")
label43_11.place(x=480, y=61)

# 第四行
label44_11 = ttk.Label(mighty3_11, text="具体名单：（人名间用空格隔开）")
label44_11.place(x=5, y=90)
createToolTip(label37_11, '这里展示自动检测后的生成的积极分子名单信息，这里不允许修改，您可以移步文件进行修改')  # Add Tooltip

# 表格文本框
scr_sheet3_11 = scrolledtext.ScrolledText(mighty3_11, width=72, height=5, wrap=WORD)
scr_sheet3_11.place(x=5, y=115)
createToolTip(scr_sheet3_11, '这里展示自动检测后的生成的积极分子名单信息，这里不允许修改，您可以移步文件进行修改')  # Add Tooltip


button9_11 = ttk.Button(mighty3_11, text="模板修改",command=zhibu_beian_model_alter) # ,
button9_11.place(x=530, y=0)

button10_11 = ttk.Button(mighty3_11, text="选择文件", command=lambda: select_file(scr3_11, pathin3_11))
button10_11.place(x=530, y=30)

button12_11 = ttk.Button(mighty3_11, text=" 自动识别\n生成信息表", command=auto_zhibu_beian_read)  #
button12_11.place(x=530, y=60, height=50)
button11_11 = ttk.Button(mighty3_11,
                        text="生成",
                        command=lambda: write_zhibu_beian(
                            party_name=number1_11.get(),
                            year=number25_11.get(),
                            month=number26_11.get(),
                            day=number29_11.get(),
                            first_people=number27_11.get(),
                            people_num=int(number28_11.get()),
                            new_workbook="积极分子备案报告名单.xlsx"
                                                         )
                        )

button11_11.place(x=530, y=115, height=70)
















'''# 12 智慧查重的控件 ############################################################'''


mighty1_12 = ttk.LabelFrame(tab12, text='数据库--比对查重')
mighty1_12.place(x=10,y=10,width=630,height=140)
mighty2_12 = ttk.LabelFrame(tab12, text='同批次--两两查重')
mighty2_12.place(x=10,y=150,width=630,height=60)
mighty3_12 = ttk.LabelFrame(tab12, text='网络联合查重')
mighty3_12.place(x=10,y=210,width=630,height=72)

# 文本框
label10_12 = ttk.Label(tab10, text="输出窗口：")
label10_12.place(x=7, y=265)
scr_12 = scrolledtext.ScrolledText(tab12, width=88, height=26, wrap=WORD)
scr_12.place(x=7,y=285)
scr_12.config(state=DISABLED)  # 关闭可写入模式

label5_12 = ttk.Label(mighty1_12, text="选择查重数据库")
label5_12.place(x=0,y=1)

number1_12 = StringVar()
number_chosen1_12 = ttk.Combobox(mighty1_12, width=8, textvariable=number1_12)
cailiao_list = ['全部材料']
# print(next(os.walk('.//mould//chachong//'))[1])
try:
    for f in next(os.walk('.//mould//chachong//'))[1]:  # 遍历文件夹下的一级文件夹
        cailiao_list.append(f)
except: pass
number_chosen1_12['values'] = cailiao_list # ['全部材料','心得体会','分组讨论','思想汇报','时政汇报']
number_chosen1_12.place(x=100,y=0,width=210)
number_chosen1_12.current(0)   # 设置初始显示值，值为元组['values']的下标
number_chosen1_12.config(state='readonly')  # 设为只读模式
createToolTip(number_chosen1_12, '这里可以选择对哪种块的材料进行查重，他将检测到./mould/chachong/下的数据库文件夹')   # Add Tooltip

number2_12 = StringVar()
number_chosen2_12 = ttk.Combobox(mighty1_12, width=8, textvariable=number2_12)
number_chosen2_12['values'] = ['全部阶段','积极分子','发展对象','预备党员','党员']
# number_chosen2_12.place(x=159,y=0)
number_chosen2_12.current(0)   # 设置初始显示值，值为元组['values']的下标
number_chosen2_12.config(state='readonly')  # 设为只读模式
createToolTip(number_chosen2_12, '这里可以更加精准选择对哪个阶段的材料进行查重，他将关系到什么阶段的数据库')   # Add Tooltip

number3_12 = StringVar()
number_chosen3_12 = ttk.Combobox(mighty1_12, width=8, textvariable=number3_12)
number_chosen3_12['values'] = ['全部支部', '第一支部','第二支部','第三支部','第四支部','第五支部','第六支部',
                               '第七支部', '第八支部', '第八支部', '研一支部', '研二支部', '研三支部']
# number_chosen3_12.place(x=240,y=0)
number_chosen3_12.current(0)   # 设置初始显示值，值为元组['values']的下标
number_chosen3_12.config(state='readonly')  # 设为只读模式
createToolTip(number_chosen3_12, '这里可以更加精准选择对哪种支部的材料进行查重，他将关系到什么支部的数据库')   # Add Tooltip

number4_12 = StringVar()
number_chosen4_12 = ttk.Combobox(mighty1_12, width=8, textvariable=number4_12)
number_chosen4_12['values'] = [ i*5 for i in range(2,20)]
number_chosen4_12.place(x=154,y=60)
number_chosen4_12.current(3)   # 设置初始显示值，值为元组['values']的下标
number_chosen4_12.config(state='readonly')  # 设为只读模式
createToolTip(number_chosen3_12, '这里可以更加精准选择判定规则，单段重复词')   # Add Tooltip

number5_12 = StringVar()
number_chosen5_12 = ttk.Combobox(mighty1_12, width=8, textvariable=number5_12)
number_chosen5_12['values'] = [ i/10 for i in range(1,10)]
number_chosen5_12.place(x=317,y=60)
number_chosen5_12.current(1)   # 设置初始显示值，值为元组['values']的下标
number_chosen5_12.config(state='readonly')  # 设为只读模式
createToolTip(number_chosen5_12, '这里可以更加精准选择判定规则，单段重复率')   # Add Tooltip

number_12_3 = IntVar()
check12_3 = Checkbutton(mighty1_12, text="选择文件-单个查重", variable=number_12_3)
check12_3.state(['selected'])  # Clears (turns off) the checkbutton.
number_12_3.set(1)
check12_3.place(x=330,y=0)
createToolTip(check12_3, '这里选择一个文件，对单个文件进行查重')   # Add Tooltip

number_12_4 = IntVar()
check12_4 = Checkbutton(mighty1_12, text="选择文件夹-批量查重", variable=number_12_4)
check12_4.state(['disabled'])  # Clears (turns off) the checkbutton.
check12_4.place(x=470,y=0)
createToolTip(check12_4, '这里选择一个文件夹，对文件夹里面所有符合条件的文件进行查重')   # Add Tooltip

number_12_2 = IntVar()
check12_2 = Checkbutton(mighty1_12, text="自定义数据库", variable=number_12_2)
check12_2.state(['selected'])  # Clears (turns off) the checkbutton.
check12_2.place(x=0,y=90)
createToolTip(check12_2, '勾选自定义，载入自建的数据库')   # Add Tooltip


label11_12 = ttk.Label(mighty1_12, text="单文件路径：")
label11_12.place(x=0,y=30)
label12_12 = ttk.Label(mighty1_12, text="个性化设置：")
label12_12.place(x=0,y=60)
label13_12 = ttk.Label(mighty1_12, text="单段重复词：")
label13_12.place(x=77,y=60)
label14_12 = ttk.Label(mighty1_12, text="单段重复率：")
label14_12.place(x=239,y=60)
createToolTip(label14_12, '单段重复词,单段重复率 不会影响输出窗口的显示，只会影响完整的查重报告')   # Add Tooltip

# GUI Callback function
def checkCallback_21(*ignoredArgs):
    if number_12_3.get():
        label11_12 = ttk.Label(mighty1_12, text="单文件路径：")
        label11_12.place(x=0,y=30)
        check12_4.configure(state='disabled')
        button2_12.place(x=455, y=28, height=55, width=90)
        # button3_12.place(x=455, y=58,width=90)
        button5_12.place(x=550, y=30, height=55, width=70)
    else:
        check12_4.configure(state='normal')
        button2_12.place_forget()
        if number_12_4.get()==0:
            label11_12 = ttk.Label(mighty1_12, text="请选择类型！")
            label11_12.place(x=0, y=30)
            # button3_12.place_forget()
            button5_12.place(x=460, y=30, height=55, width=160)
    if number_12_4.get():
        label11_12 = ttk.Label(mighty1_12, text="文件夹路径：")
        label11_12.place(x=0, y=30)
        check12_3.configure(state='disabled')
        button1_12.place(x=455, y=28, height=55, width=90)
        # button3_12.place(x=455, y=58,width=90)
        button5_12.place(x=550, y=30, height=55, width=70)
    else:
        check12_3.configure(state='normal')
        button1_12.place_forget()
        if number_12_3.get()==0:
            label11_12 = ttk.Label(mighty1_12, text="请选择类型！")
            label11_12.place(x=0, y=30)
            # button3_12.place_forget()
            button5_12.place(x=460, y=30, height=55, width=160)

# trace the state of the two checkbuttons  #？？
number_12_3.trace('w', lambda unused0, unused1, unused2: checkCallback_21())
number_12_4.trace('w', lambda unused0, unused1, unused2: checkCallback_21())

pathin3_12 = StringVar()  # 定义变量
entry_pathin3_12 = ttk.Entry(mighty1_12, textvariable=pathin3_12)  # 输入框    # entry不能和grid连写，否则会报错
entry_pathin3_12.place(x=77,y=30,width=375)
createToolTip(entry_pathin3_12, '这里不需要输入，只需要选择单一或多个文件的地址,目前支持docx后缀格式的word')   # Add Tooltip

pathin4_12 = StringVar()  # 定义变量
entry_pathin4_12 = ttk.Entry(mighty1_12, textvariable=pathin4_12)  # 输入框    # entry不能和grid连写，否则会报错
entry_pathin4_12.place(x=100,y=90,width=350)
createToolTip(entry_pathin4_12, '这里可以选择需要输入数据库文件夹的位置')   # Add Tooltip


button1_12 = ttk.Button(mighty1_12, text="选择文件夹位置", command =  lambda : select_files(scr_12, pathin3_12))
button1_12.place(x=455, y=28, height=55, width=97)
button2_12 = ttk.Button(mighty1_12, text="选择文件位置", command =  lambda : select_file(scr_12, pathin3_12))
button2_12.place(x=455,y=28, height=55, width=97)
button3_12 = ttk.Button(mighty1_12, text="  选择数据库位置  ", command = lambda : select_files(scr_12, pathin4_12))
button3_12.place(x=455, y=90,width=170)
button1_12.place_forget()

button5_12 = ttk.Button(mighty1_12, text="开始查重",
                        command = lambda : cha_chong_shujuku(path=pathin3_12.get(),
                                                count_set = int(number4_12.get()), ratio_set = float(number5_12.get()),
                                                data = [number_12_2.get(), pathin4_12.get()],
                                                type=[number1_12.get(),number2_12.get(),number3_12.get()],
                                                sole_many=str(number_12_3.get())+ str(number_12_4.get())))
button5_12.place(x=555,y=26, height=55,width=70)

# 同批次数据库两两比对 frame
label3_12 = ttk.Label(mighty2_12, text="文件夹路径：")
label3_12.place(x=0,y=2)

pathin5_12 = StringVar()  # 定义变量
entry_pathin5_12 = ttk.Entry(mighty2_12, textvariable=pathin5_12)  # 输入框    # entry不能和grid连写，否则会报错
entry_pathin5_12.place(x=77,y=2,width=375)
createToolTip(entry_pathin5_12, '这里不需要输入，只需要选择文件夹的地址,目前支持docx后缀格式的word')   # Add Tooltip

button7_12 = ttk.Button(mighty2_12, text="选择文件夹位置", command =  lambda : select_files(scr_12, pathin5_12))
button7_12.place(x=455, y=0, height=30, width=97)

button6_12 = ttk.Button(mighty2_12, text="开始查重",
                        command = lambda : cha_chong_tongpici(dirpath=pathin5_12.get(),
                                                count_set = int(number4_12.get()), ratio_set = float(number5_12.get())))
button6_12.place(x=555,y=0, height=30,width=70)




'''# 10 通用功能的控件 ############################################################'''
# 布局Frame
mighty1_10 = ttk.LabelFrame(tab10, text='强大的汇总表格文件功能')
mighty1_10.place(x=10,y=10,width=630,height=140)
mighty2_10 = ttk.LabelFrame(mighty1_10, text='参数设置')
mighty2_10.place(x=5,y=24,width=620,height=60)
mighty3_10 = ttk.LabelFrame(tab10, text='强大的PDF与word互转功能')
mighty3_10.place(x=10,y=160,width=630,height=105)

label2_10 = ttk.Label(mighty1_10, text="需汇总的文件夹路径：")
label2_10.place(x=10,y=0)
label3_10 = ttk.Label(mighty1_10, text="生成汇总后的总表名：")
label3_10.place(x=10,y=90)

label5_10 = ttk.Label(mighty2_10, text="选择工作表")
label5_10.place(x=0,y=6)
label6_10 = ttk.Label(mighty2_10, text="标题行")
label6_10.place(x=155,y=6)
label7_10 = ttk.Label(mighty2_10, text="表头行")
label7_10.place(x=275,y=6)
label8_10 = ttk.Label(mighty2_10, text="例示行")
label8_10.place(x=400,y=6)

# 工作表
# 检测列填入选项卡（自动填入，如果未打勾默认不能填写）
number1_10 = StringVar()
number_chosen1_10 = ttk.Combobox(mighty2_10, width=6, textvariable=number1_10)
number_chosen1_10['values'] = tuple('Sheet' + str(i) for i in range(1,5))
number_chosen1_10.place(x=70,y=5)
number_chosen1_10.current(0)   # 设置初始显示值，值为元组['values']的下标
number_chosen1_10.config(state='readonly')  # 设为只读模式
createToolTip(number_chosen1_10, '这里可以选择合并文件里面的哪一个工作表，默认第一个工作表，您也可以自行选择')   # Add Tooltip

# 标题行
# 检测列填入选项卡（自动填入，如果未打勾默认不能填写）
number2_10 = StringVar()
number_chosen2_10 = ttk.Combobox(mighty2_10, width=6, textvariable=number2_10)
number_chosen2_10['values'] = ('无标题','前1行','前2行','前3行')
number_chosen2_10.place(x=200,y=5)
number_chosen2_10.current(0)   # 设置初始显示值，值为元组['values']的下标
number_chosen2_10.config(state='readonly')  # 设为只读模式
createToolTip(number_chosen2_10, '展示自动检测后的的文件标题行数量，默认无标题行，您也可以自行选择')   # Add Tooltip

# 表头行
# 检测列填入选项卡（自动填入，如果未打勾默认不能填写）
number3_10 = StringVar()
number_chosen3_10 = ttk.Combobox(mighty2_10, width=6, textvariable=number3_10)
number_chosen3_10['values'] = ('无表头','第1行','第2行','第3行','第4行')
number_chosen3_10.place(x=320,y=5)
number_chosen3_10.current(1)   # 设置初始显示值，值为元组['values']的下标
number_chosen3_10.config(state='readonly')  # 设为只读模式
createToolTip(number_chosen3_10, '这里展示自动检测后的的文件表头在哪一行，默认第一行，您也可以自行选择')   # Add Tooltip

# 例示行
# 检测列填入选项卡（自动填入，如果未打勾默认不能填写）
number4_10 = StringVar()
number_chosen4_10 = ttk.Combobox(mighty2_10, width=5, textvariable=number4_10)
number_chosen4_10['values'] = ('有','没有')
number_chosen4_10.place(x=445,y=5)
number_chosen4_10.current(1)   # 设置初始显示值，值为元组['values']的下标
number_chosen4_10.config(state='readonly')  # 设为只读模式
createToolTip(number_chosen4_10, '这里展示自动检测后的的文件是否有例示行，默认没有例示行，您也可以自行选择')   # Add Tooltip


pathin_10 = StringVar()  # 定义变量
entry_pathin_10 = ttk.Entry(mighty1_10, textvariable=pathin_10)  # 输入框    # entry不能和grid连写，否则会报错
entry_pathin_10.place(x=130,y=0,width=380)
createToolTip(entry_pathin_10, '这里不需要输入')   # Add Tooltip

pathin2_10 = StringVar()  # 定义变量
entry_pathin_10 = ttk.Entry(mighty1_10, textvariable=pathin2_10)  # 输入框    # entry不能和grid连写，否则会报错
entry_pathin_10.place(x=130,y=90,width=380)
createToolTip(entry_pathin_10, '这里可以选择需要输入汇总后的总表文件名称')   # Add Tooltip
pathin2_10.set('汇总表 ' +  datetime.datetime.now().strftime("%Y_%m_%d"))


# 文本框
label10_10 = ttk.Label(tab10, text="输出窗口：")
label10_10.place(x=7, y=265)
scr_10 = scrolledtext.ScrolledText(tab10, width=88, height=26, wrap=WORD)
scr_10.place(x=7,y=285)
scr_10.config(state=DISABLED)  # 关闭可写入模式

# 按钮
button1_10 = ttk.Button(mighty2_10, text="自动检测", command = auto_general_merge_book)
button1_10.place(x=525,y=3)

button2_10 = ttk.Button(mighty1_10, text="选择文件夹", command = lambda : select_files(scr_10, pathin_10))
button2_10.place(x=530,y=0)

button3_10 = ttk.Button(mighty1_10, text="开始汇总",
                        command = lambda : general_merge_book(path=pathin_10.get(), filename=pathin2_10.get(),
                                                            sheet_what = number_chosen1_10.get(),
                                                            biaoti_row = number_chosen2_10.get(),
                                                            biaotou_row = number_chosen3_10.get(),
                                                            lishi_row = number_chosen4_10.get()
                                                            )
                        )
button3_10.place(x=530,y=90)

# 第三块frame（pdf和word互转功能）
# 选择是哪一类型
number_10_1 = IntVar()
check10_1 = Checkbutton(mighty3_10, text="pdf --> word", variable=number_10_1)
check10_1.state(['active'])  # Clears (turns off) the checkbutton.
check10_1.place(x=30,y=0)
createToolTip(check10_1, '这里选择PDF文件转换成word文件，即.pdf-->.docx')   # Add Tooltip
number_10_2 = IntVar()
check10_2 = Checkbutton(mighty3_10, text="word --> pdf", variable=number_10_2)
check10_2.state(['active'])  # Clears (turns off) the checkbutton.
check10_2.place(x=150,y=0)
createToolTip(check10_2, '这里选择word文件转换成PDF文件，即.docx或.doc-->.pdf')   # Add Tooltip
# GUI Callback function
def checkCallback_19(*ignoredArgs):
    if number_10_1.get():
        check10_2.configure(state='disabled')
    else:
        check10_2.configure(state='normal')
    if number_10_2.get():
        check10_1.configure(state='disabled')
    else:
        check10_1.configure(state='normal')
# trace the state of the two checkbuttons  #？？
number_10_1.trace('w', lambda unused0, unused1, unused2: checkCallback_19())
number_10_2.trace('w', lambda unused0, unused1, unused2: checkCallback_19())

number_10_3 = IntVar()
check10_3 = Checkbutton(mighty3_10, text="选择文件-单个转换", variable=number_10_3)
check10_3.state(['selected'])  # Clears (turns off) the checkbutton.
number_10_3.set(1)
check10_3.place(x=300,y=0)
createToolTip(check10_3, '这里选择一个文件，对单个文件进行转换')   # Add Tooltip
number_10_4 = IntVar()
check10_4 = Checkbutton(mighty3_10, text="选择文件夹-批量转换", variable=number_10_4)
check10_4.state(['disabled'])  # Clears (turns off) the checkbutton.
check10_4.place(x=450,y=0)
createToolTip(check10_4, '这里选择一个文件夹，对文件夹里面所有符合条件的文件进行转换')   # Add Tooltip

label11_10 = ttk.Label(mighty3_10, text="单文件路径：")
label11_10.place(x=0,y=30)
label12_10 = ttk.Label(mighty3_10, text="生成文件名：")
label12_10.place(x=0,y=60)

# GUI Callback function
def checkCallback_20(*ignoredArgs):
    if number_10_3.get():
        label11_10 = ttk.Label(mighty3_10, text="单文件路径：")
        label11_10.place(x=0,y=30)
        check10_4.configure(state='disabled')
        button2_10.place(x=455, y=28,width=90)
        button3_10.place(x=455, y=58,width=90)
        button5_10.place(x=550, y=30, height=55, width=70)
    else:
        check10_4.configure(state='normal')
        button2_10.place_forget()
        if number_10_4.get()==0:
            label11_10 = ttk.Label(mighty3_10, text="请选择类型！")
            label11_10.place(x=0, y=30)
            button3_10.place_forget()
            button5_10.place(x=460, y=30, height=55, width=160)
    if number_10_4.get():
        label11_10 = ttk.Label(mighty3_10, text="文件夹路径：")
        label11_10.place(x=0, y=30)
        check10_3.configure(state='disabled')
        button1_10.place(x=455, y=28,width=90)
        button3_10.place(x=455, y=58,width=90)
        button5_10.place(x=550, y=30, height=55, width=70)
    else:
        check10_3.configure(state='normal')
        button1_10.place_forget()
        if number_10_3.get()==0:
            label11_10 = ttk.Label(mighty3_10, text="请选择类型！")
            label11_10.place(x=0, y=30)
            button3_10.place_forget()
            button5_10.place(x=460, y=30, height=55, width=160)

# trace the state of the two checkbuttons  #？？
number_10_3.trace('w', lambda unused0, unused1, unused2: checkCallback_20())
number_10_4.trace('w', lambda unused0, unused1, unused2: checkCallback_20())

pathin3_10 = StringVar()  # 定义变量
entry_pathin3_10 = ttk.Entry(mighty3_10, textvariable=pathin3_10)  # 输入框    # entry不能和grid连写，否则会报错
entry_pathin3_10.place(x=77,y=30,width=375)
createToolTip(entry_pathin3_10, '这里不需要输入，只需要选择单一或多个文件的地址')   # Add Tooltip

pathin4_10 = StringVar()  # 定义变量
entry_pathin4_10 = ttk.Entry(mighty3_10, textvariable=pathin4_10)  # 输入框    # entry不能和grid连写，否则会报错
entry_pathin4_10.place(x=77,y=60,width=375)
createToolTip(entry_pathin4_10, '这里可以选择需要输入汇总后的总表文件名称')   # Add Tooltip

button1_10 = ttk.Button(mighty3_10, text="选择文件夹位置", command =  lambda : select_files(scr_10, pathin3_10))
button1_10.place(x=455, y=28,width=90)
button2_10 = ttk.Button(mighty3_10, text="选择文件位置", command =  lambda : select_file(scr_10, pathin3_10))
button2_10.place(x=455,y=28,width=90)
button3_10 = ttk.Button(mighty3_10, text="  保存位置  ", command = lambda : select_files(scr_10, pathin4_10))
button3_10.place(x=455, y=58,width=90)
button1_10.place_forget()

button5_10 = ttk.Button(mighty3_10, text="开始转换",
                        command = lambda : pdf2word_or_word2pdf(path=pathin3_10.get(),out_path=pathin4_10.get(),
                        what_to_what=str(number_10_1.get())+str(number_10_2.get()),sole=str(number_10_3.get())+str(number_10_4.get())))
button5_10.place(x=550,y=32, height=55,width=70)













'''# 9 帮助的控件 ############################################################'''

# 布局Frame
mighty1_9 = ttk.LabelFrame(tab9, text='开发者信息')
mighty1_9.place(x=10,y=0,width=630,height=130)

# 标签
label1_9 = ttk.Label(mighty1_9, text="当前版本： {} ".format(banben))
label1_9.place(x=230,y=0)
label2_9 = ttk.Label(mighty1_9, text="开发作者：  鸿武  ")
label2_9.place(x=230,y=27)
label3_9 = ttk.Label(mighty1_9, text="创作时间：2021年2月至今")
label3_9.place(x=230,y=54)
label3_9 = ttk.Label(mighty1_9, text="著作权信息：软著登字第9503465号")
label3_9.place(x=226,y=81)


label4_9 = ttk.Label(tab9, text="通用设置：")
label4_9.place(x=10,y=132)
peoplename = IntVar()
check11_7 = Checkbutton(tab9, text="请示、批复、备案等人名首字母排序", variable=peoplename)
check11_7.state(['active'])  # Clears (turns off) the checkbutton.
check11_7.place(x=10,y=152)
# active, disabled, focus, pressed, selected, background,readonly, alternate, invalid

check11_7.state(['selected'])  # 默认勾选
peoplename.set(1)
# print(peoplename.get())#没有勾选==0，勾选==1

# 按钮
button1_9 = ttk.Button(tab9, text="打开执行文件夹目录", command=open_abspath)
button1_9.place(x=10,y=210,width=630,height=100)

button1_9 = ttk.Button(tab9, text="  初始化案例文件  ", command=case_)
button1_9.place(x=10,y=315,width=315,height=100)
button1_9 = ttk.Button(tab9, text=" 检查新版本 \n密码：0911", command=version_)
button1_9.place(x=325, y=315, width=315, height=100)

button1_9 = ttk.Button(tab9, text="打开软件说明文档", command=about_)
button1_9.place(x=10,y=420,width=315,height=50)
button1_9 = ttk.Button(tab9, text="查看软件食用说明和近期更新内容", command=help_)
button1_9.place(x=10,y=470,width=315,height=50)

button1_9 = ttk.Button(tab9, text="查看视频教程", command=video_)
button1_9.place(x=325,y=420,width=315,height=100)

button1_9 = ttk.Button(tab9, text="合作开发 & BUG反馈", command=author_)
button1_9.place(x=10,y=525,width=315,height=100)
button1_9 = ttk.Button(tab9, text="致谢", command=thank_)
button1_9.place(x=325,y=525,width=315,height=100)


'''# tab13 小黑实验室 ############################################################'''
try:
    html_gif = PhotoImage(file="mould/data/me.gif")
    label = Label(tab13, image = html_gif)
    label.place(x=120,y=10)
except:pass
mighty1_13 = ttk.LabelFrame(tab13, text='小黑实验室')
mighty1_13.place(x=10,y=400,width=630,height=250)

xingxing_path = 'mould/data/xingxing'
if not os.path.exists(xingxing_path):
    print("不存在")
else:
    try:
        file = open(xingxing_path, mode='r', encoding='utf-8')
        linexx = file.readline()
        for a in range(len(linexx)):
            yushu = int(a%4)
            chushu = int(a//4)
            exec(f'xingxing_gif_{a} = PhotoImage(file="mould/data/xingxing.gif")')
            exec(f'Label(mighty1_13, image=xingxing_gif_{a}).place(x=(a // 8) * 45, y=(a % 8) * 27)')
            if a >16:
                break
    except:
        pass

label2_13 = Label(mighty1_13, text='扫描上面二维码关注小鸿武官方公众号！！！\n或者微信搜索：努力发芽的小鸿武\n记得加入本软件答疑交流QQ群:866367580\n不迷路，更早获取新版本！',compound=CENTER)
label2_13.pack()






# LICENCE_path = os.getcwd() + "\\mould\\data\\" + 'LICENCE'  # os.getcwd()返回当前文件所在的目录绝对路径
# # 若文件夹不存在则自动新建
if not os.path.exists(LICENCE_path):  # 判断是否存在文件
    # os.makedirs(folder_path) # 则创建为文件夹
    file = open(LICENCE_path, mode='w', encoding='utf-8')
    file.write("Flase")
    file.close()
    index_Toplevel = index_Toplevel()

else:
    file = open(LICENCE_path, encoding='utf-8')
    line1 = file.readline()
    line2 = file.readline()
    print("是谁：" + line1)
    print("权限：" + line2)
    file.close()
    if  line1 == str(platform.uname())+"\n" and line2 == 'True':
        tabControl.add(tab0_4, text=' 积极分子培训 ')  # Add the tab 二级菜单
        # window.mainloop()
    else:
        pass


# 显示窗口(消息循环)
window.mainloop()













'''
########################################################################################################################################################
备注：
考勤表姓名和学号列互换了，要采用最新的考勤表
学员册发放给支部填的，增加了一行推荐支部在O列，要采用最新的学员册模板
########################################################################################################################################################


考试座位表的，四个考场号
学员册B4单元格及其以下面单元格找不到学员名字

系数表格合并  给定系数，计算比例，生成名额行
学员册人员合并

时间检测
出生年月（日）
（首次）递交入党申请书时间 / 申请入党时间
确定为（确认为）积极分子时间
确定为（确认为）（列为）发展对象时间
发展为预备党员时间
预备党员转正时间

不符合  第n行 姓名 xxx xxxx 不符合



实时输出窗口
定义模块函数
每个文件都调用
一点信直接调用函数，实时写出来

'''





"""
tkinter.filedialog.asksaveasfilename():选择以什么文件名保存，返回文件名
tkinter.filedialog.asksaveasfile():选择以什么文件保存，创建文件并返回文件流对象
tkinter.filedialog.askopenfilename():选择打开什么文件，返回文件名
tkinter.filedialog.askopenfile():选择打开什么文件，返回IO流对象
tkinter.filedialog.askdirectory():选择目录，返回目录名
tkinter.filedialog.askopenfilenames():选择打开多个文件，以元组形式返回多个文件名
tkinter.filedialog.askopenfiles():选择打开多个文件，以列表形式返回多个IO流对象
"""



# frame = Frame(window,height=120, width=400, bd=1, bg="azure")
# frame.place(x=0, y=0)
# frame2 = Frame(window,height=120, width=400, bd=1, bg="light yellow1")
# frame2.place(x=0, y=120)

# # 菜单栏
# menubar = Menu(window)
#
# menubar.add_command(label="座位表", command=apass)
# menubar.add_command(label="考试表", command=apass)
# menubar.add_command(label="分组名单", command=apass)
# menubar.add_command(label="考勤表", command=apass)
#
# # 创建一个第一个菜单列表
# submenu1 = Menu(menubar, tearoff=0)
# submenu1.add_command(label='关于', command=apass)
# submenu1.add_command(label='帮助', command=apass)
# # 把菜单列表加入菜单栏
# menubar.add_cascade(label=' 菜单 ', menu=submenu1)  # 将menubar的menu属性指定为submenu1，即submenu1为menubar的下拉菜单
#
# window['menu'] = menubar  # == window.config(menu=menubar)  == window.configure(menu=menubar)
#
# # window.config(menu=menubar)
#
# style = ttk.Style()
# style.configure("EntryStyle.TEntry")
# style = ttk.Style()
# style.configure("white.TCheckbutton", foreground="white")
#



# 汇总表格功能，大规模删除空行，运行死机
# 姓名、专业班级、推荐支部，需要灵活定位，标题行和表头也需要灵活定位